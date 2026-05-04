from __future__ import annotations

import asyncio
import base64
import io
import json
from pathlib import Path

import pytest
from docx import Document
from fastmcp import Client
import mcp.types as mcp_types

from mcp_server import config as mcp_config, server as mcp_server_main
from mcp_server.bridge import InspyroBridge
from mcp_server.session_state import McpSessionState

EXPECTED_AUTHORING_TOOLS = {
    "get_system_info",
    "get_health",
    "list_component_profiles",
    "set_component_profile",
    "notebook_create",
    "notebook_load",
    "list_session_notebooks",
    "notebook_sync_cells",
    "notebook_save",
    "execute_cell",
    "execute_all_cells",
    "get_kernel_status",
    "get_run_status",
    "cancel_run",
    "resume_run",
    "list_cells",
    "get_cell",
    "find_in_notebook",
    "reset_kernel",
    "interrupt_kernel",
    "shutdown_kernel",
    "close_session_notebook",
    "get_variables",
    "check_document_quality",
    "run_document_workbench",
    "compare_document_versions",
    "manage_document_review",
    "prepare_document_delivery",
    "get_document_docx",
    "get_document_pdf",
    "export_clean_document_docx",
    "export_document_docx",
    "export_document_pdf",
    "reconvert_pdf",
    "upload_template",
    "get_template_info",
    "delete_template",
    "update_template_style",
    "convert_units",
    "get_units_catalog",
    "check_units_compatible",
}

EXPECTED_ANALYSIS_TOOLS = {
    "get_system_info",
    "get_health",
    "list_component_profiles",
    "set_component_profile",
    "analyze_dependencies",
    "analyze_impact",
    "run_sensitivity",
    "optimize_design",
    "compare_scenarios",
    "run_code_checks",
}

EXPECTED_FILES_TOOLS = {
    "get_system_info",
    "get_health",
    "list_component_profiles",
    "set_component_profile",
    "list_files",
    "read_file",
    "write_file",
    "create_file",
    "delete_file",
    "rename_file",
}


def _build_demo_docx_base64(title: str = "Black-box MCP DOCX") -> str:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Reporte generado por la suite black-box MCP.")
    buffer = io.BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


class FakeBlackBoxBridge:
    def __init__(self, root: Path) -> None:
        self.root = root.resolve()
        self.files: dict[str, object] = {}
        self.execution_count = 0
        self.activity_events: list[dict] = []
        self.mirror_events: list[dict] = []
        self.template_token = "template-123"
        self.docx_b64 = _build_demo_docx_base64()
        self.quality_summaries: dict[str, dict] = {}
        self.workbench_calls: list[dict] = []
        self.execution_observers: dict[str, list[asyncio.Queue]] = {}

    async def connect(self) -> None:
        return None

    async def disconnect(self) -> None:
        return None

    async def rest_get(self, path: str, params: dict | None = None) -> dict:
        params = params or {}
        if path == "/api/system/info":
            return {
                "service": "inspyro-backend",
                "workspace_root": str(self.root),
                "python": "3.12",
            }
        if path == "/health":
            return {"status": "healthy", "service": "inspyro-backend", "pdf_conversion_available": True}
        if path == "/metrics":
            return {"status": "ok", "ws_connections_active": 0}
        if path == "/pdf-status":
            return {
                "pdf_available": True,
                "word_available": False,
                "soffice_path": "C:/Program Files/LibreOffice/program/soffice.exe",
                "engine": "libreoffice",
            }
        if path == "/api/units/catalog":
            return {
                "count": 3,
                "units": [
                    {"category": "Force", "canonical": "kN", "display": "kN", "aliases": ["kN"]},
                    {"category": "Force", "canonical": "N", "display": "N", "aliases": ["N"]},
                    {"category": "Stress", "canonical": "MPa", "display": "MPa", "aliases": ["MPa"]},
                ],
            }
        if path == "/api/files/tree":
            requested = Path(str(params.get("path") or self.root)).resolve()
            children = []
            for stored_path in sorted(self.files):
                path_obj = Path(stored_path)
                if path_obj.parent == requested:
                    children.append({"name": path_obj.name, "path": stored_path, "type": "file"})
            return {"path": str(requested), "children": children}
        if path == "/api/files/read":
            file_path = str(Path(str(params["path"])).resolve())
            return {"content": self.files[file_path]}
        if path == "/api/docx/history":
            return {
                "items": [
                    {
                        "artifact_id": "artifact-1",
                        "source_path": params.get("source_path"),
                        "kernel_id": params.get("kernel_id") or "kernel-1",
                        "execution_id": "exec-docx",
                        "filename": "inspyro_document.docx",
                    }
                ]
            }
        if path == "/api/docx/quality":
            artifact_id = str(params.get("artifact_id") or "artifact-1")
            return self.quality_summaries[artifact_id]
        raise AssertionError(f"Unexpected REST GET path: {path}")

    async def rest_post(self, path: str, *, json_data: dict) -> dict:
        if path == "/api/mcp/activity/events":
            self.activity_events.append(json_data)
            return {"status": "accepted", "active_count": 0}
        if path == "/api/mcp/mirror-events":
            self.mirror_events.append(json_data)
            return {"status": "accepted"}
        if path == "/api/files/write":
            file_path = str(Path(str(json_data["path"])).resolve())
            self.files[file_path] = json_data["content"]
            return {"status": "ok", "path": file_path}
        if path == "/api/files/create":
            file_path = str(Path(str(json_data["path"])).joinpath(str(json_data["name"])).resolve())
            self.files[file_path] = "" if json_data.get("type") == "file" else {}
            return {"status": "ok", "path": file_path}
        if path == "/api/files/rename":
            old_path = str(Path(str(json_data["oldPath"])).resolve())
            new_path = str(Path(old_path).with_name(str(json_data["newName"])))
            self.files[new_path] = self.files.pop(old_path)
            return {"status": "ok", "oldPath": old_path, "newPath": new_path}
        if path == "/api/units/convert":
            factor = 1000.0 if json_data["from_unit"] == "kN" and json_data["to_unit"] == "N" else 1.0
            magnitude = json_data["magnitude"]
            converted = magnitude * factor if isinstance(magnitude, (int, float)) else magnitude
            return {
                "converted_magnitude": converted,
                "repr": f"{converted} {json_data['to_unit']}",
                "category": "Force",
                "dimension": "[force]",
            }
        if path == "/api/units/compatible":
            return {
                "dimension": "[force]",
                "compatible": [
                    {"canonical": "N", "display": "N", "aliases": ["N"]},
                    {"canonical": "kN", "display": "kN", "aliases": ["kN"]},
                ],
            }
        if path == "/api/docx/quality/run":
            artifact_id = str(json_data.get("artifact_id") or "artifact-1")
            summary = {
                "schema_version": 1,
                "artifact_id": artifact_id,
                "status": "warning",
                "score": 88,
                "counts": {"error": 0, "warning": 1, "info": 0},
                "sections": [
                    {
                        "id": "accessibility",
                        "status": "warning",
                        "findings": [
                            {
                                "section": "accessibility",
                                "severity": "warning",
                                "message": "Tabla sin primera fila marcada como encabezado",
                                "context": {"part": "word/document.xml"},
                            }
                        ],
                    }
                ],
                "findings": [
                    {
                        "section": "accessibility",
                        "severity": "warning",
                        "message": "Tabla sin primera fila marcada como encabezado",
                        "context": {"part": "word/document.xml"},
                    }
                ],
                "pages_rendered": None,
            }
            self.quality_summaries[artifact_id] = summary
            return summary
        if path == "/api/docx/workbench/run":
            self.workbench_calls.append(dict(json_data))
            operation = str(json_data.get("operation") or "audit")
            artifact_id = str(json_data.get("artifact_id") or "artifact-1")
            if operation == "render_manifest":
                return {
                    "status": "ok",
                    "operation": operation,
                    "artifact_id": artifact_id,
                    "workbench_id": "wb-render-manifest",
                    "visual": {
                        "status": "partial",
                        "page_count": 3,
                        "cached_pages": 1,
                        "converter_used": "word",
                        "page_resources": [
                            {
                                "name": "page-0001-z2_00.png",
                                "mime_type": "image/png",
                                "resource_uri": "/api/docx/render/resource?render_id=render-1&name=page-0001-z2_00.png",
                            }
                        ],
                    },
                    "resources": [
                        {
                            "name": "page-0001-z2_00.png",
                            "mime_type": "image/png",
                            "resource_uri": "/api/docx/render/resource?render_id=render-1&name=page-0001-z2_00.png",
                        }
                    ],
                }
            if operation == "render_all_pages":
                return {
                    "status": "ok",
                    "operation": operation,
                    "artifact_id": artifact_id,
                    "workbench_id": "wb-render-all",
                    "visual": {
                        "status": "complete",
                        "page_count": 3,
                        "cached_pages": 3,
                        "converter_used": "word",
                    },
                    "rendered_pages": [
                        {"page": 1, "resource_uri": "/api/docx/render/resource?render_id=render-1&name=page-0001-z2_00.png"}
                    ],
                    "resources": [
                        {
                            "name": "page-0001-z2_00.png",
                            "mime_type": "image/png",
                            "resource_uri": "/api/docx/render/resource?render_id=render-1&name=page-0001-z2_00.png",
                        }
                    ],
                }
            if operation == "prepare_delivery":
                return {
                    "status": "ok",
                    "operation": operation,
                    "artifact_id": artifact_id,
                    "workbench_id": "wb-delivery",
                    "summary": {
                        "artifact_id": artifact_id,
                        "status": "ok",
                        "score": 96,
                        "counts": {"error": 0, "warning": 0, "info": 1},
                        "sections": [],
                        "findings": [],
                    },
                    "variant": {
                        "parent_artifact_id": artifact_id,
                        "operation": "prepare_delivery",
                        "filename": "inspyro_document-delivery.docx",
                        "size_bytes": 1234,
                        "hash": "delivery-hash",
                        "resource_uri": "/api/docx/workbench/resource?workbench_id=wb-delivery&name=inspyro_document-delivery.docx",
                    },
                    "resources": [
                        {
                            "name": "inspyro_document-delivery.docx",
                            "size_bytes": 1234,
                            "hash": "delivery-hash",
                            "resource_uri": "/api/docx/workbench/resource?workbench_id=wb-delivery&name=inspyro_document-delivery.docx",
                        }
                    ],
                }
            return {
                "status": "ok",
                "operation": operation,
                "artifact_id": artifact_id,
                "workbench_id": f"wb-{operation}",
                "resources": [],
            }
        raise AssertionError(f"Unexpected REST POST path: {path}")

    async def rest_delete(self, path: str, *, params: dict) -> dict:
        if path == "/api/files/delete":
            file_path = str(Path(str(params["path"])).resolve())
            self.files.pop(file_path, None)
            return {"status": "ok", "path": file_path}
        raise AssertionError(f"Unexpected REST DELETE path: {path}")

    async def rest_post_files(self, path: str, *, files: dict) -> dict:
        if path == "/api/templates/upload":
            return {"template_token": self.template_token}
        raise AssertionError(f"Unexpected REST POST FILES path: {path}")

    async def rest_get_bytes(self, path: str, *, params: dict | None = None) -> bytes:
        if path == "/api/docx/download":
            return base64.b64decode(self.docx_b64)
        if path == "/api/pdf/download":
            return b"%PDF-1.4\n%fake\n"
        if path == "/api/docx/workbench/resource":
            return base64.b64decode(self.docx_b64)
        if path == "/api/docx/render/resource":
            return b"PNG"
        raise AssertionError(f"Unexpected REST GET BYTES path: {path}")

    async def rest_post_bytes(self, path: str, *, json_data: dict | None = None) -> bytes:
        if path == "/api/docx/quality/clean":
            return base64.b64decode(self.docx_b64)
        raise AssertionError(f"Unexpected REST POST BYTES path: {path}")

    def register_execution_observer(self, execution_id: str):
        queue: asyncio.Queue = asyncio.Queue()
        self.execution_observers.setdefault(execution_id, []).append(queue)

        def unregister() -> None:
            queues = self.execution_observers.get(execution_id)
            if queues is None:
                return
            try:
                queues.remove(queue)
            except ValueError:
                return
            if not queues:
                self.execution_observers.pop(execution_id, None)

        return queue, unregister

    async def ws_request(self, msg_type: str, data: dict, **kwargs) -> dict:
        if msg_type == "notebook_create":
            return {
                "type": "notebook_created",
                "kernel_id": "kernel-1",
                "notebook": {
                    "cells": [
                        {
                            "id": "welcome-cell",
                            "cell_type": "code",
                            "source": "# Bienvenido a Inspyro Notebook\nprint('Hola desde MCP')",
                            "metadata": {},
                            "outputs": [],
                            "execution_count": None,
                        }
                    ],
                    "metadata": {},
                    "nbformat": 4,
                    "nbformat_minor": 5,
                },
            }
        if msg_type == "notebook_load":
            return {
                "type": "notebook_loaded",
                "kernel_id": "kernel-1",
                "notebook": data["content"],
            }
        if msg_type == "notebook_save":
            return {"type": "notebook_saved", "status": "saved"}
        if msg_type == "notebook_set_order":
            return {"type": "notebook_order_set", "status": "ok"}
        if msg_type == "notebook_reset_kernel":
            return {"type": "notebook_kernel_reset"}
        if msg_type == "notebook_interrupt_kernel":
            return {"type": "notebook_kernel_interrupted"}
        if msg_type == "notebook_cancel_execution":
            return {
                "type": "notebook_execution_cancelled",
                "kernel_id": data["kernel_id"],
                "execution_id": data.get("execution_id"),
                "cancelled": True,
            }
        if msg_type == "notebook_shutdown_kernel":
            return {"type": "notebook_kernel_shutdown"}
        if msg_type == "force_reconvert_pdf":
            return {
                "type": "pdf_reconverted",
                "status": "ok",
                "pdf_file_token": "pdf-token-1",
                "pdf_ref": "/api/pdf/download?token=pdf-token-1",
                "pdf_file_name": "inspyro_document.pdf",
                "pdf_hash": "pdf-hash-1",
            }
        if msg_type == "template_attach":
            return {
                "type": "template_uploaded",
                "kernel_id": data["kernel_id"],
                "template_token": self.template_token,
                "template": {"styles": ["Normal"]},
            }
        if msg_type == "template_get":
            return {
                "type": "template_info",
                "template": {"styles": ["Normal"]},
            }
        if msg_type == "template_delete":
            return {
                "type": "template_deleted",
                "was_deleted": True,
            }
        if msg_type == "template_update_style":
            return {
                "type": "template_style_updated",
                "template": {"styles": [data["style_name"]]},
            }
        if msg_type == "analyze_dependencies":
            return {
                "type": "dependency_analysis_result",
                "symbol": data["symbol"],
                "graph": {"nodes": [{"name": data["symbol"]}], "edges": [], "analysis_mode": "dependencies"},
            }
        if msg_type == "analyze_impact":
            return {
                "type": "impact_analysis_result",
                "symbol": data["symbol"],
                "graph": {"nodes": [{"name": data["symbol"]}], "edges": [], "analysis_mode": "impact"},
            }
        if msg_type == "sensitivity_analyze":
            return {"type": "sensitivity_result", "success": True, "results": {"ok": True}}
        if msg_type == "optimize_design":
            return {"type": "optimization_result", "recommended_design": {"status": "ok"}}
        if msg_type == "compare_scenarios":
            return {"type": "scenario_comparison_result", "baseline": {}, "comparisons": []}
        if msg_type == "run_code_checks":
            return {"type": "code_checks_result", "summary": {"status": "ok"}, "checks": []}
        raise AssertionError(f"Unexpected WS request type: {msg_type}")

    async def ws_request_multi(self, msg_type: str, data: dict, **kwargs) -> list[dict]:
        if msg_type != "notebook_execute_cell":
            raise AssertionError(f"Unexpected WS multi request type: {msg_type}")

        self.execution_count += 1
        messages: list[dict] = []
        source = str(data.get("source") or "")
        if "print(" in source:
            messages.append(
                {
                    "type": "notebook_stream",
                    "execution_id": data["execution_id"],
                    "content": {"text": "stream\n", "name": "stdout"},
                }
            )

        terminal = {
            "type": "notebook_cell_executed",
            "execution_id": data["execution_id"],
            "execution_count": self.execution_count,
            "outputs": [],
            "variables": {},
            "kernel_id": data["kernel_id"],
            "cell_id": data["cell_id"],
        }
        if any(token in source for token in ("build_doc(", "doc_reset(", "Heading(", "Text(")):
            terminal["docx_file_b64"] = self.docx_b64
            terminal["docx_file_token"] = "docx-token-1"
            terminal["docx_artifact_id"] = "artifact-1"
            terminal["docx_ref"] = "/api/docx/download?token=docx-token-1"
            terminal["docx_download_url"] = "/api/docx/download?artifact_id=artifact-1"
            terminal["docx_file_name"] = "inspyro_document.docx"
            terminal["docx_hash"] = "docx-hash-1"
            terminal["docx_size_bytes"] = len(base64.b64decode(self.docx_b64))
            terminal["pdf_file_token"] = "pdf-token-1"
            terminal["pdf_ref"] = "/api/pdf/download?token=pdf-token-1"
            terminal["pdf_file_name"] = "inspyro_document.pdf"
            terminal["pdf_hash"] = "pdf-hash-1"
            terminal["pdf_size_bytes"] = len(b"%PDF-1.4\n%fake\n")
        messages.append(terminal)
        on_message = kwargs.get("on_message")
        if on_message is not None:
            for message in messages:
                callback_result = on_message(message)
                if asyncio.iscoroutine(callback_result):
                    await callback_result
        execution_id = data.get("execution_id")
        if execution_id:
            for queue in list(self.execution_observers.get(execution_id, [])):
                for message in messages:
                    queue.put_nowait(message)
        return messages


def _resource_text(contents: list) -> str:
    return "\n".join(
        item.text
        for item in contents
        if getattr(item, "text", None)
    )


def _prompt_text(result) -> str:
    return "\n".join(
        message.content.text
        for message in result.messages
        if getattr(message.content, "text", None)
    )


def _resource_bytes(contents: list) -> bytes:
    for item in contents:
        blob = getattr(item, "blob", None)
        if blob:
            return base64.b64decode(blob)
        text = getattr(item, "text", None)
        if text:
            return str(text).encode("utf-8")
    return b""


def _tool_tags(tool) -> list[str]:
    direct_tags = list(getattr(tool, "tags", []) or [])
    if direct_tags:
        return direct_tags
    meta = getattr(tool, "meta", None) or {}
    if isinstance(meta, dict):
        fastmcp_meta = meta.get("fastmcp") or {}
        if isinstance(fastmcp_meta, dict):
            return list(fastmcp_meta.get("tags", []) or [])
    return []


@pytest.fixture(autouse=True)
def _clear_mcp_state():
    state = McpSessionState.get()
    state.clear()
    InspyroBridge._instance = None
    InspyroBridge._instances = {}
    mcp_config.set_runtime_transport(transport="streamable-http", stateless_http=False)
    yield
    state.clear()
    InspyroBridge._instance = None
    InspyroBridge._instances = {}
    mcp_config.set_runtime_transport(transport="streamable-http", stateless_http=False)


@pytest.fixture
def fake_bridge(monkeypatch, tmp_path: Path) -> FakeBlackBoxBridge:
    bridge = FakeBlackBoxBridge(tmp_path)
    monkeypatch.setattr(InspyroBridge, "get", classmethod(lambda cls: bridge))
    return bridge


@pytest.mark.asyncio
async def test_mcp_ai_first_catalog_exposes_guides_and_prompts(fake_bridge: FakeBlackBoxBridge):
    async with Client(mcp_server_main.mcp) as client:
        instructions = client.initialize_result.instructions or ""
        resources = await client.list_resources()
        prompts = await client.list_prompts()
        tools = await client.list_tools()
        resource_templates = await client.list_resource_templates()

        resource_uris = {str(resource.uri) for resource in resources}
        prompt_names = {str(prompt.name) for prompt in prompts}
        tool_names = {str(tool.name) for tool in tools}
        resource_template_uris = {str(resource.uriTemplate) for resource in resource_templates}

        assert "inspyro://guides/start-here" in instructions
        assert len(resources) == 15
        assert len(prompts) == 7
        assert len(resource_templates) == 7
        assert tool_names == EXPECTED_AUTHORING_TOOLS
        assert "inspyro://guides/start-here" in resource_uris
        assert "inspyro://manifest" in resource_uris
        assert "inspyro://session/notebooks" in resource_uris
        assert "inspyro://examples/notebook-docx-report" in resource_uris
        assert "inspyro://workspace/tree/{path*}" in resource_template_uris
        assert "inspyro://artifacts/{kernel_id}/{kind}" in resource_template_uris
        assert "inspyro://artifacts/token/{kind}/{token}" in resource_template_uris
        assert "start_inspyro_session" in prompt_names
        assert "create_docx_report_notebook" in prompt_names
        assert "recover_mcp_notebook_session" in prompt_names
        assert "notebook_create" in tool_names
        assert "notebook_sync_cells" in tool_names
        assert "get_document_docx" in tool_names
        assert "check_document_quality" in tool_names
        assert "run_document_workbench" in tool_names
        assert "compare_document_versions" in tool_names
        assert "manage_document_review" in tool_names
        assert "prepare_document_delivery" in tool_names
        assert "export_clean_document_docx" in tool_names
        assert "export_document_docx" in tool_names
        assert "export_document_pdf" in tool_names
        assert "get_run_status" in tool_names
        assert "get_kernel_status" in tool_names
        assert "list_session_notebooks" in tool_names
        assert "list_cells" in tool_names
        assert "get_cell" in tool_names
        assert "find_in_notebook" in tool_names
        assert "cancel_run" in tool_names
        assert "resume_run" in tool_names
        assert "close_session_notebook" in tool_names
        assert "list_component_profiles" in tool_names
        assert "set_component_profile" in tool_names
        assert "create_kernel" not in tool_names
        assert "execution_status" not in tool_names
        assert "add_cell" not in tool_names
        assert "read_file" not in tool_names
        assert "get_metrics" not in tool_names
        assert "get_pdf_status" not in tool_names

        start_here = _resource_text(await client.read_resource("inspyro://guides/start-here"))
        docx_guide = _resource_text(await client.read_resource("inspyro://guides/docx-quickstart"))
        artifact_guide = _resource_text(await client.read_resource("inspyro://guides/artifact-lifecycle"))
        recovery_guide = _resource_text(await client.read_resource("inspyro://guides/error-recovery"))
        example = _resource_text(await client.read_resource("inspyro://examples/notebook-docx-report"))
        manifest = _resource_text(await client.read_resource("inspyro://manifest"))
        system_info_resource = _resource_text(await client.read_resource("inspyro://system/info"))
        system_info_resource_payload = json.loads(system_info_resource)
        prompt = _prompt_text(await client.get_prompt("start_inspyro_session", {"goal": "crear reporte"}))
        system_info = await client.call_tool("get_system_info", {})
        health_info = await client.call_tool("get_health", {})
        system_struct = system_info.structured_content
        health_struct = health_info.structured_content
        system_tool = next(tool for tool in tools if str(tool.name) == "get_system_info")
        document_tool = next(tool for tool in tools if str(tool.name) == "get_document_docx")

        assert "inspyro://guides/notebook-workflow" in start_here
        assert "`kernel_id`" in start_here
        assert "build_doc" in example
        assert "get_document_docx" in example
        assert "check_document_quality" in docx_guide
        assert "check_document_quality" in artifact_guide
        assert "missing_quality" in recovery_guide
        assert "resource_templates" in manifest
        assert "inspyro://guides/start-here" in prompt
        assert "inspyro://examples/notebook-docx-report" in prompt
        assert system_info_resource_payload["workspace_path"] == str(fake_bridge.root)
        assert system_info_resource_payload["workspace_root"] == str(fake_bridge.root)
        assert system_info_resource_payload["notebook_session_mode"] == "stateful-http"
        assert system_info_resource_payload["notebook_sessions_supported"] is True
        assert system_info_resource_payload["pdf_status"]["conversion_available"] is True
        assert system_info_resource_payload["pdf_status"]["preferred_engine"] == "libreoffice"
        assert system_struct["workspace_path"] == str(fake_bridge.root)
        assert system_struct["workspace_root"] == str(fake_bridge.root)
        assert system_struct["component_profile"] == "authoring"
        assert system_struct["notebook_session_mode"] == "stateful-http"
        assert system_struct["notebook_sessions_supported"] is True
        assert system_struct["pdf_status"]["conversion_available"] is True
        assert system_struct["pdf_status"]["preferred_engine"] == "libreoffice"
        assert health_struct["pdf_status"]["conversion_available"] is True
        assert health_struct["pdf_status"]["preferred_engine"] == "libreoffice"
        assert health_struct["notebook_session_mode"] == "stateful-http"
        assert getattr(system_tool, "annotations", None) is not None
        assert getattr(system_tool.annotations, "readOnlyHint", False) is True
        assert "group:system" in _tool_tags(system_tool)
        assert getattr(document_tool, "annotations", None) is not None
        assert getattr(document_tool.annotations, "readOnlyHint", False) is True
        assert getattr(document_tool.annotations, "idempotentHint", False) is True

        changed_analysis = await client.call_tool("set_component_profile", {"profile": "analysis"})
        analysis_tools = {str(tool.name) for tool in await client.list_tools()}
        changed_files = await client.call_tool("set_component_profile", {"profile": "files"})
        files_tools = {str(tool.name) for tool in await client.list_tools()}

        assert changed_analysis.structured_content["status"] == "ok"
        assert analysis_tools == EXPECTED_ANALYSIS_TOOLS
        assert changed_files.structured_content["status"] == "ok"
        assert files_tools == EXPECTED_FILES_TOOLS


@pytest.mark.asyncio
async def test_mcp_blackbox_notebook_to_docx_flow(fake_bridge: FakeBlackBoxBridge, tmp_path: Path):
    async with Client(mcp_server_main.mcp) as client:
        example = _resource_text(await client.read_resource("inspyro://examples/notebook-docx-report"))
        assert "notebook_sync_cells" in example

        created = await client.call_tool(
            "notebook_create",
            {"path": str(tmp_path), "name": "beam_report.ipynb"},
        )
        notebook_path = created.structured_content["path"]
        kernel_id = created.structured_content["kernel_id"]

        await client.call_tool(
            "notebook_sync_cells",
            {
                "notebook_path": notebook_path,
                "cells": [
                    {
                        "cell_type": "code",
                        "source": "L_m = 6.0\nw_kN_m = 12.5\nM_max_kNm = w_kN_m * L_m**2 / 8.0\nprint(M_max_kNm)",
                    },
                    {
                        "cell_type": "code",
                        "source": (
                            "doc_reset(hard=True)\n"
                            "with build_doc(block_id='cover', order=10) as builder:\n"
                            "    builder.heading('Informe tecnico', level=1)\n"
                            "    builder.text('Reporte generado desde MCP')\n"
                        ),
                    },
                ],
            },
        )

        executed = await client.call_tool(
            "execute_all_cells",
            {"kernel_id": kernel_id, "notebook_path": notebook_path},
        )
        docx_result = await client.call_tool(
            "get_document_docx",
            {"kernel_id": kernel_id, "inline_content": True},
        )
        assert fake_bridge.quality_summaries == {}
        quality_result = await client.call_tool(
            "check_document_quality",
            {"kernel_id": kernel_id, "run": True, "max_findings": 1},
        )
        docx_with_quality = await client.call_tool(
            "get_document_docx",
            {"kernel_id": kernel_id, "include_quality": True},
        )
        visual_result = await client.call_tool(
            "run_document_workbench",
            {"kernel_id": kernel_id, "operation": "render_all_pages"},
        )
        delivery_result = await client.call_tool(
            "prepare_document_delivery",
            {"kernel_id": kernel_id},
        )
        export_result = await client.call_tool(
            "export_document_docx",
            {"kernel_id": kernel_id, "path": str(tmp_path / "exports")},
        )
        clean_export_result = await client.call_tool(
            "export_clean_document_docx",
            {"kernel_id": kernel_id, "path": str(tmp_path / "clean-docx")},
        )

        assert executed.structured_content["status"] == "completed"
        assert executed.structured_content["failed"] == 0
        assert docx_result.structured_content["status"] == "ok"
        assert docx_result.structured_content["resource_uri"].startswith("inspyro://artifacts/")
        assert docx_result.structured_content["portable_resource_uri"].startswith("inspyro://artifacts/token/docx/")
        assert docx_result.structured_content["resource_scope"] == "portable"
        assert docx_result.structured_content["inline_content"] is True
        assert "docx_quality" not in docx_result.structured_content
        assert quality_result.structured_content["status"] == "ok"
        assert quality_result.structured_content["quality_status"] == "warning"
        assert len(quality_result.structured_content["findings"]) == 1
        assert docx_with_quality.structured_content["docx_quality"]["detail"] == "compact"
        assert visual_result.structured_content["status"] == "ok"
        assert visual_result.structured_content["visual"]["status"] == "complete"
        assert visual_result.structured_content["resources"][0]["resource_uri"].startswith("/api/docx/render/resource")
        assert "docx_base64" not in visual_result.structured_content
        assert delivery_result.structured_content["status"] == "ok"
        assert delivery_result.structured_content["visual"]["status"] == "partial"
        assert delivery_result.structured_content["variant"]["resource_uri"].startswith("/api/docx/workbench/resource")
        assert export_result.structured_content["status"] == "ok"
        assert Path(export_result.structured_content["path"]).exists()
        assert clean_export_result.structured_content["status"] == "ok"
        assert Path(clean_export_result.structured_content["path"]).exists()

        document = Document(io.BytesIO(base64.b64decode(docx_result.structured_content["docx_base64"])))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        assert any("Black-box MCP DOCX" in paragraph for paragraph in paragraphs)
        assert fake_bridge.activity_events

    async with Client(mcp_server_main.mcp) as second_client:
        portable_docx = _resource_bytes(
            await second_client.read_resource(docx_result.structured_content["portable_resource_uri"])
        )
        assert portable_docx == base64.b64decode(docx_result.structured_content["docx_base64"])


@pytest.mark.asyncio
async def test_mcp_blackbox_background_run_status_and_run_resource(
    fake_bridge: FakeBlackBoxBridge,
    tmp_path: Path,
):
    async with Client(mcp_server_main.mcp) as client:
        created = await client.call_tool(
            "notebook_create",
            {"path": str(tmp_path), "name": "background.ipynb"},
        )
        notebook_path = created.structured_content["path"]
        kernel_id = created.structured_content["kernel_id"]

        await client.call_tool(
            "notebook_sync_cells",
            {
                "notebook_path": notebook_path,
                "cells": [
                    {"cell_type": "code", "source": "print('background run')"},
                    {"cell_type": "code", "source": "result = 2 + 2"},
                ],
            },
        )

        started = await client.call_tool(
            "execute_all_cells",
            {"kernel_id": kernel_id, "notebook_path": notebook_path, "background": True},
        )
        run_id = started.structured_content["run_id"]

        assert started.structured_content["status"] == "started"
        assert started.structured_content["execution_id"] == run_id

        status = None
        for _ in range(20):
            status = await client.call_tool("get_run_status", {"run_id": run_id})
            if status.structured_content["execution_status"] != "running":
                break
            await asyncio.sleep(0.01)

        assert status is not None
        assert status.structured_content["run_id"] == run_id
        assert status.structured_content["execution_id"] == run_id
        assert status.structured_content["execution_status"] == "completed"
        assert status.structured_content["last_output_preview"] == "stream\n"

        run_resource = json.loads(_resource_text(await client.read_resource(f"inspyro://runs/{run_id}")))
        assert run_resource["run_id"] == run_id
        assert run_resource["status"] == "completed"
        assert run_resource["last_output_preview"] == "stream\n"


@pytest.mark.asyncio
async def test_mcp_blackbox_error_recovery_from_missing_artifact(fake_bridge: FakeBlackBoxBridge, tmp_path: Path):
    async with Client(mcp_server_main.mcp) as client:
        recovery_guide = _resource_text(await client.read_resource("inspyro://guides/error-recovery"))
        recovery_prompt = _prompt_text(
            await client.get_prompt(
                "recover_mcp_notebook_session",
                {"observed_error": "missing_artifact", "notebook_path": str(tmp_path / "recover.ipynb")},
            )
        )

        created = await client.call_tool(
            "notebook_create",
            {"path": str(tmp_path), "name": "recover.ipynb"},
        )
        notebook_path = created.structured_content["path"]
        kernel_id = created.structured_content["kernel_id"]

        missing_docx = await client.call_tool("get_document_docx", {"kernel_id": kernel_id})
        assert missing_docx.structured_content["status"] == "missing_artifact"

        await client.call_tool(
            "notebook_sync_cells",
            {
                "notebook_path": notebook_path,
                "cells": [
                    {
                        "cell_type": "code",
                        "source": (
                            "doc_reset(hard=True)\n"
                            "with build_doc(block_id='cover', order=10) as builder:\n"
                            "    builder.heading('Recuperacion', level=1)\n"
                        ),
                    }
                ],
            },
        )
        await client.call_tool(
            "execute_all_cells",
            {"kernel_id": kernel_id, "notebook_path": notebook_path},
        )
        recovered_docx = await client.call_tool(
            "get_document_docx",
            {"kernel_id": kernel_id, "inline_content": True},
        )

        assert "missing_artifact" in recovery_guide
        assert "notebook_load" in recovery_prompt
        assert "get_document_docx" in recovery_prompt
        assert recovered_docx.structured_content["status"] == "ok"


@pytest.mark.asyncio
async def test_mcp_blackbox_completions_cover_prompts_and_templates(
    fake_bridge: FakeBlackBoxBridge,
    tmp_path: Path,
):
    template_path = tmp_path / "beam-template.docx"
    Document().save(template_path)

    async with Client(mcp_server_main.mcp) as client:
        created = await client.call_tool(
            "notebook_create",
            {"path": str(tmp_path), "name": "beam_report.ipynb"},
        )
        notebook_path = created.structured_content["path"]
        kernel_id = created.structured_content["kernel_id"]

        await client.call_tool(
            "upload_template",
            {"kernel_id": kernel_id, "file_path": str(template_path)},
        )

        notebook_completion = await client.complete(
            mcp_types.PromptReference(type="ref/prompt", name="review_notebook"),
            {"name": "notebook_path", "value": "beam"},
        )
        kernel_completion = await client.complete(
            mcp_types.PromptReference(type="ref/prompt", name="recover_mcp_notebook_session"),
            {"name": "kernel_id", "value": "ker"},
        )
        style_completion = await client.complete(
            mcp_types.PromptReference(type="ref/prompt", name="recover_mcp_notebook_session"),
            {"name": "style_name", "value": "Nor"},
            {"kernel_id": kernel_id},
        )
        unit_completion = await client.complete(
            mcp_types.PromptReference(type="ref/prompt", name="unit_conversion_help"),
            {"name": "from_unit", "value": "k"},
        )
        cell_completion = await client.complete(
            mcp_types.ResourceTemplateReference(
                type="ref/resource",
                uri="inspyro://notebooks/{path*}/cells/{cell_id}",
            ),
            {"name": "cell_id", "value": "wel"},
            {"path": notebook_path},
        )

        assert any(value.endswith("beam_report.ipynb") for value in notebook_completion.values)
        assert "kernel-1" in kernel_completion.values
        assert "Normal" in style_completion.values
        assert "kN" in unit_completion.values
        assert "welcome-cell" in cell_completion.values


@pytest.mark.asyncio
async def test_mcp_files_profile_allows_notebook_paths_but_guidance_prefers_notebook_tools(
    fake_bridge: FakeBlackBoxBridge,
    tmp_path: Path,
):
    async with Client(mcp_server_main.mcp) as client:
        manifest = _resource_text(await client.read_resource("inspyro://manifest"))
        notebook_guide = _resource_text(await client.read_resource("inspyro://guides/notebook-workflow"))

        changed_files = await client.call_tool("set_component_profile", {"profile": "files"})
        notebook_path = str(tmp_path / "generic_notebook.ipynb")
        renamed_path = str(tmp_path / "generic_notebook_renamed.ipynb")

        created = await client.call_tool("create_file", {"path": notebook_path})
        written = await client.call_tool(
            "write_file",
            {
                "path": notebook_path,
                "content": '{"cells": [], "metadata": {}, "nbformat": 4, "nbformat_minor": 5}',
            },
        )
        read_back = await client.call_tool("read_file", {"path": notebook_path})
        renamed = await client.call_tool("rename_file", {"old_path": notebook_path, "new_path": renamed_path})
        deleted = await client.call_tool("delete_file", {"path": renamed_path})

        assert changed_files.structured_content["status"] == "ok"
        assert created.structured_content["path"].endswith("generic_notebook.ipynb")
        assert written.structured_content["path"].endswith("generic_notebook.ipynb")
        assert '"nbformat": 4' in read_back.structured_content["content"]
        assert renamed.structured_content["newPath"].endswith("generic_notebook_renamed.ipynb")
        assert deleted.structured_content["path"].endswith("generic_notebook_renamed.ipynb")
        assert "Prefer notebook editing through `notebook_load`, `notebook_sync_cells`, and `notebook_save`" in manifest
        assert "prefiere `notebook_load(include_source=True)` -> `notebook_sync_cells` -> `notebook_save`" in notebook_guide
