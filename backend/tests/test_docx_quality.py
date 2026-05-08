import asyncio
import base64
import io
import json
import os
import sys
import zipfile
from contextlib import contextmanager
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document
from fastapi import HTTPException
from lxml import etree

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.routers import docx as docx_router
from app.services import docx_artifacts
from app.services import docx_render_cache
from app.services.docx_quality.audit import audit_docx_bytes
from app.services.docx_quality.content_controls import fill_content_controls, inspect_content_controls, wrap_placeholders_as_content_controls
from app.services.docx_quality.ooxml import NS, parse_xml, qn, xml_bytes
from app.services.docx_quality.publish import clean_docx_bytes
from app.services.docx_quality.protection import set_document_protection
from app.services.docx_quality.redaction import redact_docx_bytes
from app.services.docx_quality.render import render_docx_all_pages, render_docx_manifest, render_docx_page_png
from app.services.docx_quality.review import extract_comments
from app.services.docx_quality.workbench import run_workbench_operation


def _make_docx_bytes(*paragraphs: str, author: str = "Inspyro QA") -> bytes:
    doc = Document()
    doc.core_properties.author = author
    doc.add_heading("Titulo", level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _rewrite_docx(docx_bytes: bytes, replacements: dict[str, bytes], additions: dict[str, bytes] | None = None) -> bytes:
    output = io.BytesIO()
    additions = additions or {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as source, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target:
        existing = set(source.namelist())
        for name in source.namelist():
            target.writestr(name, replacements.get(name, source.read(name)))
        for name, payload in additions.items():
            if name not in existing:
                target.writestr(name, payload)
    return output.getvalue()


def _document_xml(docx_bytes: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        return archive.read("word/document.xml")


def _add_hyperlink_text(docx_bytes: bytes, text_value: str) -> bytes:
    root = parse_xml(_document_xml(docx_bytes))
    body = root.find("w:body", namespaces=NS)
    assert body is not None
    sect_pr = body.find("w:sectPr", namespaces=NS)
    insert_at = body.index(sect_pr) if sect_pr is not None else len(body)

    paragraph = etree.Element(qn("w", "p"))
    hyperlink = etree.SubElement(paragraph, qn("w", "hyperlink"))
    run = etree.SubElement(hyperlink, qn("w", "r"))
    text = etree.SubElement(run, qn("w", "t"))
    text.text = text_value
    body.insert(insert_at, paragraph)
    return _rewrite_docx(docx_bytes, {"word/document.xml": xml_bytes(root)})


def _add_content_controls(docx_bytes: bytes) -> bytes:
    root = parse_xml(_document_xml(docx_bytes))
    body = root.find("w:body", namespaces=NS)
    assert body is not None
    sect_pr = body.find("w:sectPr", namespaces=NS)
    insert_at = body.index(sect_pr) if sect_pr is not None else len(body)

    sdt = etree.Element(qn("w", "sdt"))
    sdt_pr = etree.SubElement(sdt, qn("w", "sdtPr"))
    etree.SubElement(sdt_pr, qn("w", "tag")).set(qn("w", "val"), "CLIENTE")
    etree.SubElement(sdt_pr, qn("w", "alias")).set(qn("w", "val"), "Cliente")
    etree.SubElement(sdt_pr, qn("w", "text"))
    content = etree.SubElement(sdt, qn("w", "sdtContent"))
    paragraph = etree.SubElement(content, qn("w", "p"))
    run = etree.SubElement(paragraph, qn("w", "r"))
    text = etree.SubElement(run, qn("w", "t"))
    text.text = "{{CLIENTE}}"
    body.insert(insert_at, sdt)

    loose_paragraph = etree.Element(qn("w", "p"))
    loose_run = etree.SubElement(loose_paragraph, qn("w", "r"))
    loose_text = etree.SubElement(loose_run, qn("w", "t"))
    loose_text.text = "{{OBRA}}"
    body.insert(insert_at + 1, loose_paragraph)

    return _rewrite_docx(docx_bytes, {"word/document.xml": xml_bytes(root)})


def _add_comments_and_tracked_changes(docx_bytes: bytes) -> bytes:
    root = parse_xml(_document_xml(docx_bytes))
    body = root.find("w:body", namespaces=NS)
    assert body is not None
    sect_pr = body.find("w:sectPr", namespaces=NS)
    insert_at = body.index(sect_pr) if sect_pr is not None else len(body)

    paragraph = etree.Element(qn("w", "p"))
    insertion = etree.SubElement(paragraph, qn("w", "ins"))
    inserted_run = etree.SubElement(insertion, qn("w", "r"))
    inserted_text = etree.SubElement(inserted_run, qn("w", "t"))
    inserted_text.text = "Texto aceptado"
    deletion = etree.SubElement(paragraph, qn("w", "del"))
    deleted_run = etree.SubElement(deletion, qn("w", "r"))
    deleted_text = etree.SubElement(deleted_run, qn("w", "delText"))
    deleted_text.text = "Texto rechazado"
    body.insert(insert_at, paragraph)

    comments = b'<?xml version="1.0" encoding="UTF-8"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    return _rewrite_docx(docx_bytes, {"word/document.xml": xml_bytes(root)}, {"word/comments.xml": comments})


@contextmanager
def _patched_artifact_store(root_dir: Path):
    blobs_dir = root_dir / "blobs"
    manifests_dir = root_dir / "manifests"
    quality_dir = root_dir / "quality"
    workbench_dir = root_dir / "workbench"
    render_dir = root_dir / "render"
    for directory in (blobs_dir, manifests_dir, quality_dir, workbench_dir, render_dir):
        directory.mkdir(parents=True, exist_ok=True)
    with patch.multiple(
        docx_artifacts,
        DOCX_ARTIFACT_ROOT=root_dir,
        DOCX_ARTIFACT_BLOBS_DIR=blobs_dir,
        DOCX_ARTIFACT_MANIFESTS_DIR=manifests_dir,
        DOCX_ARTIFACT_QUALITY_DIR=quality_dir,
        DOCX_ARTIFACT_WORKBENCH_DIR=workbench_dir,
        DOCX_ARTIFACT_INDEX_PATH=root_dir / "index.json",
        get_workspace_snapshot=lambda: {"active_workspace": str(root_dir / "workspace")},
    ):
        with patch.multiple(
            docx_render_cache,
            DOCX_RENDER_CACHE_DIR=render_dir,
            _INFLIGHT_BY_KEY={},
        ):
            with docx_artifacts._ARTIFACT_LOCK:
                docx_artifacts._ARTIFACTS = None
                docx_artifacts._PROVENANCE_INDEX = None
                docx_artifacts._LAST_CLEANUP_AT = 0.0
                docx_artifacts._initialize_docx_artifact_store()
            yield


def _store_docx_artifact(docx_bytes: bytes, *, root_dir: Path) -> dict:
    payload = base64.b64encode(docx_bytes).decode("ascii")
    (root_dir / "demo.ipynb").write_text("{}", encoding="utf-8")
    return docx_artifacts.store_docx_artifact(
        payload,
        filename="quality.docx",
        docx_hash="hash-quality",
        source_kind="notebook",
        source_path=str(root_dir / "demo.ipynb"),
        kernel_id="kernel-quality",
        execution_id="exec-quality",
    )


def test_audit_detects_fields_placeholders_and_content_controls():
    docx_bytes = _add_content_controls(_make_docx_bytes("Pagina literal {PAGE}"))

    controls = inspect_content_controls(docx_bytes)
    assert controls["control_count"] == 1
    assert controls["placeholder_count"] == 2
    assert controls["unwrapped_placeholder_count"] == 1

    summary = audit_docx_bytes(docx_bytes, artifact_id="artifact-1")
    assert summary["artifact_id"] == "artifact-1"
    assert summary["status"] == "warning"
    assert summary["content_controls"]["unwrapped_placeholder_count"] == 1
    assert any(
        finding["section"] == "fields" and finding.get("context", {}).get("placeholder") == "{PAGE}"
        for finding in summary["findings"]
    )

    filled = fill_content_controls(docx_bytes, {"CLIENTE": "ACME"})
    assert b"ACME" in _document_xml(filled)

    wrapped, wrap_stats = wrap_placeholders_as_content_controls(docx_bytes)
    assert wrap_stats["wrapped"] == 1
    wrapped_controls = inspect_content_controls(wrapped)
    assert any(item["tag"] == "OBRA" for item in wrapped_controls["controls"])


def test_audit_detects_accented_generic_link_text():
    summary = audit_docx_bytes(_add_hyperlink_text(_make_docx_bytes("Base"), "aquí"))

    assert any(
        finding.get("code") == "accessibility.link_text_generic"
        and finding.get("context", {}).get("text") == "aquí"
        for finding in summary["findings"]
    )


def test_clean_docx_scrubs_metadata_comments_and_accepts_or_rejects_tracked_changes():
    source = _add_comments_and_tracked_changes(_make_docx_bytes("Base", author="Persona Interna"))

    accepted, accepted_stats = clean_docx_bytes(source, tracked_changes="accept")
    accepted_xml = _document_xml(accepted)
    assert accepted_stats["comments_removed"] == 1
    assert accepted_stats["tracked_changes_processed"] == 2
    assert "word/comments.xml" not in zipfile.ZipFile(io.BytesIO(accepted)).namelist()
    assert b"Persona Interna" not in accepted
    assert b"Texto aceptado" in accepted_xml
    assert b"Texto rechazado" not in accepted_xml
    assert b"<w:ins" not in accepted_xml and b"<w:del" not in accepted_xml

    rejected, _ = clean_docx_bytes(source, tracked_changes="reject")
    rejected_xml = _document_xml(rejected)
    assert b"Texto aceptado" not in rejected_xml
    assert b"Texto rechazado" in rejected_xml


def test_quality_routes_persist_summary_render_png_and_clean_copy():
    source = _add_comments_and_tracked_changes(_add_content_controls(_make_docx_bytes("Pagina literal {PAGE}")))
    with TemporaryDirectory() as tmp_dir:
        root_dir = Path(tmp_dir)
        with _patched_artifact_store(root_dir):
            stored = _store_docx_artifact(source, root_dir=root_dir)
            artifact_id = stored["artifact_id"]

            with pytest_raises_http_404():
                asyncio.run(docx_router.get_docx_quality(artifact_id=artifact_id))

            run_response = asyncio.run(docx_router.run_docx_quality({"artifact_id": artifact_id}))
            summary = json.loads(run_response.body.decode("utf-8"))
            assert summary["artifact_id"] == artifact_id
            assert summary["status"] in {"ok", "warning", "error"}

            get_response = asyncio.run(docx_router.get_docx_quality(artifact_id=artifact_id))
            loaded = json.loads(get_response.body.decode("utf-8"))
            assert loaded["binary_hash"] == summary["binary_hash"]

            history_response = asyncio.run(
                docx_router.get_docx_history(source_path=str(root_dir / "demo.ipynb"), kernel_id=None, limit=10)
            )
            history = json.loads(history_response.body.decode("utf-8"))
            assert history["items"][0]["docx_quality_status"] == summary["status"]
            assert history["items"][0]["docx_quality_counts"] == summary["counts"]
            assert history["items"][0]["docx_render_status"] == "missing"

            with patch.object(docx_router, "render_docx_page_png", return_value=(b"PNGDATA", {"page": 1, "page_count": 2})):
                render_response = asyncio.run(docx_router.render_docx_quality_page(artifact_id=artifact_id, page=1))
            assert render_response.media_type == "image/png"
            assert render_response.body == b"PNGDATA"
            assert render_response.headers["x-inspyro-docx-page-count"] == "2"

            clean_response = asyncio.run(docx_router.clean_docx_quality({"artifact_id": artifact_id, "tracked_changes": "accept"}))
            assert clean_response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            assert "word/comments.xml" not in zipfile.ZipFile(io.BytesIO(clean_response.body)).namelist()


def test_render_cache_persists_pdf_and_page_png_by_binary_hash():
    docx_bytes = b"fake-docx-for-render-cache"
    calls = {"convert": 0, "render": 0}

    def fake_convert(_docx_b64, timeout_s=None, progress_callback=None):
        calls["convert"] += 1
        return {
            "pdf_b64": base64.b64encode(b"%PDF-render-cache").decode("ascii"),
            "converter_used": "word",
            "duration_ms": 17,
            "stage_timings_ms": {"pdf_convert_ms": 17},
        }

    def fake_extract_pages(_pdf_bytes):
        return 2, [
            {"page": 1, "width_pt": 612.0, "height_pt": 792.0},
            {"page": 2, "width_pt": 612.0, "height_pt": 792.0},
        ]

    def fake_render_page(_pdf_path, *, page, zoom):
        calls["render"] += 1
        return f"PNG-{page}-{zoom}".encode("ascii")

    with TemporaryDirectory() as tmp_dir:
        render_dir = Path(tmp_dir) / "render-cache"
        render_dir.mkdir()
        with patch.multiple(docx_render_cache, DOCX_RENDER_CACHE_DIR=render_dir, _INFLIGHT_BY_KEY={}):
            with patch.object(docx_render_cache, "convert_docx_with_diagnostics", side_effect=fake_convert), \
                 patch.object(docx_render_cache, "_extract_pdf_pages", side_effect=fake_extract_pages), \
                 patch.object(docx_render_cache, "_render_pdf_page", side_effect=fake_render_page):
                png_1, meta_1 = render_docx_page_png(
                    docx_bytes,
                    page=1,
                    binary_hash="render-hash",
                    artifact_id="artifact-render",
                    zoom=2.0,
                )
                png_2, meta_2 = render_docx_page_png(
                    docx_bytes,
                    page=1,
                    binary_hash="render-hash",
                    artifact_id="artifact-render",
                    zoom=2.0,
                )
                manifest = render_docx_manifest(binary_hash="render-hash", artifact_id="artifact-render")
                all_pages = render_docx_all_pages(
                    docx_bytes,
                    binary_hash="render-hash",
                    artifact_id="artifact-render",
                    zoom=2.0,
                )

    assert png_1 == b"PNG-1-2.0"
    assert png_2 == b"PNG-1-2.0"
    assert meta_1["cached"] is False
    assert meta_2["cached"] is True
    assert calls["convert"] == 1
    assert calls["render"] == 2
    assert manifest["status"] == "partial"
    assert manifest["page_count"] == 2
    assert manifest["cached_pages"] == 1
    assert manifest["resources"][0]["resource_uri"].startswith("/api/docx/render/resource")
    assert all_pages["manifest"]["status"] == "complete"
    assert len(all_pages["rendered_pages"]) == 2


def test_workbench_operations_return_variants_and_compact_resources():
    source = _add_content_controls(_make_docx_bytes("Secreto RUT 12345678-9", author="Persona Interna"))
    redacted, redaction_stats = redact_docx_bytes(source, patterns=[r"12345678-9"], replacement="X")
    assert redaction_stats["redactions"] == 1
    assert b"12345678-9" not in _document_xml(redacted)

    protected, protection_stats = set_document_protection(source, mode="comments")
    assert protection_stats["protected"] is True
    with zipfile.ZipFile(io.BytesIO(protected)) as archive:
        assert b"documentProtection" in archive.read("word/settings.xml")

    audit_result, audit_resources = run_workbench_operation(source, operation="audit", payload={"profile": "agent"}, artifact={"artifact_id": "a1"})
    assert audit_result["summary"]["schema_version"] == 2
    assert audit_resources == {}

    fields_result, _ = run_workbench_operation(source, operation="fields_report", artifact={"artifact_id": "a1"})
    assert fields_result["fields"]["field_count"] == 0

    wrap_result, wrap_resources = run_workbench_operation(source, operation="content_controls_wrap", artifact={"artifact_id": "a1", "filename": "template.docx"})
    assert wrap_result["variant"]["operation"] == "content_controls_wrap"
    assert list(wrap_resources.keys()) == ["template-sdt.docx"]

    redact_result, redact_resources = run_workbench_operation(
        source,
        operation="redact",
        payload={"patterns": [r"12345678-9"]},
        artifact={"artifact_id": "a1", "filename": "demo.docx"},
    )
    assert redact_result["redaction"]["redactions"] == 1
    assert redact_result["variant"]["filename"] == "demo-redacted.docx"
    assert "demo-redacted.docx" in redact_resources

    with TemporaryDirectory() as tmp_dir:
        render_dir = Path(tmp_dir) / "render-cache"
        render_dir.mkdir()
        with patch.multiple(docx_render_cache, DOCX_RENDER_CACHE_DIR=render_dir, _INFLIGHT_BY_KEY={}):
            with patch.object(docx_render_cache, "convert_docx_with_diagnostics", return_value={
                "pdf_b64": base64.b64encode(b"%PDF-workbench").decode("ascii"),
                "converter_used": "word",
                "duration_ms": 11,
                "stage_timings_ms": {},
            }), patch.object(docx_render_cache, "_extract_pdf_pages", return_value=(1, [{"page": 1, "width_pt": 612.0, "height_pt": 792.0}])), \
                patch.object(docx_render_cache, "_render_pdf_page", return_value=b"PNG-WORKBENCH"):
                render_result, render_resources = run_workbench_operation(
                    source,
                    operation="render_page",
                    payload={"page": 1},
                    artifact={"artifact_id": "a1", "binary_hash": "workbench-render-hash"},
                )
                manifest_result, _ = run_workbench_operation(
                    source,
                    operation="render_manifest",
                    artifact={"artifact_id": "a1", "binary_hash": "workbench-render-hash"},
                )
                render_all_result, _ = run_workbench_operation(
                    source,
                    operation="render_all_pages",
                    artifact={"artifact_id": "a1", "binary_hash": "workbench-render-hash"},
                )
                render_resource_path = Path(render_result["resources"][0]["local_path"])
                manifest_pages_dir = Path(manifest_result["visual"]["pages_dir"])
                manifest_page_path = Path(manifest_result["visual"]["page_resources"][0]["local_path"])
                render_all_page_path = Path(render_all_result["rendered_pages"][0]["local_path"])
                assert render_resource_path.is_file()
                assert manifest_pages_dir.is_dir()
                assert manifest_page_path.is_file()
                assert render_all_page_path.is_file()
                clear_result, _ = run_workbench_operation(
                    source,
                    operation="clear_render_cache",
                    artifact={"artifact_id": "a1", "binary_hash": "workbench-render-hash"},
                )
    assert render_resources == {}
    assert render_result["render"]["resource_uri"].startswith("/api/docx/render/resource")
    assert render_result["resources"][0]["mime_type"] == "image/png"
    assert manifest_result["visual"]["status"] == "complete"
    assert clear_result["stats"]["removed_dirs"] == 1


def test_workbench_router_persists_results_and_resources():
    source = _add_content_controls(_make_docx_bytes("Pagina literal {PAGE}"))
    with TemporaryDirectory() as tmp_dir:
        root_dir = Path(tmp_dir)
        with _patched_artifact_store(root_dir):
            stored = _store_docx_artifact(source, root_dir=root_dir)
            artifact_id = stored["artifact_id"]

            response = asyncio.run(docx_router.run_docx_workbench({"artifact_id": artifact_id, "operation": "content_controls_wrap"}))
            payload = json.loads(response.body.decode("utf-8"))
            assert payload["operation"] == "content_controls_wrap"
            assert payload["variant"]["parent_artifact_id"] == artifact_id
            assert payload["variant"]["resource_uri"]

            loaded_response = asyncio.run(docx_router.get_docx_workbench_result_route(workbench_id=payload["workbench_id"]))
            loaded = json.loads(loaded_response.body.decode("utf-8"))
            assert loaded["workbench_id"] == payload["workbench_id"]

            resource_response = asyncio.run(
                docx_router.get_docx_workbench_resource_route(
                    workbench_id=payload["workbench_id"],
                    name=payload["variant"]["filename"],
                )
            )
            assert resource_response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            assert resource_response.body


def test_docx_resource_helpers_reject_traversal_ids():
    with TemporaryDirectory() as tmp_dir:
        root_dir = Path(tmp_dir)
        with _patched_artifact_store(root_dir):
            outside_dir = root_dir / "outside"
            outside_dir.mkdir()
            (outside_dir / "summary.json").write_text('{"leak": true}', encoding="utf-8")
            (outside_dir / "secret.docx").write_bytes(b"secret")

            assert docx_artifacts.get_docx_workbench_result("../outside") is None
            assert docx_artifacts.get_docx_workbench_resource("../outside", "secret.docx") is None

        render_dir = root_dir / "render-cache"
        render_dir.mkdir(exist_ok=True)
        outside_render_dir = root_dir / "render-outside"
        outside_render_dir.mkdir()
        (outside_render_dir / "document.pdf").write_bytes(b"secret")
        with patch.multiple(docx_render_cache, DOCX_RENDER_CACHE_DIR=render_dir):
            assert docx_render_cache.get_docx_render_resource("../render-outside", "document.pdf") is None


@contextmanager
def pytest_raises_http_404():
    try:
        yield
    except HTTPException as exc:
        assert exc.status_code == 404
    else:
        raise AssertionError("Expected HTTPException 404")
