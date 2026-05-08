import asyncio
import base64
import io
import json
import os
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.routers import docx as docx_router_module
from app.services import docx_artifacts, notebook_service
from app.services.docx_artifacts import detect_docx_body_is_empty
from app.services.jupyter_kernel import jupyter_kernel_manager
from app.services.notebook_service import (
    _build_notebook_instrumented_code,
    _process_notebook_cell_execution,
    _retrieve_docx_via_stdout,
    set_kernel_docx_source,
)
from librerias_propias.docx_builder.api import build_doc
from librerias_propias.docx_builder.session import get_session, reset_session_cache


def _make_docx_b64(*paragraphs: str) -> str:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    buffer = io.BytesIO()
    doc.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def _make_header_footer_only_docx_b64(*, header_text: str = "Header", footer_text: str = "Footer") -> str:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = header_text
    section.footer.paragraphs[0].text = footer_text
    buffer = io.BytesIO()
    doc.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def _docx_has_header_footer_refs(docx_bytes: bytes) -> bool:
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return any(
        sect.find("w:headerReference", namespaces) is not None
        or sect.find("w:footerReference", namespaces) is not None
        for sect in root.findall(".//w:sectPr", namespaces)
    )


def _read_docx_body_text(docx_bytes: bytes) -> str:
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return "".join(node.text or "" for node in root.findall(".//w:body//w:t", namespaces))


class TestDocxEmptyArtifacts:
    def _patch_artifact_store(self, root_dir: Path):
        blobs_dir = root_dir / "blobs"
        blobs_dir.mkdir(parents=True, exist_ok=True)
        manifests_dir = root_dir / "manifests"
        manifests_dir.mkdir(parents=True, exist_ok=True)
        index_path = root_dir / "index.json"
        return patch.multiple(
            docx_artifacts,
            DOCX_ARTIFACT_ROOT=root_dir,
            DOCX_ARTIFACT_BLOBS_DIR=blobs_dir,
            DOCX_ARTIFACT_MANIFESTS_DIR=manifests_dir,
            DOCX_ARTIFACT_INDEX_PATH=index_path,
        )

    def _reset_artifact_store(self):
        with docx_artifacts._ARTIFACT_LOCK:
            docx_artifacts._ARTIFACTS = None
            docx_artifacts._PROVENANCE_INDEX = None
            docx_artifacts._LAST_CLEANUP_AT = 0.0
            docx_artifacts._initialize_docx_artifact_store()

    def test_empty_artifacts_remain_in_history_but_latest_skips_them(self):
        valid_payload = _make_docx_b64("Documento válido")
        empty_payload = _make_docx_b64()
        whitespace_payload = _make_docx_b64("   ")
        with TemporaryDirectory() as tmp_dir:
            root_dir = Path(tmp_dir)
            with self._patch_artifact_store(root_dir):
                self._reset_artifact_store()
                valid = docx_artifacts.store_docx_artifact(
                    valid_payload,
                    filename="valid.docx",
                    docx_hash="hash-valid",
                    source_kind="notebook",
                    source_path="C:/workspace/demo.ipynb",
                    kernel_id="kernel-1",
                    execution_id="exec-valid",
                )
                whitespace = docx_artifacts.store_docx_artifact(
                    whitespace_payload,
                    filename="whitespace.docx",
                    docx_hash="hash-whitespace",
                    source_kind="notebook",
                    source_path="C:/workspace/demo.ipynb",
                    kernel_id="kernel-1",
                    execution_id="exec-whitespace",
                )
                empty = docx_artifacts.store_docx_artifact(
                    empty_payload,
                    filename="empty.docx",
                    docx_hash="hash-empty",
                    source_kind="notebook",
                    source_path="C:/workspace/demo.ipynb",
                    kernel_id="kernel-1",
                    execution_id="exec-empty",
                )

                history = docx_artifacts.list_docx_artifacts(source_path="C:/workspace/demo.ipynb", limit=10)
                assert len(history) == 3
                assert history[0]["artifact_id"] == empty["artifact_id"]
                assert history[0]["docx_is_empty"] is True
                assert history[0]["docx_warning"]
                assert history[1]["artifact_id"] == whitespace["artifact_id"]
                assert history[1]["docx_is_empty"] is True
                assert history[1]["docx_warning"]
                latest = docx_artifacts.get_latest_docx_artifact(source_path="C:/workspace/demo.ipynb")
                assert latest is not None
                assert latest["artifact_id"] == valid["artifact_id"]

                response = asyncio.run(
                    docx_router_module.get_docx_history(
                        source_path="C:/workspace/demo.ipynb",
                        kernel_id=None,
                        limit=10,
                    )
                )
                payload = json.loads(response.body)
                assert payload["items"][0]["docx_is_empty"] is True
                assert payload["items"][0]["docx_warning"]


@pytest.mark.parametrize("paragraphs", [(), ("   ",)])
def test_detect_docx_body_is_empty_for_blank_paragraphs(paragraphs):
    payload = _make_docx_b64(*paragraphs)
    assert detect_docx_body_is_empty(base64.b64decode(payload)) is True

    def test_header_footer_only_docx_is_not_treated_as_empty_latest(self):
        valid_payload = _make_docx_b64("Documento con body")
        header_footer_only_payload = _make_header_footer_only_docx_b64(
            header_text="Encabezado contractual",
            footer_text="Pie contractual",
        )
        with TemporaryDirectory() as tmp_dir:
            root_dir = Path(tmp_dir)
            with self._patch_artifact_store(root_dir):
                self._reset_artifact_store()
                older = docx_artifacts.store_docx_artifact(
                    valid_payload,
                    filename="older.docx",
                    docx_hash="hash-older",
                    source_kind="notebook",
                    source_path="C:/workspace/demo.ipynb",
                    kernel_id="kernel-1",
                    execution_id="exec-older",
                )
                newer = docx_artifacts.store_docx_artifact(
                    header_footer_only_payload,
                    filename="header-footer-only.docx",
                    docx_hash="hash-header-footer",
                    source_kind="notebook",
                    source_path="C:/workspace/demo.ipynb",
                    kernel_id="kernel-1",
                    execution_id="exec-header-footer",
                )

                history = docx_artifacts.list_docx_artifacts(source_path="C:/workspace/demo.ipynb", limit=10)
                assert len(history) == 2
                assert history[0]["artifact_id"] == newer["artifact_id"]
                assert history[0]["docx_is_empty"] is False

                latest = docx_artifacts.get_latest_docx_artifact(source_path="C:/workspace/demo.ipynb")
                assert latest is not None
                assert latest["artifact_id"] == newer["artifact_id"]
                assert latest["artifact_id"] != older["artifact_id"]


def test_detect_docx_body_is_not_empty_when_header_footer_have_content():
    payload = _make_header_footer_only_docx_b64(
        header_text="Encabezado visible",
        footer_text="Pie visible",
    )
    assert detect_docx_body_is_empty(base64.b64decode(payload)) is False


def test_notebook_docx_preamble_injects_host_api_into_builtins():
    instrumented = _build_notebook_instrumented_code(
        source_code="print('ok')",
        cell_id="cell-builtins",
        cell_index=1,
        emit_docx=True,
    )
    assert "import builtins as __insp_docx_builtins" in instrumented
    assert "setattr(__insp_docx_builtins, __name, globals()[__name])" in instrumented
    assert "doc_finalize" in instrumented


def test_doc_finalize_is_public_from_math_to_docx():
    from librerias_propias.math_to_docx import doc_finalize

    assert callable(doc_finalize)


def test_notebook_instrumentation_does_not_export_docx_inline_in_finally():
    instrumented = _build_notebook_instrumented_code(
        source_code="print('ok')",
        cell_id="cell-no-inline-export",
        cell_index=1,
        emit_docx=True,
    )
    assert "__INSP_NOTEBOOK_DOCX =" not in instrumented
    assert "__INSP_NOTEBOOK_DOCX_PROVENANCE =" not in instrumented
    assert "doc_export(format='docx')" not in instrumented
    assert "doc_export_provenance()" not in instrumented


def test_export_restores_template_header_footer_refs_after_raw_add_section():
    reset_session_cache()
    try:
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            template_doc = Document()
            template_section = template_doc.sections[0]
            template_section.header.paragraphs[0].text = "Encabezado EFE"
            template_section.footer.paragraphs[0].text = "Pie EFE"
            template_doc.save(template_path)

            namespace = {}
            session = get_session(namespace)
            session.set_template_path(str(template_path))
            session.reset(hard=True)

            with build_doc(namespace=namespace, block_id="cell-section", order=10) as builder:
                builder.text("Primera página")
                builder.document.add_section(WD_SECTION_START.NEW_PAGE)
                builder.text("Segunda página")

            exported = session.serialize_docx_bytes()
            assert _docx_has_header_footer_refs(exported) is True
    finally:
        reset_session_cache()


def test_template_body_is_cleared_when_template_is_used_as_runtime_base():
    reset_session_cache()
    try:
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template-with-body.docx"
            template_doc = Document()
            template_doc.sections[0].header.paragraphs[0].text = "Encabezado plantilla"
            template_doc.add_paragraph("TEMPLATE BODY PLACEHOLDER")
            template_doc.save(template_path)

            namespace = {}
            session = get_session(namespace)
            session.set_template_path(str(template_path))
            session.reset(hard=True)

            with build_doc(namespace=namespace, block_id="cell-generated", order=10) as builder:
                builder.text("GENERATED BODY CONTENT")

            exported = session.serialize_docx_bytes()
            body_text = _read_docx_body_text(exported)
            assert "GENERATED BODY CONTENT" in body_text
            assert "TEMPLATE BODY PLACEHOLDER" not in body_text
            assert _docx_has_header_footer_refs(exported) is True
    finally:
        reset_session_cache()


def test_math_to_docx_cell_tracking_uses_notebook_namespace_after_hard_reset():
    reset_session_cache()
    try:
        namespace = {}
        exec(
            """
from librerias_propias.math_to_docx import (
    build_doc,
    doc_begin,
    doc_end,
    doc_finish_cell,
    doc_reset,
    doc_start_cell,
    get_session,
)

doc_start_cell("notebook-cell-1")
doc_begin(block_id="auto-cell-1", order=1, notebook_cell_id="notebook-cell-1")
doc_reset(hard=True)
with build_doc(block_id="body-cell-1", order=10, notebook_cell_id="notebook-cell-1") as builder:
    builder.text("Generated after reset")
doc_end()
doc_finish_cell("notebook-cell-1")
SESSION_DIAG = {
    "items": sorted((get_session().ns.get("__DOCX_CELL_ITEMS") or {}).keys()),
    "groups": get_session().ns.get("__DOCX_NOTEBOOK_CELL_GROUPS") or {},
}
""",
            namespace,
            namespace,
        )

        session = get_session(namespace)
        assert "body-cell-1" in namespace["SESSION_DIAG"]["items"]
        assert namespace["SESSION_DIAG"]["groups"]["notebook-cell-1"] == ["body-cell-1"]
        assert "Generated after reset" in _read_docx_body_text(session.serialize_docx_bytes())
    finally:
        reset_session_cache()


def test_process_notebook_cell_execution_defers_docx_classification_to_document_job():
    kernel_id = "kernel-empty-docx"
    previous_docx = _make_docx_b64("Último documento válido")
    notebook_service.notebook_last_docx_b64[kernel_id] = previous_docx
    notebook_service.notebook_docx_hash[kernel_id] = "prev-hash"

    async def _run():
        with patch.object(
            notebook_service.jupyter_kernel_manager,
            "execute_cell",
            new=AsyncMock(return_value=([], 1, {}, {"doc_b64": _make_docx_b64()})),
        ):
            with patch.object(notebook_service, "store_docx_artifact") as artifact_mock:
                with patch.object(notebook_service, "store_docx_base64") as legacy_mock:
                    response = await _process_notebook_cell_execution(
                        kernel_id=kernel_id,
                        cell_id="cell-empty",
                        instrumented_code="print('ok')",
                        source_code="print('ok')",
                        enable_tracing=False,
                        emit_docx=True,
                        skip_pdf=False,
                        on_iopub=AsyncMock(),
                        websocket=object(),
                        execution_id="exec-empty-docx",
                    )
            return response, artifact_mock, legacy_mock

    try:
        response, artifact_mock, legacy_mock = asyncio.run(_run())
        assert response["docx_is_empty"] is False
        assert response["docx_ref"] is None
        assert response["docx_artifact_id"] is None
        assert response["pdf_attempted"] is True
        document_job = response["_document_job_request"]
        assert document_job is not None
        assert document_job.execution_id == "exec-empty-docx"
        assert document_job.reason == "terminal"
        assert notebook_service.notebook_last_docx_b64[kernel_id] == previous_docx
        assert notebook_service.notebook_docx_hash[kernel_id] == "prev-hash"
        assert response["docx_warnings"] is None
        assert artifact_mock.call_count == 0
        assert legacy_mock.call_count == 0
    finally:
        notebook_service.notebook_last_docx_b64.pop(kernel_id, None)
        notebook_service.notebook_docx_hash.pop(kernel_id, None)


@pytest.mark.integration
@pytest.mark.asyncio
async def test_reexecution_clears_stale_user_docx_blocks_when_second_run_emits_no_docx():
    if jupyter_kernel_manager is None:
        pytest.skip("Kernel manager no disponible")

    try:
        kernel_id = await jupyter_kernel_manager.start_kernel("python3")
    except RuntimeError as exc:
        pytest.skip(f"Kernel no disponible para integración: {exc}")

    async def _noop_iopub(*args, **kwargs):
        return None

    try:
        set_kernel_docx_source(
            kernel_id,
            source_path="C:/workspace/reexecution.ipynb",
            source_kind="notebook",
        )
        first_source = (
            "from librerias_propias.math_to_docx import build_doc\n"
            "with build_doc(order=10) as doc:\n"
            "    doc.text('Primera versión con DOCX')\n"
        )
        first_payload = await _process_notebook_cell_execution(
            kernel_id=kernel_id,
            cell_id="cell-reexec",
            instrumented_code=_build_notebook_instrumented_code(
                source_code=first_source,
                cell_id="cell-reexec",
                cell_index=10,
                emit_docx=True,
            ),
            source_code=first_source,
            enable_tracing=False,
            emit_docx=True,
            skip_pdf=False,
            on_iopub=_noop_iopub,
            websocket=object(),
            execution_id="exec-first",
        )
        assert first_payload["docx_hash"]

        second_source = "print('segunda corrida sin contenido docx')"
        second_payload = await _process_notebook_cell_execution(
            kernel_id=kernel_id,
            cell_id="cell-reexec",
            instrumented_code=_build_notebook_instrumented_code(
                source_code=second_source,
                cell_id="cell-reexec",
                cell_index=10,
                emit_docx=True,
            ),
            source_code=second_source,
            enable_tracing=False,
            emit_docx=True,
            skip_pdf=False,
            on_iopub=_noop_iopub,
            websocket=object(),
            execution_id="exec-second",
        )
        assert second_payload["docx_is_empty"] is True
        assert second_payload["docx_ref"] is None

        exported_b64 = await _retrieve_docx_via_stdout(kernel_id)
        assert exported_b64 is not None
        assert detect_docx_body_is_empty(base64.b64decode(exported_b64)) is True
    finally:
        await jupyter_kernel_manager.shutdown_kernel(kernel_id)
