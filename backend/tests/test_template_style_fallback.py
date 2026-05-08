"""Regression tests for template style precedence and safe DOCX fallbacks."""

from __future__ import annotations

import base64
import copy
import io
import os
import sys
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.shared import qn  # type: ignore
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    Document = None
    WD_STYLE_TYPE = None
    OxmlElement = None
    Pt = None
    HAS_DOCX = False

from librerias_propias.docx_builder.api import build_doc
from librerias_propias.docx_builder.session import get_session, reset_session_cache
from app.services import template_extract, template_service


DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
A_NS = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
OOXML_NS = {**DOCX_NS, **A_NS}
PNG_DOT_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0nQAAAAASUVORK5CYII="


def _append_simple_seq_field(paragraph, sequence_name: str, display_value: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f"SEQ {sequence_name} \\* ARABIC")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = display_value
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _find_style_element(styles_root: ET.Element, style_name: str) -> ET.Element | None:
    for style_elem in styles_root.findall("w:style", DOCX_NS):
        name_elem = style_elem.find("w:name", DOCX_NS)
        if name_elem is not None and name_elem.get(qn("w:val")) == style_name:
            return style_elem
    return None


def _remove_style_from_template(
    template_path: str,
    style_name: str,
    *,
    strip_numbering: bool = False,
    strip_numpr_only: bool = False,
) -> None:
    with open(template_path, "rb") as fh:
        raw = fh.read()
    src = io.BytesIO(raw)
    dst = io.BytesIO()

    with zipfile.ZipFile(src, "r") as zin:
        styles_root = ET.fromstring(zin.read("word/styles.xml"))
        style_elem = _find_style_element(styles_root, style_name)
        abstract_num_id = None
        num_id = None
        if style_elem is not None:
            p_pr = style_elem.find("w:pPr", DOCX_NS)
            num_pr = p_pr.find("w:numPr", DOCX_NS) if p_pr is not None else None
            if num_pr is not None:
                num_id_elem = num_pr.find("w:numId", DOCX_NS)
                num_id = num_id_elem.get(qn("w:val")) if num_id_elem is not None else None
            if strip_numpr_only:
                if p_pr is not None and num_pr is not None:
                    p_pr.remove(num_pr)
            else:
                styles_root.remove(style_elem)

        numbering_root = None
        if strip_numbering and num_id:
            try:
                numbering_root = ET.fromstring(zin.read("word/numbering.xml"))
            except KeyError:
                numbering_root = None
            if numbering_root is not None:
                num_elem = numbering_root.find(f".//w:num[@w:numId='{num_id}']", DOCX_NS)
                if num_elem is not None:
                    abstract_num_elem = num_elem.find("w:abstractNumId", DOCX_NS)
                    abstract_num_id = (
                        abstract_num_elem.get(qn("w:val")) if abstract_num_elem is not None else None
                    )
                    numbering_root.remove(num_elem)
                if abstract_num_id:
                    abstract_elem = numbering_root.find(
                        f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']",
                        DOCX_NS,
                    )
                    if abstract_elem is not None:
                        numbering_root.remove(abstract_elem)

        updated_styles = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
        updated_numbering = (
            ET.tostring(numbering_root, encoding="utf-8", xml_declaration=True)
            if numbering_root is not None
            else None
        )

        with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, updated_styles)
                    continue
                if updated_numbering is not None and item.filename == "word/numbering.xml":
                    zout.writestr(item, updated_numbering)
                    continue
                zout.writestr(item, zin.read(item.filename))

    with open(template_path, "wb") as fh:
        fh.write(dst.getvalue())


def _patch_theme_fonts_and_docdefaults(
    template_path: str,
    *,
    major_latin: str,
    minor_latin: str,
    doc_defaults_theme: str = "minorHAnsi",
) -> None:
    with open(template_path, "rb") as fh:
        raw = fh.read()
    src = io.BytesIO(raw)
    dst = io.BytesIO()

    with zipfile.ZipFile(src, "r") as zin:
        styles_root = ET.fromstring(zin.read("word/styles.xml"))
        theme_root = ET.fromstring(zin.read("word/theme/theme1.xml"))

        major_font = theme_root.find(".//a:fontScheme/a:majorFont/a:latin", OOXML_NS)
        minor_font = theme_root.find(".//a:fontScheme/a:minorFont/a:latin", OOXML_NS)
        if major_font is not None:
            major_font.set("typeface", major_latin)
        if minor_font is not None:
            minor_font.set("typeface", minor_latin)

        doc_defaults = styles_root.find("w:docDefaults", DOCX_NS)
        if doc_defaults is None:
            doc_defaults = ET.SubElement(styles_root, qn("w:docDefaults"))
        r_pr_default = doc_defaults.find("w:rPrDefault", DOCX_NS)
        if r_pr_default is None:
            r_pr_default = ET.SubElement(doc_defaults, qn("w:rPrDefault"))
        r_pr = r_pr_default.find("w:rPr", DOCX_NS)
        if r_pr is None:
            r_pr = ET.SubElement(r_pr_default, qn("w:rPr"))
        r_fonts = r_pr.find("w:rFonts", DOCX_NS)
        if r_fonts is None:
            r_fonts = ET.SubElement(r_pr, qn("w:rFonts"))

        for attr_name in ("ascii", "hAnsi", "cs", "eastAsia", "asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
            r_fonts.attrib.pop(qn(f"w:{attr_name}"), None)

        theme_prefix = "major" if str(doc_defaults_theme).lower().startswith("major") else "minor"
        r_fonts.set(qn("w:asciiTheme"), doc_defaults_theme)
        r_fonts.set(qn("w:hAnsiTheme"), doc_defaults_theme)
        r_fonts.set(qn("w:csTheme"), f"{theme_prefix}Bidi")
        r_fonts.set(qn("w:eastAsiaTheme"), f"{theme_prefix}EastAsia")

        updated_styles = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
        updated_theme = ET.tostring(theme_root, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, updated_styles)
                    continue
                if item.filename == "word/theme/theme1.xml":
                    zout.writestr(item, updated_theme)
                    continue
                zout.writestr(item, zin.read(item.filename))

    with open(template_path, "wb") as fh:
        fh.write(dst.getvalue())


def _patch_style_legacy_font_and_alt_name(
    template_path: str,
    *,
    style_name: str,
    style_id: str,
    font_name: str,
    alt_name: str,
) -> None:
    with open(template_path, "rb") as fh:
        raw = fh.read()
    src = io.BytesIO(raw)
    dst = io.BytesIO()

    with zipfile.ZipFile(src, "r") as zin:
        styles_root = ET.fromstring(zin.read("word/styles.xml"))
        style_elem = styles_root.find(f"w:style[@w:styleId='{style_id}']", DOCX_NS)
        if style_elem is None:
            style_elem = _find_style_element(styles_root, style_name)
        if style_elem is None:
            raise AssertionError(f"Missing style {style_name}")
        style_elem.set(qn("w:styleId"), style_id)

        r_pr = style_elem.find("w:rPr", DOCX_NS)
        if r_pr is None:
            r_pr = ET.SubElement(style_elem, qn("w:rPr"))
        r_fonts = r_pr.find("w:rFonts", DOCX_NS)
        if r_fonts is None:
            r_fonts = ET.SubElement(r_pr, qn("w:rFonts"))
        for attr_name in ("ascii", "hAnsi", "cs", "eastAsia"):
            r_fonts.set(qn(f"w:{attr_name}"), font_name)
        for attr_name in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
            r_fonts.attrib.pop(qn(f"w:{attr_name}"), None)

        try:
            font_table_root = ET.fromstring(zin.read("word/fontTable.xml"))
        except KeyError:
            font_table_root = ET.Element(qn("w:fonts"))

        font_elem = font_table_root.find(f"w:font[@w:name='{font_name}']", DOCX_NS)
        if font_elem is None:
            font_elem = ET.SubElement(font_table_root, qn("w:font"))
            font_elem.set(qn("w:name"), font_name)
        alt_elem = font_elem.find("w:altName", DOCX_NS)
        if alt_elem is None:
            alt_elem = ET.SubElement(font_elem, qn("w:altName"))
        alt_elem.set(qn("w:val"), alt_name)

        updated_styles = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
        updated_font_table = ET.tostring(font_table_root, encoding="utf-8", xml_declaration=True)

        wrote_font_table = False
        with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, updated_styles)
                    continue
                if item.filename == "word/fontTable.xml":
                    zout.writestr(item, updated_font_table)
                    wrote_font_table = True
                    continue
                zout.writestr(item, zin.read(item.filename))
            if not wrote_font_table:
                zout.writestr("word/fontTable.xml", updated_font_table)

    with open(template_path, "wb") as fh:
        fh.write(dst.getvalue())


def _get_style_numpr(style_obj):
    p_pr = style_obj.element.find(qn("w:pPr"))
    if p_pr is None:
        return None
    return p_pr.find(qn("w:numPr"))


def _read_style_rfonts_attrs(
    docx_path: str,
    style_name: str,
    style_id: str | None = None,
) -> dict[str, str | None]:
    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_root = ET.fromstring(zf.read("word/styles.xml"))

    style_elem = None
    if style_id:
        style_elem = styles_root.find(f"w:style[@w:styleId='{style_id}']", DOCX_NS)
    if style_elem is None:
        style_elem = _find_style_element(styles_root, style_name)
    if style_elem is None:
        return {}

    r_fonts = style_elem.find("w:rPr/w:rFonts", DOCX_NS)
    if r_fonts is None:
        return {}

    attrs = {}
    for attr_name in ("ascii", "hAnsi", "cs", "eastAsia", "asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        attrs[attr_name] = r_fonts.get(qn(f"w:{attr_name}"))
    return attrs


def _read_doc_defaults_props(docx_path: str) -> dict[str, dict[str, str | None]]:
    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_root = ET.fromstring(zf.read("word/styles.xml"))

    doc_defaults = styles_root.find("w:docDefaults", DOCX_NS)
    if doc_defaults is None:
        return {"rfonts": {}, "spacing": {}, "jc": {}, "ind": {}}

    r_fonts = doc_defaults.find("w:rPrDefault/w:rPr/w:rFonts", DOCX_NS)
    spacing = doc_defaults.find("w:pPrDefault/w:pPr/w:spacing", DOCX_NS)
    jc = doc_defaults.find("w:pPrDefault/w:pPr/w:jc", DOCX_NS)
    ind = doc_defaults.find("w:pPrDefault/w:pPr/w:ind", DOCX_NS)

    def _attrs(elem: ET.Element | None, names: tuple[str, ...]) -> dict[str, str | None]:
        if elem is None:
            return {}
        return {
            attr_name: elem.get(qn(f"w:{attr_name}"))
            for attr_name in names
        }

    return {
        "rfonts": _attrs(r_fonts, ("ascii", "hAnsi", "cs", "eastAsia", "asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme")),
        "spacing": _attrs(spacing, ("before", "after", "line", "lineRule")),
        "jc": _attrs(jc, ("val",)),
        "ind": _attrs(ind, ("left", "right", "firstLine", "hanging")),
    }


def _read_paragraph_style_ids(docx_bytes: bytes) -> list[str | None]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        document_root = ET.fromstring(zf.read("word/document.xml"))

    style_ids: list[str | None] = []
    for paragraph in document_root.findall(".//w:body/w:p", DOCX_NS):
        p_style = paragraph.find("w:pPr/w:pStyle", DOCX_NS)
        style_ids.append(p_style.get(qn("w:val")) if p_style is not None else None)
    return style_ids


@unittest.skipUnless(HAS_DOCX, "python-docx not available")
class TestTemplateStyleFallback(unittest.TestCase):
    def setUp(self):
        reset_session_cache()
        self.tempdir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.tempdir.cleanup()
        reset_session_cache()

    def _template_path(self, name: str = "template.docx") -> str:
        return os.path.join(self.tempdir.name, name)

    def _save_blank_template(self, name: str = "template.docx") -> str:
        path = self._template_path(name)
        doc = Document()
        doc.save(path)
        return path

    def _load_session_with_template(self, template_path: str):
        namespace = {}
        session = get_session(namespace)
        session.set_template_path(template_path)
        session.reset(hard=True)
        return namespace, session

    def test_existing_template_style_is_not_overwritten(self):
        path = self._template_path("custom-code.docx")
        doc = Document()
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(14)
        doc.save(path)

        _, session = self._load_session_with_template(path)

        style = session.doc.styles["Code"]
        self.assertEqual(style.font.name, "Courier New")
        self.assertEqual(style.font.size, Pt(14))

    def test_missing_list_bullet_style_is_recreated_with_numbering(self):
        path = self._save_blank_template("missing-list-bullet.docx")
        _remove_style_from_template(path, "List Bullet", strip_numbering=True)

        namespace, session = self._load_session_with_template(path)
        style = session.doc.styles["List Bullet"]
        num_pr = _get_style_numpr(style)

        self.assertIsNotNone(num_pr, "Fallback List Bullet style must carry numPr")
        num_id_elem = num_pr.find(qn("w:numId"))
        self.assertIsNotNone(num_id_elem, "Fallback List Bullet style must carry numId")
        num_id = num_id_elem.get(qn("w:val"))
        self.assertTrue(num_id)

        with build_doc(order=1, namespace=namespace, block_id="cell-list-fallback") as builder:
            builder.list(["Item A", "Item B"])

        raw = base64.b64decode(session.export_docx_base64())
        with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
            styles_xml = ET.fromstring(zf.read("word/styles.xml"))
            numbering_xml = ET.fromstring(zf.read("word/numbering.xml"))

        exported_style = _find_style_element(styles_xml, "List Bullet")
        self.assertIsNotNone(exported_style)
        exported_num_pr = exported_style.find("w:pPr/w:numPr", DOCX_NS)
        self.assertIsNotNone(exported_num_pr)
        exported_num_id_elem = exported_num_pr.find("w:numId", DOCX_NS)
        self.assertIsNotNone(exported_num_id_elem)
        exported_num_id = exported_num_id_elem.get(qn("w:val"))
        self.assertIsNotNone(
            numbering_xml.find(f".//w:num[@w:numId='{exported_num_id}']", DOCX_NS),
            "numbering.xml must contain the fallback list definition",
        )

    def test_missing_code_style_is_synthesized_for_template(self):
        path = self._save_blank_template("missing-code.docx")
        namespace, session = self._load_session_with_template(path)

        with build_doc(order=2, namespace=namespace, block_id="cell-code-fallback") as builder:
            builder.code("print('hello')")

        paragraph = next((p for p in session.doc.paragraphs if "hello" in p.text), None)
        self.assertIsNotNone(paragraph)
        self.assertEqual(paragraph.style.name, "Code")

    def test_broken_list_style_without_numpr_uses_visible_bullet_fallback(self):
        path = self._save_blank_template("broken-list-style.docx")
        _remove_style_from_template(path, "List Bullet", strip_numbering=True, strip_numpr_only=True)

        namespace, session = self._load_session_with_template(path)

        with build_doc(order=3, namespace=namespace, block_id="cell-broken-list") as builder:
            builder.list(["Visible bullet"])

        paragraph = next((p for p in session.doc.paragraphs if "Visible bullet" in p.text), None)
        self.assertIsNotNone(paragraph)
        self.assertTrue(
            paragraph.text.startswith("\u2022 "),
            "Broken list style should fall back to an explicit visible bullet marker",
        )

    def test_broken_numbered_list_style_uses_visible_number_fallback(self):
        path = self._save_blank_template("broken-number-style.docx")
        _remove_style_from_template(path, "List Number", strip_numbering=True, strip_numpr_only=True)

        namespace, session = self._load_session_with_template(path)

        with build_doc(order=4, namespace=namespace, block_id="cell-broken-olist") as builder:
            builder.list(["Uno", "Dos"], ordered=True)

        paragraphs = [p for p in session.doc.paragraphs if p.text]
        self.assertTrue(any(p.text.startswith("1. Uno") for p in paragraphs))
        self.assertTrue(any(p.text.startswith("2. Dos") for p in paragraphs))

    def test_hard_reset_reapplies_missing_heading_style(self):
        path = self._save_blank_template("missing-heading.docx")
        _remove_style_from_template(path, "Heading 1")

        namespace = {}
        session = get_session(namespace)
        session.set_template_path(path)
        session.reset(hard=True)

        with build_doc(order=5, namespace=namespace, block_id="cell-heading-fallback") as builder:
            builder.heading("Titulo fallback", level=1)

        paragraph = next((p for p in session.doc.paragraphs if "Titulo fallback" in p.text), None)
        self.assertIsNotNone(paragraph)
        self.assertEqual(paragraph.style.name, "Heading 1")

    def test_builder_required_style_defaults_inherit_template_default_font_for_missing_code_style(self):
        path = self._template_path("century-gothic-template.docx")
        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Century Gothic"
        normal_style.font.size = Pt(11)
        doc.save(path)

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        builder_defaults = template_service.build_builder_required_style_defaults(extracted)
        self.assertEqual(builder_defaults.get("Code", {}).get("font_name"), "Century Gothic")

    def test_extract_styles_from_docx_includes_template_font_catalog_and_default_font(self):
        path = self._template_path("template-font-catalog.docx")
        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Century Gothic"
        doc.save(path)

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        self.assertIn("Century Gothic", extracted.get("font_catalog", []))
        self.assertEqual((extracted.get("default_font") or {}).get("name"), "Century Gothic")

    def test_legacy_font_altname_is_preserved_and_can_be_replaced(self):
        path = self._template_path("legacy-font-altname.docx")
        doc = Document()
        body_style = doc.styles["Body Text"]
        body_style.font.name = "CG Times (W1)"
        body_style.font.size = Pt(13)
        doc.add_paragraph("Texto base", style="Body Text")
        doc.save(path)
        _patch_style_legacy_font_and_alt_name(
            path,
            style_name="Body Text",
            style_id="Textoindependiente",
            font_name="CG Times (W1)",
            alt_name="Times New Roman",
        )

        docx_bytes = Path(path).read_bytes()
        extracted = template_service.extract_styles_from_docx(docx_bytes)
        extracted_body = next(
            style for style in extracted.get("styles", [])
            if style.get("style_id") == "Textoindependiente" or style.get("name") == "Body Text"
        )
        self.assertEqual(
            (extracted_body.get("xml_font") or {}).get("font_name")
            or (extracted_body.get("xml_font") or {}).get("name"),
            "CG Times (W1)",
        )
        self.assertEqual(
            (extracted_body.get("resolved_font") or {}).get("font_name")
            or (extracted_body.get("resolved_font") or {}).get("name"),
            "CG Times (W1)",
        )
        font_table = ((extracted.get("xml_details") or {}).get("font_table") or {}).get("fonts") or []
        legacy_font = next(font for font in font_table if font.get("name") == "CG Times (W1)")
        self.assertEqual(legacy_font.get("alt_name"), "Times New Roman")

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-legacy-font-altname"
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated = template_service.update_template_style(
                    kernel_id,
                    "Body Text",
                    {
                        "style_id": "Textoindependiente",
                        "font_name": "Arial",
                    },
                )

                rfonts = _read_style_rfonts_attrs(
                    str(template_service._get_template_docx_path(kernel_id)),
                    "Body Text",
                    "Textoindependiente",
                )
                self.assertEqual(rfonts.get("ascii"), "Arial")
                self.assertEqual(rfonts.get("hAnsi"), "Arial")
                self.assertEqual(rfonts.get("cs"), "Arial")
                self.assertEqual(rfonts.get("eastAsia"), "Arial")
                self.assertIsNone(rfonts.get("asciiTheme"))
                self.assertIsNone(rfonts.get("hAnsiTheme"))
                self.assertIsNone(rfonts.get("csTheme"))
                self.assertIsNone(rfonts.get("eastAsiaTheme"))

                updated_body = next(
                    style for style in updated.get("styles", [])
                    if style.get("style_id") == "Textoindependiente" or style.get("name") == "Body Text"
                )
                self.assertEqual(
                    (updated_body.get("resolved_font") or {}).get("font_name")
                    or (updated_body.get("resolved_font") or {}).get("name"),
                    "Arial",
                )
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_extract_styles_from_docx_resolves_theme_minor_docdefaults_font(self):
        path = self._save_blank_template("theme-docdefaults-font.docx")
        _patch_theme_fonts_and_docdefaults(
            path,
            major_latin="Cambria",
            minor_latin="Century Gothic",
            doc_defaults_theme="minorHAnsi",
        )

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        self.assertEqual((extracted.get("default_font") or {}).get("name"), "Century Gothic")
        self.assertEqual((extracted.get("default_font_source") or {}).get("kind"), "theme")
        self.assertEqual((extracted.get("default_font_source") or {}).get("scope"), "docDefaults")

    def test_builder_required_style_defaults_use_theme_major_for_missing_heading_style(self):
        path = self._save_blank_template("theme-heading-fallback.docx")
        _patch_theme_fonts_and_docdefaults(
            path,
            major_latin="Century Gothic",
            minor_latin="Arial",
            doc_defaults_theme="minorHAnsi",
        )
        _remove_style_from_template(path, "Heading 1")

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        builder_defaults = template_service.build_builder_required_style_defaults(extracted)
        self.assertEqual(builder_defaults.get("Heading 1", {}).get("font_name"), "Century Gothic")

    def test_extract_styles_promotes_dominant_body_font_over_theme_default(self):
        path = self._template_path("body-font-hint.docx")
        doc = Document()
        for text in ("Uno", "Dos", "Tres", "Cuatro"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Century Gothic"
        for text in ("Cinco", "Seis"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
        doc.save(path)

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        self.assertEqual((extracted.get("default_font") or {}).get("name"), "Century Gothic")
        self.assertEqual((extracted.get("default_font_source") or {}).get("kind"), "explicit")
        self.assertEqual((extracted.get("default_font_source") or {}).get("scope"), "document")
        normal_style = next(
            style for style in extracted.get("styles", [])
            if style.get("style_id") == "Normal"
        )
        self.assertEqual((normal_style.get("resolved_font") or {}).get("name"), "Century Gothic")
        self.assertEqual((normal_style.get("resolved_font_source") or {}).get("scope"), "document")

    def test_get_template_refreshes_stale_font_metadata_from_docx(self):
        path = self._template_path("stale-body-font-hint.docx")
        doc = Document()
        for text in ("Uno", "Dos", "Tres", "Cuatro"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Century Gothic"
        for text in ("Cinco", "Seis"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
        doc.save(path)

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-stale-font-metadata"
                with open(path, "rb") as fh:
                    docx_bytes = fh.read()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                stale_json = copy.deepcopy(extracted)
                stale_json["default_font"] = {
                    "name": "Calibri",
                    "font_name": "Calibri",
                    "size_pt": 11.0,
                    "font_size_pt": 11.0,
                }
                stale_json["default_font_source"] = {
                    "kind": "theme",
                    "font_name": "Calibri",
                    "theme_key": "minorHAnsi",
                    "scope": "docDefaults",
                }
                stale_json_xml_details = stale_json.get("xml_details") or {}
                if isinstance(stale_json_xml_details, dict):
                    stale_json_xml_details.pop("body_font_hint", None)
                    stale_json["xml_details"] = stale_json_xml_details
                normal_style = next(
                    style for style in stale_json.get("styles", [])
                    if style.get("style_id") == "Normal"
                )
                normal_style["resolved_font"] = {
                    "name": "Calibri",
                    "font_name": "Calibri",
                    "size_pt": 11.0,
                    "font_size_pt": 11.0,
                }
                normal_style["resolved_font_source"] = {
                    "kind": "theme",
                    "font_name": "Calibri",
                    "theme_key": "minorHAnsi",
                    "scope": "docDefaults",
                }
                normal_style["font_source"] = {
                    "kind": "theme",
                    "font_name": "Calibri",
                    "theme_key": "minorHAnsi",
                    "scope": "docDefaults",
                }

                template_service._write_template_files(kernel_id, docx_bytes, stale_json)

                refreshed = template_service.get_template(kernel_id)

                self.assertEqual((refreshed.get("default_font") or {}).get("name"), "Century Gothic")
                self.assertEqual((refreshed.get("default_font_source") or {}).get("scope"), "document")
                refreshed_normal = next(
                    style for style in refreshed.get("styles", [])
                    if style.get("style_id") == "Normal"
                )
                self.assertEqual((refreshed_normal.get("resolved_font") or {}).get("name"), "Century Gothic")
                self.assertEqual((refreshed_normal.get("resolved_font_source") or {}).get("scope"), "document")
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_update_template_style_persists_explicit_font_name_in_all_rfonts_slots(self):
        path = self._save_blank_template("theme-to-explicit-font.docx")
        _patch_theme_fonts_and_docdefaults(
            path,
            major_latin="Cambria",
            minor_latin="Arial",
            doc_defaults_theme="minorHAnsi",
        )

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-explicit-font"
                with open(path, "rb") as fh:
                    docx_bytes = fh.read()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                heading_before = next(
                    style for style in extracted.get("styles", [])
                    if style.get("style_id") == "Heading1" or style.get("name") == "Heading 1"
                )
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated = template_service.update_template_style(
                    kernel_id,
                    "Normal",
                    {"font_name": "Century Gothic"},
                )

                docx_path = template_service._get_template_docx_path(kernel_id)
                rfonts = _read_style_rfonts_attrs(str(docx_path), "Normal", "Normal")

                self.assertEqual(rfonts.get("ascii"), "Century Gothic")
                self.assertEqual(rfonts.get("hAnsi"), "Century Gothic")
                self.assertEqual(rfonts.get("cs"), "Century Gothic")
                self.assertEqual(rfonts.get("eastAsia"), "Century Gothic")
                self.assertIsNone(rfonts.get("asciiTheme"))
                self.assertIsNone(rfonts.get("hAnsiTheme"))
                self.assertIsNone(rfonts.get("csTheme"))
                self.assertIsNone(rfonts.get("eastAsiaTheme"))

                normal_style = next(
                    style for style in updated.get("styles", [])
                    if style.get("style_id") == "Normal"
                )
                self.assertEqual((normal_style.get("resolved_font") or {}).get("name"), "Century Gothic")
                self.assertEqual((normal_style.get("resolved_font_source") or {}).get("kind"), "explicit")
                self.assertEqual(
                    (updated.get(template_service.BUILDER_REQUIRED_STYLE_DEFAULTS_KEY) or {}).get("Code", {}).get("font_name"),
                    "Century Gothic",
                )
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_update_document_defaults_materializes_global_font_and_preserves_explicit_heading(self):
        path = self._save_blank_template("document-defaults-global-font.docx")
        _patch_theme_fonts_and_docdefaults(
            path,
            major_latin="Cambria",
            minor_latin="Calibri",
            doc_defaults_theme="minorHAnsi",
        )

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-document-defaults"
                with open(path, "rb") as fh:
                    docx_bytes = fh.read()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                heading_before = next(
                    style for style in extracted.get("styles", [])
                    if style.get("style_id") == "Heading1" or style.get("name") == "Heading 1"
                )
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated = template_service.update_template_document_defaults(
                    kernel_id,
                    {
                        "font": {"font_name": "Century Gothic"},
                        "paragraph": {
                            "alignment": "JUSTIFY",
                            "space_before_pt": 0,
                            "space_after_pt": 6,
                            "line_spacing": 1.15,
                            "line_spacing_rule": "MULTIPLE",
                            "left_indent_inches": 0.2,
                        },
                    },
                )

                self.assertEqual(
                    (updated.get(template_service.DOCUMENT_DEFAULTS_KEY) or {}).get("font", {}).get("font_name"),
                    "Century Gothic",
                )
                self.assertEqual(
                    (updated.get(template_service.DOCUMENT_DEFAULTS_KEY) or {}).get("font_source", {}).get("scope"),
                    "docDefaults",
                )
                self.assertEqual(
                    (updated.get(template_service.DOCUMENT_DEFAULTS_KEY) or {}).get("paragraph", {}).get("alignment"),
                    "JUSTIFY",
                )

                props = _read_doc_defaults_props(str(template_service._get_template_docx_path(kernel_id)))
                self.assertEqual(props["rfonts"].get("ascii"), "Century Gothic")
                self.assertEqual(props["rfonts"].get("hAnsi"), "Century Gothic")
                self.assertEqual(props["rfonts"].get("cs"), "Century Gothic")
                self.assertEqual(props["rfonts"].get("eastAsia"), "Century Gothic")
                self.assertIsNone(props["rfonts"].get("asciiTheme"))
                self.assertEqual(props["spacing"].get("before"), "0")
                self.assertEqual(props["spacing"].get("after"), "120")
                self.assertEqual(props["spacing"].get("lineRule"), "auto")
                self.assertEqual(props["spacing"].get("line"), "276")
                self.assertEqual(props["jc"].get("val"), "both")
                self.assertEqual(props["ind"].get("left"), "288")

                template_docx_path = template_service.get_template_docx_path(kernel_id)
                namespace = {}
                session = get_session(namespace)
                session.set_template_path(template_docx_path)
                session.reset(hard=True)
                with build_doc(order=7, namespace=namespace, block_id="cell-global-defaults") as builder:
                    builder.text("Texto base generado")
                    builder.heading("Heading explícito", level=1)

                exported_bytes = base64.b64decode(session.export_docx_base64())
                exported = template_service.extract_styles_from_docx(exported_bytes)
                heading_style = next(
                    style for style in exported.get("styles", [])
                    if style.get("style_id") == "Heading1" or style.get("name") == "Heading 1"
                )

                self.assertEqual((exported.get("default_font") or {}).get("name"), "Century Gothic")
                self.assertEqual(
                    (exported.get(template_service.DOCUMENT_DEFAULTS_KEY) or {}).get("font", {}).get("font_name"),
                    "Century Gothic",
                )
                self.assertEqual(
                    (heading_style.get("resolved_font") or {}).get("name"),
                    (heading_before.get("resolved_font") or {}).get("name"),
                )
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_spanish_style_ids_survive_style_update_and_runtime_slots(self):
        path = self._template_path("spanish-style-ids.docx")
        doc = Document()
        body_style = doc.styles["Body Text"]
        body_style.element.set(qn("w:styleId"), "Textoindependiente")
        body_style.font.name = "Century Gothic"
        body_style.font.size = Pt(11)

        heading_style = doc.styles["Heading 1"]
        heading_style.element.set(qn("w:styleId"), "Ttulo1")
        heading_style.font.name = "Century Gothic"
        heading_style.font.size = Pt(16)
        p_pr = heading_style.element.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            heading_style.element.append(p_pr)
        outline_level = OxmlElement("w:outlineLvl")
        outline_level.set(qn("w:val"), "0")
        p_pr.append(outline_level)

        doc.add_paragraph("Texto base", style="Body Text")
        doc.save(path)

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-spanish-style-ids"
                docx_bytes = Path(path).read_bytes()
                extracted = template_service.extract_styles_from_docx(docx_bytes)
                extracted[template_service.SEMANTIC_STYLE_SLOTS_KEY] = {
                    "body": {
                        "slot_name": "body",
                        "category": "body",
                        "selection_key": "body|Textoindependiente|Body Text",
                        "style_id": "Textoindependiente",
                        "style_name": "Body Text",
                        "display_name": "Body Text",
                        "style_type": "paragraph",
                    },
                    "heading_1": {
                        "slot_name": "heading_1",
                        "category": "headings",
                        "selection_key": "headings|Ttulo1|Heading 1",
                        "style_id": "Ttulo1",
                        "style_name": "Heading 1",
                        "display_name": "Heading 1",
                        "style_type": "paragraph",
                    },
                }
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated_body = template_service.update_template_style(
                    kernel_id,
                    "Body Text",
                    {
                        "style_id": "Textoindependiente",
                        "font_name": "Arial",
                        "font_size_pt": 13,
                    },
                )
                updated_heading = template_service.update_template_style(
                    kernel_id,
                    "Heading 1",
                    {
                        "style_id": "Ttulo1",
                        "font_name": "Georgia",
                        "font_size_pt": 22,
                        "color_rgb": "CC0000",
                    },
                )

                template_docx_path = template_service.get_template_docx_path(kernel_id)
                namespace = {}
                session = get_session(namespace)
                session.set_template_path(template_docx_path)
                session.set_template_required_style_defaults(
                    updated_heading.get(template_service.BUILDER_REQUIRED_STYLE_DEFAULTS_KEY) or {}
                )
                session.set_template_semantic_style_slots(
                    updated_heading.get(template_service.SEMANTIC_STYLE_SLOTS_KEY) or {}
                )
                session.reset(hard=True)

                with build_doc(order=8, namespace=namespace, block_id="cell-spanish-slots") as builder:
                    builder.heading("Heading marker", level=1)
                    builder.text("Body marker")

                exported_bytes = base64.b64decode(session.export_docx_base64())
                exported_path = self._template_path("spanish-style-ids-exported.docx")
                Path(exported_path).write_bytes(exported_bytes)

                paragraph_style_ids = _read_paragraph_style_ids(exported_bytes)
                self.assertIn("Ttulo1", paragraph_style_ids)
                self.assertIn("Textoindependiente", paragraph_style_ids)

                body_rfonts = _read_style_rfonts_attrs(
                    exported_path,
                    "Body Text",
                    "Textoindependiente",
                )
                heading_rfonts = _read_style_rfonts_attrs(
                    exported_path,
                    "Heading 1",
                    "Ttulo1",
                )
                self.assertEqual(body_rfonts.get("ascii"), "Arial")
                self.assertEqual(body_rfonts.get("hAnsi"), "Arial")
                self.assertEqual(body_rfonts.get("cs"), "Arial")
                self.assertEqual(body_rfonts.get("eastAsia"), "Arial")
                self.assertEqual(heading_rfonts.get("ascii"), "Georgia")
                self.assertEqual(
                    (updated_body.get(template_service.SEMANTIC_STYLE_SLOTS_KEY) or {}).get("body", {}).get("style_id"),
                    "Textoindependiente",
                )
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_non_font_style_updates_preserve_theme_inheritance(self):
        path = self._save_blank_template("theme-inherited-font.docx")
        _patch_theme_fonts_and_docdefaults(
            path,
            major_latin="Cambria",
            minor_latin="Century Gothic",
            doc_defaults_theme="minorHAnsi",
        )

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-theme-font"
                with open(path, "rb") as fh:
                    docx_bytes = fh.read()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated = template_service.update_template_style(
                    kernel_id,
                    "Normal",
                    {"bold": True},
                )

                self.assertEqual((updated.get("default_font") or {}).get("name"), "Century Gothic")
                self.assertEqual((updated.get("default_font_source") or {}).get("kind"), "theme")
                self.assertEqual((updated.get("default_font_source") or {}).get("scope"), "docDefaults")

                normal_style = next(
                    style for style in updated.get("styles", [])
                    if style.get("style_id") == "Normal"
                )
                self.assertEqual((normal_style.get("resolved_font_source") or {}).get("kind"), "theme")
                self.assertEqual((normal_style.get("resolved_font_source") or {}).get("scope"), "docDefaults")
                self.assertEqual(
                    _read_style_rfonts_attrs(
                        str(template_service._get_template_docx_path(kernel_id)),
                        "Normal",
                        "Normal",
                    ),
                    {},
                )
            finally:
                template_service.TEMPLATE_DIR = original_template_dir


    def test_style_coverage_includes_caption_category(self):
        path = self._save_blank_template("caption-coverage.docx")
        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        coverage = template_service.get_style_coverage(extracted)
        self.assertIn("captions", coverage.get("categories", {}))
        self.assertEqual(coverage["categories"]["captions"][0]["name"], "Caption")

    def test_style_browser_classifies_detected_styles_and_auto_selects_by_score(self):
        extracted = {
            "styles": [
                {"name": "Title", "display_name": "Title", "style_id": "Title", "type": "paragraph", "xml_font": {"font_name": "Arial"}},
                {"name": "Heading 2", "display_name": "Heading 2", "style_id": "Heading2", "type": "paragraph", "resolved_paragraph_format": {"outline_level": 1}},
                {"name": "Normal", "display_name": "Normal", "style_id": "Normal", "type": "paragraph", "xml_paragraph_format": {}},
                {"name": "List Bullet", "display_name": "List Bullet", "style_id": "ListBullet", "type": "paragraph", "list_info": {"list_format": "bullet"}},
                {"name": "Table Grid", "display_name": "Table Grid", "style_id": "TableGrid", "type": "table"},
                {
                    "name": "Code",
                    "display_name": "Code",
                    "style_id": "Code",
                    "type": "paragraph",
                    "resolved_font": {"font_name": "Consolas"},
                },
                {"name": "Caption", "display_name": "Caption", "style_id": "Caption", "type": "paragraph"},
                {"name": "Emphasis", "display_name": "Emphasis", "style_id": "Emphasis", "type": "character"},
            ],
            "document_captions": [
                {"style_id": "Caption", "style_name": "Caption"},
            ],
        }

        browser = template_extract.build_style_browser(extracted)

        self.assertEqual(browser["counts"]["titles"], 1)
        self.assertEqual(browser["counts"]["headings"], 1)
        self.assertEqual(browser["counts"]["body"], 1)
        self.assertEqual(browser["counts"]["lists"], 1)
        self.assertEqual(browser["counts"]["tables"], 1)
        self.assertEqual(browser["counts"]["code"], 1)
        self.assertEqual(browser["counts"]["captions"], 1)
        self.assertEqual(browser["counts"]["other"], 1)
        self.assertEqual(browser["categories"]["other"][0]["display_name"], "Emphasis")
        self.assertEqual(browser["auto_selected"]["titles"], browser["categories"]["titles"][0]["selection_key"])
        self.assertEqual(browser["auto_selected"]["headings"], browser["categories"]["headings"][0]["selection_key"])
        self.assertEqual(browser["auto_selected"]["body"], browser["categories"]["body"][0]["selection_key"])

    def test_style_browser_uses_unique_selection_keys_for_duplicate_names(self):
        extracted = {
            "styles": [
                {"name": "Normal", "display_name": "Normal", "style_id": "Normal", "type": "paragraph"},
                {"name": "Normal", "display_name": "Normal", "style_id": "BodyText", "type": "paragraph"},
            ],
            "document_captions": [],
        }

        browser = template_extract.build_style_browser(extracted)
        selection_keys = [entry["selection_key"] for entry in browser["categories"]["body"]]

        self.assertEqual(len(selection_keys), 2)
        self.assertEqual(len(set(selection_keys)), 2)
        self.assertTrue(selection_keys[0].startswith("body|"))

    def test_build_semantic_style_slots_prefers_body_text_for_body_slot(self):
        extracted = {
            "styles": [
                {"name": "Normal", "display_name": "Normal", "style_id": "Normal", "type": "paragraph"},
                {"name": "Body Text", "display_name": "Body Text", "style_id": "BodyText", "type": "paragraph"},
                {"name": "Heading 1", "display_name": "Heading 1", "style_id": "Heading1", "type": "paragraph"},
                {"name": "Caption", "display_name": "Caption", "style_id": "Caption", "type": "paragraph"},
                {"name": "Code", "display_name": "Code", "style_id": "Code", "type": "paragraph"},
                {"name": "Table Grid", "display_name": "Table Grid", "style_id": "TableGrid", "type": "table"},
            ],
            "document_captions": [{"style_id": "Caption", "style_name": "Caption"}],
        }

        slots = template_extract.enrich_template_metadata(extracted).get(template_service.SEMANTIC_STYLE_SLOTS_KEY) or {}

        self.assertEqual(slots["body"]["style_id"], "BodyText")
        self.assertEqual(slots["body"]["style_name"], "Body Text")

    def test_build_semantic_style_slots_preserves_previous_slot_selection_when_valid(self):
        extracted = {
            "styles": [
                {"name": "Normal", "display_name": "Normal", "style_id": "Normal", "type": "paragraph"},
                {"name": "Body Text", "display_name": "Body Text", "style_id": "BodyText", "type": "paragraph"},
            ],
            "document_captions": [],
        }

        slots = template_extract.build_semantic_style_slots(
            template_extract.enrich_template_metadata(extracted),
            previous_slots={
                "body": {
                    "selection_key": "body|Normal|Normal",
                    "style_id": "Normal",
                    "style_name": "Normal",
                },
            },
        )

        self.assertEqual(slots["body"]["style_id"], "Normal")
        self.assertEqual(slots["body"]["selection_key"], "body|Normal|Normal")

    def test_extract_styles_from_docx_detects_document_captions(self):
        path = self._template_path("document-captions.docx")
        doc = Document()

        table_caption = doc.add_paragraph(style="Caption")
        table_caption.add_run("Tabla ")
        _append_simple_seq_field(table_caption, "Tabla", "1")
        table_caption.add_run(". Resultados")

        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "A"
        table.cell(1, 0).text = "1"

        figure_paragraph = doc.add_paragraph()
        figure_paragraph.add_run().add_picture(io.BytesIO(base64.b64decode(PNG_DOT_B64)))

        figure_caption = doc.add_paragraph(style="Caption")
        figure_caption.add_run("Figura ")
        _append_simple_seq_field(figure_caption, "Figura", "1")
        figure_caption.add_run(". Diagrama")

        doc.save(path)

        with open(path, "rb") as fh:
            extracted = template_service.extract_styles_from_docx(fh.read())

        self.assertEqual(len(extracted.get("document_tables", [])), 1)
        captions = extracted.get("document_captions", [])
        self.assertEqual(len(captions), 2)

        table_info = captions[0]
        self.assertEqual(table_info.get("object_type"), "table")
        self.assertEqual(table_info.get("object_index"), 0)
        self.assertEqual(table_info.get("position"), "before")
        self.assertEqual(table_info.get("style_id"), "Caption")
        self.assertTrue(table_info.get("uses_caption_style"))
        self.assertTrue(table_info.get("has_seq_field"))
        self.assertEqual(table_info.get("sequence_name"), "Tabla")

        figure_info = captions[1]
        self.assertEqual(figure_info.get("object_type"), "figure")
        self.assertEqual(figure_info.get("object_index"), 0)
        self.assertEqual(figure_info.get("position"), "after")
        self.assertEqual(figure_info.get("sequence_name"), "Figura")


if __name__ == "__main__":
    unittest.main()
