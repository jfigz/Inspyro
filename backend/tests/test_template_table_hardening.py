import asyncio
import io
import os
import sys
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import AsyncMock, patch


sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from docx import Document
    HAS_TEST_DOCX = True
except Exception:
    Document = None
    HAS_TEST_DOCX = False

from app.routers import notebook as notebook_router
from app.routers import notebook_template as notebook_template_router
from app.services import template_service


class TestTemplateTableHardening(unittest.TestCase):
    def _read_style_tblpr_tags(self, docx_path: Path, style_id: str) -> list[str]:
        with zipfile.ZipFile(docx_path, "r") as zin:
            styles_xml = zin.read("word/styles.xml")
        styles_root = ET.fromstring(styles_xml)
        style_elem = template_service._find_style_element(styles_root, None, style_id)
        self.assertIsNotNone(style_elem)
        tbl_pr = style_elem.find("w:tblPr", template_service.DOCX_NS)
        if tbl_pr is None:
            return []
        return [template_service._local_name(child.tag) for child in list(tbl_pr)]

    def _append_table_style(
        self,
        styles_root: ET.Element,
        style_id: str,
        style_name: str,
        *,
        border_color: str | None = None,
    ) -> None:
        style = ET.SubElement(styles_root, template_service._qn("w", "style"))
        style.set(template_service._qn("w", "type"), "table")
        style.set(template_service._qn("w", "styleId"), style_id)
        name = ET.SubElement(style, template_service._qn("w", "name"))
        name.set(template_service._qn("w", "val"), style_name)
        based_on = ET.SubElement(style, template_service._qn("w", "basedOn"))
        based_on.set(template_service._qn("w", "val"), "TableNormal")
        tbl_pr = ET.SubElement(style, template_service._qn("w", "tblPr"))
        if not border_color:
            return
        borders = ET.SubElement(tbl_pr, template_service._qn("w", "tblBorders"))
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = ET.SubElement(borders, template_service._qn("w", tag))
            border.set(template_service._qn("w", "val"), "single")
            border.set(template_service._qn("w", "sz"), "4")
            border.set(template_service._qn("w", "color"), border_color)

    def _build_docx_with_table_style_reference(
        self,
        tmp_dir: str,
        *,
        source_style_id: str = "SourceGrid",
        source_style_name: str = "Source Grid",
        source_border_color: str | None = "999999",
        target_style_id: str = "TargetGrid",
        target_style_name: str = "Target Grid",
        target_border_color: str = "FF0000",
    ) -> Path:
        if not HAS_TEST_DOCX:
            self.skipTest("python-docx no disponible")

        docx_path = Path(tmp_dir) / "style_source.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        doc.save(str(docx_path))

        with zipfile.ZipFile(docx_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            styles_xml = zin.read("word/styles.xml")

        document_root = ET.fromstring(document_xml)
        tbl = document_root.find(".//" + template_service._qn("w", "tbl"))
        self.assertIsNotNone(tbl)
        tbl_pr = tbl.find(template_service._qn("w", "tblPr"))
        if tbl_pr is None:
            tbl_pr = ET.Element(template_service._qn("w", "tblPr"))
            tbl.insert(0, tbl_pr)
        for tag in ("tblBorders",):
            for old in list(tbl_pr.findall(template_service._qn("w", tag))):
                tbl_pr.remove(old)
        tbl_style = tbl_pr.find(template_service._qn("w", "tblStyle"))
        if tbl_style is None:
            tbl_style = ET.SubElement(tbl_pr, template_service._qn("w", "tblStyle"))
        tbl_style.set(template_service._qn("w", "val"), source_style_id)

        styles_root = ET.fromstring(styles_xml)
        if source_border_color is not None:
            self._append_table_style(
                styles_root,
                source_style_id,
                source_style_name,
                border_color=source_border_color,
            )
        self._append_table_style(
            styles_root,
            target_style_id,
            target_style_name,
            border_color=target_border_color,
        )

        namespace_hints = template_service._collect_docx_namespace_hints_from_path(docx_path)
        updated_document = template_service._serialize_ooxml_part(
            document_root,
            document_xml,
            namespace_hints=namespace_hints,
        )
        updated_styles = template_service._serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        template_service._write_docx_parts(
            docx_path,
            {
                "word/document.xml": updated_document,
                "word/styles.xml": updated_styles,
            },
        )
        return docx_path

    def _read_style_border_colors(self, docx_path: Path, style_id: str) -> dict[str, str | None]:
        with zipfile.ZipFile(docx_path, "r") as zin:
            styles_xml = zin.read("word/styles.xml")
        styles_root = ET.fromstring(styles_xml)
        style_elem = template_service._find_style_element(styles_root, None, style_id)
        self.assertIsNotNone(style_elem)
        tbl_pr = style_elem.find("w:tblPr", template_service.DOCX_NS)
        borders = tbl_pr.find("w:tblBorders", template_service.DOCX_NS) if tbl_pr is not None else None
        self.assertIsNotNone(borders)
        colors: dict[str, str | None] = {}
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.find(template_service._qn("w", tag))
            colors[tag] = border.get(template_service._qn("w", "color")) if border is not None else None
        return colors

    def _inject_legacy_runtime_nodes(self, docx_bytes: bytes, style_id: str = "TableGrid") -> bytes:
        parts = template_service._read_docx_parts(docx_bytes, ["word/styles.xml"])
        styles_root = ET.fromstring(parts["word/styles.xml"])
        style_elem = template_service._find_style_element(styles_root, "Table Grid", style_id)
        self.assertIsNotNone(style_elem)
        tbl_pr = style_elem.find("w:tblPr", template_service.DOCX_NS)
        if tbl_pr is None:
            tbl_pr = ET.SubElement(style_elem, template_service._qn("w", "tblPr"))

        tbl_layout = ET.SubElement(tbl_pr, template_service._qn("w", "tblLayout"))
        tbl_layout.set(template_service._qn("w", "type"), "fixed")
        tbl_w = ET.SubElement(tbl_pr, template_service._qn("w", "tblW"))
        tbl_w.set(template_service._qn("w", "type"), "pct")
        tbl_w.set(template_service._qn("w", "w"), "5000")
        tbl_look = ET.SubElement(tbl_pr, template_service._qn("w", "tblLook"))
        tbl_look.set(template_service._qn("w", "firstRow"), "1")
        tbl_look.set(template_service._qn("w", "noVBand"), "1")

        updated_styles = template_service._serialize_ooxml_part(
            styles_root,
            parts["word/styles.xml"],
            namespace_hints=template_service._collect_docx_namespace_hints_from_bytes(docx_bytes),
        )
        return template_service._rewrite_docx_bytes(docx_bytes, {"word/styles.xml": updated_styles})

    def _assert_ignorable_prefixes_declared(self, xml_text: str) -> None:
        root_tag = template_service._extract_root_start_tag(xml_text)
        self.assertIsNotNone(root_tag)
        declared = template_service._extract_namespace_declarations_from_root_tag(root_tag)
        ignorable = template_service._extract_root_ignorable_prefixes(root_tag)
        for prefix in ignorable:
            self.assertIn(prefix, declared, f"Missing xmlns declaration for ignorable prefix '{prefix}'")

    def _build_style_with_first_row(self, fill: str = "111111") -> ET.Element:
        style_elem = ET.Element(template_service._qn("w", "style"))
        first_row = ET.SubElement(style_elem, template_service._qn("w", "tblStylePr"))
        first_row.set(template_service._qn("w", "type"), "firstRow")
        tc_pr = ET.SubElement(first_row, template_service._qn("w", "tcPr"))
        shd = ET.SubElement(tc_pr, template_service._qn("w", "shd"))
        shd.set(template_service._qn("w", "fill"), fill)
        return style_elem

    def _get_first_row_variants(self, style_elem: ET.Element):
        return [
            node
            for node in style_elem.findall(template_service._qn("w", "tblStylePr"))
            if node.get(template_service._qn("w", "type")) == "firstRow"
        ]

    def test_parse_table_borders_tolerates_invalid_size(self):
        borders = template_service._parse_table_borders(
            [{"tag": "top", "attrs": {"val": "single", "sz": "not-a-number", "color": "00AAFF"}}]
        )
        self.assertIn("top", borders)
        self.assertEqual(borders["top"]["style"], "single")
        self.assertIsNone(borders["top"]["size_pt"])
        self.assertEqual(borders["top"]["color"], "00AAFF")

    def test_has_effective_preview_updates_ignores_meta_only_payload(self):
        meta_only = {"style_type": "table", "category": "tables", "style_id": "TableGrid"}
        self.assertFalse(template_service._has_effective_preview_updates(meta_only))

        with_effective_change = dict(meta_only, table_border_color="00AAFF")
        self.assertTrue(template_service._has_effective_preview_updates(with_effective_change))

    def test_resolve_preview_table_look_prefers_style_then_props(self):
        style_elem = ET.Element(template_service._qn("w", "style"))
        tbl_pr = ET.SubElement(style_elem, template_service._qn("w", "tblPr"))
        tbl_look = ET.SubElement(tbl_pr, template_service._qn("w", "tblLook"))
        tbl_look.set(template_service._qn("w", "firstRow"), "0")
        tbl_look.set(template_service._qn("w", "firstColumn"), "1")
        tbl_look.set(template_service._qn("w", "noHBand"), "1")

        style_obj = type("TableStyleStub", (), {"_element": style_elem})()
        resolved = template_service._resolve_preview_table_look({}, style_obj)
        self.assertFalse(resolved["firstRow"])
        self.assertTrue(resolved["firstColumn"])
        self.assertTrue(resolved["noHBand"])

        overridden = template_service._resolve_preview_table_look(
            {"look_first_row": True, "look_no_v_band": False},
            style_obj,
        )
        self.assertTrue(overridden["firstRow"])
        self.assertFalse(overridden["noVBand"])

    def test_apply_table_runtime_defaults_to_preview_table_uses_docx_oxml(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        runtime_defaults = {
            "look": {
                "firstRow": True,
                "firstColumn": False,
                "noHBand": False,
                "noVBand": True,
            },
            "layout_type": "fixed",
            "width_type": "pct",
            "width_value": 5000,
        }

        with patch.object(template_service._logger, "warning") as warning_mock:
            template_service._apply_table_runtime_defaults_to_preview_table(table, runtime_defaults)

        warning_mock.assert_not_called()
        tbl_pr = table._tbl.find(template_service.docx_qn("w:tblPr"))
        self.assertIsNotNone(tbl_pr)
        tbl_look = tbl_pr.find(template_service.docx_qn("w:tblLook"))
        tbl_layout = tbl_pr.find(template_service.docx_qn("w:tblLayout"))
        tbl_w = tbl_pr.find(template_service.docx_qn("w:tblW"))

        self.assertIsNotNone(tbl_look)
        self.assertIsNotNone(tbl_layout)
        self.assertIsNotNone(tbl_w)
        self.assertEqual(tbl_look.get(template_service.docx_qn("w:firstRow")), "1")
        self.assertEqual(tbl_look.get(template_service.docx_qn("w:noVBand")), "1")
        self.assertEqual(tbl_layout.get(template_service.docx_qn("w:type")), "fixed")
        self.assertEqual(tbl_w.get(template_service.docx_qn("w:type")), "pct")
        self.assertEqual(tbl_w.get(template_service.docx_qn("w:w")), "5000")

    def test_apply_table_format_preserves_target_borders_when_source_has_no_direct_borders(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir) / "templates"
            try:
                kernel_id = "kernel-preserve-target-borders"
                docx_path = self._build_docx_with_table_style_reference(
                    tmp_dir,
                    source_style_id="MissingSourceGrid",
                    source_border_color=None,
                    target_border_color="FF0000",
                )
                docx_bytes = docx_path.read_bytes()
                template_service.save_template(
                    kernel_id,
                    docx_bytes,
                    template_service.extract_styles_from_docx(docx_bytes),
                )

                updated = template_service.apply_table_format_to_style(
                    kernel_id,
                    0,
                    "Target Grid",
                    "TargetGrid",
                )

                self.assertIsNotNone(updated)
                saved_path = template_service._get_template_docx_path(kernel_id)
                border_colors = self._read_style_border_colors(saved_path, "TargetGrid")
                self.assertEqual(set(border_colors.values()), {"FF0000"})
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_apply_table_format_copies_effective_source_table_style_borders(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir) / "templates"
            try:
                kernel_id = "kernel-copy-source-style"
                docx_path = self._build_docx_with_table_style_reference(
                    tmp_dir,
                    source_style_id="SourceGrid",
                    source_style_name="Source Grid",
                    source_border_color="999999",
                    target_border_color="FF0000",
                )
                docx_bytes = docx_path.read_bytes()
                template_service.save_template(
                    kernel_id,
                    docx_bytes,
                    template_service.extract_styles_from_docx(docx_bytes),
                )

                updated = template_service.apply_table_format_to_style(
                    kernel_id,
                    0,
                    "Target Grid",
                    "TargetGrid",
                )

                self.assertIsNotNone(updated)
                saved_path = template_service._get_template_docx_path(kernel_id)
                border_colors = self._read_style_border_colors(saved_path, "TargetGrid")
                self.assertEqual(set(border_colors.values()), {"999999"})
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_extract_styles_from_docx_reads_header_footer_text_from_ooxml_parts(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "header_footer_extract.docx"
            doc = Document()
            section = doc.sections[0]
            section.header.paragraphs[0].text = "placeholder"
            section.footer.paragraphs[0].text = "placeholder"
            doc.add_paragraph("Body text")
            doc.save(str(docx_path))

            W = template_service.DOCX_NS["w"]
            header_xml = (
                f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:hdr xmlns:w="{W}">'
                f'  <w:tbl>'
                f'    <w:tr><w:tc><w:p>'
                f'      <w:r><w:t>CRITERIOS DE DISE</w:t></w:r>'
                f'      <w:r><w:t>ÑO DE DESNIVELACIÓN FERROVIARIA</w:t></w:r>'
                f'    </w:p></w:tc></w:tr>'
                f'  </w:tbl>'
                f'</w:hdr>'
            )
            footer_xml = (
                f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:ftr xmlns:w="{W}">'
                f'  <w:p>'
                f'    <w:r><w:t>GIN-DI-226492-1002650-GEN-INF-0009_E</w:t></w:r>'
                f'  </w:p>'
                f'  <w:p>'
                f'    <w:r><w:t>Página 1 de 2</w:t></w:r>'
                f'  </w:p>'
                f'</w:ftr>'
            )
            template_service._write_docx_parts(
                docx_path,
                {
                    "word/header1.xml": header_xml.encode("utf-8"),
                    "word/footer1.xml": footer_xml.encode("utf-8"),
                },
            )

            extracted = template_service.extract_styles_from_docx(docx_path.read_bytes())

            self.assertTrue(any("CRITERIOS DE DISEÑO DE DESNIVELACIÓN FERROVIARIA" in item for item in extracted.get("headers", [])))
            self.assertTrue(any("GIN-DI-226492-1002650-GEN-INF-0009_E" in item for item in extracted.get("footers", [])))
            self.assertTrue(any("Página 1 de 2" in item for item in extracted.get("footers", [])))

    def test_generate_document_table_preview_rejects_negative_index(self):
        result = template_service.generate_document_table_preview("kernel-test", -1)
        self.assertFalse(result.get("success"))
        self.assertIn("non-negative integer", result.get("error", ""))

    def test_create_table_style_from_format_rejects_negative_index(self):
        result = template_service.create_table_style_from_format("kernel-test", -1, "NuevoEstilo")
        self.assertIsNone(result)

    def test_apply_table_format_to_style_rejects_negative_index(self):
        result = template_service.apply_table_format_to_style("kernel-test", -1, "Table Grid")
        self.assertIsNone(result)

    def test_apply_style_to_docx_resolves_by_style_id_before_creating_new_style(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("sample")
            doc.save(str(docx_path))

            template_service._apply_style_to_docx(
                docx_path,
                "Missing Display Name",
                {
                    "style_id": "Normal",
                    "style_type": "paragraph",
                    "category": "body",
                    "bold": True,
                },
            )

            reloaded = Document(str(docx_path))
            style_names = {style.name for style in reloaded.styles}
            self.assertNotIn("Missing Display Name", style_names)

    def test_apply_first_row_variant_keeps_existing_when_not_applicable(self):
        style_elem = self._build_style_with_first_row(fill="111111")
        table_format = {
            "has_distinct_header": False,
            "first_row_format": {},
        }

        template_service._apply_first_row_variant(style_elem, table_format)

        first_rows = self._get_first_row_variants(style_elem)
        self.assertEqual(len(first_rows), 1)
        shd = first_rows[0].find("w:tcPr/w:shd", template_service.DOCX_NS)
        self.assertIsNotNone(shd)
        self.assertEqual(shd.get(template_service._qn("w", "fill")), "111111")

    def test_apply_first_row_variant_replaces_existing_when_applicable(self):
        style_elem = self._build_style_with_first_row(fill="111111")
        table_format = {
            "has_distinct_header": True,
            "first_row_format": {
                "shading_fill": "FF0000",
                "font_properties": {},
                "sample_cell": {},
            },
        }

        template_service._apply_first_row_variant(style_elem, table_format)

        first_rows = self._get_first_row_variants(style_elem)
        self.assertEqual(len(first_rows), 1)
        shd = first_rows[0].find("w:tcPr/w:shd", template_service.DOCX_NS)
        self.assertIsNotNone(shd)
        self.assertEqual(shd.get(template_service._qn("w", "fill")), "FF0000")

    def test_normalize_style_updates_supports_nested_table_block(self):
        updates = {
            "font": {"bold": True},
            "paragraph": {"alignment": "CENTER"},
            "table": {
                "border_style": "single",
                "look_first_row": False,
                "width_type": "pct",
                "width_value": 5000,
            },
        }
        normalized = template_service._normalize_style_updates(updates)
        self.assertTrue(normalized.get("bold"))
        self.assertEqual(normalized.get("alignment"), "CENTER")
        self.assertEqual(normalized.get("table_border_style"), "single")
        self.assertFalse(normalized.get("table_look_first_row"))
        self.assertEqual(normalized.get("table_width_type"), "pct")
        self.assertEqual(normalized.get("table_width_value"), 5000)

    def test_apply_table_properties_to_style_tblpr_omits_runtime_only_nodes(self):
        tbl_pr = ET.Element(template_service._qn("w", "tblPr"))
        template_service._apply_table_properties_to_style_tblpr(
            tbl_pr,
            {
                "borders": {"top": {"style": "single", "size_pt": 0.5, "color": "000000"}},
                "layout_type": "fixed",
                "width_type": "pct",
                "width_value": 5000,
                "look": {"firstRow": True, "noVBand": True},
                "cell_spacing_pt": 2.0,
            },
        )

        child_tags = {template_service._local_name(child.tag) for child in list(tbl_pr)}
        self.assertIn("tblBorders", child_tags)
        self.assertIn("tblCellSpacing", child_tags)
        self.assertNotIn("tblLayout", child_tags)
        self.assertNotIn("tblW", child_tags)
        self.assertNotIn("tblLook", child_tags)

    def test_sanitize_advanced_props_for_table_style_extracts_runtime_nodes(self):
        sanitized, runtime_patch = template_service._sanitize_advanced_props_for_table_style(
            {
                "tbl_pr": [
                    {"tag": "tblLayout", "attrs": {"type": "fixed"}},
                    {"tag": "tblW", "attrs": {"type": "pct", "w": "5000"}},
                    {"tag": "tblLook", "attrs": {"firstRow": "1", "noVBand": "1"}},
                    {"tag": "tblBorders", "attrs": {}},
                ]
            }
        )

        self.assertEqual(runtime_patch.get("layout_type"), "fixed")
        self.assertEqual(runtime_patch.get("width_type"), "pct")
        self.assertEqual(runtime_patch.get("width_value"), "5000")
        self.assertEqual(runtime_patch.get("look"), {"firstRow": "1", "noVBand": "1"})
        self.assertEqual(sanitized.get("tbl_pr"), [{"tag": "tblBorders", "attrs": {}}])

    def test_update_template_style_persists_runtime_defaults_outside_styles_xml(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-runtime-style"
                doc = Document()
                doc.add_paragraph("demo")
                buffer = io.BytesIO()
                doc.save(buffer)
                docx_bytes = buffer.getvalue()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                template_service.save_template(kernel_id, docx_bytes, extracted)

                updated = template_service.update_template_style(
                    kernel_id,
                    "Table Grid",
                    {
                        "style_id": "TableGrid",
                        "style_type": "table",
                        "category": "tables",
                        "table_layout_type": "fixed",
                        "table_width_type": "dxa",
                        "table_width_value": 7200,
                        "table_look_first_row": True,
                        "table_look_no_v_band": True,
                    },
                )

                runtime_defaults = updated.get(template_service.TABLE_STYLE_RUNTIME_DEFAULTS_KEY) or {}
                entry = runtime_defaults.get("TableGrid") or {}
                self.assertEqual(entry.get("layout_type"), "fixed")
                self.assertEqual(entry.get("width_type"), "dxa")
                self.assertEqual(entry.get("width_value"), 7200)
                self.assertEqual(entry.get("look"), {"firstRow": True, "noVBand": True})

                docx_path = template_service._get_template_docx_path(kernel_id)
                self.assertNotIn("tblLayout", self._read_style_tblpr_tags(docx_path, "TableGrid"))
                self.assertNotIn("tblW", self._read_style_tblpr_tags(docx_path, "TableGrid"))
                self.assertNotIn("tblLook", self._read_style_tblpr_tags(docx_path, "TableGrid"))
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_apply_table_format_to_style_persists_runtime_defaults_without_invalid_style_nodes(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-apply-style"
                doc = Document()
                doc.add_paragraph("demo")
                buffer = io.BytesIO()
                doc.save(buffer)
                docx_bytes = buffer.getvalue()

                extracted = template_service.extract_styles_from_docx(docx_bytes)
                template_service.save_template(kernel_id, docx_bytes, extracted)

                fake_table = {
                    "parsed_properties": {
                        "borders": {},
                        "alignment": "center",
                        "width_type": "pct",
                        "width_value": 5000,
                        "layout_type": "fixed",
                        "look": {"firstRow": True, "noVBand": True},
                    }
                }

                with patch.object(template_service, "_extract_document_tables", return_value=[fake_table]):
                    updated = template_service.apply_table_format_to_style(
                        kernel_id,
                        0,
                        "Table Grid",
                        "TableGrid",
                    )

                self.assertIsNotNone(updated)
                runtime_defaults = updated.get(template_service.TABLE_STYLE_RUNTIME_DEFAULTS_KEY) or {}
                entry = runtime_defaults.get("TableGrid") or {}
                self.assertEqual(entry.get("layout_type"), "fixed")
                self.assertEqual(entry.get("width_type"), "pct")
                self.assertEqual(entry.get("width_value"), 5000)
                self.assertEqual(entry.get("look"), {"firstRow": True, "noVBand": True})

                docx_path = template_service._get_template_docx_path(kernel_id)
                self.assertNotIn("tblLayout", self._read_style_tblpr_tags(docx_path, "TableGrid"))
                self.assertNotIn("tblW", self._read_style_tblpr_tags(docx_path, "TableGrid"))
                self.assertNotIn("tblLook", self._read_style_tblpr_tags(docx_path, "TableGrid"))
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_get_template_docx_path_sanitizes_legacy_runtime_nodes_and_keeps_backup(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-legacy-sanitize"
                doc = Document()
                doc.add_paragraph("demo")
                buffer = io.BytesIO()
                doc.save(buffer)
                legacy_docx_bytes = self._inject_legacy_runtime_nodes(buffer.getvalue())
                legacy_template = template_service.extract_styles_from_docx(legacy_docx_bytes)

                template_service._write_template_files(kernel_id, legacy_docx_bytes, legacy_template)

                docx_path = template_service.get_template_docx_path(kernel_id)
                self.assertIsNotNone(docx_path)

                sanitized_template = template_service.get_template(kernel_id) or {}
                runtime_defaults = sanitized_template.get(template_service.TABLE_STYLE_RUNTIME_DEFAULTS_KEY) or {}
                entry = runtime_defaults.get("TableGrid") or {}
                self.assertEqual(entry.get("layout_type"), "fixed")
                self.assertEqual(entry.get("width_type"), "pct")
                self.assertEqual(entry.get("width_value"), 5000)
                self.assertEqual(entry.get("look"), {"firstRow": True, "noVBand": True})

                sanitized_path = template_service._get_template_docx_path(kernel_id)
                self.assertNotIn("tblLayout", self._read_style_tblpr_tags(sanitized_path, "TableGrid"))
                self.assertNotIn("tblW", self._read_style_tblpr_tags(sanitized_path, "TableGrid"))
                self.assertNotIn("tblLook", self._read_style_tblpr_tags(sanitized_path, "TableGrid"))

                backup_files = list((template_service.TEMPLATE_DIR / kernel_id).glob("*.backup_*.docx"))
                self.assertTrue(backup_files)
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_serialize_ooxml_part_preserves_ignorable_prefix_declarations(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "roundtrip.docx"
            doc = Document()
            doc.add_paragraph("demo")
            doc.save(str(docx_path))

            with zipfile.ZipFile(docx_path, "r") as zin:
                styles_xml = zin.read("word/styles.xml")
                namespace_hints = template_service._collect_docx_namespace_hints_from_zip(zin)

            styles_root = ET.fromstring(styles_xml)
            roundtripped = template_service._serialize_ooxml_part(
                styles_root,
                styles_xml,
                namespace_hints=namespace_hints,
            ).decode("utf-8")

            self._assert_ignorable_prefixes_declared(roundtripped)

    def test_sanitize_table_runtime_defaults_repairs_ignorable_prefixes_even_without_runtime_nodes(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "corrupt_styles.docx"
            doc = Document()
            doc.add_paragraph("demo")
            doc.save(str(docx_path))

            with zipfile.ZipFile(docx_path, "r") as zin:
                styles_xml = zin.read("word/styles.xml").decode("utf-8")
                docx_bytes = Path(docx_path).read_bytes()

            corrupted_styles = styles_xml.replace(
                ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"',
                "",
            )
            corrupted_docx = template_service._rewrite_docx_bytes(
                docx_bytes,
                {"word/styles.xml": corrupted_styles.encode("utf-8")},
            )

            sanitized_docx, _, changed = template_service._sanitize_table_style_runtime_defaults_in_docx(corrupted_docx)
            self.assertTrue(changed)

            parts = template_service._read_docx_parts(sanitized_docx, ["word/styles.xml"])
            self._assert_ignorable_prefixes_declared(parts["word/styles.xml"])

    def test_repair_ooxml_namespace_declarations_repairs_footer_root(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "footer_corrupt.docx"
            doc = Document()
            doc.sections[0].footer.paragraphs[0].text = "footer"
            doc.save(str(docx_path))

            docx_bytes = docx_path.read_bytes()
            parts = template_service._read_docx_parts(docx_bytes, ["word/footer1.xml"])
            footer_xml = parts["word/footer1.xml"]
            root_tag = template_service._extract_root_start_tag(footer_xml)
            ignorable = sorted(template_service._extract_root_ignorable_prefixes(root_tag))
            declared = template_service._extract_namespace_declarations_from_root_tag(root_tag)

            missing_prefix = next((prefix for prefix in ignorable if prefix in declared), None)
            if not missing_prefix:
                self.skipTest("footer1.xml sin prefixes ignorable declarados")

            corrupted_footer = footer_xml.replace(
                f' xmlns:{missing_prefix}="{declared[missing_prefix]}"',
                "",
                1,
            )
            corrupted_docx = template_service._rewrite_docx_bytes(
                docx_bytes,
                {"word/footer1.xml": corrupted_footer.encode("utf-8")},
            )

            repaired_docx, changed = template_service._repair_ooxml_namespace_declarations_in_docx(corrupted_docx)
            self.assertTrue(changed)

            repaired_parts = template_service._read_docx_parts(repaired_docx, ["word/footer1.xml"])
            self._assert_ignorable_prefixes_declared(repaired_parts["word/footer1.xml"])

    def test_delete_template_sanitizes_kernel_id(self):
        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                safe_dir = template_service.TEMPLATE_DIR / "evil"
                safe_dir.mkdir(parents=True, exist_ok=True)
                (safe_dir / "template.docx").write_bytes(b"demo")

                deleted = template_service.delete_template("../evil")

                self.assertTrue(deleted)
                self.assertFalse(safe_dir.exists())
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_update_template_style_invalidates_preview_cache(self):
        with TemporaryDirectory() as tmp_dir:
            original_template_dir = template_service.TEMPLATE_DIR
            template_service.TEMPLATE_DIR = Path(tmp_dir)
            try:
                kernel_id = "kernel-cache"
                kernel_dir = template_service.TEMPLATE_DIR / kernel_id
                kernel_dir.mkdir(parents=True, exist_ok=True)
                (kernel_dir / "template.docx").write_bytes(b"fake-docx-bytes")
                (kernel_dir / "template.json").write_text("{}", encoding="utf-8")

                with patch.object(template_service, "HAS_DOCX", True):
                    with patch.object(template_service, "_apply_style_to_docx") as apply_mock:
                        with patch.object(template_service, "extract_styles_from_docx", return_value={"styles": {}}):
                            with patch.object(template_service, "get_style_coverage", return_value={"summary": {}}):
                                with patch.object(template_service, "clear_preview_cache") as clear_cache_mock:
                                    template_service.update_template_style(
                                        kernel_id,
                                        "Normal",
                                        {"font": {"bold": True}},
                                    )

                apply_mock.assert_called_once()
                self.assertGreaterEqual(clear_cache_mock.call_count, 1)
                clear_cache_mock.assert_any_call(kernel_id)
            finally:
                template_service.TEMPLATE_DIR = original_template_dir

    def test_prepare_preview_docx_preserves_header_footer_references(self):
        if not (template_service.HAS_DOCX and HAS_TEST_DOCX):
            self.skipTest("python-docx no disponible")

        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            preview_path = Path(tmp_dir) / "preview.docx"

            doc = Document()
            section = doc.sections[0]
            section.header.paragraphs[0].text = "HEADER DEMO"
            section.footer.paragraphs[0].text = "FOOTER DEMO"
            doc.add_paragraph("Body text")
            doc.save(str(template_path))

            template_service._prepare_preview_docx(
                preview_path,
                template_path=str(template_path),
                clear_body=True,
                clear_header_footer=False,
                compact_page_setup=True,
            )

            with zipfile.ZipFile(preview_path, "r") as zin:
                self.assertIn("word/header1.xml", zin.namelist())
                self.assertIn("word/footer1.xml", zin.namelist())
                document_xml = zin.read("word/document.xml").decode("utf-8")
                header_xml = zin.read("word/header1.xml").decode("utf-8")
                footer_xml = zin.read("word/footer1.xml").decode("utf-8")

            self.assertIn("headerReference", document_xml)
            self.assertIn("footerReference", document_xml)
            self.assertIn("HEADER DEMO", header_xml)
            self.assertIn("FOOTER DEMO", footer_xml)


class TestTemplateTableRouterValidation(unittest.IsolatedAsyncioTestCase):
    async def test_template_table_preview_rejects_non_integer_index(self):
        message = {
            "kernel_id": "kernel-router",
            "table_index": "abc",
            "request_id": "req-preview-1",
        }
        websocket = object()

        with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
            await notebook_router.handle_template_table_preview(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_table_preview_error")
        self.assertEqual(payload.get("kernel_id"), "kernel-router")
        self.assertEqual(payload.get("table_index"), "abc")
        self.assertIn("Invalid table_index", payload.get("error", ""))

    async def test_template_apply_table_format_rejects_negative_index(self):
        message = {
            "kernel_id": "kernel-router",
            "table_index": -1,
            "target_style_name": "Table Grid",
            "target_style_id": None,
        }
        websocket = object()

        with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
            await notebook_router.handle_template_apply_table_format(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_error")
        self.assertEqual(payload.get("kernel_id"), "kernel-router")
        self.assertEqual(payload.get("table_index"), -1)
        self.assertIn("Invalid table_index", payload.get("error", ""))

    async def test_template_create_style_from_table_invalid_range_returns_template_error(self):
        message = {
            "kernel_id": "kernel-router",
            "table_index": 9999,
            "style_name": "NuevoEstilo",
        }
        websocket = object()

        with patch.object(notebook_router.template_table_format, "create_table_style_from_format", return_value=None):
            with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
                await notebook_router.handle_template_create_style_from_table(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_error")
        self.assertIn("Failed to create table style from format", payload.get("error", ""))

    async def test_template_upload_invalid_base64_returns_typed_error(self):
        message = {
            "type": "template_upload",
            "kernel_id": "kernel-router",
            "request_id": "req-upload-1",
            "docx_base64": "invalid@@base64",
        }
        websocket = object()

        with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
            await notebook_router.handle_template_upload(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_error")
        self.assertEqual(payload.get("request_id"), "req-upload-1")
        self.assertEqual(payload.get("error_code"), "invalid_docx_base64")

    async def test_template_upload_rejects_oversized_payload_before_decode(self):
        message = {
            "type": "template_upload",
            "kernel_id": "kernel-router",
            "request_id": "req-upload-oversize",
            "docx_base64": "A" * 200,
        }
        websocket = object()

        with patch("app.services.template_logic.TEMPLATE_UPLOAD_MAX_BYTES", 16):
            with patch("app.services.template_logic.apply_template_bytes_to_kernel", new=AsyncMock()) as apply_mock:
                with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
                    await notebook_router.handle_template_upload(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_error")
        self.assertEqual(payload.get("request_id"), "req-upload-oversize")
        self.assertEqual(payload.get("error_code"), "template_upload_too_large")
        details = payload.get("details", {})
        self.assertEqual(details.get("max_bytes"), 16)
        self.assertGreater(details.get("received_estimated_bytes", 0), 16)
        apply_mock.assert_not_awaited()

    async def test_template_update_style_success_echoes_request_id(self):
        message = {
            "type": "template_update_style",
            "kernel_id": "kernel-router",
            "request_id": "req-update-1",
            "style_name": "Normal",
            "updates": {"bold": True},
        }
        websocket = object()
        lock = asyncio.Lock()

        with patch.object(notebook_router, "_get_kernel_lock", return_value=lock):
            with patch.object(notebook_router.template_style_apply, "update_template_style", return_value={"styles": {}}):
                with patch.object(notebook_router, "jupyter_kernel_manager", None):
                    with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
                        await notebook_router.handle_template_update_style(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_style_updated")
        self.assertEqual(payload.get("kernel_id"), "kernel-router")
        self.assertEqual(payload.get("request_id"), "req-update-1")

    async def test_template_update_semantic_slots_success_echoes_request_id(self):
        message = {
            "type": "template_update_semantic_slots",
            "kernel_id": "kernel-router",
            "request_id": "req-slots-1",
            "semantic_style_slots": {
                "body": {
                    "selection_key": "body|BodyText|Body Text",
                    "style_id": "BodyText",
                    "style_name": "Body Text",
                },
            },
        }
        websocket = object()
        lock = asyncio.Lock()

        with patch.object(notebook_router, "_get_kernel_lock", return_value=lock):
            with patch.object(notebook_router.template_service, "update_template_semantic_style_slots", return_value={"styles": []}):
                with patch.object(notebook_router, "jupyter_kernel_manager", None):
                    with patch.object(notebook_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
                        await notebook_router.handle_template_update_semantic_slots(message, websocket)

        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_semantic_slots_updated")
        self.assertEqual(payload.get("kernel_id"), "kernel-router")
        self.assertEqual(payload.get("request_id"), "req-slots-1")

    async def test_template_preview_style_marks_native_word_request(self):
        message = {
            "type": "template_preview_style",
            "kernel_id": "kernel-router",
            "request_id": "req-native-preview",
            "preview_key": "preview-native",
            "preview_engine": "word_native",
            "native_word_preview": True,
            "style_name": "Body Text",
            "style_props": {"style_id": "BodyText"},
        }
        websocket = object()

        async def current(*_args, **_kwargs):
            return True

        with patch.object(notebook_template_router, "_bind_template_kernel_connection", new=AsyncMock()):
            with patch.object(notebook_template_router, "_register_preview_request", new=AsyncMock(return_value=("kernel-router", "preview-native"))):
                with patch.object(notebook_template_router, "_is_preview_request_current", new=AsyncMock(side_effect=current)):
                    with patch.object(notebook_template_router, "_complete_preview_request", new=AsyncMock()):
                        with patch.object(notebook_template_router.template_preview, "get_preview_cache", return_value=None):
                            with patch.object(notebook_template_router.template_preview, "set_preview_cache") as cache_mock:
                                with patch.object(notebook_template_router.template_service, "run_template_executor", new=AsyncMock(return_value="PNG_BASE64")) as run_mock:
                                    with patch.object(notebook_template_router.manager, "send_personal_message", new=AsyncMock()) as send_mock:
                                        await notebook_template_router.handle_template_preview_style(message, websocket)

        run_mock.assert_awaited_once()
        style_props = run_mock.await_args.args[3]
        self.assertEqual(style_props.get("_preview_engine"), "word_native")
        self.assertTrue(style_props.get("native_word_preview"))
        cache_mock.assert_called_once_with("preview-native", "PNG_BASE64", "kernel-router")
        send_mock.assert_awaited_once()
        payload = send_mock.await_args.args[0]
        self.assertEqual(payload.get("type"), "template_preview_ready")
        self.assertEqual(payload.get("request_id"), "req-native-preview")


class TestFreezeHeaderFooterTableStyles(unittest.TestCase):
    """Tests for _freeze_header_footer_table_styles."""

    def _make_docx_with_header_table(self, tmp_dir, style_id="TableGrid"):
        """Create a minimal DOCX with a table in the header referencing a table style."""
        if not HAS_TEST_DOCX:
            self.skipTest("python-docx not available")

        docx_path = Path(tmp_dir) / "sample.docx"
        doc = Document()
        doc.add_paragraph("Body text")
        doc.save(str(docx_path))

        # Build a minimal header1.xml with a table referencing the style
        W = template_service.DOCX_NS["w"]
        header_xml = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:hdr xmlns:w="{W}">'
            f'  <w:tbl>'
            f'    <w:tblPr>'
            f'      <w:tblStyle w:val="{style_id}"/>'
            f'    </w:tblPr>'
            f'    <w:tr><w:tc><w:p><w:r><w:t>Header</w:t></w:r></w:p></w:tc></w:tr>'
            f'  </w:tbl>'
            f'</w:hdr>'
        )

        # Add a table style to styles.xml with borders
        with zipfile.ZipFile(docx_path, "r") as zin:
            styles_raw = zin.read("word/styles.xml")

        styles_root = ET.fromstring(styles_raw)
        new_style = ET.SubElement(styles_root, template_service._qn("w", "style"))
        new_style.set(template_service._qn("w", "type"), "table")
        new_style.set(template_service._qn("w", "styleId"), style_id)
        name_el = ET.SubElement(new_style, template_service._qn("w", "name"))
        name_el.set(template_service._qn("w", "val"), "Table Grid")
        tbl_pr = ET.SubElement(new_style, template_service._qn("w", "tblPr"))
        borders = ET.SubElement(tbl_pr, template_service._qn("w", "tblBorders"))
        for tag in ("top", "bottom", "left", "right", "insideH", "insideV"):
            b = ET.SubElement(borders, template_service._qn("w", tag))
            b.set(template_service._qn("w", "val"), "single")
            b.set(template_service._qn("w", "sz"), "4")
            b.set(template_service._qn("w", "color"), "000000")

        updated_styles = template_service._serialize_ooxml_part(
            styles_root,
            styles_raw,
            namespace_hints=template_service._collect_docx_namespace_hints_from_path(docx_path),
        )

        # Write header and updated styles into the DOCX
        template_service._write_docx_parts(docx_path, {
            "word/header1.xml": header_xml.encode("utf-8"),
            "word/styles.xml": updated_styles,
        })
        return docx_path

    def test_freeze_detaches_header_table_from_style(self):
        """After freeze, the header table no longer references the style and has inline tblBorders."""
        with TemporaryDirectory() as tmp_dir:
            docx_path = self._make_docx_with_header_table(tmp_dir, style_id="TableGrid")

            template_service._freeze_header_footer_table_styles(
                docx_path, "Table Grid", "TableGrid"
            )

            # Read back header1.xml and verify
            with zipfile.ZipFile(docx_path, "r") as zin:
                header_bytes = zin.read("word/header1.xml")

            root = ET.fromstring(header_bytes)
            tbl = root.find(".//" + template_service._qn("w", "tbl"))
            self.assertIsNotNone(tbl, "Table should exist in header")
            tbl_pr = tbl.find(template_service._qn("w", "tblPr"))
            self.assertIsNotNone(tbl_pr, "tblPr should exist")

            # Style reference should be removed
            style_ref = tbl_pr.find(template_service._qn("w", "tblStyle"))
            self.assertIsNone(style_ref, "tblStyle reference should have been removed")

            # Borders should be inlined
            borders = tbl_pr.find(template_service._qn("w", "tblBorders"))
            self.assertIsNotNone(borders, "tblBorders should be inlined from the style")
            top = borders.find(template_service._qn("w", "top"))
            self.assertIsNotNone(top, "top border should exist")
            self.assertEqual(
                top.get(template_service._qn("w", "val")), "single"
            )

    def test_freeze_handles_docx_without_headers(self):
        """Should not fail on a DOCX with no header/footer parts."""
        if not HAS_TEST_DOCX:
            self.skipTest("python-docx not available")

        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "no_headers.docx"
            doc = Document()
            doc.add_paragraph("Body only")
            doc.save(str(docx_path))

            # Should complete without error
            template_service._freeze_header_footer_table_styles(
                docx_path, "Table Grid", "TableGrid"
            )

    def test_freeze_preserves_existing_inline_overrides(self):
        """If the header table already has an explicit tblBorders, the freeze should not overwrite it."""
        if not HAS_TEST_DOCX:
            self.skipTest("python-docx not available")

        with TemporaryDirectory() as tmp_dir:
            docx_path = self._make_docx_with_header_table(tmp_dir, style_id="TableGrid")

            # Add an explicit tblBorders to the header table before freezing
            W = template_service.DOCX_NS["w"]
            with zipfile.ZipFile(docx_path, "r") as zin:
                header_bytes = zin.read("word/header1.xml")

            root = ET.fromstring(header_bytes)
            tbl_pr = root.find(".//" + template_service._qn("w", "tblPr"))
            explicit_borders = ET.SubElement(tbl_pr, template_service._qn("w", "tblBorders"))
            top_border = ET.SubElement(explicit_borders, template_service._qn("w", "top"))
            top_border.set(template_service._qn("w", "val"), "double")
            top_border.set(template_service._qn("w", "sz"), "8")

            updated_header = template_service._serialize_ooxml_part(
                root,
                header_bytes,
                namespace_hints=template_service._collect_docx_namespace_hints_from_path(docx_path),
            )
            template_service._write_docx_parts(docx_path, {"word/header1.xml": updated_header})

            # Freeze should keep the explicit "double" border, not replace with style's "single"
            template_service._freeze_header_footer_table_styles(
                docx_path, "Table Grid", "TableGrid"
            )

            with zipfile.ZipFile(docx_path, "r") as zin:
                header_bytes = zin.read("word/header1.xml")

            root = ET.fromstring(header_bytes)
            tbl_pr = root.find(".//" + template_service._qn("w", "tblPr"))
            borders = tbl_pr.find(template_service._qn("w", "tblBorders"))
            top = borders.find(template_service._qn("w", "top"))
            self.assertEqual(
                top.get(template_service._qn("w", "val")), "double",
                "Existing inline border override should be preserved, not replaced by style"
            )


if __name__ == "__main__":
    unittest.main()
