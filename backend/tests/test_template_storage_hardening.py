from __future__ import annotations

import io
import json
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

from app.services import template_service

DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _qn(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _minimal_docx_bytes() -> bytes:
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_paragraph("template")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_write_template_files_uses_complete_atomic_payloads(monkeypatch, tmp_path: Path):
    monkeypatch.setattr(template_service, "TEMPLATE_DIR", tmp_path)
    kernel_id = "kernel-template-storage"
    docx_bytes = _minimal_docx_bytes()
    extracted = {"styles": [], "metadata": {"revision": 1}}

    docx_path, json_path = template_service._write_template_files(kernel_id, docx_bytes, extracted)

    assert docx_path.read_bytes() == docx_bytes
    assert json.loads(json_path.read_text(encoding="utf-8")) == extracted
    assert list(docx_path.parent.glob("*.tmp")) == []


def test_sanitize_persisted_template_quarantines_corrupt_docx(monkeypatch, tmp_path: Path):
    monkeypatch.setattr(template_service, "TEMPLATE_DIR", tmp_path)
    kernel_id = "kernel-corrupt-template"
    template_dir = template_service._ensure_template_dir(kernel_id)
    docx_path = template_service._get_template_docx_path(kernel_id)
    json_path = template_service._get_template_json_path(kernel_id)
    docx_path.write_bytes(b"not-a-docx")
    json_path.write_text('{"styles": []}', encoding="utf-8")

    result = template_service._sanitize_persisted_template_if_needed(kernel_id)

    assert result
    assert result.get("metadata", {}).get("recovered_from_corrupt_docx") is True
    assert result.get("styles")
    assert docx_path.exists()
    assert docx_path.read_bytes() != b"not-a-docx"
    persisted = json.loads(json_path.read_text(encoding="utf-8"))
    assert persisted.get("metadata", {}).get("recovered_from_corrupt_docx") is True
    assert persisted.get("styles")
    quarantined = list(template_dir.glob("template.quarantine_*.docx"))
    assert len(quarantined) == 1
    assert quarantined[0].read_bytes() == b"not-a-docx"


def test_extract_styles_includes_hidden_styles_with_visibility_metadata():
    docx = pytest.importorskip("docx")
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement

    document = docx.Document()
    hidden_style = document.styles.add_style("Hidden Internal", WD_STYLE_TYPE.PARAGRAPH)
    hidden_style.element.append(OxmlElement("w:hidden"))
    buffer = io.BytesIO()
    document.save(buffer)

    extracted = template_service.extract_styles_from_docx(buffer.getvalue())
    style = next((item for item in extracted.get("styles", []) if item.get("name") == "Hidden Internal"), None)

    assert style is not None
    assert style.get("hidden") is True
    assert style.get("style_visibility", {}).get("hidden") is True


def test_word_style_roundtrip_applies_structured_ooxml(monkeypatch, tmp_path: Path):
    docx = pytest.importorskip("docx")

    monkeypatch.setattr(template_service, "TEMPLATE_DIR", tmp_path)
    kernel_id = "kernel-word-style-roundtrip"
    document = docx.Document()
    document.add_paragraph("template")
    buffer = io.BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()
    template_service.save_template(
        kernel_id,
        docx_bytes,
        template_service.extract_styles_from_docx(docx_bytes),
    )

    updated = template_service.update_template_style(
        kernel_id,
        "Normal",
        {
            "style_id": "Normal",
            "word_style": {
                "metadata": {"ui_priority": 7, "aliases": ["Body Copy"], "locked": True},
                "visibility": {"q_format": True, "hidden": False},
                "font": {
                    "complex_script_font_name": "Aptos",
                    "east_asia_font_name": "MS Gothic",
                    "character_spacing_twips": "20",
                    "vanish": True,
                },
                "paragraph": {
                    "contextual_spacing": True,
                    "tabs": [{"val": "right", "leader": "dot", "pos_twips": "4320"}],
                    "shading": {"fill": "F2F2F2"},
                    "borders": {
                        "bottom": {"style": "single", "size_pt": 1.5, "color": "1B4965"},
                    },
                },
            },
        },
    )

    docx_path = template_service._get_template_docx_path(kernel_id)
    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_root = ET.fromstring(zf.read("word/styles.xml"))

    normal = styles_root.find("w:style[@w:styleId='Normal']", DOCX_NS)
    assert normal is not None
    assert normal.find("w:uiPriority", DOCX_NS).get(_qn("val")) == "7"
    assert normal.find("w:aliases", DOCX_NS).get(_qn("val")) == "Body Copy"
    assert normal.find("w:locked", DOCX_NS) is not None
    assert normal.find("w:qFormat", DOCX_NS) is not None
    assert normal.find("w:hidden", DOCX_NS) is None
    r_fonts = normal.find("w:rPr/w:rFonts", DOCX_NS)
    assert r_fonts.get(_qn("cs")) == "Aptos"
    assert r_fonts.get(_qn("eastAsia")) == "MS Gothic"
    assert normal.find("w:rPr/w:spacing", DOCX_NS).get(_qn("val")) == "20"
    assert normal.find("w:rPr/w:vanish", DOCX_NS) is not None
    assert normal.find("w:pPr/w:contextualSpacing", DOCX_NS) is not None
    tab = normal.find("w:pPr/w:tabs/w:tab", DOCX_NS)
    assert tab.get(_qn("val")) == "right"
    assert tab.get(_qn("leader")) == "dot"
    assert tab.get(_qn("pos")) == "4320"
    assert normal.find("w:pPr/w:shd", DOCX_NS).get(_qn("fill")) == "F2F2F2"
    bottom_border = normal.find("w:pPr/w:pBdr/w:bottom", DOCX_NS)
    assert bottom_border.get(_qn("val")) == "single"
    assert bottom_border.get(_qn("sz")) == "12"
    assert bottom_border.get(_qn("color")) == "1B4965"

    updated_normal = next(item for item in updated.get("styles", []) if item.get("style_id") == "Normal")
    assert updated.get("word_capabilities", {}).get("paragraph")
    assert updated_normal.get("word_style", {}).get("metadata", {}).get("aliases") == ["Body Copy"]
    assert updated_normal.get("word_style", {}).get("font", {}).get("complex_script_font_name") == "Aptos"
    assert updated_normal.get("word_style", {}).get("paragraph", {}).get("contextual_spacing") is True
    assert updated_normal.get("word_style", {}).get("paragraph", {}).get("shading", {}).get("fill") == "F2F2F2"


def test_word_style_rejects_malformed_raw_ooxml(monkeypatch, tmp_path: Path):
    docx = pytest.importorskip("docx")

    monkeypatch.setattr(template_service, "TEMPLATE_DIR", tmp_path)
    kernel_id = "kernel-word-style-invalid-raw"
    document = docx.Document()
    document.add_paragraph("template")
    buffer = io.BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()
    template_service.save_template(
        kernel_id,
        docx_bytes,
        template_service.extract_styles_from_docx(docx_bytes),
    )

    with pytest.raises(ValueError, match="invalid_advanced_props.r_pr"):
        template_service.update_template_style(
            kernel_id,
            "Normal",
            {
                "style_id": "Normal",
                "advanced_props": {"r_pr": {"tag": "b"}},
            },
        )
