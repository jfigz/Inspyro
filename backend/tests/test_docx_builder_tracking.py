import os
import sys
import unittest
import base64
import io
import json
import zipfile
import xml.etree.ElementTree as ET
import tempfile
import importlib.util
import warnings
from types import SimpleNamespace
from unittest.mock import patch

try:
    import pandas as pd
    HAS_PANDAS = True
except Exception:
    pd = None
    HAS_PANDAS = False

try:
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except Exception:
    plt = None
    HAS_MATPLOTLIB = False

from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.shared import qn  # type: ignore

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from librerias_propias.docx_builder.api import build_doc
from librerias_propias.docx_builder.proxies import FontProxy, ParagraphFormatProxy
from librerias_propias.docx_builder import session as session_module
from librerias_propias.docx_builder.session import get_session, reset_session_cache
from librerias_propias.docx_builder.utils import validate_docx_package_bytes

PNG_DOT_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0nQAAAAASUVORK5CYII="
DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _export_document_root(namespace):
    session = get_session(namespace)
    raw = base64.b64decode(session.export_docx_base64())
    with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
        document_xml = zf.read("word/document.xml")
    return ET.fromstring(document_xml)


class _MutationRecorder:
    def __init__(self):
        self.calls = []

    def record_visible_mutation(self, block_id, root_element, target, **kwargs):
        self.calls.append(
            {
                "block_id": block_id,
                "root_element": root_element,
                "target": target,
                "kwargs": kwargs,
            }
        )


class _FakeFont:
    def __init__(self, *, element=None, rpr=None):
        self._element = element
        self._rPr = rpr
        self.name = None
        self.size = None


class _FakeParagraphFormat:
    def __init__(self, *, element=None, ppr=None):
        self._element = element
        self._pPr = ppr
        self.keep_with_next = None
        self.space_after = None


def _proxy_context(block_id: str):
    return _MutationRecorder(), SimpleNamespace(block_id=block_id), OxmlElement("w:p")


class TestDocxBuilderTracking(unittest.TestCase):
    def setUp(self):
        reset_session_cache()

    def tearDown(self):
        reset_session_cache()

    def test_table_builder_tracks_single_table_once(self):
        namespace = {}
        with build_doc(order=1, namespace=namespace, block_id="cell-table-1") as builder:
            builder.table([["A", "B"]], headers=["c1", "c2"])

        session = get_session(namespace)
        snapshot = session.snapshot_cell("cell-table-1")
        tracked_tables = [item for item in snapshot.get("elements", []) if item.get("type") == "Table"]
        self.assertEqual(len(tracked_tables), 1)

    def test_table_builder_tracks_multiple_tables_without_duplicates(self):
        namespace = {}
        with build_doc(order=3, namespace=namespace, block_id="cell-table-3") as builder:
            builder.table([[1]], headers=["h"])
            builder.table([[2]], headers=["h"])

        session = get_session(namespace)
        snapshot = session.snapshot_cell("cell-table-3")
        tracked_tables = [item for item in snapshot.get("elements", []) if item.get("type") == "Table"]
        self.assertEqual(len(tracked_tables), 2)

    def test_font_proxy_prefers_element_without_futurewarning(self):
        session, handle, root = _proxy_context("cell-font-primary")
        element = OxmlElement("w:r")
        fallback = OxmlElement("w:rPr")
        font = _FakeFont(element=element, rpr=fallback)
        proxy = FontProxy(font, session, handle, root)

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            proxy.name = "Aptos"

        self.assertFalse(any(issubclass(item.category, FutureWarning) for item in caught))
        self.assertEqual(font.name, "Aptos")
        self.assertEqual(session.calls[-1]["target"], element)
        self.assertEqual(session.calls[-1]["kwargs"]["api_name"], "document.run.font.name")

    def test_font_proxy_falls_back_to_rpr_without_futurewarning(self):
        session, handle, root = _proxy_context("cell-font-fallback")
        fallback = OxmlElement("w:rPr")
        font = _FakeFont(element=None, rpr=fallback)
        proxy = FontProxy(font, session, handle, root)

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            proxy.size = 12

        self.assertFalse(any(issubclass(item.category, FutureWarning) for item in caught))
        self.assertEqual(font.size, 12)
        self.assertEqual(session.calls[-1]["target"], fallback)
        self.assertEqual(session.calls[-1]["kwargs"]["api_name"], "document.run.font.size")

    def test_paragraph_format_proxy_prefers_element_without_futurewarning(self):
        session, handle, root = _proxy_context("cell-paragraph-primary")
        element = OxmlElement("w:p")
        fallback = OxmlElement("w:pPr")
        paragraph_format = _FakeParagraphFormat(element=element, ppr=fallback)
        proxy = ParagraphFormatProxy(paragraph_format, session, handle, root)

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            proxy.keep_with_next = True

        self.assertFalse(any(issubclass(item.category, FutureWarning) for item in caught))
        self.assertTrue(paragraph_format.keep_with_next)
        self.assertEqual(session.calls[-1]["target"], element)
        self.assertEqual(
            session.calls[-1]["kwargs"]["api_name"],
            "document.paragraph_format.keep_with_next",
        )

    def test_paragraph_format_proxy_falls_back_to_ppr_without_futurewarning(self):
        session, handle, root = _proxy_context("cell-paragraph-fallback")
        fallback = OxmlElement("w:pPr")
        paragraph_format = _FakeParagraphFormat(element=None, ppr=fallback)
        proxy = ParagraphFormatProxy(paragraph_format, session, handle, root)

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            proxy.space_after = 120

        self.assertFalse(any(issubclass(item.category, FutureWarning) for item in caught))
        self.assertEqual(paragraph_format.space_after, 120)
        self.assertEqual(session.calls[-1]["target"], fallback)
        self.assertEqual(
            session.calls[-1]["kwargs"]["api_name"],
            "document.paragraph_format.space_after",
        )

    def test_table_builder_reapplies_runtime_defaults_on_table_instance(self):
        namespace = {}
        session = get_session(namespace)
        session.set_template_table_style_defaults({
            "TableGrid": {
                "style_id": "TableGrid",
                "style_name": "Table Grid",
                "layout_type": "fixed",
                "width_type": "dxa",
                "width_value": 7200,
                "look": {
                    "firstRow": True,
                    "lastColumn": True,
                    "noHBand": False,
                    "noVBand": True,
                },
            },
        })

        with build_doc(order=4, namespace=namespace, block_id="cell-table-runtime") as builder:
            builder.table([["A", "B"]], headers=["h1", "h2"], style="Table Grid")

        table = get_session(namespace).doc.tables[0]
        tbl_pr = table._tbl.tblPr
        self.assertIsNotNone(tbl_pr)

        tbl_look = tbl_pr.find(qn("w:tblLook"))
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        tbl_width = tbl_pr.find(qn("w:tblW"))

        self.assertIsNotNone(tbl_look)
        self.assertEqual(tbl_look.get(qn("w:firstRow")), "1")
        self.assertEqual(tbl_look.get(qn("w:lastColumn")), "1")
        self.assertEqual(tbl_look.get(qn("w:noVBand")), "1")
        self.assertIsNotNone(tbl_layout)
        self.assertEqual(tbl_layout.get(qn("w:type")), "fixed")
        self.assertIsNotNone(tbl_width)
        self.assertEqual(tbl_width.get(qn("w:type")), "dxa")
        self.assertEqual(tbl_width.get(qn("w:w")), "7200")

    def test_dataframe_builder_inherits_runtime_defaults(self):
        if not HAS_PANDAS:
            self.skipTest("pandas no disponible")

        namespace = {}
        session = get_session(namespace)
        session.set_template_table_style_defaults({
            "TableGrid": {
                "style_id": "TableGrid",
                "style_name": "Table Grid",
                "layout_type": "fixed",
                "width_type": "pct",
                "width_value": 5000,
            },
        })

        with build_doc(order=5, namespace=namespace, block_id="cell-df-runtime") as builder:
            builder.dataframe(pd.DataFrame([{"a": 1, "b": 2}]), style="Table Grid")

        table = get_session(namespace).doc.tables[0]
        tbl_pr = table._tbl.tblPr
        self.assertIsNotNone(tbl_pr)
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        tbl_width = tbl_pr.find(qn("w:tblW"))
        self.assertIsNotNone(tbl_layout)
        self.assertEqual(tbl_layout.get(qn("w:type")), "fixed")
        self.assertIsNotNone(tbl_width)
        self.assertEqual(tbl_width.get(qn("w:type")), "pct")
        self.assertEqual(tbl_width.get(qn("w:w")), "5000")

    def test_fast_rebuild_restores_image_relationships_from_snapshot(self):
        namespace = {}
        with build_doc(order=1, namespace=namespace, block_id="cell-image-1") as builder:
            builder.image(base64.b64decode(PNG_DOT_B64))

        session = get_session(namespace)
        session._persist_cell_snapshot("cell-image-1")
        session.ns[session._CELL_ITEMS_KEY] = {}
        session._mark_dirty(reason="force_snapshot_restore")

        docx_b64 = session.export_docx_base64()
        raw = base64.b64decode(docx_b64)
        is_valid, errors = validate_docx_package_bytes(raw)
        self.assertTrue(is_valid, errors)

        with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
            media = [name for name in zf.namelist() if name.startswith("word/media/")]
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
        self.assertGreaterEqual(len(media), 1)
        self.assertIn("relationships/image", rels)

    def test_fast_rebuild_restores_figure_relationships_from_snapshot(self):
        if not HAS_MATPLOTLIB:
            self.skipTest("matplotlib no disponible")

        namespace = {}
        fig = plt.figure()
        try:
            ax = fig.add_subplot(111)
            ax.plot([0, 1], [0, 1])
            with build_doc(order=2, namespace=namespace, block_id="cell-figure-1") as builder:
                builder.figure(fig, caption="Figura de prueba")
        finally:
            plt.close(fig)

        session = get_session(namespace)
        session._persist_cell_snapshot("cell-figure-1")
        session.ns[session._CELL_ITEMS_KEY] = {}
        session._mark_dirty(reason="force_snapshot_restore")

        docx_b64 = session.export_docx_base64()
        raw = base64.b64decode(docx_b64)
        is_valid, errors = validate_docx_package_bytes(raw)
        self.assertTrue(is_valid, errors)

        with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
            media = [name for name in zf.namelist() if name.startswith("word/media/")]
        self.assertGreaterEqual(len(media), 1)

    def test_validate_docx_package_bytes_detects_missing_image_relationship_target(self):
        namespace = {}
        with build_doc(order=6, namespace=namespace, block_id="cell-image-broken") as builder:
            builder.image(base64.b64decode(PNG_DOT_B64))

        session = get_session(namespace)
        raw = base64.b64decode(session.export_docx_base64())
        src = io.BytesIO(raw)
        dst = io.BytesIO()
        with zipfile.ZipFile(src, "r") as zin:
            with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename.startswith("word/media/"):
                        continue
                    zout.writestr(item, zin.read(item.filename))

        is_valid, errors = validate_docx_package_bytes(dst.getvalue())
        self.assertFalse(is_valid)
        self.assertTrue(any("parte faltante" in error or "Relación rota" in error for error in errors))


    def test_image_caption_uses_seq_field_and_number_bookmark(self):
        namespace = {}
        with build_doc(order=7, namespace=namespace, block_id="cell-image-caption") as builder:
            builder.image(
                base64.b64decode(PNG_DOT_B64),
                caption="Plano general",
                label="fig:plano",
            )

        document_root = _export_document_root(namespace)
        fld_simple = document_root.find(".//w:fldSimple", DOCX_NS)
        self.assertIsNotNone(fld_simple)
        self.assertIn("SEQ Figura", fld_simple.get(qn("w:instr")))

        bookmark = fld_simple.find("w:bookmarkStart", DOCX_NS)
        self.assertIsNotNone(bookmark)
        self.assertEqual(bookmark.get(qn("w:name")), "fig:plano")

        text = "".join(node.text or "" for node in document_root.findall(".//w:t", DOCX_NS))
        self.assertIn("Figura 1. Plano general", text)

    def test_table_caption_is_inserted_before_table_with_seq_field(self):
        namespace = {}
        with build_doc(order=8, namespace=namespace, block_id="cell-table-caption") as builder:
            builder.table(
                [[1, 2]],
                headers=["A", "B"],
                caption="Resultados",
                label="tbl:resultados",
            )

        document_root = _export_document_root(namespace)
        body_children = [
            child for child in list(document_root.find("w:body", DOCX_NS))
            if child.tag != qn("w:sectPr")
        ]
        self.assertGreaterEqual(len(body_children), 2)
        self.assertEqual(body_children[0].tag, qn("w:p"))
        self.assertEqual(body_children[1].tag, qn("w:tbl"))

        caption_field = body_children[0].find(".//w:fldSimple", DOCX_NS)
        self.assertIsNotNone(caption_field)
        self.assertIn("SEQ Tabla", caption_field.get(qn("w:instr")))

        caption_text = "".join(node.text or "" for node in body_children[0].findall(".//w:t", DOCX_NS))
        self.assertIn("Tabla 1. Resultados", caption_text)

    def test_plain_caption_without_number_keeps_text_only(self):
        namespace = {}
        with build_doc(order=9, namespace=namespace, block_id="cell-caption-plain") as builder:
            builder.caption("Leyenda manual", label="cap:manual")

        document_root = _export_document_root(namespace)
        self.assertIsNone(document_root.find(".//w:fldSimple", DOCX_NS))
        text = "".join(node.text or "" for node in document_root.findall(".//w:t", DOCX_NS))
        self.assertIn("Leyenda manual", text)

    def test_provenance_manifest_marks_builder_exact_and_raw_document_fallback(self):
        namespace = {}
        with build_doc(order=10, namespace=namespace, block_id="cell-provenance") as builder:
            builder.heading("Titulo de prueba")
            builder.document.add_paragraph("Parrafo raw")

        session = get_session(namespace)
        manifest = json.loads(session.export_provenance_manifest_json())
        items = manifest.get("items") or []

        heading_item = next(item for item in items if item.get("api_name") == "heading")
        raw_item = next(item for item in items if "Parrafo raw" in str(item.get("text_preview") or ""))

        self.assertEqual(heading_item.get("precision"), "exact")
        self.assertEqual(raw_item.get("precision"), "exact")
        self.assertTrue(heading_item.get("clickable"))
        self.assertTrue(raw_item.get("clickable"))

    def test_provenance_manifest_captures_external_file_path(self):
        namespace = {}
        with tempfile.TemporaryDirectory() as tmp_dir:
            helper_path = os.path.join(tmp_dir, "helper_docx_builder_tracking.py")
            with open(helper_path, "w", encoding="utf-8") as fh:
                fh.write("def add_heading(builder):\n")
                fh.write("    builder.heading('Titulo externo')\n")

            spec = importlib.util.spec_from_file_location("helper_docx_builder_tracking", helper_path)
            module = importlib.util.module_from_spec(spec)
            assert spec and spec.loader
            spec.loader.exec_module(module)

            with build_doc(order=11, namespace=namespace, block_id="cell-provenance-file") as builder:
                module.add_heading(builder)

        session = get_session(namespace)
        manifest = json.loads(session.export_provenance_manifest_json())
        heading_item = next(
            item for item in (manifest.get("items") or [])
            if item.get("api_name") == "heading" and item.get("text_preview") == "Titulo externo"
        )

        self.assertNotEqual(os.path.normcase(heading_item.get("file_path")), os.path.normcase(os.path.abspath(helper_path)))
        self.assertEqual(os.path.normcase(heading_item.get("exact_file_path")), os.path.normcase(os.path.abspath(helper_path)))
        self.assertEqual(heading_item.get("exact_line"), 2)

    def test_provenance_callsite_ignores_ipykernel_temp_frame_when_notebook_frames_exist(self):
        namespace = {}
        session = get_session(namespace)
        fake_stack = [
            SimpleNamespace(filename=__file__, lineno=1),
            SimpleNamespace(filename=__file__, lineno=2),
            SimpleNamespace(filename="<inspyro-notebook:cell-exact>", lineno=21),
            SimpleNamespace(filename="<inspyro-notebook:cell-callsite>", lineno=9),
            SimpleNamespace(
                filename=r"C:\tmp\ipykernel_35784\2485563350.py",
                lineno=113,
            ),
        ]

        with patch.object(session_module.inspect, "stack", return_value=fake_stack):
            location = session._capture_provenance_location()

        self.assertEqual(location.get("exact_notebook_cell_id"), "cell-exact")
        self.assertEqual(location.get("exact_line"), 21)
        self.assertEqual(location.get("notebook_cell_id"), "cell-callsite")
        self.assertEqual(location.get("line"), 9)
        self.assertIsNone(location.get("file_path"))
        self.assertEqual(len(location.get("user_stack") or []), 3)
        self.assertIn("ipykernel_", str((location.get("user_stack") or [])[-1].get("file_path") or ""))

    def test_provenance_open_url_prefers_backend_url_then_backend_port(self):
        with patch.dict(os.environ, {"INSPYRO_BACKEND_URL": "http://127.0.0.1:19000", "INSPYRO_BACKEND_PORT": "18000"}):
            self.assertEqual(
                session_module._build_provenance_open_url("prov-123"),
                "http://127.0.0.1:19000/api/docx/provenance/open?provenance_id=prov-123",
            )

        with patch.dict(os.environ, {"INSPYRO_BACKEND_URL": "", "INSPYRO_BACKEND_PORT": "18000", "INSPYRO_BACKEND_HOST": "127.0.0.1"}):
            self.assertEqual(
                session_module._build_provenance_open_url("prov-123"),
                "http://127.0.0.1:18000/api/docx/provenance/open?provenance_id=prov-123",
            )

    def test_python_docx_public_proxies_track_visible_mutations_exactly(self):
        namespace = {}
        with build_doc(order=12, namespace=namespace, block_id="cell-provenance-python-docx") as builder:
            paragraph_before = builder.document.add_paragraph("Base")
            paragraph_before.insert_paragraph_before("Antes")

            paragraph_bold = builder.document.add_paragraph()
            run_bold = paragraph_bold.add_run("Negrita")
            run_bold.bold = True

            paragraph_text = builder.document.add_paragraph()
            run_text = paragraph_text.add_run("Temporal")
            run_text.text = "Texto final"

            table = builder.document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Celda directa"
            table.rows[0].cells[1].paragraphs[0].add_run("Celda run")

        session = get_session(namespace)
        manifest = json.loads(session.export_provenance_manifest_json())
        items = manifest.get("items") or []

        before_item = next(item for item in items if item.get("text_preview") == "Antes")
        bold_item = next(item for item in items if item.get("text_preview") == "Negrita")
        text_item = next(item for item in items if item.get("text_preview") == "Texto final")
        cell_text_item = next(item for item in items if item.get("text_preview") == "Celda directa")
        cell_run_item = next(item for item in items if item.get("text_preview") == "Celda run")

        self.assertEqual(before_item.get("api_name"), "document.insert_paragraph_before")
        self.assertEqual(bold_item.get("api_name"), "document.run.bold")
        self.assertEqual(text_item.get("api_name"), "document.run.text")
        self.assertEqual(cell_text_item.get("api_name"), "document.cell.text")
        self.assertEqual(cell_run_item.get("api_name"), "document.add_run")

        for item in (before_item, bold_item, text_item, cell_text_item, cell_run_item):
            self.assertEqual(item.get("precision"), "exact")
            self.assertTrue(item.get("clickable"))
            self.assertIsNotNone(item.get("line"))

    def test_python_docx_manual_xml_mutations_track_visible_nodes(self):
        namespace = {}
        with build_doc(order=13, namespace=namespace, block_id="cell-provenance-xml") as builder:
            paragraph = builder.document.add_paragraph()
            xml_run = OxmlElement("w:r")
            xml_text = OxmlElement("w:t")
            xml_text.text = "Parrafo xml"
            xml_run.append(xml_text)
            paragraph._p.append(xml_run)

            table = builder.document.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell_paragraph = OxmlElement("w:p")
            cell_run = OxmlElement("w:r")
            cell_text = OxmlElement("w:t")
            cell_text.text = "Celda xml"
            cell_run.append(cell_text)
            cell_paragraph.append(cell_run)
            cell._tc.append(cell_paragraph)

        session = get_session(namespace)
        manifest = json.loads(session.export_provenance_manifest_json())
        items = manifest.get("items") or []

        paragraph_item = next(item for item in items if item.get("text_preview") == "Parrafo xml")
        cell_item = next(item for item in items if item.get("text_preview") == "Celda xml")

        self.assertEqual(paragraph_item.get("api_name"), "document.xml.append")
        self.assertEqual(cell_item.get("api_name"), "document.xml.append")
        self.assertEqual(paragraph_item.get("precision"), "exact")
        self.assertEqual(cell_item.get("precision"), "exact")
        self.assertTrue(paragraph_item.get("clickable"))
        self.assertTrue(cell_item.get("clickable"))

    def test_snapshot_rebuild_preserves_fine_grained_provenance_fragments(self):
        namespace = {}
        with build_doc(order=14, namespace=namespace, block_id="cell-provenance-rebuild") as builder:
            builder.document.add_paragraph("Persistido run")

            table = builder.document.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            xml_paragraph = OxmlElement("w:p")
            xml_run = OxmlElement("w:r")
            xml_text = OxmlElement("w:t")
            xml_text.text = "Persistido xml"
            xml_run.append(xml_text)
            xml_paragraph.append(xml_run)
            cell._tc.append(xml_paragraph)

        session = get_session(namespace)
        before_manifest = json.loads(session.export_provenance_manifest_json())
        before_items = before_manifest.get("items") or []
        before_previews = {item.get("text_preview") for item in before_items}
        self.assertIn("Persistido run", before_previews)
        self.assertIn("Persistido xml", before_previews)

        session._persist_cell_snapshot("cell-provenance-rebuild")
        session.ns[session._CELL_ITEMS_KEY] = {}
        session._mark_dirty(reason="force_provenance_snapshot_restore")
        session.export_docx_base64()

        after_manifest = json.loads(session.export_provenance_manifest_json())
        after_items = after_manifest.get("items") or []
        after_by_preview = {item.get("text_preview"): item for item in after_items}

        self.assertIn("Persistido run", after_by_preview)
        self.assertIn("Persistido xml", after_by_preview)
        self.assertEqual(after_by_preview["Persistido run"].get("precision"), "exact")
        self.assertEqual(after_by_preview["Persistido xml"].get("precision"), "exact")
        self.assertTrue(after_by_preview["Persistido run"].get("clickable"))
        self.assertTrue(after_by_preview["Persistido xml"].get("clickable"))


if __name__ == "__main__":
    unittest.main()
