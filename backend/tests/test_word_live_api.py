import io
import os
import sys
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from main import app
from app.services import word_live

WORD_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_NS = {"w": WORD_W_NS}


def _qn(tag: str) -> str:
    return f"{{{WORD_W_NS}}}{tag}"


def _rewrite_docx_part(docx_bytes: bytes, updates: dict[str, bytes]) -> bytes:
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            written = set()
            for item in zin.infolist():
                if item.filename in updates:
                    zout.writestr(item, updates[item.filename])
                    written.add(item.filename)
                else:
                    zout.writestr(item, zin.read(item.filename))
            for name, payload in updates.items():
                if name not in written:
                    zout.writestr(name, payload)
    return dst.getvalue()


def _wrap_body_child_in_sdt(node: ET.Element, *, tag: str, title: str) -> ET.Element:
    sdt = ET.Element(_qn("sdt"))
    sdt_pr = ET.SubElement(sdt, _qn("sdtPr"))
    tag_node = ET.SubElement(sdt_pr, _qn("tag"))
    tag_node.set(_qn("val"), tag)
    alias = ET.SubElement(sdt_pr, _qn("alias"))
    alias.set(_qn("val"), title)
    sdt_content = ET.SubElement(sdt, _qn("sdtContent"))
    sdt_content.append(node)
    return sdt


def _make_docx_with_content_controls() -> bytes:
    doc = Document()
    doc.add_paragraph("Editable introduction")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "42"
    buf = io.BytesIO()
    doc.save(buf)

    with zipfile.ZipFile(io.BytesIO(buf.getvalue()), "r") as zin:
        document_xml = zin.read("word/document.xml")

    root = ET.fromstring(document_xml)
    body = root.find("w:body", WORD_NS)
    assert body is not None
    paragraph = next(child for child in list(body) if child.tag == _qn("p"))
    table_node = next(child for child in list(body) if child.tag == _qn("tbl"))

    body.remove(paragraph)
    body.remove(table_node)
    body.insert(0, _wrap_body_child_in_sdt(paragraph, tag="region-intro", title="Introduction"))
    body.insert(1, _wrap_body_child_in_sdt(table_node, tag="region-table", title="Results Table"))

    updated_document = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return _rewrite_docx_part(buf.getvalue(), {"word/document.xml": updated_document})


def test_open_session_extracts_regions_from_docx_path(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("INSPYRO_APP_STATE_DIR", str(tmp_path / "state"))
    word_live.reset_word_live_cache()
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(_make_docx_with_content_controls())

    client = TestClient(app)
    response = client.post(
        "/api/word-live/session/open",
        json={"document_path": str(docx_path), "addin_version": "0.1.0"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["document_path"] == str(docx_path.resolve())
    assert payload["validation_status"] == "ok"
    assert payload["document_version"] == 1
    kinds = {item["region_id"]: item["kind"] for item in payload["regions"]}
    assert kinds["region-intro"] == "text"
    assert kinds["region-table"] == "table"


def test_update_text_rejects_non_text_region(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("INSPYRO_APP_STATE_DIR", str(tmp_path / "state"))
    word_live.reset_word_live_cache()
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(_make_docx_with_content_controls())

    client = TestClient(app)
    open_response = client.post("/api/word-live/session/open", json={"document_path": str(docx_path)})
    session_id = open_response.json()["session_id"]

    response = client.post(
        "/api/word-live/region/update-text",
        json={
            "session_id": session_id,
            "region_id": "region-table",
            "text": "new value",
            "document_version": 1,
        },
    )

    assert response.status_code == 400
    assert "no permite update_text" in response.json()["detail"]


def test_replace_fragment_requires_kind_compatible_ooxml(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("INSPYRO_APP_STATE_DIR", str(tmp_path / "state"))
    word_live.reset_word_live_cache()
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(_make_docx_with_content_controls())

    client = TestClient(app)
    open_response = client.post("/api/word-live/session/open", json={"document_path": str(docx_path)})
    session_id = open_response.json()["session_id"]

    response = client.post(
        "/api/word-live/region/replace-fragment",
        json={
            "session_id": session_id,
            "region_id": "region-table",
            "fragment_ooxml": (
                f'<?xml version="1.0" encoding="UTF-8"?>'
                f'<w:p xmlns:w="{WORD_W_NS}"><w:r><w:t>Bad table replacement</w:t></w:r></w:p>'
            ),
            "document_version": 1,
        },
    )

    assert response.status_code == 400
    assert "debe contener w:tbl" in response.json()["detail"]


def test_resync_merges_live_word_region_metadata(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("INSPYRO_APP_STATE_DIR", str(tmp_path / "state"))
    word_live.reset_word_live_cache()
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(_make_docx_with_content_controls())

    client = TestClient(app)
    open_response = client.post("/api/word-live/session/open", json={"document_path": str(docx_path)})
    session_payload = open_response.json()

    response = client.post(
        "/api/word-live/session/resync",
        json={
            "session_id": session_payload["session_id"],
            "document_version": 2,
            "regions": [
                {
                    "region_id": "region-intro",
                    "kind": "text",
                    "content_control_tag": "region-intro",
                    "title": "Introduction",
                    "allowed_ops": ["update_text"],
                    "supports_ooxml_replace": False,
                    "word_control_id": 99,
                    "text_preview": "Changed from Word",
                }
            ],
        },
    )

    assert response.status_code == 200
    payload = response.json()
    intro = next(item for item in payload["regions"] if item["region_id"] == "region-intro")
    assert intro["word_control_id"] == 99
    assert intro["text_preview"] == "Changed from Word"
    assert payload["document_version"] == 2


def test_word_addin_assets_are_served() -> None:
    client = TestClient(app)
    manifest = client.get("/word-addin/manifest.xml")
    taskpane = client.get("/word-addin/taskpane.html")

    if not word_live._manifest_path().exists():
        assert manifest.status_code == 404
        assert taskpane.status_code == 404
        return

    assert manifest.status_code == 200
    assert "Inspyro Word Live" in manifest.text
    assert taskpane.status_code == 200
    assert "Structured editing over content controls" in taskpane.text


def test_word_live_launcher_install_endpoint(monkeypatch) -> None:
    client = TestClient(app)
    monkeypatch.setattr(
        word_live,
        "install_word_live_addin",
        lambda: {
            "installed": True,
            "registered": True,
            "addin_id": "addin-demo",
            "manifest_path": "C:\\demo\\manifest.xml",
            "platform": "win32",
        },
    )

    response = client.post("/api/word-live/launcher/install")

    assert response.status_code == 200
    payload = response.json()
    assert payload["installed"] is True
    assert payload["registered"] is True
    assert payload["addin_id"] == "addin-demo"


def test_word_live_launcher_open_endpoint_returns_selected_document(monkeypatch) -> None:
    client = TestClient(app)
    monkeypatch.setattr(
        word_live,
        "open_word_live_launcher",
        lambda **_: {
            "cancelled": False,
            "launched": True,
            "installed": True,
            "registered": True,
            "document_path": "C:\\Temp\\Inspyro\\demo.docx",
            "target_document_name": "demo.docx",
            "host_document_name": "Word add-in demo.docx",
        },
    )

    response = client.post("/api/word-live/launcher/open", json={"use_picker": True})

    assert response.status_code == 200
    payload = response.json()
    assert payload["launched"] is True
    assert payload["document_path"].endswith("demo.docx")
    assert payload["target_document_name"] == "demo.docx"
