"""Service for extracting and managing DOCX templates.

This service handles:
- Extracting styles, page setup, headers/footers from a DOCX file
- Persisting template data (JSON + original DOCX)
- Retrieving active templates for a kernel
- Deleting templates
"""

from __future__ import annotations

import asyncio
import base64
import copy
import functools
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
import io
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

from app.services.docx_quality.content_controls import inspect_content_controls
from app.services.workspace_service import get_app_storage_dir

# File locking imports
if sys.platform == "win32":
    import msvcrt
    try:
        import winreg
    except ImportError:  # pragma: no cover - defensive on non-standard runtimes
        winreg = None  # type: ignore[assignment]
else:
    import fcntl


try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
    from docx.enum.text import WD_LINE_SPACING
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

# FIX #3: Process tracking for zombie cleanup
import atexit
import threading
try:
    import psutil
    HAS_PSUTIL = True
except ImportError:
    psutil = None
    HAS_PSUTIL = False

_LEGACY_TEMPLATE_DIR = Path(__file__).parent.parent.parent / ".templates"


def _migrate_legacy_template_dir(target_dir: Path) -> None:
    try:
        legacy_dir = _LEGACY_TEMPLATE_DIR.resolve()
        resolved_target = target_dir.resolve()
    except OSError:
        return

    if legacy_dir == resolved_target or not legacy_dir.exists() or not legacy_dir.is_dir():
        return

    try:
        target_has_content = any(resolved_target.iterdir())
    except OSError:
        target_has_content = False

    if target_has_content:
        return

    try:
        shutil.copytree(legacy_dir, resolved_target, dirs_exist_ok=True)
        logging.getLogger(__name__).info(
            "[Template] Migrated default template storage from %s to %s",
            legacy_dir,
            resolved_target,
        )
    except Exception as exc:
        logging.getLogger(__name__).warning(
            "[Template] Could not migrate legacy template storage %s -> %s: %s",
            legacy_dir,
            resolved_target,
            exc,
        )


def _resolve_template_base_dir() -> Path:
    override = os.getenv("INSPYRO_TEMPLATE_DIR")
    if override:
        template_dir = Path(override).expanduser().resolve()
        template_dir.mkdir(parents=True, exist_ok=True)
        return template_dir

    template_dir = get_app_storage_dir("templates")
    _migrate_legacy_template_dir(template_dir)
    template_dir.mkdir(parents=True, exist_ok=True)
    return template_dir


TEMPLATE_DIR = _resolve_template_base_dir()

# Required styles that docx-builder API uses
# These are checked for coverage when displaying template status
REQUIRED_STYLES = {
    "titles": {
        "Title": {"description": "Título principal del documento"},
        "Subtitle": {"description": "Subtítulo del documento"},
    },
    "headings": {
        "Heading 1": {"description": "Encabezado nivel 1"},
        "Heading 2": {"description": "Encabezado nivel 2"},
        "Heading 3": {"description": "Encabezado nivel 3"},
        "Heading 4": {"description": "Encabezado nivel 4"},
        "Heading 5": {"description": "Encabezado nivel 5"},
        "Heading 6": {"description": "Encabezado nivel 6"},
    },
    "body": {
        "Normal": {"description": "Texto normal del documento"},
        "Quote": {"description": "Citas o bloques de texto destacado"},
    },
    "captions": {
        "Caption": {"description": "Títulos y leyendas de figuras o tablas"},
    },
    "lists": {
        "List Bullet": {"description": "Lista con viñetas"},
        "List Number": {"description": "Lista numerada"},
    },
    "tables": {
        "Table Grid": {"description": "Tabla con bordes visibles"},
    },
    "code": {
        "Code": {"description": "Bloque de código o texto monoespaciado"},
    },
}

REQUIRED_STYLE_IDS = {
    "Title": "Title",
    "Subtitle": "Subtitle",
    "Heading 1": "Heading1",
    "Heading 2": "Heading2",
    "Heading 3": "Heading3",
    "Heading 4": "Heading4",
    "Heading 5": "Heading5",
    "Heading 6": "Heading6",
    "Normal": "Normal",
    "Caption": "Caption",
    "Quote": "Quote",
    "List Bullet": "ListBullet",
    "List Number": "ListNumber",
    "Table Grid": "TableGrid",
    "Code": "Code",
}

TABLE_STYLE_RUNTIME_DEFAULTS_KEY = "table_style_runtime_defaults"
BUILDER_REQUIRED_STYLE_DEFAULTS_KEY = "builder_required_style_defaults"
SYSTEM_FONT_CATALOG_KEY = "system_font_catalog"
DOCUMENT_DEFAULTS_KEY = "document_defaults"
SEMANTIC_STYLE_SLOTS_KEY = "semantic_style_slots"
_TABLE_STYLE_RUNTIME_TAGS = {"tblLayout", "tblLook", "tblW"}
_TABLE_STYLE_RUNTIME_LOOK_KEYS = (
    "firstRow",
    "lastRow",
    "firstColumn",
    "lastColumn",
    "noHBand",
    "noVBand",
)
_XML_DECLARATION_RE = re.compile(r"^\ufeff?\s*(<\?xml[^>]+\?>)", re.IGNORECASE)
_ROOT_START_TAG_RE = re.compile(r"<(?P<tag>[A-Za-z_][\w:.-]*)(?P<attrs>(?:\s+[^<>]*)?)>", re.DOTALL)
_XMLNS_DECLARATION_RE = re.compile(
    r"""\sxmlns(?::(?P<prefix>[A-Za-z_][\w.-]*))?=(?P<quote>["'])(?P<uri>.*?)(?P=quote)""",
    re.DOTALL,
)
_IGNORABLE_ATTR_RE = re.compile(
    r"""\s(?:[A-Za-z_][\w.-]*:)?Ignorable=(?P<quote>["'])(?P<value>.*?)(?P=quote)""",
    re.DOTALL,
)

# FIX #3: Global process tracking
_active_preview_processes: Dict[str, Set[int]] = {}
_process_lock = threading.Lock()
_xml_namespace_lock = threading.Lock()


DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

ET.register_namespace("w", DOCX_NS["w"])
ET.register_namespace("a", DOCX_NS["a"])

_logger = logging.getLogger(__name__)
_template_executor_workers = max(1, int(os.getenv("INSPYRO_TEMPLATE_EXECUTOR_WORKERS", "4")))
_template_executor = ThreadPoolExecutor(
    max_workers=_template_executor_workers,
    thread_name_prefix="template-worker",
)
_template_executor_shutdown = False
_template_executor_state_lock = threading.Lock()
_system_font_catalog_lock = threading.Lock()
_system_font_catalog_cache: Optional[List[str]] = None


def _coerce_optional_int(value: Any) -> Optional[int]:
    if value in (None, ""):
        return None
    try:
        return int(round(float(value)))
    except (TypeError, ValueError):
        return None


def _normalize_table_runtime_look(look: Any) -> Optional[Dict[str, bool]]:
    if not isinstance(look, dict):
        return None
    normalized: Dict[str, bool] = {}
    for key in _TABLE_STYLE_RUNTIME_LOOK_KEYS:
        bool_value = _coerce_boolish(look.get(key))
        if bool_value is not None:
            normalized[key] = bool_value
    return normalized or None


def _normalize_runtime_defaults_entry(
    style_id: Any,
    *,
    style_name: Any = None,
    layout_type: Any = None,
    width_type: Any = None,
    width_value: Any = None,
    look: Any = None,
) -> Optional[Dict[str, Any]]:
    style_id_text = str(style_id).strip() if style_id not in (None, "") else ""
    if not style_id_text:
        return None

    entry: Dict[str, Any] = {"style_id": style_id_text}
    if style_name not in (None, ""):
        entry["style_name"] = str(style_name)

    layout_text = str(layout_type).strip().lower() if layout_type not in (None, "") else ""
    if layout_text:
        entry["layout_type"] = layout_text

    width_type_text = str(width_type).strip().lower() if width_type not in (None, "") else ""
    if width_type_text:
        entry["width_type"] = width_type_text

    width_int = _coerce_optional_int(width_value)
    if width_int is not None:
        entry["width_value"] = width_int

    normalized_look = _normalize_table_runtime_look(look)
    if normalized_look:
        entry["look"] = normalized_look

    if len(entry) == 1 and "style_name" not in entry:
        return None
    if not any(key in entry for key in ("layout_type", "width_type", "width_value", "look")):
        return None
    return entry


def _normalize_table_style_runtime_defaults(data: Any) -> Dict[str, Dict[str, Any]]:
    if not isinstance(data, dict):
        return {}
    normalized: Dict[str, Dict[str, Any]] = {}
    for key, raw_entry in data.items():
        entry_dict = raw_entry if isinstance(raw_entry, dict) else {}
        normalized_entry = _normalize_runtime_defaults_entry(
            entry_dict.get("style_id") or key,
            style_name=entry_dict.get("style_name"),
            layout_type=entry_dict.get("layout_type"),
            width_type=entry_dict.get("width_type"),
            width_value=entry_dict.get("width_value"),
            look=entry_dict.get("look"),
        )
        if normalized_entry:
            normalized[normalized_entry["style_id"]] = normalized_entry
    return normalized


def _merge_table_style_runtime_defaults(
    current_defaults: Any,
    incoming_defaults: Any,
) -> Dict[str, Dict[str, Any]]:
    merged = _normalize_table_style_runtime_defaults(current_defaults)
    incoming = _normalize_table_style_runtime_defaults(incoming_defaults)
    for style_id, incoming_entry in incoming.items():
        entry = dict(merged.get(style_id, {}))
        entry["style_id"] = style_id
        if incoming_entry.get("style_name"):
            entry["style_name"] = incoming_entry["style_name"]
        for key in ("layout_type", "width_type", "width_value", "look"):
            if key in incoming_entry:
                entry[key] = copy.deepcopy(incoming_entry[key])
        normalized_entry = _normalize_runtime_defaults_entry(
            style_id,
            style_name=entry.get("style_name"),
            layout_type=entry.get("layout_type"),
            width_type=entry.get("width_type"),
            width_value=entry.get("width_value"),
            look=entry.get("look"),
        )
        if normalized_entry:
            merged[style_id] = normalized_entry
        else:
            merged.pop(style_id, None)
    return merged


def _apply_runtime_defaults_patch(
    current_defaults: Any,
    *,
    style_id: Any,
    style_name: Any = None,
    patch: Optional[Dict[str, Any]] = None,
) -> Dict[str, Dict[str, Any]]:
    merged = _normalize_table_style_runtime_defaults(current_defaults)
    style_id_text = str(style_id).strip() if style_id not in (None, "") else ""
    if not style_id_text:
        return merged

    patch = patch if isinstance(patch, dict) else {}
    entry = dict(merged.get(style_id_text, {}))
    entry["style_id"] = style_id_text
    if style_name not in (None, ""):
        entry["style_name"] = str(style_name)

    for key in ("layout_type", "width_type", "width_value", "look"):
        if key not in patch:
            continue
        value = patch.get(key)
        if key == "look":
            normalized_value = _normalize_table_runtime_look(value)
            if normalized_value:
                entry[key] = normalized_value
            else:
                entry.pop(key, None)
            continue
        if value in (None, ""):
            entry.pop(key, None)
        elif key == "width_value":
            coerced = _coerce_optional_int(value)
            if coerced is None:
                entry.pop(key, None)
            else:
                entry[key] = coerced
        else:
            entry[key] = str(value).strip().lower()

    normalized_entry = _normalize_runtime_defaults_entry(
        style_id_text,
        style_name=entry.get("style_name"),
        layout_type=entry.get("layout_type"),
        width_type=entry.get("width_type"),
        width_value=entry.get("width_value"),
        look=entry.get("look"),
    )
    if normalized_entry:
        merged[style_id_text] = normalized_entry
    else:
        merged.pop(style_id_text, None)
    return merged


def _extract_runtime_defaults_from_table_format(table_format: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(table_format, dict):
        return {}
    patch: Dict[str, Any] = {}
    if "layout_type" in table_format:
        patch["layout_type"] = table_format.get("layout_type")
    if "width_type" in table_format:
        patch["width_type"] = table_format.get("width_type")
    if "width_value" in table_format:
        patch["width_value"] = table_format.get("width_value")
    if "look" in table_format:
        patch["look"] = table_format.get("look")
    return patch


def _extract_runtime_defaults_from_updates(updates: Dict[str, Any]) -> Dict[str, Any]:
    patch: Dict[str, Any] = {}
    if "table_layout_type" in updates:
        patch["layout_type"] = updates.get("table_layout_type")
    width_type_present = "table_width_type" in updates
    width_value_present = "table_width_value" in updates
    if width_type_present:
        patch["width_type"] = updates.get("table_width_type")
    if width_value_present:
        patch["width_value"] = updates.get("table_width_value")
    look_patch: Dict[str, Any] = {}
    look_key_map = {
        "firstRow": "table_look_first_row",
        "lastRow": "table_look_last_row",
        "firstColumn": "table_look_first_column",
        "lastColumn": "table_look_last_column",
        "noHBand": "table_look_no_h_band",
        "noVBand": "table_look_no_v_band",
    }
    for look_key, update_key in look_key_map.items():
        if update_key in updates:
            look_patch[look_key] = updates.get(update_key)
    if look_patch:
        patch["look"] = look_patch
    return patch


def _extract_runtime_defaults_from_tbl_pr_props(props: Any) -> tuple[Any, Dict[str, Any]]:
    if not isinstance(props, list):
        return props, {}

    filtered: List[Dict[str, Any]] = []
    runtime_patch: Dict[str, Any] = {}
    look_patch: Dict[str, Any] = {}

    for node in props:
        if not isinstance(node, dict):
            filtered.append(node)
            continue
        tag = node.get("tag")
        attrs = node.get("attrs") or {}
        if tag == "tblLayout":
            runtime_patch["layout_type"] = attrs.get("type")
            continue
        if tag == "tblW":
            runtime_patch["width_type"] = attrs.get("type")
            runtime_patch["width_value"] = attrs.get("w")
            continue
        if tag == "tblLook":
            for key in _TABLE_STYLE_RUNTIME_LOOK_KEYS:
                if key in attrs:
                    look_patch[key] = attrs.get(key)
            continue
        filtered.append(copy.deepcopy(node))

    if look_patch:
        runtime_patch["look"] = look_patch
    return filtered, runtime_patch


def _sanitize_advanced_props_for_table_style(advanced_props: Any) -> tuple[Any, Dict[str, Any]]:
    if not isinstance(advanced_props, dict):
        return advanced_props, {}

    sanitized = copy.deepcopy(advanced_props)
    runtime_patch: Dict[str, Any] = {}

    filtered_tbl_pr, tbl_pr_patch = _extract_runtime_defaults_from_tbl_pr_props(sanitized.get("tbl_pr"))
    if "tbl_pr" in sanitized:
        sanitized["tbl_pr"] = filtered_tbl_pr
    if tbl_pr_patch:
        runtime_patch.update(tbl_pr_patch)

    return sanitized, runtime_patch


def _merge_runtime_defaults_into_extracted(
    extracted: Dict[str, Any],
    previous_slots: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    runtime_defaults = _normalize_table_style_runtime_defaults(extracted.get(TABLE_STYLE_RUNTIME_DEFAULTS_KEY))
    extracted[TABLE_STYLE_RUNTIME_DEFAULTS_KEY] = runtime_defaults
    if runtime_defaults:
        runtime_by_name = {
            entry.get("style_name"): entry
            for entry in runtime_defaults.values()
            if entry.get("style_name")
        }

        for style_info in extracted.get("styles", []):
            if style_info.get("type") != "table":
                continue
            runtime_entry = (
                runtime_defaults.get(style_info.get("style_id"))
                or runtime_by_name.get(style_info.get("display_name"))
                or runtime_by_name.get(style_info.get("name"))
            )
            if not runtime_entry:
                continue

            resolved_table = dict(style_info.get("resolved_table_format") or style_info.get("xml_table_format") or {})
            if runtime_entry.get("layout_type"):
                resolved_table["layout_type"] = runtime_entry["layout_type"]
            if runtime_entry.get("width_type"):
                resolved_table["width_type"] = runtime_entry["width_type"]
            if "width_value" in runtime_entry:
                resolved_table["width_value"] = runtime_entry["width_value"]
            if runtime_entry.get("look"):
                merged_look = dict(resolved_table.get("look") or {})
                merged_look.update(runtime_entry["look"])
                resolved_table["look"] = merged_look

            style_info["resolved_table_format"] = resolved_table
            style_info["table_runtime_defaults"] = copy.deepcopy(runtime_entry)

    from app.services import template_extract

    return template_extract.enrich_template_metadata(extracted, previous_slots=previous_slots)


def _merge_runtime_patch_dicts(*patches: Any) -> Dict[str, Any]:
    merged: Dict[str, Any] = {}
    for patch in patches:
        if not isinstance(patch, dict):
            continue
        for key in ("layout_type", "width_type", "width_value"):
            if key in patch:
                merged[key] = patch.get(key)
        if "look" in patch:
            look_value = patch.get("look")
            if isinstance(look_value, dict):
                merged["look"] = dict(merged.get("look") or {})
                merged["look"].update(look_value)
            else:
                merged["look"] = look_value
    return merged


def _find_template_style_info(
    template_info: Optional[Dict[str, Any]],
    style_name: Optional[str] = None,
    style_id: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    if not isinstance(template_info, dict):
        return None

    styles = template_info.get("styles")
    if not isinstance(styles, list):
        return None

    style_name_text = str(style_name).strip() if style_name not in (None, "") else None
    style_id_text = str(style_id).strip() if style_id not in (None, "") else None

    for style_info in styles:
        if not isinstance(style_info, dict):
            continue
        if style_id_text and style_info.get("style_id") == style_id_text:
            return style_info

    if not style_name_text:
        return None

    for style_info in styles:
        if not isinstance(style_info, dict):
            continue
        if style_info.get("display_name") == style_name_text or style_info.get("name") == style_name_text:
            return style_info
    return None


def _rewrite_docx_bytes(docx_bytes: bytes, updates: Dict[str, bytes]) -> bytes:
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            written: Set[str] = set()
            for item in zin.infolist():
                if item.filename in updates:
                    zout.writestr(item, updates[item.filename])
                    written.add(item.filename)
                else:
                    zout.writestr(item, zin.read(item.filename))
            for filename, data in updates.items():
                if filename not in written:
                    zout.writestr(filename, data)
    return dst.getvalue()


def _coerce_xml_text(xml_source: Any) -> str:
    if isinstance(xml_source, bytes):
        return xml_source.decode("utf-8", errors="ignore")
    if isinstance(xml_source, str):
        return xml_source
    return str(xml_source or "")


def _extract_xml_declaration(xml_source: Any) -> Optional[str]:
    xml_text = _coerce_xml_text(xml_source)
    match = _XML_DECLARATION_RE.match(xml_text)
    if not match:
        return None
    return match.group(1)


def _strip_xml_declaration(xml_source: Any) -> str:
    xml_text = _coerce_xml_text(xml_source)
    return _XML_DECLARATION_RE.sub("", xml_text, count=1)


def _extract_root_start_tag(xml_source: Any) -> Optional[str]:
    xml_body = _strip_xml_declaration(xml_source)
    match = _ROOT_START_TAG_RE.search(xml_body)
    if not match:
        return None
    return match.group(0)


def _extract_namespace_declarations_from_root_tag(root_tag: Optional[str]) -> Dict[str, str]:
    declarations: Dict[str, str] = {}
    if not root_tag:
        return declarations
    for match in _XMLNS_DECLARATION_RE.finditer(root_tag):
        prefix = match.group("prefix") or ""
        uri = match.group("uri")
        if uri:
            declarations[prefix] = uri
    return declarations


def _extract_root_namespace_declarations(xml_source: Any) -> Dict[str, str]:
    return _extract_namespace_declarations_from_root_tag(_extract_root_start_tag(xml_source))


def _extract_root_ignorable_prefixes(xml_source: Any) -> Set[str]:
    root_tag = _extract_root_start_tag(xml_source)
    if not root_tag:
        return set()
    match = _IGNORABLE_ATTR_RE.search(root_tag)
    if not match:
        return set()
    return {token for token in match.group("value").split() if token}


def _is_generated_namespace_prefix(prefix: str) -> bool:
    return bool(prefix) and prefix.startswith("ns") and prefix[2:].isdigit()


def _normalize_namespace_hints(namespace_hints: Any) -> Dict[str, str]:
    if not isinstance(namespace_hints, dict):
        return {}
    normalized: Dict[str, str] = {}
    for raw_prefix, raw_uri in namespace_hints.items():
        if raw_uri in (None, ""):
            continue
        prefix = "" if raw_prefix in (None, "") else str(raw_prefix)
        normalized[prefix] = str(raw_uri)
    return normalized


def _filter_required_namespace_declarations(
    required_declarations: Dict[str, str],
    namespace_hints: Dict[str, str],
) -> Dict[str, str]:
    if not required_declarations:
        return {}

    preferred_prefix_by_uri = {
        uri: prefix
        for prefix, uri in namespace_hints.items()
        if prefix and not _is_generated_namespace_prefix(prefix)
    }

    filtered: Dict[str, str] = {}
    for prefix, uri in required_declarations.items():
        if _is_generated_namespace_prefix(prefix) and preferred_prefix_by_uri.get(uri):
            continue
        filtered[prefix] = uri
    return filtered


def _inject_namespace_declarations_into_root(
    xml_text: str,
    declarations: List[tuple[str, str]],
) -> str:
    root_tag = _extract_root_start_tag(xml_text)
    if not root_tag or not declarations:
        return xml_text

    suffix = "/>" if root_tag.endswith("/>") else ">"
    insertion = "".join(
        f' xmlns:{prefix}="{uri}"' if prefix else f' xmlns="{uri}"'
        for prefix, uri in declarations
    )
    updated_root = root_tag[:-len(suffix)] + insertion + suffix
    return xml_text.replace(root_tag, updated_root, 1)


def _ensure_root_namespace_declarations(
    xml_source: Any,
    *,
    required_declarations: Any = None,
    namespace_hints: Any = None,
) -> tuple[bytes, bool]:
    xml_text = _coerce_xml_text(xml_source)
    root_tag = _extract_root_start_tag(xml_text)
    if not root_tag:
        return xml_text.encode("utf-8"), False

    current_declarations = _extract_namespace_declarations_from_root_tag(root_tag)
    hint_declarations = _normalize_namespace_hints(namespace_hints)
    required = _filter_required_namespace_declarations(
        _normalize_namespace_hints(required_declarations),
        hint_declarations,
    )

    missing: List[tuple[str, str]] = []
    for prefix, uri in required.items():
        if prefix not in current_declarations and uri:
            missing.append((prefix, uri))

    missing_prefixes = {prefix for prefix, _ in missing}
    combined_hints = dict(hint_declarations)
    combined_hints.update(required)
    for prefix in sorted(_extract_root_ignorable_prefixes(root_tag)):
        if prefix in current_declarations or prefix in missing_prefixes:
            continue
        uri = combined_hints.get(prefix)
        if uri:
            missing.append((prefix, uri))
            missing_prefixes.add(prefix)

    if not missing:
        return xml_text.encode("utf-8"), False

    updated_xml = _inject_namespace_declarations_into_root(xml_text, missing)
    return updated_xml.encode("utf-8"), True


def _collect_docx_namespace_hints_from_zip(zin: zipfile.ZipFile) -> Dict[str, str]:
    hints: Dict[str, str] = dict(DOCX_NS)
    hints.setdefault("mc", "http://schemas.openxmlformats.org/markup-compatibility/2006")

    for name in zin.namelist():
        if not name.endswith(".xml"):
            continue
        try:
            xml_text = zin.read(name).decode("utf-8", errors="ignore")
        except Exception:
            continue
        declarations = _extract_root_namespace_declarations(xml_text)
        for prefix, uri in declarations.items():
            if prefix in hints and hints[prefix] == uri:
                continue
            existing_uri = hints.get(prefix)
            if prefix not in hints or not existing_uri:
                hints[prefix] = uri
                continue
            if _is_generated_namespace_prefix(prefix):
                continue
            hints[prefix] = uri
    return hints


def _collect_docx_namespace_hints_from_bytes(docx_bytes: bytes) -> Dict[str, str]:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
            return _collect_docx_namespace_hints_from_zip(zin)
    except Exception:
        return {**DOCX_NS, "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006"}


def _collect_docx_namespace_hints_from_path(docx_path: Path) -> Dict[str, str]:
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            return _collect_docx_namespace_hints_from_zip(zin)
    except Exception:
        return {**DOCX_NS, "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006"}


def _register_namespace_hints(namespace_hints: Any) -> None:
    for prefix, uri in _normalize_namespace_hints(namespace_hints).items():
        if prefix in {"xml", "xmlns"} or _is_generated_namespace_prefix(prefix):
            continue
        try:
            ET.register_namespace(prefix, uri)
        except ValueError:
            continue


def _serialize_ooxml_part(
    element: ET.Element,
    original_xml: Any,
    *,
    namespace_hints: Any = None,
) -> bytes:
    original_text = _coerce_xml_text(original_xml)
    declaration = _extract_xml_declaration(original_text) or "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
    original_declarations = _extract_root_namespace_declarations(original_text)
    combined_hints = _normalize_namespace_hints(namespace_hints)
    combined_hints.update(original_declarations)

    with _xml_namespace_lock:
        _register_namespace_hints(combined_hints)
        serialized_body = ET.tostring(element, encoding="unicode")

    ensured_body, _ = _ensure_root_namespace_declarations(
        serialized_body,
        required_declarations=original_declarations,
        namespace_hints=combined_hints,
    )
    return f"{declaration}{ensured_body.decode('utf-8')}".encode("utf-8")


def _repair_ooxml_namespace_declarations_in_docx(docx_bytes: bytes) -> tuple[bytes, bool]:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)
            updates: Dict[str, bytes] = {}
            for name in zin.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    xml_text = zin.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue
                repaired_bytes, changed = _ensure_root_namespace_declarations(
                    xml_text,
                    required_declarations=_extract_root_namespace_declarations(xml_text),
                    namespace_hints=namespace_hints,
                )
                if changed:
                    updates[name] = repaired_bytes

        if not updates:
            return docx_bytes, False
        return _rewrite_docx_bytes(docx_bytes, updates), True
    except Exception as exc:
        _logger.warning("[Template] Could not repair OOXML namespace declarations: %s", exc)
        return docx_bytes, False


def _sanitize_table_style_runtime_defaults_in_docx(
    docx_bytes: bytes,
    existing_defaults: Any = None,
) -> tuple[bytes, Dict[str, Dict[str, Any]], bool]:
    runtime_defaults = _normalize_table_style_runtime_defaults(existing_defaults)
    try:
        namespace_hints = _collect_docx_namespace_hints_from_bytes(docx_bytes)
        parts = _read_docx_parts(docx_bytes, ["word/styles.xml"])
        styles_xml = parts.get("word/styles.xml")
        if not styles_xml:
            return docx_bytes, runtime_defaults, False

        repaired_styles, namespace_changed = _ensure_root_namespace_declarations(
            styles_xml,
            required_declarations=_extract_root_namespace_declarations(styles_xml),
            namespace_hints=namespace_hints,
        )
        styles_source = repaired_styles.decode("utf-8")
        styles_root = ET.fromstring(styles_source)
        changed = False
        discovered_defaults: Dict[str, Dict[str, Any]] = {}

        for style_elem in styles_root.findall("w:style", DOCX_NS):
            if style_elem.get(_qn("w", "type")) != "table":
                continue
            style_id = style_elem.get(_qn("w", "styleId"))
            if not style_id:
                continue
            style_name = _find_val(style_elem, "w:name") or style_id
            tbl_pr = style_elem.find("w:tblPr", DOCX_NS)
            if tbl_pr is None:
                continue

            runtime_patch: Dict[str, Any] = {}
            for child in list(tbl_pr):
                local_tag = _local_name(child.tag)
                if local_tag == "tblLayout":
                    runtime_patch["layout_type"] = child.get(_qn("w", "type"))
                    tbl_pr.remove(child)
                    changed = True
                elif local_tag == "tblW":
                    runtime_patch["width_type"] = child.get(_qn("w", "type"))
                    runtime_patch["width_value"] = child.get(_qn("w", "w"))
                    tbl_pr.remove(child)
                    changed = True
                elif local_tag == "tblLook":
                    look: Dict[str, Any] = {}
                    for look_key in _TABLE_STYLE_RUNTIME_LOOK_KEYS:
                        attr_value = child.get(_qn("w", look_key))
                        if attr_value is not None:
                            look[look_key] = attr_value
                    if look:
                        runtime_patch["look"] = look
                    tbl_pr.remove(child)
                    changed = True

            if len(list(tbl_pr)) == 0:
                style_elem.remove(tbl_pr)

            normalized_entry = _normalize_runtime_defaults_entry(
                style_id,
                style_name=style_name,
                layout_type=runtime_patch.get("layout_type"),
                width_type=runtime_patch.get("width_type"),
                width_value=runtime_patch.get("width_value"),
                look=runtime_patch.get("look"),
            )
            if normalized_entry:
                discovered_defaults[style_id] = normalized_entry

        if discovered_defaults:
            runtime_defaults = _merge_table_style_runtime_defaults(runtime_defaults, discovered_defaults)

        if not changed and not namespace_changed:
            return docx_bytes, runtime_defaults, False

        if changed:
            updated_styles = _serialize_ooxml_part(
                styles_root,
                styles_source,
                namespace_hints=namespace_hints,
            )
        else:
            updated_styles = repaired_styles
        updated_docx = _rewrite_docx_bytes(docx_bytes, {"word/styles.xml": updated_styles})
        return updated_docx, runtime_defaults, True
    except Exception as exc:
        _logger.warning("[Template] Could not sanitize table runtime defaults from DOCX: %s", exc)
        return docx_bytes, runtime_defaults, False


def _needs_font_metadata_refresh(extracted_json: Optional[Dict[str, Any]]) -> bool:
    if not isinstance(extracted_json, dict):
        return True

    required_top_level = (
        "default_font",
        "default_font_source",
        "font_catalog",
        SYSTEM_FONT_CATALOG_KEY,
        BUILDER_REQUIRED_STYLE_DEFAULTS_KEY,
        DOCUMENT_DEFAULTS_KEY,
        SEMANTIC_STYLE_SLOTS_KEY,
    )
    if any(key not in extracted_json for key in required_top_level):
        return True

    xml_details = extracted_json.get("xml_details")
    if not isinstance(xml_details, dict):
        return True
    if "body_font_hint" not in xml_details:
        return True

    for style_info in extracted_json.get("styles") or []:
        if not isinstance(style_info, dict):
            continue
        if style_info.get("resolved_font") and "resolved_font_source" not in style_info:
            return True
        if style_info.get("xml_font") and "xml_font_source" not in style_info:
            return True
    return False


def _prepare_template_payload(
    docx_bytes: bytes,
    extracted_json: Optional[Dict[str, Any]] = None,
) -> tuple[bytes, Dict[str, Any], bool]:
    extracted = copy.deepcopy(extracted_json) if isinstance(extracted_json, dict) else extract_styles_from_docx(docx_bytes)
    previous_slots = (
        copy.deepcopy(extracted.get(SEMANTIC_STYLE_SLOTS_KEY))
        if isinstance(extracted.get(SEMANTIC_STYLE_SLOTS_KEY), dict)
        else None
    )
    runtime_defaults = _normalize_table_style_runtime_defaults(extracted.get(TABLE_STYLE_RUNTIME_DEFAULTS_KEY))
    namespace_repaired_bytes, namespace_changed = _repair_ooxml_namespace_declarations_in_docx(docx_bytes)
    sanitized_bytes, runtime_defaults, docx_changed = _sanitize_table_style_runtime_defaults_in_docx(
        namespace_repaired_bytes,
        runtime_defaults,
    )
    if docx_changed or namespace_changed or _needs_font_metadata_refresh(extracted):
        extracted = extract_styles_from_docx(sanitized_bytes)
    extracted[TABLE_STYLE_RUNTIME_DEFAULTS_KEY] = runtime_defaults
    _merge_runtime_defaults_into_extracted(extracted, previous_slots=previous_slots)
    return sanitized_bytes, extracted, (docx_changed or namespace_changed)


async def run_template_executor(func, *args, **kwargs):
    """Run template-heavy work on a dedicated executor."""
    loop = asyncio.get_running_loop()
    call = functools.partial(func, *args, **kwargs)
    return await loop.run_in_executor(_template_executor, call)


def get_template_executor_stats() -> dict:
    with _template_executor_state_lock:
        shutdown = _template_executor_shutdown
    return {
        "template_executor_max_workers": _template_executor_workers,
        "template_executor_shutdown": shutdown,
    }


def shutdown_template_executor() -> None:
    global _template_executor_shutdown
    with _template_executor_state_lock:
        if _template_executor_shutdown:
            return
        _template_executor_shutdown = True
    _template_executor.shutdown(wait=False, cancel_futures=True)


# ==============================================================================
# FIX #3: Process cleanup functions
# ==============================================================================

def _register_preview_process(kernel_id, pid):
    if not HAS_PSUTIL:  
        return
    with _process_lock:
        if kernel_id not in _active_preview_processes:
            _active_preview_processes[kernel_id] = set()
        _active_preview_processes[kernel_id].add(pid)
        _logger.debug(f'[Template] Registered process {pid} for kernel {kernel_id}')


def _unregister_preview_process(kernel_id, pid):
    if not HAS_PSUTIL:
        return
    with _process_lock:
        if kernel_id in _active_preview_processes:
            _active_preview_processes[kernel_id].discard(pid)
            if not _active_preview_processes[kernel_id]:
                del _active_preview_processes[kernel_id]
            _logger.debug(f'[Template] Unregistered process {pid} for kernel {kernel_id}')


def _kill_process_tree(pid):
    if not HAS_PSUTIL:
        return
    try:
        parent = psutil.Process(pid)
        children = parent.children(recursive=True)
        for child in children:
            try:
                child.terminate()
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        gone, alive = psutil.wait_procs(children, timeout=3)
        for p in alive:
            try:
                p.kill()
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        try:
            parent.terminate()
            parent.wait(timeout=3)
        except psutil.TimeoutExpired:
            parent.kill()
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            pass
        _logger.info(f'[Template] Killed process tree for PID {pid}')
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        _logger.debug(f'[Template] Process {pid} already gone: {e}')
    except Exception as e:
        _logger.warning(f'[Template] Error killing process {pid}: {e}')


def cleanup_kernel_processes(kernel_id):
    if not HAS_PSUTIL:
        return
    with _process_lock:
        pids = _active_preview_processes.get(kernel_id, set()).copy()
    if not pids:
        return
    _logger.info(f'[Template] Cleaning up {len(pids)} process(es) for kernel {kernel_id}')
    for pid in pids:
        _kill_process_tree(pid)
        _unregister_preview_process(kernel_id, pid)


def _cleanup_all_processes():
    if not HAS_PSUTIL:
        return
    with _process_lock:
        all_pids = []
        for kernel_id, pids in _active_preview_processes.items():
            all_pids.extend((kernel_id, pid) for pid in pids)
    if all_pids:
        _logger.info(f'[Template] Shutdown: cleaning up {len(all_pids)} process(es)')
        for kernel_id, pid in all_pids:
            _kill_process_tree(pid)


if HAS_PSUTIL:
    atexit.register(_cleanup_all_processes)

ET.register_namespace("r", DOCX_NS["r"])


def _qn(prefix: str, tag: str) -> str:
    return f"{{{DOCX_NS[prefix]}}}{tag}"


def _local_name(tag: str) -> str:
    return tag.split("}", 1)[1] if "}" in tag else tag


def _as_bool(value: Optional[str]) -> Optional[bool]:
    if value is None:
        return None
    return str(value).lower() in {"1", "true", "yes", "on"}


def _serialize_xml(element: Optional[ET.Element]) -> Optional[str]:
    if element is None:
        return None
    try:
        return ET.tostring(element, encoding="unicode")
    except Exception:
        return None


def _collect_props(element: Optional[ET.Element]) -> Optional[List[Dict[str, Any]]]:
    if element is None:
        return None
    props: List[Dict[str, Any]] = []
    for child in list(element):
        entry: Dict[str, Any] = {
            "tag": _local_name(child.tag),
        }
        if child.attrib:
            entry["attrs"] = {_local_name(k): v for k, v in child.attrib.items()}
        text = (child.text or "").strip()
        if text:
            entry["text"] = text
        if list(child):
            entry["children"] = _collect_props(child)
        props.append(entry)
    return props


def _append_props(parent: ET.Element, props: List[Dict[str, Any]]) -> None:
    for prop in props:
        tag = prop.get("tag")
        if not tag:
            continue
        child = ET.Element(_qn("w", tag))
        attrs = prop.get("attrs") or {}
        for attr_name, attr_value in attrs.items():
            child.set(_qn("w", attr_name), str(attr_value))
        text_value = prop.get("text")
        if text_value is not None:
            child.text = str(text_value)
        children = prop.get("children") or []
        if children:
            _append_props(child, children)
        parent.append(child)


def _build_props_element(tag: str, props: List[Dict[str, Any]]) -> ET.Element:
    elem = ET.Element(_qn("w", tag))
    _append_props(elem, props)
    return elem


def _find_style_element(
    styles_root: ET.Element,
    style_name: Optional[str],
    style_id: Optional[str]
) -> Optional[ET.Element]:
    if style_id:
        style = styles_root.find(f"w:style[@w:styleId='{style_id}']", DOCX_NS)
        if style is not None:
            return style
    if style_name:
        for style in styles_root.findall("w:style", DOCX_NS):
            name_elem = style.find("w:name", DOCX_NS)
            if name_elem is not None and name_elem.get(_qn("w", "val")) == style_name:
                return style
        fallback_id = REQUIRED_STYLE_IDS.get(style_name) or style_name.replace(" ", "")
        style = styles_root.find(f"w:style[@w:styleId='{fallback_id}']", DOCX_NS)
        if style is not None:
            return style
    return None


def _replace_child(parent: ET.Element, tag: str, new_elem: Optional[ET.Element]) -> None:
    existing = parent.find(f"w:{tag}", DOCX_NS)
    if existing is not None:
        parent.remove(existing)
    if new_elem is not None:
        parent.append(new_elem)


def _ensure_child(parent: ET.Element, tag: str) -> ET.Element:
    child = parent.find(f"w:{tag}", DOCX_NS)
    if child is None:
        child = ET.Element(_qn("w", tag))
        parent.append(child)
    return child


_RFONTS_EXPLICIT_ATTRS = ("ascii", "hAnsi", "cs", "eastAsia")
_RFONTS_THEME_ATTRS = ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme")


def _set_explicit_family_on_rfonts(r_fonts: ET.Element, font_name: Optional[str]) -> None:
    normalized_name = str(font_name).strip() if font_name not in (None, "") else ""

    for attr_name in (*_RFONTS_EXPLICIT_ATTRS, *_RFONTS_THEME_ATTRS):
        r_fonts.attrib.pop(_qn("w", attr_name), None)

    if not normalized_name:
        return

    for attr_name in _RFONTS_EXPLICIT_ATTRS:
        r_fonts.set(_qn("w", attr_name), normalized_name)


def _sync_style_explicit_font_name(
    docx_path: Path,
    style_name: str,
    style_id: Optional[str],
    font_name: Optional[str],
) -> None:
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                _logger.warning("[Template] styles.xml not found in DOCX")
                return
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)

        styles_root = ET.fromstring(styles_xml)
        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            _logger.warning("[Template] Style '%s' not found in styles.xml for font sync", style_name)
            return

        normalized_name = str(font_name).strip() if font_name not in (None, "") else ""
        r_pr = style_elem.find("w:rPr", DOCX_NS)
        if r_pr is None and not normalized_name:
            return
        if r_pr is None:
            r_pr = ET.Element(_qn("w", "rPr"))
            style_elem.append(r_pr)

        r_fonts = r_pr.find("w:rFonts", DOCX_NS)
        if r_fonts is None:
            r_fonts = ET.Element(_qn("w", "rFonts"))
            r_pr.insert(0, r_fonts)

        _set_explicit_family_on_rfonts(r_fonts, normalized_name)

        if not r_fonts.attrib and not list(r_fonts) and not (r_fonts.text or "").strip():
            r_pr.remove(r_fonts)
        if not list(r_pr) and not r_pr.attrib and not (r_pr.text or "").strip():
            style_elem.remove(r_pr)

        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/styles.xml": updated_xml})
    except Exception as exc:
        _logger.error("[Template] Failed to sync explicit font name for style '%s': %s", style_name, exc)
        raise


def _sync_document_defaults_in_docx(
    docx_path: Path,
    document_defaults: Optional[Dict[str, Any]],
) -> None:
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                _logger.warning("[Template] styles.xml not found in DOCX")
                return
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)

        styles_root = ET.fromstring(styles_xml)
        font_props = _filter_document_default_font_props((document_defaults or {}).get("font"))
        paragraph_props = _filter_document_default_paragraph_props((document_defaults or {}).get("paragraph"))

        doc_defaults = styles_root.find("w:docDefaults", DOCX_NS)
        should_keep_doc_defaults = _has_meaningful_font_props(font_props) or _has_meaningful_paragraph_props(paragraph_props)
        if doc_defaults is None and not should_keep_doc_defaults:
            return
        if doc_defaults is None:
            doc_defaults = ET.Element(_qn("w", "docDefaults"))
            styles_root.insert(0, doc_defaults)

        for child_tag in ("rPrDefault", "pPrDefault"):
            existing = doc_defaults.find(f"w:{child_tag}", DOCX_NS)
            if existing is not None:
                doc_defaults.remove(existing)

        if _has_meaningful_font_props(font_props):
            r_pr_default = ET.SubElement(doc_defaults, _qn("w", "rPrDefault"))
            r_pr = ET.SubElement(r_pr_default, _qn("w", "rPr"))
            _append_rpr_from_font_props(r_pr, font_props)
            if not list(r_pr) and not r_pr.attrib and not (r_pr.text or "").strip():
                r_pr_default.remove(r_pr)
                if not list(r_pr_default) and not r_pr_default.attrib and not (r_pr_default.text or "").strip():
                    doc_defaults.remove(r_pr_default)

        if _has_meaningful_paragraph_props(paragraph_props):
            p_pr_default = ET.SubElement(doc_defaults, _qn("w", "pPrDefault"))
            p_pr = ET.SubElement(p_pr_default, _qn("w", "pPr"))
            _append_ppr_from_paragraph_props(p_pr, paragraph_props)
            if not list(p_pr) and not p_pr.attrib and not (p_pr.text or "").strip():
                p_pr_default.remove(p_pr)
                if not list(p_pr_default) and not p_pr_default.attrib and not (p_pr_default.text or "").strip():
                    doc_defaults.remove(p_pr_default)

        if not list(doc_defaults) and not doc_defaults.attrib and not (doc_defaults.text or "").strip():
            styles_root.remove(doc_defaults)

        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/styles.xml": updated_xml})
    except Exception as exc:
        _logger.error("[Template] Failed to sync document defaults: %s", exc)
        raise


def _write_docx_parts(docx_path: Path, updates: Dict[str, bytes]) -> None:
    """Safely write updated parts to DOCX file with locking to prevent race conditions.
    
    This function:
    1. Acquires an exclusive lock on the file to prevent concurrent modifications
    2. Writes changes to a temporary file in the same directory
    3. Atomically replaces the original file (Unix) or uses backup strategy (Windows)
    4. Cleans up temporary files even if errors occur
    
    Args:
        docx_path: Path to the DOCX file to modify
        updates: Dictionary mapping part paths to their new content bytes
    
    Raises:
        Exception: If file operations fail (lock will be released automatically)
    """
    if not updates:
        return
    
    # Create lock file path
    lock_path = docx_path.with_suffix(docx_path.suffix + ".lock")
    temp_path = None
    backup_path = None
    
    try:
        # Ensure parent directory exists
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Acquire exclusive lock
        with open(lock_path, "w") as lock_file:
            if sys.platform == "win32":
                # Windows: lock 1 byte at position 0
                msvcrt.locking(lock_file.fileno(), msvcrt.LK_LOCK, 1)
            else:
                # Unix: exclusive lock on entire file
                fcntl.flock(lock_file.fileno(), fcntl.LOCK_EX)
            
            try:
                # Generate unique temp file in same directory (required for atomic move)
                temp_path = docx_path.with_suffix(f".{uuid.uuid4().hex[:8]}.tmp")
                
                # Write to temp file
                with zipfile.ZipFile(docx_path, "r") as zin:
                    with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                        updated = set()
                        for item in zin.infolist():
                            if item.filename in updates:
                                zout.writestr(item, updates[item.filename])
                                updated.add(item.filename)
                            else:
                                zout.writestr(item, zin.read(item.filename))
                        # Add any new parts not in original
                        for filename, data in updates.items():
                            if filename not in updated:
                                zout.writestr(filename, data)
                
                # Atomic replacement (platform-specific)
                if sys.platform == "win32":
                    # Windows doesn't support atomic replace, use backup strategy
                    backup_path = docx_path.with_suffix(".backup")
                    
                    # Remove old backup if exists
                    if backup_path.exists():
                        try:
                            backup_path.unlink()
                        except OSError:
                            pass
                    
                    # Move current file to backup
                    if docx_path.exists():
                        try:
                            docx_path.rename(backup_path)
                        except OSError as e:
                            _logger.warning(f"Could not create backup: {e}")
                            backup_path = None
                    
                    # Move temp to target
                    try:
                        temp_path.rename(docx_path)
                        # Success, remove backup
                        if backup_path and backup_path.exists():
                            backup_path.unlink()
                    except Exception as e:
                        # Rollback: restore from backup
                        _logger.error(f"Failed to replace DOCX, rolling back: {e}")
                        if backup_path and backup_path.exists():
                            backup_path.rename(docx_path)
                        raise
                else:
                    # Unix: atomic replace (Python 3.3+)
                    temp_path.replace(docx_path)
                
                temp_path = None  # Successfully moved, don't cleanup
                
            finally:
                # Lock is automatically released when file closes
                pass
    
    finally:
        # Cleanup lock file
        try:
            if lock_path.exists():
                lock_path.unlink()
        except Exception:
            pass
        
        # Cleanup temp file if still exists (error occurred)
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except Exception:
                pass
        
        # Cleanup backup if still exists (shouldn't happen normally)
        if backup_path and backup_path.exists():
            try:
                backup_path.unlink()
            except Exception:
                pass


def _apply_style_xml_updates(
    docx_path: Path,
    style_name: str,
    style_id: Optional[str],
    advanced_props: Dict[str, Any]
) -> Dict[str, Any]:
    if not advanced_props or not isinstance(advanced_props, dict):
        return {}
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                _logger.warning("[Template] styles.xml not found in DOCX")
                return {}
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)
        styles_root = ET.fromstring(styles_xml)
        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            _logger.warning(f"[Template] Style '{style_name}' not found in styles.xml")
            return {}

        runtime_patch: Dict[str, Any] = {}
        effective_props = advanced_props
        if style_elem.get(_qn("w", "type")) == "table":
            effective_props, runtime_patch = _sanitize_advanced_props_for_table_style(advanced_props)

        for key, tag in ("r_pr", "rPr"), ("p_pr", "pPr"), ("tbl_pr", "tblPr"), ("tc_pr", "tcPr"):
            if key not in effective_props:
                continue
            props = effective_props.get(key)
            if props is None:
                continue
            if isinstance(props, list) and len(props) == 0:
                _replace_child(style_elem, tag, None)
                continue
            if isinstance(props, list):
                _replace_child(style_elem, tag, _build_props_element(tag, props))

        if "tbl_style_pr" in effective_props:
            tbl_style_pr = effective_props.get("tbl_style_pr")
            if tbl_style_pr is not None:
                for elem in style_elem.findall("w:tblStylePr", DOCX_NS):
                    style_elem.remove(elem)
                if isinstance(tbl_style_pr, list):
                    for entry in tbl_style_pr:
                        tbl_elem = ET.Element(_qn("w", "tblStylePr"))
                        if isinstance(entry, dict):
                            tbl_type = entry.get("type")
                            if tbl_type:
                                tbl_elem.set(_qn("w", "type"), str(tbl_type))
                            props = entry.get("properties") or []
                            if isinstance(props, list):
                                _append_props(tbl_elem, props)
                        style_elem.append(tbl_elem)

        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/styles.xml": updated_xml})
        return runtime_patch
    except Exception as exc:
        _logger.error(f"[Template] Failed to apply XML updates: {exc}")
        raise


def _apply_outline_level(
    docx_path: Path,
    style_name: str,
    style_id: Optional[str],
    outline_level: Optional[int]
) -> None:
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                _logger.warning("[Template] styles.xml not found in DOCX")
                return
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)
        styles_root = ET.fromstring(styles_xml)
        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            _logger.warning(f"[Template] Style '{style_name}' not found in styles.xml")
            return

        p_pr = style_elem.find("w:pPr", DOCX_NS)
        if p_pr is None:
            p_pr = ET.Element(_qn("w", "pPr"))
            style_elem.append(p_pr)

        outline_elem = p_pr.find("w:outlineLvl", DOCX_NS)
        if outline_elem is not None:
            p_pr.remove(outline_elem)

        if outline_level is not None:
            outline_elem = ET.Element(_qn("w", "outlineLvl"))
            outline_elem.set(_qn("w", "val"), str(outline_level))
            p_pr.append(outline_elem)

        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/styles.xml": updated_xml})
    except Exception as exc:
        _logger.error(f"[Template] Failed to apply outline level: {exc}")
        raise


def _freeze_header_footer_table_styles(
    docx_path: Path,
    style_name: Optional[str],
    style_id: Optional[str],
) -> None:
    """Detach header/footer tables from a table style before it is modified.

    For each header/footer XML part, find tables that reference the target
    style.  Copy the style's current ``<w:tblPr>`` children as **direct
    formatting** on the table's own ``<w:tblPr>`` and remove the
    ``<w:tblStyle>`` reference so subsequent style changes do not
    propagate to header/footer content.
    """
    if not style_name and not style_id:
        return

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            part_names = zin.namelist()
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                return  # no styles.xml, nothing to do
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)

            # Identify which header/footer parts exist
            hf_parts: Dict[str, bytes] = {}
            for name in part_names:
                basename = name.lower()
                if basename.startswith("word/header") or basename.startswith("word/footer"):
                    if basename.endswith(".xml"):
                        hf_parts[name] = zin.read(name)

        if not hf_parts:
            return  # no header/footer parts

        # Parse styles.xml to find the target style's tblPr
        styles_root = ET.fromstring(styles_xml)
        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            return  # style not found; will be caught by caller

        # Gather the style IDs this style is known by so we can match
        # <w:tblStyle w:val="..."> references in header/footer tables.
        known_ids: Set[str] = set()
        sid = style_elem.get(_qn("w", "styleId"))
        if sid:
            known_ids.add(sid)
        name_elem = style_elem.find("w:name", DOCX_NS)
        if name_elem is not None:
            name_val = name_elem.get(_qn("w", "val"))
            if name_val:
                known_ids.add(name_val)
        if style_name:
            known_ids.add(style_name)
        if style_id:
            known_ids.add(style_id)

        # Collect the style's tblPr children (to be inlined in each table)
        style_tblpr = style_elem.find("w:tblPr", DOCX_NS)

        # Process each header/footer part
        updates_to_write: Dict[str, bytes] = {}
        for part_name, part_bytes in hf_parts.items():
            try:
                part_root = ET.fromstring(part_bytes)
            except Exception:
                continue  # skip unparseable parts

            modified = False
            for tbl in part_root.findall(".//" + _qn("w", "tbl")):
                tbl_pr = tbl.find(_qn("w", "tblPr"))
                if tbl_pr is None:
                    continue
                tbl_style_ref = tbl_pr.find(_qn("w", "tblStyle"))
                if tbl_style_ref is None:
                    continue
                ref_val = tbl_style_ref.get(_qn("w", "val")) or ""
                if ref_val not in known_ids:
                    continue

                # ---- Inline the style's tblPr children ----
                if style_tblpr is not None:
                    # Copy each property child from the style
                    existing_tags = {child.tag for child in tbl_pr}
                    for child in list(style_tblpr):
                        # Skip the tblStyle element itself
                        if _local_name(child.tag) == "tblStyle":
                            continue
                        # Only add properties that don't already have explicit overrides
                        if child.tag not in existing_tags:
                            tbl_pr.append(copy.deepcopy(child))

                # Remove the style reference so the table is detached
                tbl_pr.remove(tbl_style_ref)
                modified = True
                _logger.info(
                    f"[Template] Detached table from style '{ref_val}' in {part_name}"
                )

            if modified:
                updates_to_write[part_name] = _serialize_ooxml_part(
                    part_root,
                    part_bytes,
                    namespace_hints=namespace_hints,
                )

        if updates_to_write:
            _write_docx_parts(docx_path, updates_to_write)
            _logger.info(
                f"[Template] Froze {len(updates_to_write)} header/footer part(s) "
                f"before modifying style '{style_name or style_id}'"
            )

    except Exception as exc:
        _logger.warning(f"[Template] Failed to freeze header/footer tables: {exc}")
        # Non-fatal: proceed with style modification even if freeze fails


def _apply_table_style_updates(
    docx_path: Path,
    style_name: str,
    style_id: Optional[str],
    updates: Dict[str, Any]
) -> None:
    table_keys = {
        "table_border_style",
        "table_border_size_pt",
        "table_border_color",
        "table_shading_color",
        "table_alignment",
        "table_cell_spacing_pt",
        "table_cell_margin_top_pt",
        "table_cell_margin_bottom_pt",
        "table_cell_margin_left_pt",
        "table_cell_margin_right_pt",
        "table_cell_shading_color",
        "table_cell_vertical_align",
    }
    if not any(key in updates for key in table_keys):
        return

    # Protect header/footer tables from the upcoming style change
    _freeze_header_footer_table_styles(docx_path, style_name, style_id)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
            except KeyError:
                _logger.warning("[Template] styles.xml not found in DOCX")
                return
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)

        styles_root = ET.fromstring(styles_xml)
        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            _logger.warning(f"[Template] Style '{style_name}' not found for table updates")
            return

        def _remove_child(parent: ET.Element, tag: str) -> None:
            existing = parent.find(f"w:{tag}", DOCX_NS)
            if existing is not None:
                parent.remove(existing)

        tbl_pr = style_elem.find("w:tblPr", DOCX_NS)
        if tbl_pr is None:
            tbl_pr = ET.Element(_qn("w", "tblPr"))
            style_elem.append(tbl_pr)
        else:
            for runtime_tag in _TABLE_STYLE_RUNTIME_TAGS:
                _remove_child(tbl_pr, runtime_tag)

        border_style = updates.get("table_border_style")
        border_size_pt = updates.get("table_border_size_pt")
        border_color = updates.get("table_border_color")

        if border_style or border_size_pt or border_color:
            borders = _ensure_child(tbl_pr, "tblBorders")
            border_tags = ["top", "bottom", "left", "right", "insideH", "insideV"]
            for tag in border_tags:
                border = _ensure_child(borders, tag)
                if border_style:
                    val = "nil" if str(border_style).lower() == "none" else str(border_style)
                    border.set(_qn("w", "val"), val)
                if border_size_pt not in (None, ""):
                    try:
                        size_val = int(float(border_size_pt) * 8)
                        border.set(_qn("w", "sz"), str(size_val))
                    except (TypeError, ValueError):
                        pass
                if border_color:
                    border.set(_qn("w", "color"), str(border_color).replace("#", ""))

        shading_color = updates.get("table_shading_color")
        if shading_color is not None:
            if shading_color == "":
                existing = tbl_pr.find("w:shd", DOCX_NS)
                if existing is not None:
                    tbl_pr.remove(existing)
            else:
                shd = _ensure_child(tbl_pr, "shd")
                shd.set(_qn("w", "val"), "clear")
                shd.set(_qn("w", "fill"), str(shading_color).replace("#", ""))

        table_alignment = updates.get("table_alignment")
        if table_alignment is not None:
            if table_alignment == "":
                existing = tbl_pr.find("w:jc", DOCX_NS)
                if existing is not None:
                    tbl_pr.remove(existing)
            else:
                jc = _ensure_child(tbl_pr, "jc")
                jc.set(_qn("w", "val"), str(table_alignment).lower())

        if "table_cell_spacing_pt" in updates:
            spacing_value = updates.get("table_cell_spacing_pt")
            if spacing_value in (None, ""):
                _remove_child(tbl_pr, "tblCellSpacing")
            else:
                try:
                    spacing_twips = int(round(float(spacing_value) * 20))
                    spacing = _ensure_child(tbl_pr, "tblCellSpacing")
                    spacing.set(_qn("w", "w"), str(max(0, spacing_twips)))
                    spacing.set(_qn("w", "type"), "dxa")
                except (TypeError, ValueError):
                    pass

        margin_map = {
            "top": updates.get("table_cell_margin_top_pt"),
            "bottom": updates.get("table_cell_margin_bottom_pt"),
            "left": updates.get("table_cell_margin_left_pt"),
            "right": updates.get("table_cell_margin_right_pt"),
        }
        if any(value not in (None, "") for value in margin_map.values()):
            cell_mar = _ensure_child(tbl_pr, "tblCellMar")
            for tag, value in margin_map.items():
                if value in (None, ""):
                    continue
                mar = _ensure_child(cell_mar, tag)
                try:
                    twips = int(float(value) * 20)
                    mar.set(_qn("w", "w"), str(twips))
                    mar.set(_qn("w", "type"), "dxa")
                except (TypeError, ValueError):
                    continue

        look_key_map = {
            "firstRow": "table_look_first_row",
            "lastRow": "table_look_last_row",
            "firstColumn": "table_look_first_column",
            "lastColumn": "table_look_last_column",
            "noHBand": "table_look_no_h_band",
            "noVBand": "table_look_no_v_band",
        }
        look_updates: Dict[str, Optional[bool]] = {}
        for xml_key, update_key in look_key_map.items():
            if update_key not in updates:
                continue
            look_updates[xml_key] = _coerce_optional_bool(updates.get(update_key))

        if look_updates:
            _logger.debug(
                "[Template] Captured table look runtime defaults for style '%s' without writing invalid OOXML",
                style_name,
            )

        cell_shading_color = updates.get("table_cell_shading_color")
        cell_vertical_align = updates.get("table_cell_vertical_align")
        if "table_cell_shading_color" in updates or "table_cell_vertical_align" in updates:
            tc_pr = style_elem.find("w:tcPr", DOCX_NS)
            if tc_pr is None and (cell_shading_color not in (None, "") or cell_vertical_align not in (None, "")):
                tc_pr = ET.SubElement(style_elem, _qn("w", "tcPr"))

            if tc_pr is not None:
                if "table_cell_shading_color" in updates:
                    if cell_shading_color in (None, ""):
                        _remove_child(tc_pr, "shd")
                    else:
                        normalized_fill = _normalize_ooxml_color(cell_shading_color)
                        if normalized_fill:
                            shd = _ensure_child(tc_pr, "shd")
                            shd.set(_qn("w", "val"), "clear")
                            shd.set(_qn("w", "color"), "auto")
                            shd.set(_qn("w", "fill"), normalized_fill)
                        else:
                            _remove_child(tc_pr, "shd")

                if "table_cell_vertical_align" in updates:
                    if cell_vertical_align in (None, ""):
                        _remove_child(tc_pr, "vAlign")
                    else:
                        v_align = _ensure_child(tc_pr, "vAlign")
                        v_align.set(_qn("w", "val"), str(cell_vertical_align))

                if len(list(tc_pr)) == 0:
                    style_elem.remove(tc_pr)

        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/styles.xml": updated_xml})
    except Exception as exc:
        _logger.error(f"[Template] Failed table style updates: {exc}")
        raise


def _apply_list_style_updates(
    docx_path: Path,
    style_name: str,
    style_id: Optional[str],
    updates: Dict[str, Any]
) -> None:
    list_keys = {
        "list_format",
        "list_bullet_char",
        "list_start",
        "list_level",
        "list_alignment",
        "list_left_indent_inches",
        "list_hanging_indent_inches",
    }
    if not any(key in updates for key in list_keys):
        return

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            try:
                styles_xml = zin.read("word/styles.xml")
                numbering_xml = zin.read("word/numbering.xml")
            except KeyError:
                _logger.warning("[Template] numbering.xml/styles.xml missing for list updates")
                return
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)

        styles_root = ET.fromstring(styles_xml)
        numbering_root = ET.fromstring(numbering_xml)

        style_elem = _find_style_element(styles_root, style_name, style_id)
        if style_elem is None:
            _logger.warning(f"[Template] Style '{style_name}' not found for list updates")
            return

        p_pr = style_elem.find("w:pPr", DOCX_NS)
        if p_pr is None:
            p_pr = ET.Element(_qn("w", "pPr"))
            style_elem.append(p_pr)

        num_pr = p_pr.find("w:numPr", DOCX_NS)
        if num_pr is None:
            num_pr = ET.Element(_qn("w", "numPr"))
            p_pr.append(num_pr)

        num_id_elem = num_pr.find("w:numId", DOCX_NS)
        ilvl_elem = num_pr.find("w:ilvl", DOCX_NS)

        list_level = updates.get("list_level")
        if list_level not in (None, ""):
            try:
                level_value = int(list_level)
            except (TypeError, ValueError):
                level_value = 0
        else:
            level_value = int(ilvl_elem.get(_qn("w", "val"))) if ilvl_elem is not None else 0

        if ilvl_elem is None:
            ilvl_elem = ET.Element(_qn("w", "ilvl"))
            num_pr.append(ilvl_elem)
        ilvl_elem.set(_qn("w", "val"), str(level_value))

        if num_id_elem is None:
            num_id_elem = ET.Element(_qn("w", "numId"))
            num_pr.append(num_id_elem)

        num_id_val = num_id_elem.get(_qn("w", "val"))
        if not num_id_val:
            existing_nums = [
                int(elem.get(_qn("w", "numId")))
                for elem in numbering_root.findall("w:num", DOCX_NS)
                if elem.get(_qn("w", "numId"))
            ]
            new_num_id = max(existing_nums) + 1 if existing_nums else 1
            num_id_val = str(new_num_id)
            num_id_elem.set(_qn("w", "val"), num_id_val)

            existing_abs = [
                int(elem.get(_qn("w", "abstractNumId")))
                for elem in numbering_root.findall("w:abstractNum", DOCX_NS)
                if elem.get(_qn("w", "abstractNumId"))
            ]
            new_abs_id = max(existing_abs) + 1 if existing_abs else 1

            abstract = ET.Element(_qn("w", "abstractNum"))
            abstract.set(_qn("w", "abstractNumId"), str(new_abs_id))
            lvl = ET.SubElement(abstract, _qn("w", "lvl"))
            lvl.set(_qn("w", "ilvl"), str(level_value))
            numbering_root.append(abstract)

            num = ET.Element(_qn("w", "num"))
            num.set(_qn("w", "numId"), num_id_val)
            abstract_ref = ET.SubElement(num, _qn("w", "abstractNumId"))
            abstract_ref.set(_qn("w", "val"), str(new_abs_id))
            numbering_root.append(num)

        num_elem = numbering_root.find(f"w:num[@w:numId='{num_id_val}']", DOCX_NS)
        if num_elem is None:
            _logger.warning("[Template] list numId missing in numbering.xml")
            return
        abstract_id = _find_val(num_elem, "w:abstractNumId")
        if abstract_id is None:
            _logger.warning("[Template] abstractNumId missing for list num")
            return
        abstract_elem = numbering_root.find(f"w:abstractNum[@w:abstractNumId='{abstract_id}']", DOCX_NS)
        if abstract_elem is None:
            _logger.warning("[Template] abstractNum not found")
            return

        lvl_elem = abstract_elem.find(f"w:lvl[@w:ilvl='{level_value}']", DOCX_NS)
        if lvl_elem is None:
            lvl_elem = ET.Element(_qn("w", "lvl"))
            lvl_elem.set(_qn("w", "ilvl"), str(level_value))
            abstract_elem.append(lvl_elem)

        list_format = updates.get("list_format")
        list_bullet_char = updates.get("list_bullet_char")
        list_start = updates.get("list_start")
        list_alignment = updates.get("list_alignment")
        list_left_indent = updates.get("list_left_indent_inches")
        list_hanging_indent = updates.get("list_hanging_indent_inches")

        if list_start not in (None, ""):
            start_elem = _ensure_child(lvl_elem, "start")
            start_elem.set(_qn("w", "val"), str(list_start))

        current_fmt_elem = lvl_elem.find("w:numFmt", DOCX_NS)
        current_fmt = current_fmt_elem.get(_qn("w", "val")) if current_fmt_elem is not None else None
        if list_format:
            num_fmt = _ensure_child(lvl_elem, "numFmt")
            num_fmt.set(_qn("w", "val"), str(list_format))
            lvl_text = _ensure_child(lvl_elem, "lvlText")
            if str(list_format).lower() == "bullet":
                bullet_char = list_bullet_char or "•"
                lvl_text.set(_qn("w", "val"), bullet_char)
            else:
                lvl_text_val = f"%{level_value + 1}."
                lvl_text.set(_qn("w", "val"), lvl_text_val)
        elif list_bullet_char and (current_fmt or "").lower() == "bullet":
            lvl_text = _ensure_child(lvl_elem, "lvlText")
            lvl_text.set(_qn("w", "val"), str(list_bullet_char))

        if list_alignment:
            lvl_jc = _ensure_child(lvl_elem, "lvlJc")
            lvl_jc.set(_qn("w", "val"), str(list_alignment).lower())

        if list_left_indent not in (None, "") or list_hanging_indent not in (None, ""):
            ppr = _ensure_child(lvl_elem, "pPr")
            ind = _ensure_child(ppr, "ind")
            if list_left_indent not in (None, ""):
                try:
                    ind.set(_qn("w", "left"), str(int(float(list_left_indent) * 1440)))
                except (TypeError, ValueError):
                    pass
            if list_hanging_indent not in (None, ""):
                try:
                    ind.set(_qn("w", "hanging"), str(int(float(list_hanging_indent) * 1440)))
                except (TypeError, ValueError):
                    pass

        updated_styles = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        updated_numbering = _serialize_ooxml_part(
            numbering_root,
            numbering_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {
            "word/styles.xml": updated_styles,
            "word/numbering.xml": updated_numbering,
        })
    except Exception as exc:
        _logger.error(f"[Template] Failed list style updates: {exc}")
        raise


def _read_docx_parts(docx_bytes: bytes, part_paths: List[str]) -> Dict[str, str]:
    parts: Dict[str, str] = {}
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
            for path in part_paths:
                try:
                    with docx_zip.open(path) as part:
                        parts[path] = part.read().decode("utf-8", errors="ignore")
                except KeyError:
                    continue
    except Exception as exc:
        _logger.warning(f"[Template] Failed to read DOCX parts: {exc}")
    return parts


def _twips_to_inches(value: Optional[str]) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(float(value) / 1440.0, 3)
    except (TypeError, ValueError):
        return None


def _find_val(element: ET.Element, path: str) -> Optional[str]:
    target = element.find(path, DOCX_NS)
    if target is None:
        return None
    return target.get(_qn("w", "val"))


def _parse_style_element(style_elem: ET.Element) -> Dict[str, Any]:
    style_id = style_elem.get(_qn("w", "styleId"))
    style_type = style_elem.get(_qn("w", "type"))
    display_name = _find_val(style_elem, "w:name")
    based_on = _find_val(style_elem, "w:basedOn")
    next_style = _find_val(style_elem, "w:next")
    link_style = _find_val(style_elem, "w:link")
    ui_priority = _find_val(style_elem, "w:uiPriority")
    outline_val = None
    outline_elem = style_elem.find("w:pPr/w:outlineLvl", DOCX_NS)
    if outline_elem is not None:
        outline_val = outline_elem.get(_qn("w", "val"))
    outline_level = None
    if outline_val is not None:
        try:
            outline_level = int(outline_val)
        except (TypeError, ValueError):
            outline_level = None

    r_pr = style_elem.find("w:rPr", DOCX_NS)
    p_pr = style_elem.find("w:pPr", DOCX_NS)
    tbl_pr = style_elem.find("w:tblPr", DOCX_NS)
    tc_pr = style_elem.find("w:tcPr", DOCX_NS)

    tbl_style_props: List[Dict[str, Any]] = []
    for tbl_style in style_elem.findall("w:tblStylePr", DOCX_NS):
        tbl_style_props.append({
            "type": tbl_style.get(_qn("w", "type")),
            "properties": _collect_props(tbl_style),
            "raw_xml": _serialize_xml(tbl_style),
        })

    return {
        "style_id": style_id,
        "display_name": display_name,
        "type": style_type,
        "based_on": based_on,
        "next": next_style,
        "link": link_style,
        "ui_priority": int(ui_priority) if ui_priority and ui_priority.isdigit() else ui_priority,
        "default": _as_bool(style_elem.get(_qn("w", "default"))),
        "custom": _as_bool(style_elem.get(_qn("w", "customStyle"))),
        "hidden": style_elem.find("w:hidden", DOCX_NS) is not None,
        "semi_hidden": style_elem.find("w:semiHidden", DOCX_NS) is not None,
        "q_format": style_elem.find("w:qFormat", DOCX_NS) is not None,
        "unhide_when_used": style_elem.find("w:unhideWhenUsed", DOCX_NS) is not None,
        "outline_level": outline_level,
        "r_pr": _collect_props(r_pr),
        "p_pr": _collect_props(p_pr),
        "tbl_pr": _collect_props(tbl_pr),
        "tc_pr": _collect_props(tc_pr),
        "tbl_style_pr": tbl_style_props or None,
        "raw_xml": _serialize_xml(style_elem),
    }


def _parse_doc_defaults(doc_defaults_elem: ET.Element) -> Dict[str, Any]:
    r_pr_default = doc_defaults_elem.find("w:rPrDefault/w:rPr", DOCX_NS)
    p_pr_default = doc_defaults_elem.find("w:pPrDefault/w:pPr", DOCX_NS)
    return {
        "r_pr": _collect_props(r_pr_default),
        "p_pr": _collect_props(p_pr_default),
        "r_pr_xml": _serialize_xml(r_pr_default),
        "p_pr_xml": _serialize_xml(p_pr_default),
        "raw_xml": _serialize_xml(doc_defaults_elem),
    }


def _parse_latent_styles(latent_elem: ET.Element) -> Dict[str, Any]:
    attrs = {_local_name(k): v for k, v in latent_elem.attrib.items()}
    exceptions: List[Dict[str, Any]] = []
    for exc in latent_elem.findall("w:lsdException", DOCX_NS):
        exceptions.append({
            "name": exc.get(_qn("w", "name")),
            "locked": _as_bool(exc.get(_qn("w", "locked"))),
            "ui_priority": exc.get(_qn("w", "uiPriority")),
            "q_format": _as_bool(exc.get(_qn("w", "qFormat"))),
            "semi_hidden": _as_bool(exc.get(_qn("w", "semiHidden"))),
            "unhide_when_used": _as_bool(exc.get(_qn("w", "unhideWhenUsed"))),
            "hidden": _as_bool(exc.get(_qn("w", "hidden"))),
        })
    return {
        "attrs": attrs,
        "exceptions": exceptions,
        "raw_xml": _serialize_xml(latent_elem),
    }


def _parse_numbering(numbering_root: ET.Element) -> Dict[str, Any]:
    abstract_nums: List[Dict[str, Any]] = []
    for abstract in numbering_root.findall("w:abstractNum", DOCX_NS):
        abstract_id = abstract.get(_qn("w", "abstractNumId"))
        levels: List[Dict[str, Any]] = []
        for lvl in abstract.findall("w:lvl", DOCX_NS):
            ilvl = lvl.get(_qn("w", "ilvl"))
            levels.append({
                "ilvl": int(ilvl) if ilvl and ilvl.isdigit() else ilvl,
                "start": _find_val(lvl, "w:start"),
                "num_fmt": _find_val(lvl, "w:numFmt"),
                "lvl_text": _find_val(lvl, "w:lvlText"),
                "lvl_jc": _find_val(lvl, "w:lvlJc"),
                "p_style": _find_val(lvl, "w:pStyle"),
                "p_pr": _collect_props(lvl.find("w:pPr", DOCX_NS)),
                "r_pr": _collect_props(lvl.find("w:rPr", DOCX_NS)),
                "raw_xml": _serialize_xml(lvl),
            })
        abstract_nums.append({
            "abstract_num_id": abstract_id,
            "nsid": _find_val(abstract, "w:nsid"),
            "multilevel_type": _find_val(abstract, "w:multiLevelType"),
            "tmpl": _find_val(abstract, "w:tmpl"),
            "levels": levels,
            "raw_xml": _serialize_xml(abstract),
        })

    nums: List[Dict[str, Any]] = []
    for num in numbering_root.findall("w:num", DOCX_NS):
        num_id = num.get(_qn("w", "numId"))
        abstract_ref = _find_val(num, "w:abstractNumId")
        lvl_overrides: List[Dict[str, Any]] = []
        for override in num.findall("w:lvlOverride", DOCX_NS):
            ilvl = override.get(_qn("w", "ilvl"))
            lvl_overrides.append({
                "ilvl": int(ilvl) if ilvl and ilvl.isdigit() else ilvl,
                "start_override": _find_val(override, "w:startOverride"),
                "lvl": _collect_props(override.find("w:lvl", DOCX_NS)),
                "raw_xml": _serialize_xml(override),
            })
        nums.append({
            "num_id": int(num_id) if num_id and num_id.isdigit() else num_id,
            "abstract_num_id": int(abstract_ref) if abstract_ref and abstract_ref.isdigit() else abstract_ref,
            "lvl_overrides": lvl_overrides or None,
            "raw_xml": _serialize_xml(num),
        })

    return {
        "abstract_nums": abstract_nums,
        "nums": nums,
        "raw_xml": _serialize_xml(numbering_root),
    }


def _parse_theme(theme_root: ET.Element) -> Dict[str, Any]:
    theme_name = theme_root.get("name")
    font_scheme_elem = theme_root.find(".//a:fontScheme", DOCX_NS)
    color_scheme_elem = theme_root.find(".//a:clrScheme", DOCX_NS)

    def _parse_font_group(group_elem: Optional[ET.Element]) -> Dict[str, Any]:
        if group_elem is None:
            return {}
        fonts: Dict[str, Any] = {}
        for node in list(group_elem):
            tag = _local_name(node.tag)
            if tag in {"latin", "ea", "cs"}:
                fonts[tag] = node.get("typeface")
            elif tag == "font":
                script = node.get("script")
                fonts.setdefault("script", {})[script] = node.get("typeface")
        return fonts

    def _parse_color_node(node: ET.Element) -> Dict[str, Any]:
        if len(node) == 0:
            return {"attrs": node.attrib}
        child = list(node)[0]
        return {
            "type": _local_name(child.tag),
            "attrs": child.attrib,
        }

    font_scheme = {}
    if font_scheme_elem is not None:
        font_scheme = {
            "major": _parse_font_group(font_scheme_elem.find("a:majorFont", DOCX_NS)),
            "minor": _parse_font_group(font_scheme_elem.find("a:minorFont", DOCX_NS)),
            "raw_xml": _serialize_xml(font_scheme_elem),
        }

    color_scheme = {}
    if color_scheme_elem is not None:
        colors: Dict[str, Any] = {}
        for node in list(color_scheme_elem):
            colors[_local_name(node.tag)] = _parse_color_node(node)
        color_scheme = {
            "name": color_scheme_elem.get("name"),
            "colors": colors,
            "raw_xml": _serialize_xml(color_scheme_elem),
        }

    return {
        "name": theme_name,
        "font_scheme": font_scheme or None,
        "color_scheme": color_scheme or None,
        "raw_xml": _serialize_xml(theme_root),
    }


def _normalize_header_footer_paragraph_text(paragraph_text: str) -> str:
    return " ".join(str(paragraph_text or "").replace("\xa0", " ").split()).strip()


def _extract_text_from_header_footer_part_xml(xml_bytes: bytes) -> str:
    try:
        root = ET.fromstring(xml_bytes)
    except Exception:
        return ""

    paragraphs: List[str] = []
    for paragraph in root.findall(".//w:p", DOCX_NS):
        text_parts = [node.text or "" for node in paragraph.findall(".//w:t", DOCX_NS)]
        paragraph_text = _normalize_header_footer_paragraph_text("".join(text_parts))
        if paragraph_text:
            paragraphs.append(paragraph_text)

    unique_paragraphs = list(dict.fromkeys(paragraphs))
    if unique_paragraphs:
        return "\n".join(unique_paragraphs)

    has_table = root.find(".//w:tbl", DOCX_NS) is not None
    has_drawing = (
        root.find(".//w:drawing", DOCX_NS) is not None
        or root.find(".//w:pict", DOCX_NS) is not None
    )
    if has_table and has_drawing:
        return "[contenido en tabla con imagen]"
    if has_table:
        return "[contenido en tabla]"
    if has_drawing:
        return "[contenido gráfico]"
    return ""


def _extract_header_footer_texts_from_docx_bytes(docx_bytes: bytes) -> tuple[List[str], List[str]]:
    headers: List[str] = []
    footers: List[str] = []

    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
            for part_name in sorted(docx_zip.namelist()):
                if not part_name.endswith(".xml"):
                    continue
                if part_name.startswith("word/header"):
                    text = _extract_text_from_header_footer_part_xml(docx_zip.read(part_name))
                    if text:
                        headers.append(text)
                elif part_name.startswith("word/footer"):
                    text = _extract_text_from_header_footer_part_xml(docx_zip.read(part_name))
                    if text:
                        footers.append(text)
    except Exception as exc:
        _logger.warning(f"[Template] Failed extracting header/footer OOXML text: {exc}")

    return list(dict.fromkeys(headers)), list(dict.fromkeys(footers))


def _parse_font_table(fonts_root: ET.Element) -> Dict[str, Any]:
    fonts: List[Dict[str, Any]] = []
    for font_elem in fonts_root.findall("w:font", DOCX_NS):
        name = font_elem.get(_qn("w", "name"))
        if not name:
            continue
        entry: Dict[str, Any] = {"name": name}
        alt_name = font_elem.find("w:altName", DOCX_NS)
        if alt_name is not None and alt_name.get(_qn("w", "val")):
            entry["alt_name"] = alt_name.get(_qn("w", "val"))
        family = font_elem.find("w:family", DOCX_NS)
        if family is not None and family.get(_qn("w", "val")):
            entry["family"] = family.get(_qn("w", "val"))
        pitch = font_elem.find("w:pitch", DOCX_NS)
        if pitch is not None and pitch.get(_qn("w", "val")):
            entry["pitch"] = pitch.get(_qn("w", "val"))
        fonts.append(entry)
    return {
        "fonts": fonts,
        "raw_xml": _serialize_xml(fonts_root),
    }


def _parse_settings(settings_root: ET.Element) -> Dict[str, Any]:
    default_tab = settings_root.find("w:defaultTabStop", DOCX_NS)
    return {
        "track_revisions": settings_root.find("w:trackRevisions", DOCX_NS) is not None,
        "default_tab_stop_twips": default_tab.get(_qn("w", "val")) if default_tab is not None else None,
        "raw_xml": _serialize_xml(settings_root),
    }


def _parse_sections(document_root: ET.Element) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    for idx, sect in enumerate(document_root.findall(".//w:sectPr", DOCX_NS)):
        pg_sz = sect.find("w:pgSz", DOCX_NS)
        pg_mar = sect.find("w:pgMar", DOCX_NS)
        sections.append({
            "index": idx,
            "page_size": {
                "width_twips": pg_sz.get(_qn("w", "w")) if pg_sz is not None else None,
                "height_twips": pg_sz.get(_qn("w", "h")) if pg_sz is not None else None,
                "width_inches": _twips_to_inches(pg_sz.get(_qn("w", "w")) if pg_sz is not None else None),
                "height_inches": _twips_to_inches(pg_sz.get(_qn("w", "h")) if pg_sz is not None else None),
                "orient": pg_sz.get(_qn("w", "orient")) if pg_sz is not None else None,
            },
            "margins": {
                "top_twips": pg_mar.get(_qn("w", "top")) if pg_mar is not None else None,
                "bottom_twips": pg_mar.get(_qn("w", "bottom")) if pg_mar is not None else None,
                "left_twips": pg_mar.get(_qn("w", "left")) if pg_mar is not None else None,
                "right_twips": pg_mar.get(_qn("w", "right")) if pg_mar is not None else None,
                "top_inches": _twips_to_inches(pg_mar.get(_qn("w", "top")) if pg_mar is not None else None),
                "bottom_inches": _twips_to_inches(pg_mar.get(_qn("w", "bottom")) if pg_mar is not None else None),
                "left_inches": _twips_to_inches(pg_mar.get(_qn("w", "left")) if pg_mar is not None else None),
                "right_inches": _twips_to_inches(pg_mar.get(_qn("w", "right")) if pg_mar is not None else None),
            },
            "properties": _collect_props(sect),
            "raw_xml": _serialize_xml(sect),
        })
    return sections


def _extract_table_info_from_element(tbl_idx: int, tbl: ET.Element) -> Dict[str, Any]:
    table_info: Dict[str, Any] = {
        "index": tbl_idx,
        "rows": 0,
        "cols": 0,
        "style_name": None,
        "style_id": None,
        "table_properties": {},
        "first_row_format": {},
        "sample_cells": [],
        "has_distinct_header": False,
        "raw_xml": _serialize_xml(tbl),
    }

    tbl_pr = tbl.find("w:tblPr", DOCX_NS)
    if tbl_pr is not None:
        table_info["table_properties"] = _collect_props(tbl_pr)
        tbl_style = tbl_pr.find("w:tblStyle", DOCX_NS)
        if tbl_style is not None:
            table_info["style_id"] = tbl_style.get(_qn("w", "val"))
            table_info["style_name"] = table_info["style_id"]

        parsed_tbl_pr = _parse_tblpr_nodes(_collect_props(tbl_pr))
        table_info["parsed_properties"] = parsed_tbl_pr
        table_info["borders"] = parsed_tbl_pr.get("borders")
        table_info["shading_fill"] = parsed_tbl_pr.get("shading_color")
        table_info["margins"] = parsed_tbl_pr.get("cell_margins")
        table_info["alignment"] = parsed_tbl_pr.get("alignment")
        table_info["look"] = parsed_tbl_pr.get("look")

    rows = tbl.findall("w:tr", DOCX_NS)
    table_info["rows"] = len(rows)

    first_row_cells = []
    other_row_cells = []
    for row_idx, row in enumerate(rows):
        cells = row.findall("w:tc", DOCX_NS)
        if row_idx == 0:
            table_info["cols"] = len(cells)

        for cell_idx, cell in enumerate(cells):
            cell_info: Dict[str, Any] = {
                "row": row_idx,
                "col": cell_idx,
                "properties": {},
                "text": "",
            }
            tc_pr = cell.find("w:tcPr", DOCX_NS)
            if tc_pr is not None:
                cell_info["properties"] = _collect_props(tc_pr)
                cell_info["parsed_properties"] = _parse_tcpr_nodes(_collect_props(tc_pr))

            paragraphs = cell.findall("w:p", DOCX_NS)
            text_parts = []
            for p in paragraphs[:1]:
                for r in p.findall("w:r", DOCX_NS):
                    r_pr = r.find("w:rPr", DOCX_NS)
                    if r_pr is not None:
                        font_props = _parse_rpr_nodes(_collect_props(r_pr))
                        if font_props:
                            cell_info["font_properties"] = font_props
                            break
                for t in p.findall(".//w:t", DOCX_NS):
                    if t.text:
                        text_parts.append(t.text)

            cell_info["text"] = "".join(text_parts)[:50]
            if row_idx == 0:
                first_row_cells.append(cell_info)
            elif row_idx == 1 and len(other_row_cells) < 4:
                other_row_cells.append(cell_info)

    table_info["sample_cells"] = first_row_cells + other_row_cells

    if first_row_cells:
        first_sample_cell = first_row_cells[0].get("parsed_properties", {}) or {}
        first_font = first_row_cells[0].get("font_properties", {}) or {}
        other_sample_cell = other_row_cells[0].get("parsed_properties", {}) if other_row_cells else {}
        other_font = other_row_cells[0].get("font_properties", {}) if other_row_cells else {}

        first_bg = next(
            (
                (c.get("parsed_properties", {}) or {}).get("shading_color")
                for c in first_row_cells
                if (c.get("parsed_properties", {}) or {}).get("shading_color")
            ),
            None,
        )
        other_bg = next(
            (
                (c.get("parsed_properties", {}) or {}).get("shading_color")
                for c in other_row_cells
                if (c.get("parsed_properties", {}) or {}).get("shading_color")
            ),
            None,
        )

        first_borders = first_sample_cell.get("borders") or {}
        other_borders = (other_sample_cell or {}).get("borders") or {}
        meaningful_first_font = {k: v for k, v in first_font.items() if v not in (None, "", False)}
        meaningful_other_font = {k: v for k, v in (other_font or {}).items() if v not in (None, "", False)}

        shading_diff = (first_bg or None) != (other_bg or None)
        border_diff = bool(first_borders) and first_borders != other_borders
        font_diff = bool(meaningful_first_font) and meaningful_first_font != meaningful_other_font

        table_info["has_distinct_header"] = bool(shading_diff or border_diff or font_diff)
        table_info["first_row_format"] = {
            "sample_cell": first_sample_cell,
            "font_properties": first_font,
            "shading_fill": first_bg,
            "borders": first_sample_cell.get("borders"),
            "margins": first_sample_cell.get("margins"),
            "vertical_align": first_sample_cell.get("vertical_align"),
        }

    return table_info


def _extract_style_name_map(styles_xml: Optional[bytes]) -> Dict[str, str]:
    if not styles_xml:
        return {}
    try:
        styles_root = ET.fromstring(styles_xml)
    except Exception:
        return {}

    style_map: Dict[str, str] = {}
    for style_elem in styles_root.findall("w:style", DOCX_NS):
        style_id = style_elem.get(_qn("w", "styleId"))
        name_elem = style_elem.find("w:name", DOCX_NS)
        style_name = name_elem.get(_qn("w", "val")) if name_elem is not None else None
        if style_id and style_name:
            style_map[str(style_id)] = str(style_name)
    return style_map


def _extract_paragraph_text(paragraph: ET.Element) -> tuple[str, str]:
    text = "".join((node.text or "") for node in paragraph.findall(".//w:t", DOCX_NS))
    plain_text = re.sub(r"\s+", " ", text).strip()
    return text, plain_text


def _paragraph_contains_drawing(paragraph: ET.Element) -> bool:
    return (
        paragraph.find(".//w:drawing", DOCX_NS) is not None
        or paragraph.find(".//w:pict", DOCX_NS) is not None
    )


def _paragraph_field_instructions(paragraph: ET.Element) -> List[str]:
    instructions: List[str] = []
    for field in paragraph.findall(".//w:fldSimple", DOCX_NS):
        instruction = field.get(_qn("w", "instr"))
        if instruction:
            instructions.append(str(instruction))

    complex_chunks = [
        (node.text or "").strip()
        for node in paragraph.findall(".//w:instrText", DOCX_NS)
        if (node.text or "").strip()
    ]
    if complex_chunks:
        instructions.append(" ".join(complex_chunks))
    return instructions


def _extract_sequence_name(instruction: str) -> Optional[str]:
    if not instruction:
        return None
    match = re.search(r'\bSEQ\s+(?:"([^"]+)"|([^\s\\]+))', str(instruction), flags=re.IGNORECASE)
    if not match:
        return None
    return (match.group(1) or match.group(2) or "").strip() or None


def _caption_object_type_hint(sequence_name: Optional[str], plain_text: str) -> Optional[str]:
    if sequence_name:
        normalized_sequence = str(sequence_name).strip().lower()
        if normalized_sequence.startswith("figura"):
            return "figure"
        if normalized_sequence.startswith("tabla"):
            return "table"

    prefix_match = re.match(r"^\s*(figura|tabla)\b", plain_text, flags=re.IGNORECASE)
    if not prefix_match:
        return None
    return "figure" if prefix_match.group(1).lower().startswith("figura") else "table"


def _extract_caption_candidate(
    paragraph: ET.Element,
    style_name_map: Dict[str, str],
) -> Optional[Dict[str, Any]]:
    text, plain_text = _extract_paragraph_text(paragraph)
    instructions = _paragraph_field_instructions(paragraph)
    sequence_name = next(
        (
            extracted
            for extracted in (_extract_sequence_name(instruction) for instruction in instructions)
            if extracted
        ),
        None,
    )

    p_style = paragraph.find("w:pPr/w:pStyle", DOCX_NS)
    style_id = p_style.get(_qn("w", "val")) if p_style is not None else None
    style_name = style_name_map.get(style_id, style_id) if style_id else None
    uses_caption_style = style_id == "Caption" or style_name == "Caption"
    has_seq_field = any("SEQ" in str(instruction).upper() for instruction in instructions)
    object_type_hint = _caption_object_type_hint(sequence_name, plain_text)

    if not (uses_caption_style or has_seq_field or object_type_hint):
        return None

    return {
        "text": text,
        "plain_text": plain_text,
        "style_name": style_name,
        "style_id": style_id,
        "uses_caption_style": uses_caption_style,
        "has_seq_field": has_seq_field,
        "sequence_name": sequence_name,
        "object_type_hint": object_type_hint,
    }


def _is_empty_body_paragraph(node: ET.Element) -> bool:
    if _local_name(node.tag) != "p":
        return False
    _, plain_text = _extract_paragraph_text(node)
    return not plain_text and not _paragraph_contains_drawing(node)


def _find_adjacent_object(
    entries: List[Dict[str, Any]],
    start_idx: int,
    *,
    step: int,
) -> Optional[Dict[str, Any]]:
    idx = start_idx + step
    while 0 <= idx < len(entries):
        entry = entries[idx]
        if entry.get("object") is not None:
            return entry["object"]
        if _is_empty_body_paragraph(entry["element"]):
            idx += step
            continue
        return None
    return None


def _pick_caption_anchor(
    prev_object: Optional[Dict[str, Any]],
    next_object: Optional[Dict[str, Any]],
    object_type_hint: Optional[str],
) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    if object_type_hint == "table":
        if next_object and next_object.get("type") == "table":
            return next_object, "before"
        if prev_object and prev_object.get("type") == "table":
            return prev_object, "after"
    if object_type_hint == "figure":
        if prev_object and prev_object.get("type") == "figure":
            return prev_object, "after"
        if next_object and next_object.get("type") == "figure":
            return next_object, "before"

    if prev_object is not None:
        return prev_object, "after"
    if next_object is not None:
        return next_object, "before"
    return None, None


def _extract_document_tables_and_captions(docx_bytes: bytes) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    parts = _read_docx_parts(docx_bytes, ["word/document.xml", "word/styles.xml"])
    document_xml = parts.get("word/document.xml")
    if not document_xml:
        return [], []

    try:
        doc_root = ET.fromstring(document_xml)
    except Exception as exc:
        _logger.warning(f"[Template] Failed parsing document.xml for structure: {exc}")
        return [], []

    body = doc_root.find("w:body", DOCX_NS)
    if body is None:
        return [], []

    style_name_map = _extract_style_name_map(parts.get("word/styles.xml"))
    tables: List[Dict[str, Any]] = []
    captions: List[Dict[str, Any]] = []
    entries: List[Dict[str, Any]] = []
    figure_index = 0

    for child in list(body):
        tag = _local_name(child.tag)
        if tag == "sectPr":
            continue
        entry: Dict[str, Any] = {"element": child, "tag": tag, "object": None}
        if tag == "tbl":
            table_info = _extract_table_info_from_element(len(tables), child)
            tables.append(table_info)
            entry["object"] = {"type": "table", "index": table_info["index"]}
        elif tag == "p" and _paragraph_contains_drawing(child):
            entry["object"] = {"type": "figure", "index": figure_index}
            figure_index += 1
        entries.append(entry)

    for idx, entry in enumerate(entries):
        if entry.get("tag") != "p":
            continue
        candidate = _extract_caption_candidate(entry["element"], style_name_map)
        if candidate is None:
            continue

        prev_object = _find_adjacent_object(entries, idx, step=-1)
        next_object = _find_adjacent_object(entries, idx, step=1)
        anchor, position = _pick_caption_anchor(
            prev_object,
            next_object,
            candidate.get("object_type_hint"),
        )
        object_type = (anchor or {}).get("type") or candidate.get("object_type_hint")
        object_index = (anchor or {}).get("index")

        candidate.pop("object_type_hint", None)
        captions.append({
            "index": len(captions),
            "object_type": object_type,
            "object_index": object_index,
            "position": position,
            **candidate,
        })

    _logger.info(
        "[Template] Extracted %d tables and %d captions from document",
        len(tables),
        len(captions),
    )
    return tables, captions


def _extract_document_tables(docx_bytes: bytes) -> List[Dict[str, Any]]:
    tables, _ = _extract_document_tables_and_captions(docx_bytes)
    return tables


def _extract_body_font_hint(
    document_root: ET.Element,
    theme_details: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    body = document_root.find("w:body", DOCX_NS)
    if body is None:
        return None

    font_counts: Dict[str, int] = {}

    def _register_rpr_font(r_pr: Optional[ET.Element]) -> None:
        if r_pr is None:
            return
        props = _collect_props(r_pr)
        font_source = _extract_rpr_font_source(
            props,
            theme_details,
            source_style="document.xml",
            scope="document",
        )
        if not isinstance(font_source, dict) or font_source.get("kind") != "explicit":
            return
        font_name = str(font_source.get("font_name") or "").strip()
        if not font_name:
            return
        font_counts[font_name] = font_counts.get(font_name, 0) + 1

    stack: List[tuple[ET.Element, bool]] = [
        (child, False) for child in reversed(list(body))
    ]
    while stack:
        node, inside_table = stack.pop()
        now_inside_table = inside_table or node.tag in {_qn("w", "tbl"), _qn("w", "tc")}

        if node.tag == _qn("w", "p") and not now_inside_table:
            has_text = any((text_node.text or "").strip() for text_node in node.findall(".//w:t", DOCX_NS))
            if has_text:
                _register_rpr_font(node.find("w:pPr/w:rPr", DOCX_NS))
                for run in node.findall("w:r", DOCX_NS):
                    if any((text_node.text or "").strip() for text_node in run.findall(".//w:t", DOCX_NS)):
                        _register_rpr_font(run.find("w:rPr", DOCX_NS))

        for child in reversed(list(node)):
            stack.append((child, now_inside_table))

    if not font_counts:
        return None

    dominant_font, dominant_count = max(
        font_counts.items(),
        key=lambda item: (item[1], item[0].lower()),
    )
    total = sum(font_counts.values())
    if dominant_count < 3 or (dominant_count / max(total, 1)) < 0.55:
        return None

    return {
        "font": {
            "font_name": dominant_font,
            "name": dominant_font,
        },
        "source": _build_font_source(
            "explicit",
            font_name=dominant_font,
            source_style="document.xml",
            scope="document",
        ),
        "count": dominant_count,
        "total": total,
    }


def _extract_docx_xml_details(docx_bytes: bytes) -> Dict[str, Any]:
    parts = _read_docx_parts(
        docx_bytes,
        [
            "word/styles.xml",
            "word/numbering.xml",
            "word/theme/theme1.xml",
            "word/fontTable.xml",
            "word/settings.xml",
            "word/document.xml",
        ],
    )

    details: Dict[str, Any] = {
        "styles": [],
        "doc_defaults": None,
        "latent_styles": None,
        "numbering": None,
        "theme": None,
        "font_table": None,
        "settings": None,
        "sections": [],
        "body_font_hint": None,
        "heading_styles": [],
        "meta": {},
    }

    styles_xml = parts.get("word/styles.xml")
    if styles_xml:
        try:
            styles_root = ET.fromstring(styles_xml)
            doc_defaults_elem = styles_root.find("w:docDefaults", DOCX_NS)
            if doc_defaults_elem is not None:
                details["doc_defaults"] = _parse_doc_defaults(doc_defaults_elem)
            latent_elem = styles_root.find("w:latentStyles", DOCX_NS)
            if latent_elem is not None:
                details["latent_styles"] = _parse_latent_styles(latent_elem)
            styles: List[Dict[str, Any]] = []
            for style_elem in styles_root.findall("w:style", DOCX_NS):
                styles.append(_parse_style_element(style_elem))
            details["styles"] = styles
            details["heading_styles"] = [
                {
                    "style_id": style.get("style_id"),
                    "display_name": style.get("display_name"),
                    "outline_level": style.get("outline_level"),
                }
                for style in styles
                if style.get("outline_level") is not None
            ]
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing styles.xml: {exc}")

    numbering_xml = parts.get("word/numbering.xml")
    if numbering_xml:
        try:
            numbering_root = ET.fromstring(numbering_xml)
            details["numbering"] = _parse_numbering(numbering_root)
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing numbering.xml: {exc}")

    theme_xml = parts.get("word/theme/theme1.xml")
    if theme_xml:
        try:
            theme_root = ET.fromstring(theme_xml)
            details["theme"] = _parse_theme(theme_root)
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing theme1.xml: {exc}")

    font_table_xml = parts.get("word/fontTable.xml")
    if font_table_xml:
        try:
            font_table_root = ET.fromstring(font_table_xml)
            details["font_table"] = _parse_font_table(font_table_root)
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing fontTable.xml: {exc}")

    settings_xml = parts.get("word/settings.xml")
    if settings_xml:
        try:
            settings_root = ET.fromstring(settings_xml)
            details["settings"] = _parse_settings(settings_root)
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing settings.xml: {exc}")

    document_xml = parts.get("word/document.xml")
    if document_xml:
        try:
            document_root = ET.fromstring(document_xml)
            details["sections"] = _parse_sections(document_root)
            details["body_font_hint"] = _extract_body_font_hint(
                document_root,
                details.get("theme") or {},
            )
        except Exception as exc:
            _logger.warning(f"[Template] Failed parsing document.xml: {exc}")

    styles_count = len(details.get("styles", []))
    details["meta"] = {
        "style_count": styles_count,
        "table_style_count": len([s for s in details.get("styles", []) if s.get("type") == "table"]),
        "numbering_abstract_count": len(details.get("numbering", {}).get("abstract_nums", [])) if details.get("numbering") else 0,
        "numbering_num_count": len(details.get("numbering", {}).get("nums", [])) if details.get("numbering") else 0,
        "section_count": len(details.get("sections", [])),
    }

    return details


def _to_int(value: Optional[str]) -> Optional[int]:
    try:
        return int(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _validate_table_index_value(table_index: Any) -> Optional[int]:
    """Return a valid non-negative table index or None when invalid."""
    if isinstance(table_index, bool) or not isinstance(table_index, int):
        return None
    if table_index < 0:
        return None
    return table_index


def _table_index_error_message(table_index: Any, table_count: Optional[int] = None) -> str:
    if table_count is None:
        return f"Invalid table_index {table_index!r}. Expected a non-negative integer."
    return (
        f"Invalid table_index {table_index!r}. "
        f"Expected 0 <= table_index < {table_count}."
    )


def _twips_to_points(value: Optional[str]) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(float(value) / 20.0, 2)
    except (TypeError, ValueError):
        return None


def _twips_to_lines(value: Optional[str]) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(float(value) / 240.0, 2)
    except (TypeError, ValueError):
        return None


def _normalize_bool_attr(value: Optional[str]) -> bool:
    if value is None:
        return True
    return str(value).lower() not in {"0", "false", "off", "none"}


def _coerce_optional_bool(value: Any) -> Optional[bool]:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return None


def _resolve_theme_font(theme_details: Optional[Dict[str, Any]], theme_key: Optional[str]) -> Optional[str]:
    if not theme_details or not theme_key:
        return None
    font_scheme = (theme_details.get("font_scheme") or {})
    key = str(theme_key).lower()
    group = None
    if key.startswith("major"):
        group = font_scheme.get("major") or {}
    elif key.startswith("minor"):
        group = font_scheme.get("minor") or {}
    if not group:
        return None
    if "cs" in key:
        return group.get("cs") or group.get("latin")
    if "ea" in key or "eastasia" in key:
        return group.get("ea") or group.get("latin")
    return group.get("latin")


def _resolve_theme_color(theme_details: Optional[Dict[str, Any]], color_key: Optional[str]) -> Optional[str]:
    if not theme_details or not color_key:
        return None
    scheme = theme_details.get("color_scheme") or {}
    colors = scheme.get("colors") or {}
    entry = colors.get(color_key)
    if not entry:
        return None
    attrs = entry.get("attrs") or {}
    if entry.get("type") == "srgbClr":
        return attrs.get("val")
    if entry.get("type") == "sysClr":
        return attrs.get("lastClr") or attrs.get("val")
    return None


def _parse_rpr_nodes(nodes: Optional[List[Dict[str, Any]]], theme_details: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    if not nodes:
        return {}
    font: Dict[str, Any] = {}
    underline_map = {
        "single": "SINGLE",
        "double": "DOUBLE",
        "dotted": "DOTTED",
        "dash": "DASH",
        "wave": "WAVY",
        "wavy": "WAVY",
    }
    for node in nodes:
        tag = node.get("tag")
        if not tag:
            continue
        attrs = node.get("attrs") or {}
        if tag == "rFonts":
            font_name = (
                attrs.get("ascii")
                or attrs.get("hAnsi")
                or attrs.get("cs")
                or attrs.get("eastAsia")
            )
            if not font_name:
                theme_key = (
                    attrs.get("asciiTheme")
                    or attrs.get("hAnsiTheme")
                    or attrs.get("csTheme")
                    or attrs.get("eastAsiaTheme")
                )
                font_name = _resolve_theme_font(theme_details, theme_key)
            if font_name:
                font["font_name"] = font_name
        elif tag == "sz":
            size_val = _to_int(attrs.get("val"))
            if size_val is not None:
                font["font_size_pt"] = round(size_val / 2.0, 1)
        elif tag == "b":
            font["bold"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "i":
            font["italic"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "u":
            val = attrs.get("val")
            if val and str(val).lower() in {"none", "0", "false"}:
                font["underline"] = False
                font["underline_style"] = None
            else:
                font["underline"] = True
                if val:
                    font["underline_style"] = underline_map.get(str(val).lower(), str(val).upper())
        elif tag == "color":
            val = attrs.get("val")
            if val and str(val).lower() != "auto":
                font["color_rgb"] = str(val).upper()
            else:
                theme_color = _resolve_theme_color(theme_details, attrs.get("themeColor"))
                if theme_color:
                    font["color_rgb"] = str(theme_color).upper()
        elif tag == "highlight":
            val = attrs.get("val")
            if val:
                font["highlight_color"] = str(val).upper()
        elif tag == "strike":
            font["strike"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "dstrike":
            font["double_strike"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "caps":
            font["all_caps"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "smallCaps":
            font["small_caps"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "vertAlign":
            val = str(attrs.get("val") or "").lower()
            if val == "superscript":
                font["superscript"] = True
            elif val == "subscript":
                font["subscript"] = True
    return font


def _parse_ppr_nodes(nodes: Optional[List[Dict[str, Any]]]) -> Dict[str, Any]:
    if not nodes:
        return {}
    para: Dict[str, Any] = {}
    line_rule_map = {
        "auto": "MULTIPLE",
        "exact": "EXACTLY",
        "atleast": "AT_LEAST",
    }
    for node in nodes:
        tag = node.get("tag")
        if not tag:
            continue
        attrs = node.get("attrs") or {}
        if tag == "jc":
            val = attrs.get("val")
            if val:
                normalized_val = str(val).upper()
                para["alignment"] = "JUSTIFY" if normalized_val == "BOTH" else normalized_val
        elif tag == "spacing":
            before = _twips_to_points(attrs.get("before"))
            after = _twips_to_points(attrs.get("after"))
            if before is not None:
                para["space_before_pt"] = before
            if after is not None:
                para["space_after_pt"] = after
            line_val = attrs.get("line")
            line_rule_raw = attrs.get("lineRule")
            if line_rule_raw:
                line_rule_key = str(line_rule_raw).lower()
                para["line_spacing_rule"] = line_rule_map.get(line_rule_key, str(line_rule_raw).upper())
            if line_val is not None:
                if line_rule_raw and str(line_rule_raw).lower() in {"exact", "atleast"}:
                    para["line_spacing"] = _twips_to_points(line_val)
                else:
                    para["line_spacing"] = _twips_to_lines(line_val)
        elif tag == "ind":
            left = _twips_to_inches(attrs.get("left"))
            right = _twips_to_inches(attrs.get("right"))
            first = _twips_to_inches(attrs.get("firstLine"))
            hanging = _twips_to_inches(attrs.get("hanging"))
            if left is not None:
                para["left_indent_inches"] = left
            if right is not None:
                para["right_indent_inches"] = right
            if hanging is not None:
                para["first_line_indent_inches"] = -hanging
                para["hanging_indent_inches"] = hanging
            elif first is not None:
                para["first_line_indent_inches"] = first
        elif tag == "keepNext":
            para["keep_with_next"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "keepLines":
            para["keep_together"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "pageBreakBefore":
            para["page_break_before"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "widowControl":
            para["widow_control"] = _normalize_bool_attr(attrs.get("val"))
        elif tag == "outlineLvl":
            outline_val = _to_int(attrs.get("val"))
            if outline_val is not None:
                para["outline_level"] = outline_val
    return para


def _parse_numpr_from_ppr(nodes: Optional[List[Dict[str, Any]]]) -> Optional[Dict[str, Any]]:
    if not nodes:
        return None
    for node in nodes:
        if node.get("tag") != "numPr":
            continue
        num_id = None
        level = None
        for child in node.get("children") or []:
            if child.get("tag") == "numId":
                num_id = child.get("attrs", {}).get("val")
            elif child.get("tag") == "ilvl":
                level = child.get("attrs", {}).get("val")
        if num_id is None and level is None:
            return None
        parsed_num_id = _to_int(num_id) if num_id is not None else None
        parsed_level = _to_int(level) if level is not None else None
        return {
            "num_id": parsed_num_id if parsed_num_id is not None else num_id,
            "level": parsed_level if parsed_level is not None else level,
        }
    return None


def _parse_list_indents(nodes: Optional[List[Dict[str, Any]]]) -> Dict[str, Optional[float]]:
    if not nodes:
        return {"left": None, "hanging": None, "first": None}
    for node in nodes:
        if node.get("tag") != "ind":
            continue
        attrs = node.get("attrs") or {}
        return {
            "left": _twips_to_inches(attrs.get("left")),
            "hanging": _twips_to_inches(attrs.get("hanging")),
            "first": _twips_to_inches(attrs.get("firstLine")),
        }
    return {"left": None, "hanging": None, "first": None}


# =============================================================================
# Table property parsing functions
# =============================================================================

def _parse_table_borders(children: Optional[List[Dict[str, Any]]]) -> Dict[str, Any]:
    """Parse tblBorders children into structured border properties."""
    if not children:
        return {}
    borders = {}
    border_tags = ["top", "bottom", "left", "right", "insideH", "insideV", "start", "end"]
    for child in children:
        tag = child.get("tag")
        if tag not in border_tags:
            continue
        attrs = child.get("attrs") or {}
        # Normalize start/end to left/right
        key = "left" if tag == "start" else ("right" if tag == "end" else tag)
        val = attrs.get("val")
        sz = attrs.get("sz")
        color = attrs.get("color")
        size_pt = None
        if sz not in (None, ""):
            try:
                size_pt = round(float(sz) / 8, 2)  # OOXML sz is 1/8 of pt
            except (TypeError, ValueError):
                size_pt = None
        borders[key] = {
            "style": val if val != "nil" else "none",
            "size_pt": size_pt,
            "color": color.upper() if color and color.lower() != "auto" else None,
        }
    return borders


def _parse_cell_margins(children: Optional[List[Dict[str, Any]]]) -> Dict[str, Optional[float]]:
    """Parse tblCellMar children into margin values in pt."""
    if not children:
        return {}
    margins = {}
    margin_tags = {"top", "bottom", "left", "right", "start", "end"}
    for child in children:
        tag = child.get("tag")
        if tag not in margin_tags:
            continue
        attrs = child.get("attrs") or {}
        # Normalize start/end to left/right
        key = "left" if tag == "start" else ("right" if tag == "end" else tag)
        w = attrs.get("w")
        w_type = attrs.get("type", "dxa")
        if w is not None:
            try:
                if w_type == "dxa":
                    margins[key] = round(float(w) / 20, 2)  # twips to pt
                else:
                    margins[key] = float(w)
            except (TypeError, ValueError):
                pass
    return margins


def _parse_tblpr_nodes(nodes: Optional[List[Dict[str, Any]]]) -> Dict[str, Any]:
    """Parse tblPr nodes into structured table format properties."""
    if not nodes:
        return {}
    result: Dict[str, Any] = {}
    for node in nodes:
        tag = node.get("tag")
        if not tag:
            continue
        attrs = node.get("attrs") or {}
        children = node.get("children")
        
        if tag == "tblBorders":
            result["borders"] = _parse_table_borders(children)
        elif tag == "shd":
            fill = attrs.get("fill")
            val = attrs.get("val")
            result["shading_color"] = fill.upper() if fill and fill.lower() not in ("auto", "none") else None
            result["shading_pattern"] = val
        elif tag == "tblCellMar":
            result["cell_margins"] = _parse_cell_margins(children)
        elif tag == "jc":
            result["alignment"] = attrs.get("val")
        elif tag == "tblW":
            w = attrs.get("w")
            w_type = attrs.get("type")
            result["width_value"] = _to_int(w) if w else None
            result["width_type"] = w_type  # auto, dxa, pct
        elif tag == "tblInd":
            w = attrs.get("w")
            w_type = attrs.get("type", "dxa")
            if w is not None:
                try:
                    if w_type == "dxa":
                        result["indent_pt"] = round(float(w) / 20, 2)
                    else:
                        result["indent_pt"] = float(w)
                except (TypeError, ValueError):
                    pass
        elif tag == "tblLayout":
            result["layout_type"] = attrs.get("type")  # fixed or autofit
        elif tag == "tblCellSpacing":
            w = attrs.get("w")
            w_type = attrs.get("type", "dxa")
            if w is not None:
                try:
                    if w_type == "dxa":
                        result["cell_spacing_pt"] = round(float(w) / 20, 2)
                    else:
                        result["cell_spacing_pt"] = float(w)
                except (TypeError, ValueError):
                    pass
        elif tag == "tblLook":
            # Table look options for conditional formatting
            result["look"] = {
                "firstRow": attrs.get("firstRow") == "1" or attrs.get("val", "")[:1] == "1",
                "lastRow": attrs.get("lastRow") == "1",
                "firstColumn": attrs.get("firstColumn") == "1",
                "lastColumn": attrs.get("lastColumn") == "1",
                "noHBand": attrs.get("noHBand") == "1",
                "noVBand": attrs.get("noVBand") == "1",
            }
        elif tag == "tblStyle":
            result["style_ref"] = attrs.get("val")
        elif tag == "tblPrChange":
            # Track changes - skip for now
            pass
    return result


def _parse_tcpr_nodes(nodes: Optional[List[Dict[str, Any]]]) -> Dict[str, Any]:
    """Parse tcPr (table cell properties) nodes into structured format."""
    if not nodes:
        return {}
    result: Dict[str, Any] = {}
    for node in nodes:
        tag = node.get("tag")
        if not tag:
            continue
        attrs = node.get("attrs") or {}
        children = node.get("children")
        
        if tag == "tcW":
            w = attrs.get("w")
            w_type = attrs.get("type")
            result["width_value"] = _to_int(w) if w else None
            result["width_type"] = w_type
        elif tag == "tcBorders":
            result["borders"] = _parse_table_borders(children)
        elif tag == "shd":
            fill = attrs.get("fill")
            val = attrs.get("val")
            result["shading_color"] = fill.upper() if fill and fill.lower() not in ("auto", "none") else None
            result["shading_pattern"] = val
        elif tag == "vAlign":
            result["vertical_align"] = attrs.get("val")  # top, center, bottom
        elif tag == "noWrap":
            result["no_wrap"] = True
        elif tag == "tcMar":
            result["margins"] = _parse_cell_margins(children)
        elif tag == "gridSpan":
            val = attrs.get("val")
            result["grid_span"] = _to_int(val) if val else None
        elif tag == "vMerge":
            val = attrs.get("val")
            result["v_merge"] = val if val else "continue"  # restart or continue
        elif tag == "hMerge":
            val = attrs.get("val")
            result["h_merge"] = val if val else "continue"
        elif tag == "textDirection":
            result["text_direction"] = attrs.get("val")
    return result


def _parse_tbl_style_variant(variant: Dict[str, Any]) -> Dict[str, Any]:
    """Parse a tblStylePr variant (firstRow, lastRow, bands, etc.)."""
    result: Dict[str, Any] = {
        "type": variant.get("type"),
    }
    props = variant.get("properties") or []
    
    # Parse rPr, pPr, tblPr, tcPr from the variant
    for prop in props:
        tag = prop.get("tag")
        if tag == "rPr":
            result["font"] = _parse_rpr_nodes(prop.get("children"))
        elif tag == "pPr":
            result["paragraph"] = _parse_ppr_nodes(prop.get("children"))
        elif tag == "tblPr":
            result["table"] = _parse_tblpr_nodes(prop.get("children"))
        elif tag == "tcPr":
            result["cell"] = _parse_tcpr_nodes(prop.get("children"))
    
    return result


def _merge_props(base: Optional[Dict[str, Any]], override: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    merged = dict(base) if base else {}
    if not override:
        return merged
    for key, value in override.items():
        if value is None or value == "":
            continue
        merged[key] = value
    return merged


def _copy_font_source(source: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    return copy.deepcopy(source) if isinstance(source, dict) else None


def _build_font_source(
    kind: str,
    *,
    font_name: Optional[str] = None,
    theme_key: Optional[str] = None,
    source_style: Optional[str] = None,
    scope: Optional[str] = None,
    inherited_from: Optional[str] = None,
) -> Dict[str, Any]:
    source: Dict[str, Any] = {"kind": str(kind)}
    if font_name not in (None, ""):
        source["font_name"] = str(font_name)
    if theme_key not in (None, ""):
        source["theme_key"] = str(theme_key)
    if source_style not in (None, ""):
        source["style_name"] = str(source_style)
    if scope not in (None, ""):
        source["scope"] = str(scope)
    if inherited_from not in (None, ""):
        source["inherited_from"] = str(inherited_from)
    return source


def _font_name_from_props(font_props: Optional[Dict[str, Any]]) -> Optional[str]:
    if not isinstance(font_props, dict):
        return None
    value = font_props.get("name") or font_props.get("font_name")
    if value in (None, ""):
        return None
    return str(value)


def _font_size_from_props(font_props: Optional[Dict[str, Any]]) -> Optional[float]:
    if not isinstance(font_props, dict):
        return None
    value = font_props.get("size_pt")
    if value in (None, ""):
        value = font_props.get("font_size_pt")
    if value in (None, ""):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _font_defaults_from_props(font_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(font_props, dict):
        return {}

    defaults: Dict[str, Any] = {}
    font_name = _font_name_from_props(font_props)
    if font_name:
        defaults["font_name"] = font_name

    font_size_pt = _font_size_from_props(font_props)
    if font_size_pt is not None:
        defaults["font_size_pt"] = font_size_pt

    font_color_rgb = _hex_to_rgb_tuple(font_props.get("color_rgb"))
    if font_color_rgb:
        defaults["font_color_rgb"] = font_color_rgb

    for bool_key in ("bold", "italic", "underline", "strike"):
        if font_props.get(bool_key) is not None:
            defaults[bool_key] = bool(font_props.get(bool_key))

    underline_style = font_props.get("underline_style")
    if underline_style not in (None, ""):
        defaults["underline_style"] = str(underline_style)

    return defaults


def _extract_rpr_font_source(
    nodes: Optional[List[Dict[str, Any]]],
    theme_details: Optional[Dict[str, Any]] = None,
    *,
    source_style: Optional[str] = None,
    scope: str = "style",
) -> Optional[Dict[str, Any]]:
    if not nodes:
        return None

    for node in nodes:
        if node.get("tag") != "rFonts":
            continue
        attrs = node.get("attrs") or {}
        explicit_font = (
            attrs.get("ascii")
            or attrs.get("hAnsi")
            or attrs.get("cs")
            or attrs.get("eastAsia")
        )
        if explicit_font:
            return _build_font_source(
                "explicit",
                font_name=str(explicit_font),
                source_style=source_style,
                scope=scope,
            )

        theme_key = (
            attrs.get("asciiTheme")
            or attrs.get("hAnsiTheme")
            or attrs.get("csTheme")
            or attrs.get("eastAsiaTheme")
        )
        if theme_key:
            resolved_font = _resolve_theme_font(theme_details, str(theme_key))
            return _build_font_source(
                "theme",
                font_name=resolved_font,
                theme_key=str(theme_key),
                source_style=source_style,
                scope=scope,
            )
    return None


def _normalize_system_font_name(value: Any) -> Optional[str]:
    if value in (None, ""):
        return None
    text = str(value).strip()
    if not text:
        return None
    text = re.sub(r"\s*\([^)]*\)\s*$", "", text).strip()
    return text or None


def _collect_windows_system_fonts() -> List[str]:
    fonts: Set[str] = set()
    if winreg is None:  # pragma: no cover - depends on host platform
        return []

    registry_paths = [
        (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    ]
    for hive, key_path in registry_paths:
        try:
            with winreg.OpenKey(hive, key_path) as key:  # type: ignore[arg-type]
                index = 0
                while True:
                    try:
                        value_name, _, _ = winreg.EnumValue(key, index)
                    except OSError:
                        break
                    index += 1
                    normalized = _normalize_system_font_name(value_name)
                    if normalized:
                        fonts.add(normalized)
        except OSError:
            continue
    return sorted(fonts, key=lambda item: item.lower())


def _collect_fontconfig_system_fonts() -> List[str]:
    try:
        result = subprocess.run(
            ["fc-list", "--format=%{family}\n"],
            capture_output=True,
            text=True,
            timeout=5,
            check=False,
        )
    except Exception:
        return []
    if result.returncode != 0:
        return []

    fonts: Set[str] = set()
    for line in result.stdout.splitlines():
        families = [segment.strip() for segment in str(line).split(",")]
        for family in families:
            normalized = _normalize_system_font_name(family)
            if normalized:
                fonts.add(normalized)
    return sorted(fonts, key=lambda item: item.lower())


def get_system_font_catalog(*, force_refresh: bool = False) -> List[str]:
    global _system_font_catalog_cache
    with _system_font_catalog_lock:
        if _system_font_catalog_cache is not None and not force_refresh:
            return list(_system_font_catalog_cache)

        if sys.platform == "win32":
            fonts = _collect_windows_system_fonts()
        else:
            fonts = _collect_fontconfig_system_fonts()

        _system_font_catalog_cache = list(fonts)
        return list(_system_font_catalog_cache)


def _map_xml_font_props(font_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not font_props:
        return {}
    mapped = dict(font_props)
    if "font_name" in font_props and "name" not in mapped:
        mapped["name"] = font_props.get("font_name")
    if "font_size_pt" in font_props and "size_pt" not in mapped:
        mapped["size_pt"] = font_props.get("font_size_pt")
    return mapped


def _filter_document_default_font_props(font_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(font_props, dict):
        return {}
    allowed_keys = (
        "font_name",
        "name",
        "font_size_pt",
        "size_pt",
        "bold",
        "italic",
        "underline",
        "underline_style",
        "color_rgb",
        "highlight_color",
        "strike",
        "double_strike",
        "all_caps",
        "small_caps",
        "superscript",
        "subscript",
    )
    normalized = {}
    for key in allowed_keys:
        if key in font_props:
            normalized[key] = copy.deepcopy(font_props.get(key))
    if "font_name" not in normalized and normalized.get("name") not in (None, ""):
        normalized["font_name"] = normalized.get("name")
    if "font_size_pt" not in normalized and normalized.get("size_pt") not in (None, ""):
        normalized["font_size_pt"] = normalized.get("size_pt")
    return normalized


def _filter_document_default_paragraph_props(paragraph_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(paragraph_props, dict):
        return {}
    allowed_keys = (
        "space_before_pt",
        "space_after_pt",
        "line_spacing",
        "line_spacing_rule",
        "first_line_indent_inches",
        "left_indent_inches",
        "right_indent_inches",
        "keep_with_next",
        "keep_together",
        "page_break_before",
        "widow_control",
        "alignment",
    )
    return {
        key: copy.deepcopy(paragraph_props.get(key))
        for key in allowed_keys
        if key in paragraph_props
    }


def _has_meaningful_paragraph_props(paragraph_props: Optional[Dict[str, Any]]) -> bool:
    if not isinstance(paragraph_props, dict) or not paragraph_props:
        return False
    for value in paragraph_props.values():
        if value is None or value == "":
            continue
        if isinstance(value, bool):
            if value:
                return True
            continue
        return True
    return False


def _build_document_defaults_payload(
    font_props: Optional[Dict[str, Any]],
    paragraph_props: Optional[Dict[str, Any]],
    *,
    font_source: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    normalized_font = _map_xml_font_props(_filter_document_default_font_props(font_props))
    normalized_paragraph = _filter_document_default_paragraph_props(paragraph_props)
    paragraph_source = None
    if _has_meaningful_paragraph_props(normalized_paragraph):
        paragraph_source = {
            "kind": "explicit",
            "scope": "docDefaults",
        }
    return {
        "font": normalized_font,
        "paragraph": normalized_paragraph,
        "font_source": _copy_font_source(font_source),
        "paragraph_source": paragraph_source,
    }


def _hex_to_rgb_tuple(color_value: Any) -> Optional[tuple[int, int, int]]:
    if color_value in (None, ""):
        return None
    text = str(color_value).strip().lstrip("#")
    if len(text) < 6:
        return None
    text = text[:6]
    try:
        return tuple(int(text[i:i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return None


def _collect_font_catalog(extracted: Dict[str, Any]) -> List[str]:
    fonts: Set[str] = set()

    def _add_font(value: Any) -> None:
        if value in (None, ""):
            return
        text = str(value).strip()
        if text:
            fonts.add(text)

    default_font = extracted.get("default_font") or {}
    _add_font(default_font.get("name"))
    _add_font(default_font.get("font_name"))

    for style in extracted.get("styles") or []:
        if not isinstance(style, dict):
            continue
        for font_block_key in ("font", "resolved_font", "xml_font"):
            font_block = style.get(font_block_key) or {}
            if not isinstance(font_block, dict):
                continue
            _add_font(font_block.get("name"))
            _add_font(font_block.get("font_name"))

    xml_details = extracted.get("xml_details") or {}
    theme = xml_details.get("theme") or {}
    font_scheme = theme.get("font_scheme") or {}
    for group_name in ("major", "minor"):
        group = font_scheme.get(group_name) or {}
        if not isinstance(group, dict):
            continue
        for key, value in group.items():
            if key == "script" and isinstance(value, dict):
                for script_font in value.values():
                    _add_font(script_font)
                continue
            _add_font(value)

    font_table = xml_details.get("font_table") or {}
    for font_entry in font_table.get("fonts") or []:
        if not isinstance(font_entry, dict):
            continue
        _add_font(font_entry.get("name"))
        _add_font(font_entry.get("alt_name"))

    return sorted(fonts, key=lambda item: item.lower())


def build_builder_required_style_defaults(extracted: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    if not isinstance(extracted, dict):
        return {}

    styles = extracted.get("styles") or []
    style_lookup: Dict[str, Dict[str, Any]] = {}
    for style in styles:
        if not isinstance(style, dict):
            continue
        for key in (style.get("style_id"), style.get("name"), style.get("display_name")):
            if key not in (None, ""):
                style_lookup[str(key)] = style

    def _find_style(*candidates: str) -> Optional[Dict[str, Any]]:
        for candidate in candidates:
            style = style_lookup.get(candidate)
            if isinstance(style, dict):
                return style
        return None

    def _font_defaults_from_style(*candidates: str) -> Dict[str, Any]:
        style = _find_style(*candidates)
        if not isinstance(style, dict):
            return {}
        font_props = style.get("resolved_font") or style.get("font") or {}
        return _font_defaults_from_props(font_props)

    theme = (extracted.get("xml_details") or {}).get("theme") or {}
    theme_major_font = _resolve_theme_font(theme, "majorHAnsi")
    theme_minor_font = _resolve_theme_font(theme, "minorHAnsi")

    default_font = extracted.get("default_font") or {}
    normal_style = _find_style("Normal")
    resolved_normal_font = (
        normal_style.get("resolved_font")
        if isinstance(normal_style, dict)
        else {}
    ) or {}
    body_seed = _font_defaults_from_props(_merge_props(default_font, resolved_normal_font))
    if not body_seed and theme_minor_font:
        body_seed = {"font_name": str(theme_minor_font)}

    heading_seed = (
        _font_defaults_from_style(
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Heading 4",
            "Heading 5",
            "Heading 6",
            "Title",
            "Subtitle",
        )
        or ({"font_name": str(theme_major_font)} if theme_major_font else {})
        or dict(body_seed)
    )
    list_seed = (
        _font_defaults_from_style("List Bullet", "List Number", "List Paragraph", "Normal")
        or dict(body_seed)
    )
    caption_seed = _font_defaults_from_style("Caption", "Normal") or dict(body_seed)
    code_seed = _font_defaults_from_style("Code", "Normal") or dict(body_seed)

    builder_defaults: Dict[str, Dict[str, Any]] = {}
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"):
        style_defaults = _font_defaults_from_style(style_name) or dict(heading_seed)
        if style_defaults:
            builder_defaults[style_name] = style_defaults

    for style_name in ("List Bullet", "List Number"):
        style_defaults = _font_defaults_from_style(style_name) or dict(list_seed)
        if style_defaults:
            builder_defaults[style_name] = style_defaults

    caption_defaults = _font_defaults_from_style("Caption") or dict(caption_seed)
    if caption_defaults:
        builder_defaults["Caption"] = caption_defaults

    code_defaults = _font_defaults_from_style("Code") or dict(code_seed)
    if code_defaults:
        builder_defaults["Code"] = code_defaults

    return builder_defaults


_HIGHLIGHT_MAP: Optional[Dict[str, Any]] = None


def _get_highlight_map() -> Dict[str, Any]:
    global _HIGHLIGHT_MAP
    if _HIGHLIGHT_MAP is not None:
        return _HIGHLIGHT_MAP
    names = [
        "YELLOW",
        "BRIGHT_GREEN",
        "TURQUOISE",
        "PINK",
        "BLUE",
        "RED",
        "DARK_BLUE",
        "DARK_RED",
        "DARK_YELLOW",
        "DARK_GREEN",
        "DARK_CYAN",
        "DARK_MAGENTA",
        "GRAY_50",
        "GRAY_25",
        "BLACK",
        "WHITE",
    ]
    result: Dict[str, Any] = {}
    try:
        for name in names:
            value = getattr(WD_COLOR_INDEX, name, None)
            if value is not None:
                result[name] = value
    except Exception:
        result = {}
    _HIGHLIGHT_MAP = result
    return result


def _resolve_highlight_color(value: Any) -> Optional[Any]:
    if not value:
        return None
    name = str(value).upper()
    highlight_map = _get_highlight_map()
    if name in highlight_map:
        return highlight_map[name]
    alias_map = {
        "DARK_GREEN": "GREEN",
        "DARK_CYAN": "TEAL",
        "DARK_MAGENTA": "VIOLET",
    }
    alias = alias_map.get(name)
    if alias:
        try:
            return getattr(WD_COLOR_INDEX, alias, None)
        except Exception:
            return None
    try:
        return getattr(WD_COLOR_INDEX, name, None)
    except Exception:
        return None


def _build_list_info_map(xml_details: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    styles = xml_details.get("styles") or []
    numbering = xml_details.get("numbering") or {}
    abstract_nums = numbering.get("abstract_nums") or []
    nums = numbering.get("nums") or []

    abstract_by_id = {
        str(item.get("abstract_num_id")): item
        for item in abstract_nums
        if item.get("abstract_num_id") is not None
    }
    num_to_abstract = {
        str(item.get("num_id")): str(item.get("abstract_num_id"))
        for item in nums
        if item.get("num_id") is not None and item.get("abstract_num_id") is not None
    }
    abstract_to_num: Dict[str, str] = {}
    for item in nums:
        abs_id = item.get("abstract_num_id")
        num_id = item.get("num_id")
        if abs_id is None or num_id is None:
            continue
        abs_key = str(abs_id)
        if abs_key not in abstract_to_num:
            abstract_to_num[abs_key] = str(num_id)

    style_name_by_id = {
        style.get("style_id"): style.get("display_name")
        for style in styles
        if style.get("style_id")
    }

    def build_from_level(num_id: Optional[str], abstract_id: Optional[str], lvl: Dict[str, Any], source: str) -> Dict[str, Any]:
        level = lvl.get("ilvl")
        num_fmt = lvl.get("num_fmt")
        lvl_text = lvl.get("lvl_text")
        start_val = _to_int(lvl.get("start"))
        list_alignment = lvl.get("lvl_jc")
        indents = _parse_list_indents(lvl.get("p_pr"))
        list_info = {
            "source": source,
            "num_id": _to_int(num_id) if num_id is not None else num_id,
            "abstract_num_id": _to_int(abstract_id) if abstract_id is not None else abstract_id,
            "level": level,
            "list_level": level,
            "list_format": num_fmt,
            "list_bullet_char": lvl_text if str(num_fmt).lower() == "bullet" else None,
            "lvl_text": lvl_text,
            "num_fmt": num_fmt,
            "list_start": start_val,
            "list_alignment": str(list_alignment).upper() if list_alignment else None,
            "list_left_indent_inches": indents.get("left"),
            "list_hanging_indent_inches": indents.get("hanging"),
        }
        return list_info

    def build_from_numpr(numpr: Dict[str, Any], source: str) -> Optional[Dict[str, Any]]:
        num_id = numpr.get("num_id")
        level = numpr.get("level")
        if num_id is None:
            return None
        abstract_id = num_to_abstract.get(str(num_id))
        abstract = abstract_by_id.get(str(abstract_id)) if abstract_id is not None else None
        if not abstract:
            return {
                "source": source,
                "num_id": num_id,
                "abstract_num_id": abstract_id,
                "level": level,
                "list_level": level,
            }
        lvl = None
        for candidate in abstract.get("levels", []):
            if level is None or candidate.get("ilvl") == level:
                lvl = candidate
                break
        if not lvl:
            return {
                "source": source,
                "num_id": num_id,
                "abstract_num_id": abstract_id,
                "level": level,
                "list_level": level,
            }
        return build_from_level(str(num_id), str(abstract_id), lvl, source)

    style_map: Dict[str, Dict[str, Any]] = {}

    for style in styles:
        style_id = style.get("style_id")
        p_pr_nodes = style.get("p_pr")
        numpr = _parse_numpr_from_ppr(p_pr_nodes)
        if numpr and style_id:
            info = build_from_numpr(numpr, "style_numpr")
            if info:
                style_map[style_id] = info
                display_name = style.get("display_name")
                if display_name:
                    style_map[display_name] = info

    for abstract in abstract_nums:
        abstract_id = abstract.get("abstract_num_id")
        abstract_key = str(abstract_id) if abstract_id is not None else None
        num_id = abstract_to_num.get(abstract_key) if abstract_key else None
        for lvl in abstract.get("levels", []) or []:
            p_style = lvl.get("p_style")
            if not p_style:
                continue
            info = build_from_level(num_id, abstract_key, lvl, "numbering_pstyle")
            style_map[p_style] = info
            display_name = style_name_by_id.get(p_style)
            if display_name:
                style_map[display_name] = info

    return style_map


def _build_style_prop_maps(xml_details: Dict[str, Any]) -> Dict[str, Any]:
    """Build explicit and resolved property maps for all styles.
    
    FIX #7: Enhanced with cycle detection and max depth protection.
    """
    styles = xml_details.get("styles") or []
    theme = xml_details.get("theme") or {}
    defaults = xml_details.get("doc_defaults") or {}
    default_font = _parse_rpr_nodes(defaults.get("r_pr"), theme)
    default_para = _parse_ppr_nodes(defaults.get("p_pr"))
    default_font_source = _extract_rpr_font_source(
        defaults.get("r_pr"),
        theme,
        scope="docDefaults",
    )

    style_by_id = {style.get("style_id"): style for style in styles if style.get("style_id")}
    style_by_name = {style.get("display_name"): style for style in styles if style.get("display_name")}

    explicit_map: Dict[str, Dict[str, Any]] = {}
    resolved_map: Dict[str, Dict[str, Any]] = {}
    
    # FIX #7: Constants for cycle/depth protection
    MAX_INHERITANCE_DEPTH = 20
    
    def resolve_style(style_ref: str, stack: Optional[set] = None, depth: int = 0) -> Dict[str, Any]:
        """Resolve style with cycle detection and max depth protection.
        
        FIX #7: Added stack parameter to track inheritance chain and detect cycles.
        Added depth parameter to prevent excessive recursion.
        """
        # Base case: no reference
        if not style_ref:
            return {
                "font": dict(default_font),
                "paragraph": dict(default_para),
                "font_source": _copy_font_source(default_font_source),
            }
        
        # FIX #7: Check max depth to prevent stack overflow
        if depth >= MAX_INHERITANCE_DEPTH:
            # Assuming _logger is defined elsewhere or will be added by the user
            # from ._logger import _logger
            _logger.warning(
                f"[Template] Max inheritance depth ({MAX_INHERITANCE_DEPTH}) reached for style '{style_ref}'. "
                "Possible cycle or excessive nesting. Returning defaults."
            )
            return {
                "font": dict(default_font),
                "paragraph": dict(default_para),
                "font_source": _copy_font_source(default_font_source),
            }
        
        # Check cache
        cache_key = style_ref
        if style_ref in resolved_map:
            return resolved_map[style_ref]
        
        # Lookup style
        style = style_by_id.get(style_ref) or style_by_name.get(style_ref)
        if not style:
            return {
                "font": dict(default_font),
                "paragraph": dict(default_para),
                "font_source": _copy_font_source(default_font_source),
            }
        
        style_id = style.get("style_id") or style_ref
        cache_key = style_id
        
        # Check cache again with style_id
        if cache_key in resolved_map:
            return resolved_map[cache_key]
        
        # FIX #7: Initialize stack for cycle detection
        if stack is None:
            stack = set()
        
        # FIX #7: Detect cycle
        if cache_key in stack:
            # _logger.warning(
            #     f"[Template] Cycle detected in style inheritance: {' -> '.join(stack)} -> {cache_key}. "
            #     "Breaking cycle and returning defaults."
            # )
            _logger.warning(
                f"[Template] Cycle detected in style inheritance: {' -> '.join(stack)} -> {cache_key}. "
                "Breaking cycle and returning defaults."
            )
            return {
                "font": dict(default_font),
                "paragraph": dict(default_para),
                "font_source": _copy_font_source(default_font_source),
            }
        
        # Add current style to stack
        stack.add(cache_key)
        
        try:
            # Resolve based-on style recursively
            base_ref = style.get("based_on") or style.get("link")
            if base_ref:
                base_props = resolve_style(base_ref, stack, depth + 1)
            else:
                base_props = {
                    "font": dict(default_font),
                    "paragraph": dict(default_para),
                    "font_source": _copy_font_source(default_font_source),
                }
            
            # Parse explicit properties
            explicit_font = _parse_rpr_nodes(style.get("r_pr"), theme)
            explicit_para = _parse_ppr_nodes(style.get("p_pr"))
            explicit_font_source = _extract_rpr_font_source(
                style.get("r_pr"),
                theme,
                source_style=style.get("display_name") or style.get("style_id") or style_ref,
                scope="style",
            )
            
            # Merge with base
            resolved_font = _merge_props(base_props.get("font"), explicit_font)
            resolved_para = _merge_props(base_props.get("paragraph"), explicit_para)
            resolved_font_source = _copy_font_source(base_props.get("font_source"))
            if explicit_font_source and explicit_font_source.get("font_name"):
                resolved_font_source = explicit_font_source
            elif resolved_font_source and base_ref:
                resolved_font_source.setdefault("inherited_from", str(base_ref))
             
            # Cache result
            result = {
                "font": resolved_font,
                "paragraph": resolved_para,
                "font_source": resolved_font_source,
            }
            resolved_map[cache_key] = result
            
            return result
            
        finally:
            # FIX #7: Always remove from stack when unwinding
            stack.discard(cache_key)

    # Resolve all styles
    for style in styles:
        style_id = style.get("style_id")
        display_name = style.get("display_name")
        style_type = style.get("type")
        explicit_font = _parse_rpr_nodes(style.get("r_pr"), theme)
        explicit_para = _parse_ppr_nodes(style.get("p_pr"))
        explicit_table = _parse_tblpr_nodes(style.get("tbl_pr"))
        explicit_cell = _parse_tcpr_nodes(style.get("tc_pr"))
        explicit_font_source = _extract_rpr_font_source(
            style.get("r_pr"),
            theme,
            source_style=display_name or style_id,
            scope="style",
        )
        
        # Parse table style variants (tblStylePr)
        tbl_style_variants = None
        if style.get("tbl_style_pr"):
            tbl_style_variants = {}
            for variant in style.get("tbl_style_pr") or []:
                parsed = _parse_tbl_style_variant(variant)
                variant_type = parsed.get("type")
                if variant_type:
                    tbl_style_variants[variant_type] = parsed
        
        if style_id:
            explicit_map[style_id] = {
                "font": explicit_font, 
                "paragraph": explicit_para,
                "table": explicit_table,
                "cell": explicit_cell,
                "table_variants": tbl_style_variants,
                "font_source": explicit_font_source,
            }
            resolved_map[style_id] = resolve_style(style_id)
            # Add table properties to resolved map for table styles
            if style_type == "table" and style_id in resolved_map:
                resolved_map[style_id]["table"] = explicit_table
                resolved_map[style_id]["cell"] = explicit_cell
                resolved_map[style_id]["table_variants"] = tbl_style_variants
        
        if display_name:
            explicit_map[display_name] = {
                "font": explicit_font,
                "paragraph": explicit_para,
                "table": explicit_table,
                "cell": explicit_cell,
                "table_variants": tbl_style_variants,
                "font_source": explicit_font_source,
            }
            if style_id and style_id in resolved_map:
                resolved_map[display_name] = resolved_map[style_id]
            else:
                resolved_map[display_name] = resolve_style(display_name)
            # Add table properties for table styles
            if style_type == "table" and display_name in resolved_map:
                resolved_map[display_name]["table"] = explicit_table
                resolved_map[display_name]["cell"] = explicit_cell
                resolved_map[display_name]["table_variants"] = tbl_style_variants

    return {
        "defaults": {
            "font": default_font,
            "paragraph": default_para,
            "font_source": _copy_font_source(default_font_source),
        },
        "explicit": explicit_map,
        "resolved": resolved_map,
    }


def _sanitize_kernel_id(kernel_id: str) -> str:
    """Sanitize kernel_id to prevent path traversal attacks.
    
    FIX #11: Only allows alphanumeric characters, hyphens, and underscores.
    This prevents directory traversal attacks like "../../../etc/passwd".
    
    Args:
        kernel_id: Raw kernel ID from client
        
    Returns:
        Sanitized kernel ID safe for use in file paths
        
    Raises:
        ValueError: If kernel_id is invalid or becomes empty after sanitization
    """
    if not kernel_id or not isinstance(kernel_id, str):
        raise ValueError("kernel_id must be a non-empty string")
    
    # Only allow alphanumeric, hyphen, and underscore
    safe_id = "".join(c for c in kernel_id if c.isalnum() or c in "-_")
    
    if not safe_id:
        raise ValueError(f"kernel_id '{kernel_id}' contains no valid characters")
    
    # Additional check: ensure it doesn't start with '.' (hidden files)
    if safe_id.startswith("."):
        safe_id = safe_id.lstrip(".")
    
    if not safe_id:
        raise ValueError(f"kernel_id '{kernel_id}' is invalid after sanitization")
    
    return safe_id


def _ensure_template_dir(kernel_id: str) -> Path:
    """Ensure the template directory exists for the given kernel.
    
    FIX #11: Uses sanitized kernel_id to prevent path traversal.
    """
    safe_id = _sanitize_kernel_id(kernel_id)
    template_dir = TEMPLATE_DIR / safe_id
    template_dir.mkdir(parents=True, exist_ok=True)
    return template_dir


def _get_template_json_path(kernel_id: str) -> Path:
    """Get the path to the template JSON file.
    
    FIX #11: Uses sanitized kernel_id.
    """
    safe_id = _sanitize_kernel_id(kernel_id)
    return TEMPLATE_DIR / safe_id / "template.json"


def _get_template_docx_path(kernel_id: str) -> Path:
    """Get the path to the template DOCX file.
    
    FIX #11: Uses sanitized kernel_id.
    """
    safe_id = _sanitize_kernel_id(kernel_id)
    return TEMPLATE_DIR / safe_id / "template.docx"


def _load_template_json_file(json_path: Path) -> Optional[Dict[str, Any]]:
    if not json_path.exists():
        return None
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as exc:
        _logger.warning("[Template] Could not read template JSON %s: %s", json_path, exc)
        return None
    return data if isinstance(data, dict) else None


def _write_template_files(kernel_id: str, docx_bytes: bytes, extracted_json: Dict[str, Any]) -> tuple[Path, Path]:
    _ensure_template_dir(kernel_id)
    docx_path = _get_template_docx_path(kernel_id)
    json_path = _get_template_json_path(kernel_id)

    with open(docx_path, "wb") as f:
        f.write(docx_bytes)
    if not docx_path.exists() or docx_path.stat().st_size == 0:
        raise IOError(f"Failed to write DOCX file: {docx_path}")

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(extracted_json, f, indent=2, ensure_ascii=False)
    if not json_path.exists() or json_path.stat().st_size == 0:
        raise IOError(f"Failed to write JSON file: {json_path}")

    return docx_path, json_path


def _create_template_backups(
    kernel_id: str,
    *,
    reason: Optional[str] = None,
) -> tuple[Optional[Path], Optional[Path]]:
    docx_path = _get_template_docx_path(kernel_id)
    json_path = _get_template_json_path(kernel_id)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_docx = docx_path.with_suffix(f".backup_{timestamp}.docx")
    backup_json = json_path.with_suffix(f".backup_{timestamp}.json")

    created_docx: Optional[Path] = None
    created_json: Optional[Path] = None

    if docx_path.exists():
        shutil.copy2(docx_path, backup_docx)
        created_docx = backup_docx
    if json_path.exists():
        shutil.copy2(json_path, backup_json)
        created_json = backup_json

    if created_docx or created_json:
        _logger.info(
            "[Template] Created backup for kernel %s%s",
            kernel_id,
            f" ({reason})" if reason else "",
        )
    return created_docx, created_json


def _sanitize_persisted_template_if_needed(kernel_id: str) -> Optional[Dict[str, Any]]:
    docx_path = _get_template_docx_path(kernel_id)
    if not docx_path.exists():
        return None

    json_path = _get_template_json_path(kernel_id)
    existing_json = _load_template_json_file(json_path)

    try:
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
    except Exception as exc:
        _logger.warning("[Template] Could not read template DOCX %s: %s", docx_path, exc)
        return existing_json

    try:
        prepared_docx, prepared_json, docx_changed = _prepare_template_payload(docx_bytes, existing_json)
    except Exception as exc:
        _logger.warning("[Template] Could not sanitize persisted template for kernel %s: %s", kernel_id, exc)
        return existing_json

    json_changed = not isinstance(existing_json, dict) or existing_json != prepared_json
    if docx_changed or json_changed:
        if docx_changed:
            _create_template_backups(kernel_id, reason="legacy table style runtime defaults sanitization")
        _write_template_files(kernel_id, prepared_docx, prepared_json)
        _cleanup_old_backups(kernel_id, max_backups=5)
        clear_preview_cache(kernel_id)

    return prepared_json




def _cleanup_old_backups(kernel_id: str, max_backups: int = 5) -> None:
    """Clean up old backup files, keeping only the N most recent.
    
    FIX #17: Prevents disk space exhaustion from accumulating backups.
    
    Args:
        kernel_id: The kernel ID
        max_backups: Maximum number of backups to keep (default: 5)
    """
    try:
        safe_id = _sanitize_kernel_id(kernel_id)
        template_dir = TEMPLATE_DIR / safe_id
        
        if not template_dir.exists():
            return
        
        # Find all backup files
        backup_files = []
        for pattern in ["*.backup_*.docx", "*.backup_*.json"]:
            backup_files.extend(template_dir.glob(pattern))
        
        if len(backup_files) <= max_backups:
            return  # Nothing to clean up
        
        # Sort by modification time (newest first)
        backup_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        
        # Delete old backups
        files_to_delete = backup_files[max_backups:]
        for backup_file in files_to_delete:
            try:
                backup_file.unlink()
                _logger.debug(f"[Template] Deleted old backup: {backup_file.name}")
            except Exception as e:
                _logger.warning(f"[Template] Could not delete backup {backup_file.name}: {e}")
        
        if files_to_delete:
            _logger.info(
                f"[Template] Cleaned up {len(files_to_delete)} old backup(s) for kernel {kernel_id}, "
                f"kept {max_backups} most recent"
            )
    
    except Exception as e:
        _logger.error(f"[Template] Error during backup cleanup: {e}")


def extract_styles_from_docx(docx_bytes: bytes) -> Dict[str, Any]:
    """Extract styles, page setup, headers, and footers from a DOCX file.
    
    Args:
        docx_bytes: The raw bytes of the DOCX file.
        
    Returns:
        A dictionary containing:
        - page_setup: margins, orientation, page size
        - styles: list of paragraph/character styles with font info
        - headers: list of header texts
        - footers: list of footer texts
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed")
    
    doc = Document(io.BytesIO(docx_bytes))
    result: Dict[str, Any] = {
        "page_setup": {},
        "styles": [],
        "headers": [],
        "footers": [],
        DOCUMENT_DEFAULTS_KEY: {
            "font": {},
            "paragraph": {},
            "font_source": None,
            "paragraph_source": None,
        },
        "default_font": None,
        "default_font_source": None,
        "font_catalog": [],
        SYSTEM_FONT_CATALOG_KEY: [],
        "xml_details": None,
        "document_tables": [],
        "document_captions": [],
        "content_controls": {"controls": [], "placeholders": [], "control_count": 0, "placeholder_count": 0, "unwrapped_placeholder_count": 0},
        TABLE_STYLE_RUNTIME_DEFAULTS_KEY: {},
        BUILDER_REQUIRED_STYLE_DEFAULTS_KEY: {},
    }
    
    # -------------------------------------------------------------------------
    # Page Setup (from last section, which is typically the main one)
    # -------------------------------------------------------------------------
    if doc.sections:
        section = doc.sections[-1]
        result["page_setup"] = {
            "page_width_inches": round(section.page_width.inches, 2) if section.page_width else None,
            "page_height_inches": round(section.page_height.inches, 2) if section.page_height else None,
            "left_margin_inches": round(section.left_margin.inches, 2) if section.left_margin else None,
            "right_margin_inches": round(section.right_margin.inches, 2) if section.right_margin else None,
            "top_margin_inches": round(section.top_margin.inches, 2) if section.top_margin else None,
            "bottom_margin_inches": round(section.bottom_margin.inches, 2) if section.bottom_margin else None,
            "orientation": "landscape" if section.orientation and section.orientation.value == 1 else "portrait",
        }
    
    # -------------------------------------------------------------------------
    # Styles (paragraph, character, and table styles)
    # -------------------------------------------------------------------------
    priority_styles = {"Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle", "Caption"}
    table_priority_styles = {"Table Grid", "TableGrid"}  # Common table styles
    
    for style in doc.styles:
        # Process paragraph, character, AND table styles
        if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            continue
        
        # Skip hidden/internal styles unless they're priority ones or table styles
        # Table styles are often marked hidden but should be included for editing
        is_priority = style.name in priority_styles or style.name in table_priority_styles
        is_table_style = style.type == WD_STYLE_TYPE.TABLE
        if style.hidden and not is_priority and not is_table_style:
            continue
        
        # Determine style type string
        if style.type == WD_STYLE_TYPE.TABLE:
            style_type_str = "table"
        elif style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_type_str = "paragraph"
        else:
            style_type_str = "character"
        
        style_info: Dict[str, Any] = {
            "name": style.name,
            "display_name": style.name,
            "type": style_type_str,
            "base_style": style.base_style.name if style.base_style else None,
            "priority": style.name in priority_styles or style.name in table_priority_styles,
            "style_id": getattr(style, "style_id", None),
        }
        
        # Font information (table styles may not have font property)
        font = None
        try:
            if hasattr(style, 'font'):
                font = style.font
        except Exception:
            pass
        
        if font:
            underline_value = font.underline
            underline_style = None
            underline_bool = None
            if underline_value is not None:
                if isinstance(underline_value, bool):
                    underline_bool = underline_value
                else:
                    underline_style = underline_value.name if hasattr(underline_value, "name") else str(underline_value)
                    underline_bool = True

            highlight_value = getattr(font, "highlight_color", None)
            highlight_name = highlight_value.name if highlight_value is not None and hasattr(highlight_value, "name") else None

            style_info["font"] = {
                "name": font.name,
                "size_pt": round(font.size.pt, 1) if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
                "underline": underline_bool,
                "underline_style": underline_style,
                "color_rgb": str(font.color.rgb) if font.color and font.color.rgb else None,
                "highlight_color": highlight_name,
                "strike": getattr(font, "strike", None),
                "double_strike": getattr(font, "double_strike", None),
                "all_caps": getattr(font, "all_caps", None),
                "small_caps": getattr(font, "small_caps", None),
                "superscript": getattr(font, "superscript", None),
                "subscript": getattr(font, "subscript", None),
            }
        
        # Paragraph format (only for paragraph styles)
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            pf = style.paragraph_format
            if pf:
                style_info["paragraph_format"] = {
                    "alignment": pf.alignment.name if pf.alignment else None,
                    "space_before_pt": round(pf.space_before.pt, 1) if pf.space_before else None,
                    "space_after_pt": round(pf.space_after.pt, 1) if pf.space_after else None,
                    "line_spacing": pf.line_spacing if pf.line_spacing else None,
                    "line_spacing_rule": pf.line_spacing_rule.name if pf.line_spacing_rule else None,
                    "first_line_indent_inches": round(pf.first_line_indent.inches, 2) if pf.first_line_indent else None,
                    "left_indent_inches": round(pf.left_indent.inches, 2) if pf.left_indent else None,
                    "right_indent_inches": round(pf.right_indent.inches, 2) if pf.right_indent else None,
                    "keep_with_next": pf.keep_with_next,
                    "keep_together": pf.keep_together,
                    "page_break_before": pf.page_break_before,
                    "widow_control": pf.widow_control,
                }
        
        result["styles"].append(style_info)
    
    # Sort: priority styles first, then alphabetically
    result["styles"].sort(key=lambda s: (not s.get("priority", False), s["name"]))

    # -------------------------------------------------------------------------
    # XML Details (styles, numbering, themes, sections, defaults)
    # -------------------------------------------------------------------------
    xml_details = _extract_docx_xml_details(docx_bytes)
    result["xml_details"] = xml_details
    outline_by_id = {
        style.get("style_id"): style.get("outline_level")
        for style in (xml_details.get("styles") or [])
        if style.get("style_id")
    }
    outline_by_name = {
        style.get("display_name"): style.get("outline_level")
        for style in (xml_details.get("styles") or [])
        if style.get("display_name")
    }
    for style_info in result["styles"]:
        style_id = style_info.get("style_id")
        outline = outline_by_id.get(style_id) if style_id else None
        if outline is None:
            outline = outline_by_name.get(style_info.get("name"))
        if outline is not None:
            style_info["outline_level"] = outline
            if isinstance(style_info.get("paragraph_format"), dict):
                style_info["paragraph_format"]["outline_level"] = outline

    # -------------------------------------------------------------------------
    # Include table styles from XML details (python-docx skips table styles)
    # -------------------------------------------------------------------------
    table_styles = [
        style for style in (xml_details.get("styles") or [])
        if style.get("type") == "table"
    ]
    if table_styles:
        existing_ids = {style.get("style_id") for style in result["styles"] if style.get("style_id")}
        existing_names = {style.get("name") for style in result["styles"]}
        for table_style in table_styles:
            style_id = table_style.get("style_id")
            display_name = table_style.get("display_name") or style_id
            if style_id in existing_ids or display_name in existing_names:
                continue
            result["styles"].append({
                "name": display_name,
                "display_name": display_name,
                "type": "table",
                "base_style": table_style.get("based_on"),
                "priority": style_id == "TableGrid",
                "style_id": style_id,
                "font": {},
                "paragraph_format": {},
            })

    # -------------------------------------------------------------------------
    # Resolve XML style details (explicit + effective) and numbering
    # -------------------------------------------------------------------------
    style_prop_maps = _build_style_prop_maps(xml_details)
    list_info_map = _build_list_info_map(xml_details)

    for style_info in result["styles"]:
        keys = [
            style_info.get("style_id"),
            style_info.get("display_name"),
            style_info.get("name"),
        ]
        explicit = None
        resolved = None
        list_info = None
        for key in keys:
            if not key:
                continue
            if explicit is None:
                explicit = style_prop_maps.get("explicit", {}).get(key)
            if resolved is None:
                resolved = style_prop_maps.get("resolved", {}).get(key)
            if list_info is None:
                list_info = list_info_map.get(key)
        if explicit:
            xml_font = explicit.get("font")
            mapped_xml_font = _map_xml_font_props(xml_font)
            style_info["xml_font"] = xml_font
            style_info["xml_font_source"] = _copy_font_source(explicit.get("font_source"))
            style_info["xml_paragraph_format"] = explicit.get("paragraph")
            style_info["font"] = _merge_props(style_info.get("font"), mapped_xml_font)
            style_info["paragraph_format"] = _merge_props(style_info.get("paragraph_format"), explicit.get("paragraph"))
            # Include table/cell properties for table styles
            if explicit.get("table"):
                style_info["xml_table_format"] = explicit.get("table")
            if explicit.get("cell"):
                style_info["xml_cell_format"] = explicit.get("cell")
            if explicit.get("table_variants"):
                style_info["xml_table_variants"] = explicit.get("table_variants")
        if resolved:
            resolved_font = _map_xml_font_props(resolved.get("font"))
            style_info["resolved_font"] = resolved_font
            style_info["resolved_font_source"] = _copy_font_source(resolved.get("font_source"))
            style_info["font_source"] = _copy_font_source(resolved.get("font_source"))
            style_info["resolved_paragraph_format"] = resolved.get("paragraph")
            # Include resolved table/cell properties
            if resolved.get("table"):
                style_info["resolved_table_format"] = resolved.get("table")
            if resolved.get("cell"):
                style_info["resolved_cell_format"] = resolved.get("cell")
            if resolved.get("table_variants"):
                style_info["table_variants"] = resolved.get("table_variants")
        if list_info:
            style_info["list_info"] = list_info
    
    # -------------------------------------------------------------------------
    # Default document font (effective Normal/docDefaults/theme resolution)
    # -------------------------------------------------------------------------
    default_font_from_xml = _map_xml_font_props((style_prop_maps.get("defaults") or {}).get("font"))
    default_font_source = _copy_font_source((style_prop_maps.get("defaults") or {}).get("font_source"))
    default_paragraph_from_xml = _filter_document_default_paragraph_props(
        (style_prop_maps.get("defaults") or {}).get("paragraph")
    )
    result[DOCUMENT_DEFAULTS_KEY] = _build_document_defaults_payload(
        default_font_from_xml,
        default_paragraph_from_xml,
        font_source=default_font_source,
    )
    body_font_hint = (xml_details or {}).get("body_font_hint") or {}
    body_hint_font = _map_xml_font_props(body_font_hint.get("font"))
    body_hint_source = _copy_font_source(body_font_hint.get("source"))
    try:
        normal_style = next(
            (
                style for style in result["styles"]
                if style.get("style_id") == "Normal" or style.get("name") == "Normal"
            ),
            None,
        )
        normal_resolved_font = (
            normal_style.get("resolved_font")
            if isinstance(normal_style, dict)
            else {}
        ) or {}
        result["default_font"] = _merge_props(default_font_from_xml, normal_resolved_font) or None
    except Exception:
        normal_style = None
        normal_resolved_font = {}
        result["default_font"] = default_font_from_xml or None
    result["default_font_source"] = default_font_source

    normal_explicit_font = (
        _font_name_from_props((normal_style or {}).get("xml_font"))
        if isinstance(normal_style, dict)
        else None
    )
    current_default_font_name = _font_name_from_props(result.get("default_font"))
    body_hint_font_name = _font_name_from_props(body_hint_font)
    should_promote_body_font = (
        body_hint_font_name not in (None, "")
        and body_hint_font_name != current_default_font_name
        and not normal_explicit_font
        and (default_font_source or {}).get("kind") == "theme"
    )
    if should_promote_body_font:
        result["default_font"] = _merge_props(result.get("default_font"), body_hint_font) or body_hint_font
        result["default_font_source"] = body_hint_source
        if isinstance(normal_style, dict):
            normal_style["resolved_font"] = _merge_props(normal_style.get("resolved_font"), body_hint_font)
            normal_style["resolved_font_source"] = _copy_font_source(body_hint_source)
            normal_style["font_source"] = _copy_font_source(body_hint_source)

    result["font_catalog"] = _collect_font_catalog(result)
    result[SYSTEM_FONT_CATALOG_KEY] = get_system_font_catalog()
    result[BUILDER_REQUIRED_STYLE_DEFAULTS_KEY] = build_builder_required_style_defaults(result)
    
    # -------------------------------------------------------------------------
    # Headers and Footers
    # -------------------------------------------------------------------------
    headers_from_ooxml, footers_from_ooxml = _extract_header_footer_texts_from_docx_bytes(docx_bytes)
    result["headers"] = headers_from_ooxml
    result["footers"] = footers_from_ooxml
    
    _logger.info(
        f"[Template] Extracted {len(result['styles'])} styles, "
        f"{len(result['headers'])} headers, {len(result['footers'])} footers"
    )
    
    # -------------------------------------------------------------------------
    # Document Tables (actual table instances with direct formatting)
    # -------------------------------------------------------------------------
    document_tables, document_captions = _extract_document_tables_and_captions(docx_bytes)
    result["document_tables"] = document_tables
    result["document_captions"] = document_captions
    try:
        result["content_controls"] = inspect_content_controls(docx_bytes)
    except Exception as exc:
        result["content_controls"] = {
            "controls": [],
            "placeholders": [],
            "control_count": 0,
            "placeholder_count": 0,
            "unwrapped_placeholder_count": 0,
            "error": f"{type(exc).__name__}",
        }

    return _merge_runtime_defaults_into_extracted(result)


def get_style_coverage(extracted: Dict[str, Any]) -> Dict[str, Any]:
    from app.services import template_extract

    return template_extract.get_style_coverage(extracted)


_TABLE_BLOCK_KEY_MAP = {
    "border_style": "table_border_style",
    "border_size_pt": "table_border_size_pt",
    "border_color": "table_border_color",
    "shading_color": "table_shading_color",
    "alignment": "table_alignment",
    "cell_margin_top_pt": "table_cell_margin_top_pt",
    "cell_margin_bottom_pt": "table_cell_margin_bottom_pt",
    "cell_margin_left_pt": "table_cell_margin_left_pt",
    "cell_margin_right_pt": "table_cell_margin_right_pt",
    "layout_type": "table_layout_type",
    "cell_spacing_pt": "table_cell_spacing_pt",
    "width_type": "table_width_type",
    "width_value": "table_width_value",
    "look_first_row": "table_look_first_row",
    "look_last_row": "table_look_last_row",
    "look_first_column": "table_look_first_column",
    "look_last_column": "table_look_last_column",
    "look_no_h_band": "table_look_no_h_band",
    "look_no_v_band": "table_look_no_v_band",
    "cell_shading_color": "table_cell_shading_color",
    "cell_vertical_align": "table_cell_vertical_align",
}


def _normalize_style_updates(updates: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize legacy flat updates and nested block updates into flat keys."""
    if not isinstance(updates, dict):
        return {}

    normalized: Dict[str, Any] = dict(updates)

    for block_name in ("font", "paragraph"):
        block = updates.get(block_name)
        if isinstance(block, dict):
            for key, value in block.items():
                if isinstance(key, str) and key:
                    normalized[key] = value

    table_block = updates.get("table")
    if isinstance(table_block, dict):
        for key, value in table_block.items():
            if not isinstance(key, str) or not key:
                continue
            mapped_key = _TABLE_BLOCK_KEY_MAP.get(key)
            if mapped_key:
                normalized[mapped_key] = value
            elif key.startswith("table_"):
                normalized[key] = value

    return normalized


def _normalize_document_defaults_updates(updates: Dict[str, Any]) -> Dict[str, Any]:
    normalized = _normalize_style_updates(updates)
    allowed_keys = set(_filter_document_default_font_props({}).keys()) | {
        "font_name",
        "font_size_pt",
        "bold",
        "italic",
        "underline",
        "underline_style",
        "color_rgb",
        "highlight_color",
        "strike",
        "double_strike",
        "all_caps",
        "small_caps",
        "superscript",
        "subscript",
        "space_before_pt",
        "space_after_pt",
        "line_spacing",
        "line_spacing_rule",
        "first_line_indent_inches",
        "left_indent_inches",
        "right_indent_inches",
        "keep_with_next",
        "keep_together",
        "page_break_before",
        "widow_control",
        "alignment",
    }
    return {
        key: value
        for key, value in normalized.items()
        if key in allowed_keys
    }


def update_template_document_defaults(
    kernel_id: str,
    updates: Dict[str, Any],
) -> Dict[str, Any]:
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed")

    normalized_updates = _normalize_document_defaults_updates(updates)
    if not normalized_updates:
        return get_template(kernel_id) or {}

    current_template = get_template(kernel_id) or {}
    current_document_defaults = current_template.get(DOCUMENT_DEFAULTS_KEY) or {}
    current_font = _filter_document_default_font_props(current_document_defaults.get("font"))
    current_paragraph = _filter_document_default_paragraph_props(current_document_defaults.get("paragraph"))

    merged_font = dict(current_font)
    for key in (
        "font_name",
        "name",
        "font_size_pt",
        "size_pt",
        "bold",
        "italic",
        "underline",
        "underline_style",
        "color_rgb",
        "highlight_color",
        "strike",
        "double_strike",
        "all_caps",
        "small_caps",
        "superscript",
        "subscript",
    ):
        if key in normalized_updates:
            merged_font[key] = copy.deepcopy(normalized_updates.get(key))

    if "font_name" in merged_font and merged_font.get("font_name") in (None, ""):
        merged_font.pop("font_name", None)
        merged_font.pop("name", None)
    elif "font_name" in merged_font:
        merged_font["name"] = merged_font.get("font_name")

    if "font_size_pt" in merged_font and merged_font.get("font_size_pt") in (None, ""):
        merged_font.pop("font_size_pt", None)
        merged_font.pop("size_pt", None)
    elif "font_size_pt" in merged_font:
        merged_font["size_pt"] = merged_font.get("font_size_pt")

    merged_paragraph = dict(current_paragraph)
    for key in (
        "space_before_pt",
        "space_after_pt",
        "line_spacing",
        "line_spacing_rule",
        "first_line_indent_inches",
        "left_indent_inches",
        "right_indent_inches",
        "keep_with_next",
        "keep_together",
        "page_break_before",
        "widow_control",
        "alignment",
    ):
        if key in normalized_updates:
            merged_paragraph[key] = copy.deepcopy(normalized_updates.get(key))

    document_defaults = {
        "font": _filter_document_default_font_props(merged_font),
        "paragraph": _filter_document_default_paragraph_props(merged_paragraph),
    }

    json_path = _get_template_json_path(kernel_id)
    docx_path = _get_template_docx_path(kernel_id)
    if not json_path.exists() or not docx_path.exists():
        raise ValueError(f"No template found for kernel {kernel_id}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_docx = docx_path.with_suffix(f".backup_{timestamp}.docx")
    backup_json = json_path.with_suffix(f".backup_{timestamp}.json")

    try:
        shutil.copy2(docx_path, backup_docx)
        if json_path.exists():
            shutil.copy2(json_path, backup_json)

        _logger.info(
            "[Template] Created backup for document defaults update (kernel %s): %s",
            kernel_id,
            backup_docx.name,
        )

        _sync_document_defaults_in_docx(docx_path, document_defaults)

        with open(docx_path, "rb") as f:
            new_docx_bytes = f.read()
        extracted = extract_styles_from_docx(new_docx_bytes)
        extracted["style_coverage"] = get_style_coverage(extracted)
        save_template(kernel_id, new_docx_bytes, extracted)

        try:
            backup_docx.unlink()
            if backup_json.exists():
                backup_json.unlink()
        except Exception as cleanup_err:
            _logger.warning("[Template] Could not cleanup backup: %s", cleanup_err)

        _logger.info("[Template] Successfully updated document defaults for kernel %s", kernel_id)
        return extracted
    except Exception as exc:
        _logger.error(
            "[Template] Error updating document defaults for kernel %s: %s. Rolling back to backup.",
            kernel_id,
            exc,
        )
        try:
            if backup_docx.exists():
                shutil.copy2(backup_docx, docx_path)
            if backup_json.exists():
                shutil.copy2(backup_json, json_path)
        except Exception as rollback_err:
            _logger.critical(
                "[Template] CRITICAL: document defaults rollback failed for kernel %s: %s",
                kernel_id,
                rollback_err,
            )
        raise


def update_template_semantic_style_slots(
    kernel_id: str,
    semantic_style_slots: Dict[str, Any],
) -> Dict[str, Any]:
    if not isinstance(semantic_style_slots, dict):
        raise ValueError("semantic_style_slots must be a mapping")

    current_template = get_template(kernel_id) or {}
    json_path = _get_template_json_path(kernel_id)
    docx_path = _get_template_docx_path(kernel_id)
    if not json_path.exists() or not docx_path.exists():
        raise ValueError(f"No template found for kernel {kernel_id}")

    from app.services import template_extract

    current_slots = current_template.get(SEMANTIC_STYLE_SLOTS_KEY)
    merged_slots = copy.deepcopy(current_slots) if isinstance(current_slots, dict) else {}
    for slot_name, slot_payload in semantic_style_slots.items():
        if slot_name in template_extract.SEMANTIC_STYLE_SLOT_ORDER:
            merged_slots[slot_name] = copy.deepcopy(slot_payload)

    normalized_slots = template_extract.build_semantic_style_slots(
        current_template,
        previous_slots=merged_slots,
    )

    with open(docx_path, "rb") as f:
        docx_bytes = f.read()

    prepared_template = copy.deepcopy(current_template)
    prepared_template[SEMANTIC_STYLE_SLOTS_KEY] = normalized_slots
    save_template(kernel_id, docx_bytes, prepared_template)
    updated = get_template(kernel_id) or prepared_template
    if not isinstance(updated.get(SEMANTIC_STYLE_SLOTS_KEY), dict):
        updated[SEMANTIC_STYLE_SLOTS_KEY] = normalized_slots
    return updated


def update_template_style(
    kernel_id: str, 
    style_name: str, 
    updates: Dict[str, Any]
) -> Dict[str, Any]:
    """Update a specific style in the template and regenerate the DOCX.
    
    FIX #8: Enhanced with transactional update - creates backup before changes
    and rolls back if any error occurs.
    
    Args:
        kernel_id: The kernel ID.
        style_name: Name of the style to update.
        updates: Dict with style properties to update:
            - font_name: str
            - font_size_pt: float
            - bold: bool
            - italic: bool
            - underline: bool
            - underline_style: SINGLE|DOUBLE|DOTTED|DASH|WAVY
            - color_rgb: str (hex like "FF0000")
            - highlight_color: WD_COLOR_INDEX name
            - strike: bool
            - double_strike: bool
            - all_caps: bool
            - small_caps: bool
            - superscript: bool
            - subscript: bool
            - space_before_pt: float
            - space_after_pt: float
            - alignment: LEFT|CENTER|RIGHT|JUSTIFY
            - line_spacing: float
            - line_spacing_rule: SINGLE|ONE_POINT_FIVE|DOUBLE|AT_LEAST|EXACTLY|MULTIPLE
            - first_line_indent_inches: float
            - left_indent_inches: float
            - right_indent_inches: float
            - keep_with_next: bool
            - keep_together: bool
            - page_break_before: bool
            - widow_control: bool
            - style_id: optional styleId override
            - outline_level: int
            - advanced_props: dict with XML lists for r_pr/p_pr/tbl_pr/tc_pr/tbl_style_pr
            - list_format: bullet|decimal|lowerLetter|upperLetter|lowerRoman|upperRoman
            - list_bullet_char: str
            - list_start: int
            - list_level: int
            - list_alignment: LEFT|CENTER|RIGHT
            - list_left_indent_inches: float
            - list_hanging_indent_inches: float
            - table_border_style: single|none
            - table_border_size_pt: float
            - table_border_color: hex
            - table_shading_color: hex
            - table_alignment: LEFT|CENTER|RIGHT
            - table_cell_margin_*_pt: float
    
    Returns:
        Updated template JSON.
        
    Raises:
        RuntimeError: If python-docx is not installed
        ValueError: If no template exists for the kernel
        Exception: If update fails (after rollback is attempted)
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed")

    updates = _normalize_style_updates(updates)
    current_template = get_template(kernel_id) or {}
    
    json_path = _get_template_json_path(kernel_id)
    docx_path = _get_template_docx_path(kernel_id)

    if not json_path.exists() or not docx_path.exists():
        raise ValueError(f"No template found for kernel {kernel_id}")

    requested_style_id = updates.get("style_id") if isinstance(updates, dict) else None
    current_style_info = _find_template_style_info(current_template, style_name, requested_style_id)
    current_style_id = requested_style_id or (current_style_info or {}).get("style_id")
    current_style_name = (
        (current_style_info or {}).get("display_name")
        or (current_style_info or {}).get("name")
        or style_name
    )
    is_table_style = bool(
        (current_style_info or {}).get("type") == "table"
        or updates.get("style_type") == "table"
        or updates.get("category") == "tables"
        or any(key.startswith("table_") for key in updates.keys())
    )

    runtime_defaults = _normalize_table_style_runtime_defaults(
        current_template.get(TABLE_STYLE_RUNTIME_DEFAULTS_KEY)
    )
    runtime_patch = _extract_runtime_defaults_from_updates(updates) if is_table_style else {}
    if is_table_style and isinstance(updates.get("advanced_props"), dict):
        sanitized_advanced_props, advanced_runtime_patch = _sanitize_advanced_props_for_table_style(
            updates.get("advanced_props")
        )
        updates["advanced_props"] = sanitized_advanced_props
        runtime_patch = _merge_runtime_patch_dicts(runtime_patch, advanced_runtime_patch)

    # FIX #8: Create backup before applying changes
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_docx = docx_path.with_suffix(f".backup_{timestamp}.docx")
    backup_json = json_path.with_suffix(f".backup_{timestamp}.json")
    
    try:
        # Create backups
        shutil.copy2(docx_path, backup_docx)
        if json_path.exists():
            shutil.copy2(json_path, backup_json)
        
        _logger.info(f"[Template] Created backup for kernel {kernel_id}: {backup_docx.name}")
        
        # Apply changes
        applied_runtime_patch = _apply_style_to_docx(docx_path, style_name, updates)
        runtime_patch = _merge_runtime_patch_dicts(runtime_patch, applied_runtime_patch)

        # Re-extract and save JSON
        with open(docx_path, "rb") as f:
            new_docx_bytes = f.read()
        extracted = extract_styles_from_docx(new_docx_bytes)

        updated_style_info = _find_template_style_info(extracted, style_name, current_style_id)
        effective_style_id = current_style_id or (updated_style_info or {}).get("style_id")
        effective_style_name = (
            (updated_style_info or {}).get("display_name")
            or (updated_style_info or {}).get("name")
            or current_style_name
        )

        if is_table_style or (updated_style_info or {}).get("type") == "table":
            runtime_defaults = _apply_runtime_defaults_patch(
                runtime_defaults,
                style_id=effective_style_id,
                style_name=effective_style_name,
                patch=runtime_patch,
            )

        extracted[TABLE_STYLE_RUNTIME_DEFAULTS_KEY] = runtime_defaults

        extracted["style_coverage"] = get_style_coverage(extracted)
        save_template(kernel_id, new_docx_bytes, extracted)

        # Success - cleanup backup
        try:
            backup_docx.unlink()
            if backup_json.exists():
                backup_json.unlink()
        except Exception as cleanup_err:
            _logger.warning(f"[Template] Could not cleanup backup: {cleanup_err}")

        _logger.info(f"[Template] Successfully updated style '{style_name}' for kernel {kernel_id}")
        return extracted
        
    except Exception as exc:
        # FIX #8: Rollback on error
        _logger.error(
            f"[Template] Error updating style '{style_name}' for kernel {kernel_id}: {exc}. "
            "Rolling back to backup."
        )
        
        try:
            # Restore from backup
            if backup_docx.exists():
                shutil.copy2(backup_docx, docx_path)
                _logger.info(f"[Template] Restored DOCX from backup: {backup_docx.name}")
                
            if backup_json.exists():
                shutil.copy2(backup_json, json_path)
                _logger.info(f"[Template] Restored JSON from backup")
            
            # Keep backup files for debugging
            _logger.info(f"[Template] Backup files preserved for debugging: {backup_docx.name}")
            
        except Exception as rollback_err:
            _logger.critical(
                f"[Template] CRITICAL: Rollback failed for kernel {kernel_id}: {rollback_err}. "
                f"Manual intervention may be required. Backup: {backup_docx}"
            )
        
        # Re-raise original exception
        raise


def _apply_style_to_docx(docx_path: Path, style_name: str, updates: Dict[str, Any]) -> Dict[str, Any]:
    """Apply style updates to the actual DOCX file."""
    if not HAS_DOCX:
        return {}

    advanced_props = updates.get("advanced_props") if isinstance(updates, dict) else None
    style_id = updates.get("style_id") if isinstance(updates, dict) else None

    try:
        advanced_runtime_patch: Dict[str, Any] = {}
        doc = Document(str(docx_path))

        # Try to resolve existing style by display name first, then style_id.
        style = None
        lookup_candidates: List[str] = []
        if style_name:
            lookup_candidates.append(str(style_name))
        if style_id and style_id not in lookup_candidates:
            lookup_candidates.append(str(style_id))

        for candidate in lookup_candidates:
            try:
                style = doc.styles[candidate]
                break
            except KeyError:
                continue

        # Style doesn't exist, try to add it.
        if style is None:
            try:
                style_type_name = updates.get("style_type") if isinstance(updates, dict) else None
                if not style_type_name and isinstance(updates, dict):
                    category = updates.get("category")
                    if category == "tables":
                        style_type_name = "table"
                    elif category == "lists":
                        style_type_name = "paragraph"

                style_type_map = {
                    "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                    "character": WD_STYLE_TYPE.CHARACTER,
                    "table": WD_STYLE_TYPE.TABLE,
                }
                style_type = style_type_map.get(str(style_type_name).lower(), WD_STYLE_TYPE.PARAGRAPH)
                style_label = str(style_name or style_id or "CustomStyle")

                style = doc.styles.add_style(style_label, style_type)
                if style_type == WD_STYLE_TYPE.PARAGRAPH:
                    style.base_style = doc.styles["Normal"]
            except Exception as e:
                _logger.warning(f"Could not create style '{style_name or style_id}': {e}")
                return {}
        
        def _to_float(value: Any) -> Optional[float]:
            if value is None or value == "":
                return None
            try:
                return float(value)
            except (TypeError, ValueError):
                return None

        def _to_bool(value: Any) -> Optional[bool]:
            if value is None:
                return None
            return bool(value)

        underline_map = {
            "SINGLE": WD_UNDERLINE.SINGLE,
            "DOUBLE": WD_UNDERLINE.DOUBLE,
            "DOTTED": WD_UNDERLINE.DOTTED,
            "DASH": WD_UNDERLINE.DASH,
            "WAVY": WD_UNDERLINE.WAVY,
        }
        highlight_map = _get_highlight_map()

        # Apply font updates
        if style.font:
            if "font_name" in updates:
                style.font.name = updates["font_name"] or None
            if "font_size_pt" in updates:
                size_pt = _to_float(updates.get("font_size_pt"))
                style.font.size = Pt(size_pt) if size_pt is not None else None
            if "bold" in updates:
                style.font.bold = _to_bool(updates.get("bold"))
            if "italic" in updates:
                style.font.italic = _to_bool(updates.get("italic"))
            if "underline_style" in updates:
                underline_style = updates.get("underline_style")
                if underline_style:
                    style.font.underline = underline_map.get(str(underline_style).upper(), True)
                else:
                    style.font.underline = None
            elif "underline" in updates:
                underline = updates.get("underline")
                style.font.underline = None if underline is None else bool(underline)
            if "color_rgb" in updates:
                color_value = updates.get("color_rgb")
                if color_value:
                    try:
                        hex_color = str(color_value).lstrip("#")
                        r = int(hex_color[0:2], 16)
                        g = int(hex_color[2:4], 16)
                        b = int(hex_color[4:6], 16)
                        style.font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, IndexError):
                        _logger.warning(f"Invalid color: {color_value}")
                else:
                    try:
                        style.font.color.rgb = None
                    except Exception:
                        pass
            if "highlight_color" in updates:
                highlight_value = updates.get("highlight_color")
                if highlight_value:
                    style.font.highlight_color = _resolve_highlight_color(highlight_value)
                else:
                    style.font.highlight_color = None
            if "strike" in updates:
                style.font.strike = _to_bool(updates.get("strike"))
            if "double_strike" in updates:
                style.font.double_strike = _to_bool(updates.get("double_strike"))
            if "all_caps" in updates:
                style.font.all_caps = _to_bool(updates.get("all_caps"))
            if "small_caps" in updates:
                style.font.small_caps = _to_bool(updates.get("small_caps"))
            if "superscript" in updates:
                style.font.superscript = _to_bool(updates.get("superscript"))
            if "subscript" in updates:
                style.font.subscript = _to_bool(updates.get("subscript"))
        
        # Apply paragraph format updates
        if hasattr(style, 'paragraph_format') and style.paragraph_format:
            if "space_before_pt" in updates:
                before_pt = _to_float(updates.get("space_before_pt"))
                style.paragraph_format.space_before = Pt(before_pt) if before_pt is not None else None
            if "space_after_pt" in updates:
                after_pt = _to_float(updates.get("space_after_pt"))
                style.paragraph_format.space_after = Pt(after_pt) if after_pt is not None else None
            if "line_spacing" in updates:
                line_spacing = _to_float(updates.get("line_spacing"))
                style.paragraph_format.line_spacing = line_spacing
            if "line_spacing_rule" in updates:
                rule_value = updates.get("line_spacing_rule")
                rule_map = {
                    "SINGLE": WD_LINE_SPACING.SINGLE,
                    "ONE_POINT_FIVE": WD_LINE_SPACING.ONE_POINT_FIVE,
                    "DOUBLE": WD_LINE_SPACING.DOUBLE,
                    "AT_LEAST": WD_LINE_SPACING.AT_LEAST,
                    "EXACTLY": WD_LINE_SPACING.EXACTLY,
                    "MULTIPLE": WD_LINE_SPACING.MULTIPLE,
                }
                if not rule_value:
                    style.paragraph_format.line_spacing_rule = None
                else:
                    style.paragraph_format.line_spacing_rule = rule_map.get(str(rule_value).upper(), None)
            if "first_line_indent_inches" in updates:
                first_indent = _to_float(updates.get("first_line_indent_inches"))
                style.paragraph_format.first_line_indent = Inches(first_indent) if first_indent is not None else None
            if "left_indent_inches" in updates:
                left_indent = _to_float(updates.get("left_indent_inches"))
                style.paragraph_format.left_indent = Inches(left_indent) if left_indent is not None else None
            if "right_indent_inches" in updates:
                right_indent = _to_float(updates.get("right_indent_inches"))
                style.paragraph_format.right_indent = Inches(right_indent) if right_indent is not None else None
            if "keep_with_next" in updates:
                style.paragraph_format.keep_with_next = _to_bool(updates.get("keep_with_next"))
            if "keep_together" in updates:
                style.paragraph_format.keep_together = _to_bool(updates.get("keep_together"))
            if "page_break_before" in updates:
                style.paragraph_format.page_break_before = _to_bool(updates.get("page_break_before"))
            if "widow_control" in updates:
                style.paragraph_format.widow_control = _to_bool(updates.get("widow_control"))
            if "alignment" in updates:
                align_value = updates.get("alignment")
                align_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if not align_value:
                    style.paragraph_format.alignment = None
                else:
                    align_key = str(align_value).upper().split()[0]
                    style.paragraph_format.alignment = align_map.get(align_key, None)
        
        # Save the modified document
        doc.save(str(docx_path))

        outline_level = updates.get("outline_level") if isinstance(updates, dict) else None
        if outline_level is not None and outline_level != "":
            try:
                outline_value = int(outline_level)
            except (TypeError, ValueError):
                outline_value = None
        else:
            outline_value = None

        _apply_list_style_updates(docx_path, style_name, style_id, updates)
        _apply_table_style_updates(docx_path, style_name, style_id, updates)

        if advanced_props:
            advanced_runtime_patch = _apply_style_xml_updates(docx_path, style_name, style_id, advanced_props)

        if outline_level is not None:
            _apply_outline_level(docx_path, style_name, style_id, outline_value)

        if "font_name" in updates:
            _sync_style_explicit_font_name(
                docx_path,
                style_name,
                style_id,
                updates.get("font_name"),
            )

        _logger.info(f"[Template] Applied updates to DOCX for style '{style_name}'")
        return advanced_runtime_patch
        
    except Exception as e:
        _logger.error(f"[Template] Error applying style to DOCX: {e}")
        raise


def save_template(kernel_id: str, docx_bytes: bytes, extracted_json: Dict[str, Any]) -> bool:
    """Save the template DOCX and extracted JSON for a kernel.
    
    FIX #16: Added validation for file write success.
    FIX #17: Cleanup old backups after save.
    
    Args:
        kernel_id: The kernel ID to associate the template with.
        docx_bytes: The raw bytes of the DOCX file.
        extracted_json: The extracted template information.
        
    Returns:
        True if successful.
        
    Raises:
        ValueError: If docx_bytes is empty
        IOError: If file write fails
    """
    # FIX #16: Validate input
    if not docx_bytes:
        raise ValueError("Cannot save empty DOCX file")
    
    try:
        prepared_docx, prepared_json, docx_changed = _prepare_template_payload(docx_bytes, extracted_json)
        docx_path = _get_template_docx_path(kernel_id)

        if docx_changed and docx_path.exists():
            _create_template_backups(kernel_id, reason="table style runtime defaults sanitization")

        _write_template_files(kernel_id, prepared_docx, prepared_json)
        
        # FIX #17: Cleanup old backups (keep only 5 most recent)
        _cleanup_old_backups(kernel_id, max_backups=5)

        # Invalidate cached style previews for this kernel; styles/template changed.
        clear_preview_cache(kernel_id)
        
        _logger.info(f"[Template] Saved template for kernel {kernel_id}")
        return True
    except Exception as e:
        _logger.error(f"[Template] Error saving template: {e}")
        raise


def get_template(kernel_id: str) -> Optional[Dict[str, Any]]:
    """Get the extracted template JSON for a kernel.
    
    FIX #16: Enhanced validation and error handling.
    
    Args:
        kernel_id: The kernel ID.
        
    Returns:
        The template JSON, or None if no template exists.
    """
    sanitized_template = _sanitize_persisted_template_if_needed(kernel_id)
    if isinstance(sanitized_template, dict):
        return _merge_runtime_defaults_into_extracted(sanitized_template)

    json_path = _get_template_json_path(kernel_id)
    
    # FIX #16: Check existence
    if not json_path.exists():
        return None
    
    # FIX #16: Check if file is readable and not empty
    try:
        file_size = json_path.stat().st_size
        if file_size == 0:
            _logger.warning(f"[Template] JSON file is empty for kernel {kernel_id}")
            return None
    except OSError as e:
        _logger.error(f"[Template] Cannot access JSON file: {e}")
        return None
    
    # Load JSON
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            # FIX #16: Validate it's a dict
            if not isinstance(data, dict):
                _logger.error(f"[Template] Invalid JSON structure for kernel {kernel_id}")
                return None
            return _merge_runtime_defaults_into_extracted(data)
    except json.JSONDecodeError as e:
        _logger.error(f"[Template] JSON decode error for kernel {kernel_id}: {e}")
        return None
    except Exception as e:
        _logger.error(f"[Template] Error loading template JSON: {e}")
        return None


def get_template_docx_path(kernel_id: str) -> Optional[str]:
    """Get the path to the template DOCX file for a kernel.
    
    This path is used by DocxSession to initialize documents with the template.
    
    Args:
        kernel_id: The kernel ID.
        
    Returns:
        The absolute path to the template DOCX, or None if no template exists.
    """
    _sanitize_persisted_template_if_needed(kernel_id)
    docx_path = _get_template_docx_path(kernel_id)
    if docx_path.exists():
        return str(docx_path.absolute())
    return None


def delete_template(kernel_id: str) -> bool:
    """Delete the template for a kernel.
    
    Args:
        kernel_id: The kernel ID.
        
    Returns:
        True if deleted, False if no template existed.
    """
    safe_id = _sanitize_kernel_id(kernel_id)
    template_dir = TEMPLATE_DIR / safe_id
    if template_dir.exists():
        try:
            shutil.rmtree(template_dir)
            clear_preview_cache(kernel_id)
            _logger.info(f"[Template] Deleted template for kernel {kernel_id}")
            return True
        except Exception as e:
            _logger.error(f"[Template] Error deleting template: {e}")
            raise
    return False


def list_templates() -> List[str]:
    """List all kernel IDs that have templates.
    
    Returns:
        List of kernel IDs.
    """
    if not TEMPLATE_DIR.exists():
        return []
    
    return [
        d.name for d in TEMPLATE_DIR.iterdir()
        if d.is_dir() and (d / "template.docx").exists()
    ]


# =============================================================================
# STYLE PREVIEW (Rendered with Word)
# =============================================================================

# Lazy import for PyMuPDF (optional dependency for preview)
_fitz = None


def get_preview_cache(preview_key: str | None, kernel_id: str | None = None) -> Optional[str]:
    from app.services.template import preview as template_preview

    return template_preview.get_preview_cache(preview_key, kernel_id)


def set_preview_cache(preview_key: str, preview_b64: str, kernel_id: str | None = None) -> None:
    from app.services.template import preview as template_preview

    template_preview.set_preview_cache(preview_key, preview_b64, kernel_id)


def clear_preview_cache(kernel_id: str | None = None) -> None:
    from app.services.template import preview as template_preview

    template_preview.clear_preview_cache(kernel_id)

def _get_fitz():
    """Lazy-load PyMuPDF (fitz) module."""
    global _fitz
    if _fitz is None:
        try:
            import fitz
            _fitz = fitz
        except ImportError:
            _logger.warning("[Template] PyMuPDF not installed, style preview unavailable")
            return None
    return _fitz


def _convert_pdf_to_png(pdf_bytes: bytes, dpi: int = 150) -> Optional[bytes]:
    """Convert first page of PDF to PNG image.
    
    Args:
        pdf_bytes: Raw PDF bytes
        dpi: Resolution for rendering (default 150 for fast preview)
    
    Returns:
        PNG image bytes, or None if conversion fails
    """
    fitz = _get_fitz()
    if fitz is None:
        return None
    
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
        
        page = doc[0]
        
        # Scale for desired DPI (72 is PDF default)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        # Render to pixmap
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PNG bytes
        png_bytes = pix.tobytes("png")
        
        doc.close()
        return png_bytes
    except Exception as e:
        _logger.error(f"[Template] Error converting PDF to PNG: {e}")
        return None


def _hex_to_fitz_color(hex_color: Optional[str], default: tuple[float, float, float]) -> tuple[float, float, float]:
    """Convert a hex color (RRGGBB) to a PyMuPDF RGB tuple (0..1)."""
    if not hex_color:
        return default
    value = str(hex_color).strip().replace("#", "")
    if len(value) < 6:
        return default
    try:
        r = int(value[0:2], 16) / 255.0
        g = int(value[2:4], 16) / 255.0
        b = int(value[4:6], 16) / 255.0
        return (r, g, b)
    except ValueError:
        return default


def _extract_table_preview_rows(raw_xml: Optional[str], max_rows: int, default_cols: int) -> List[List[Dict[str, Any]]]:
    """Extract lightweight preview rows/cells from a table raw XML fragment."""
    if not raw_xml:
        return []
    try:
        root = ET.fromstring(raw_xml)
    except Exception:
        return []

    rows: List[List[Dict[str, Any]]] = []
    tr_nodes = root.findall("w:tr", DOCX_NS)
    for tr in tr_nodes[:max(1, max_rows)]:
        row_cells: List[Dict[str, Any]] = []
        tc_nodes = tr.findall("w:tc", DOCX_NS)
        for tc in tc_nodes:
            text = "".join((t.text or "") for t in tc.findall(".//w:t", DOCX_NS)).strip()
            tc_pr = tc.find("w:tcPr", DOCX_NS)
            shading = None
            span = 1
            if tc_pr is not None:
                shd = tc_pr.find("w:shd", DOCX_NS)
                if shd is not None:
                    fill = shd.get(_qn("w", "fill"))
                    if fill and str(fill).lower() not in {"auto", "none"}:
                        shading = str(fill).upper()
                grid_span = tc_pr.find("w:gridSpan", DOCX_NS)
                if grid_span is not None:
                    try:
                        span = max(1, int(grid_span.get(_qn("w", "val")) or "1"))
                    except (TypeError, ValueError):
                        span = 1

            row_cells.append({"text": text, "shading_color": shading})
            for _ in range(span - 1):
                row_cells.append({"text": "", "shading_color": shading})

        if not row_cells and default_cols > 0:
            row_cells = [{"text": "", "shading_color": None} for _ in range(default_cols)]
        rows.append(row_cells)

    return rows


def _render_table_preview_with_fitz(table_info: Dict[str, Any], max_rows: int = 4) -> Optional[str]:
    """Render a table preview directly with PyMuPDF (no DOCX->PDF conversion)."""
    fitz = _get_fitz()
    if fitz is None:
        return None

    try:
        table_index = int(table_info.get("index") or 0)
        rows_total = int(table_info.get("rows") or 0)
        cols_total = int(table_info.get("cols") or 0)
        rows_preview = max(1, min(max_rows, rows_total if rows_total > 0 else max_rows))
        cols_total = max(1, cols_total)

        rows = _extract_table_preview_rows(table_info.get("raw_xml"), rows_preview, cols_total)
        if not rows:
            rows = [[{"text": f"Tabla {table_index + 1}", "shading_color": None}] for _ in range(rows_preview)]

        cols_preview = max(cols_total, max((len(r) for r in rows), default=1))
        for row in rows:
            if len(row) < cols_preview:
                row.extend([{"text": "", "shading_color": None}] * (cols_preview - len(row)))

        table_props = table_info.get("parsed_properties") or {}
        borders = table_props.get("borders") or table_info.get("borders") or {}
        border_sample = next((v for v in borders.values() if isinstance(v, dict)), {}) or {}
        border_pt = border_sample.get("size_pt")
        try:
            border_width = float(border_pt) if border_pt not in (None, "") else 0.8
        except (TypeError, ValueError):
            border_width = 0.8
        border_width = max(0.6, min(2.0, border_width))
        border_color = _hex_to_fitz_color(border_sample.get("color"), (0.22, 0.24, 0.28))

        first_row = table_info.get("first_row_format") or {}
        header_fill = first_row.get("shading_fill") or table_props.get("shading_color") or table_info.get("shading_fill")
        body_fill = table_props.get("shading_color") or table_info.get("shading_fill")
        header_fill_color = _hex_to_fitz_color(header_fill, (0.20, 0.40, 0.62))
        body_fill_color = _hex_to_fitz_color(body_fill, (0.96, 0.97, 0.98))

        header_font = first_row.get("font_properties") or {}
        header_text_color = _hex_to_fitz_color(header_font.get("color_rgb"), (1.0, 1.0, 1.0))
        body_text_color = (0.12, 0.12, 0.14)

        margin_x = 18.0
        margin_y = 18.0
        title_h = 18.0
        cell_h = 34.0
        if cols_preview <= 4:
            cell_w = 140.0
        elif cols_preview <= 7:
            cell_w = 108.0
        else:
            cell_w = 86.0

        page_w = max(340.0, min(1180.0, margin_x * 2 + cols_preview * cell_w))
        page_h = max(180.0, min(1600.0, margin_y * 2 + title_h + rows_preview * cell_h))

        doc = fitz.open()
        page = doc.new_page(width=page_w, height=page_h)
        usable_w = page_w - margin_x * 2
        if cols_preview > 0:
            cell_w = usable_w / cols_preview

        title = f"Tabla {table_index + 1}  |  {rows_total} x {cols_total}"
        page.insert_text((margin_x, margin_y + 11), title, fontsize=9, color=(0.55, 0.57, 0.60))

        top_y = margin_y + title_h
        for r_idx in range(rows_preview):
            row = rows[r_idx] if r_idx < len(rows) else [{"text": "", "shading_color": None}] * cols_preview
            for c_idx in range(cols_preview):
                cell = row[c_idx] if c_idx < len(row) else {"text": "", "shading_color": None}
                x0 = margin_x + c_idx * cell_w
                y0 = top_y + r_idx * cell_h
                x1 = x0 + cell_w
                y1 = y0 + cell_h
                rect = fitz.Rect(x0, y0, x1, y1)

                raw_cell_fill = cell.get("shading_color")
                if raw_cell_fill:
                    fill_color = _hex_to_fitz_color(raw_cell_fill, header_fill_color if r_idx == 0 else body_fill_color)
                else:
                    fill_color = header_fill_color if r_idx == 0 else body_fill_color

                page.draw_rect(rect, color=border_color, fill=fill_color, width=border_width)

                text = str(cell.get("text") or "")
                if not text:
                    text = f"R{r_idx + 1}C{c_idx + 1}"
                text_rect = fitz.Rect(x0 + 4, y0 + 4, x1 - 4, y1 - 4)
                page.insert_textbox(
                    text_rect,
                    text,
                    fontsize=9 if r_idx == 0 else 8.5,
                    fontname="helv",
                    color=header_text_color if r_idx == 0 else body_text_color,
                    align=0,
                )

        pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
        png_bytes = pix.tobytes("png")
        doc.close()
        return base64.b64encode(png_bytes).decode("utf-8")
    except Exception as exc:
        _logger.warning(f"[Template] fitz fallback preview failed: {exc}")
        return None


def _build_preview_updates(style_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not style_props or not isinstance(style_props, dict):
        return {}
    explicit_keys = style_props.get("changed_keys") if "changed_keys" in style_props else style_props.get("explicit_keys")
    explicit_keys_provided = "changed_keys" in style_props or "explicit_keys" in style_props
    meta_keys = {"style_type", "category", "style_id"}
    updates: Dict[str, Any] = {}
    for key, value in style_props.items():
        if key in {"changed_keys", "explicit_keys"}:
            continue
        if key in meta_keys:
            if value is not None and value != "":
                updates[key] = value
            continue
        if explicit_keys_provided:
            if key in (explicit_keys or []):
                updates[key] = value
            continue
        if value is None or value == "":
            continue
        updates[key] = value
    return updates


_PREVIEW_META_KEYS = {"style_type", "category", "style_id", "is_table_style"}
_TABLE_LOOK_DEFAULTS = {
    "firstRow": True,
    "lastRow": False,
    "firstColumn": False,
    "lastColumn": False,
    "noHBand": False,
    "noVBand": True,
}
_TABLE_LOOK_PROP_MAP = {
    "firstRow": "look_first_row",
    "lastRow": "look_last_row",
    "firstColumn": "look_first_column",
    "lastColumn": "look_last_column",
    "noHBand": "look_no_h_band",
    "noVBand": "look_no_v_band",
}


def _has_effective_preview_updates(updates: Optional[Dict[str, Any]]) -> bool:
    if not updates or not isinstance(updates, dict):
        return False
    return any(key not in _PREVIEW_META_KEYS for key in updates.keys())


def _coerce_boolish(value: Any) -> Optional[bool]:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return None


def _extract_table_look_from_style_props(style_props: Optional[Dict[str, Any]]) -> Dict[str, bool]:
    if not isinstance(style_props, dict):
        return {}

    look: Dict[str, bool] = {}
    nested_look = style_props.get("look")
    if isinstance(nested_look, dict):
        for key in _TABLE_LOOK_DEFAULTS.keys():
            parsed = _coerce_boolish(nested_look.get(key))
            if parsed is not None:
                look[key] = parsed

    for look_key, prop_key in _TABLE_LOOK_PROP_MAP.items():
        parsed = _coerce_boolish(style_props.get(prop_key))
        if parsed is None:
            parsed = _coerce_boolish(style_props.get(f"table_{prop_key}"))
        if parsed is not None:
            look[look_key] = parsed

    return look


def _extract_table_look_from_style(style_obj: Any) -> Dict[str, bool]:
    style_elem = getattr(style_obj, "_element", None)
    if style_elem is None:
        return {}

    tbl_look = style_elem.find("w:tblPr/w:tblLook", DOCX_NS)
    if tbl_look is None:
        return {}

    look: Dict[str, bool] = {}
    for key in _TABLE_LOOK_DEFAULTS.keys():
        parsed = _coerce_boolish(tbl_look.get(_qn("w", key)))
        if parsed is not None:
            look[key] = parsed
    return look


def _resolve_preview_table_look(style_props: Optional[Dict[str, Any]], style_obj: Any) -> Dict[str, bool]:
    resolved = dict(_TABLE_LOOK_DEFAULTS)
    resolved.update(_extract_table_look_from_style(style_obj))
    resolved.update(_extract_table_look_from_style_props(style_props))
    return resolved


def _extract_preview_table_runtime_defaults(style_props: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(style_props, dict):
        return {}

    runtime_defaults: Dict[str, Any] = {}
    layout_value = style_props.get("table_layout_type", style_props.get("layout_type"))
    if layout_value not in (None, ""):
        runtime_defaults["layout_type"] = str(layout_value).strip().lower()

    width_type = style_props.get("table_width_type", style_props.get("width_type"))
    if width_type not in (None, ""):
        runtime_defaults["width_type"] = str(width_type).strip().lower()

    width_value = style_props.get("table_width_value", style_props.get("width_value"))
    coerced_width = _coerce_optional_int(width_value)
    if coerced_width is not None:
        runtime_defaults["width_value"] = coerced_width

    look = _extract_table_look_from_style_props(style_props)
    if look:
        runtime_defaults["look"] = look

    return runtime_defaults


def _apply_table_runtime_defaults_to_preview_table(table: Any, runtime_defaults: Optional[Dict[str, Any]]) -> None:
    if not runtime_defaults or not isinstance(runtime_defaults, dict):
        return

    try:
        tbl = table._tbl
        tbl_pr = tbl.find(_qn("w", "tblPr"))
        if tbl_pr is None:
            tbl_pr = ET.SubElement(tbl, _qn("w", "tblPr"))
            tbl.insert(0, tbl_pr)

        for tag in _TABLE_STYLE_RUNTIME_TAGS:
            existing = tbl_pr.find(_qn("w", tag))
            if existing is not None:
                tbl_pr.remove(existing)

        look = runtime_defaults.get("look")
        if isinstance(look, dict) and look:
            tbl_look = ET.SubElement(tbl_pr, _qn("w", "tblLook"))
            for key, default_value in _TABLE_LOOK_DEFAULTS.items():
                bool_value = _coerce_boolish(look.get(key))
                resolved_bool = default_value if bool_value is None else bool_value
                tbl_look.set(_qn("w", key), "1" if resolved_bool else "0")

        layout_type = runtime_defaults.get("layout_type")
        if layout_type not in (None, ""):
            tbl_layout = ET.SubElement(tbl_pr, _qn("w", "tblLayout"))
            tbl_layout.set(_qn("w", "type"), str(layout_type).strip().lower())

        width_type = runtime_defaults.get("width_type")
        width_value = runtime_defaults.get("width_value")
        if width_type not in (None, "") or width_value not in (None, ""):
            tbl_w = ET.SubElement(tbl_pr, _qn("w", "tblW"))
            resolved_width_type = str(width_type or "auto").strip().lower()
            tbl_w.set(_qn("w", "type"), resolved_width_type)
            if resolved_width_type == "auto":
                tbl_w.set(_qn("w", "w"), "0")
            else:
                tbl_w.set(_qn("w", "w"), str(_coerce_optional_int(width_value) or 0))
    except Exception as exc:
        _logger.warning("[Template] Could not apply preview table runtime defaults: %s", exc)


def _clear_document_body(target_doc: Document) -> None:
    body = target_doc.element.body
    for child in list(body):
        if _local_name(child.tag) == "sectPr":
            continue
        body.remove(child)


def _clear_header_footer(target_doc: Document) -> None:
    for section in target_doc.sections:
        try:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
        except Exception:
            pass
        for container in (section.header, section.footer):
            try:
                element = container._element
                for child in list(element):
                    element.remove(child)
            except Exception:
                continue


def _apply_compact_preview_page_setup(target_doc: Document) -> None:
    for section in target_doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(1.6)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.1)


def _prepare_preview_docx(
    target_path: Path,
    *,
    template_path: Optional[str] = None,
    clear_body: bool = False,
    clear_header_footer: bool = False,
    compact_page_setup: bool = False,
) -> Document:
    if template_path and os.path.exists(template_path):
        shutil.copy2(template_path, target_path)
        doc = Document(str(target_path))
    else:
        doc = Document()

    if clear_body:
        _clear_document_body(doc)
    if clear_header_footer:
        _clear_header_footer(doc)
    if compact_page_setup:
        _apply_compact_preview_page_setup(doc)

    doc.save(str(target_path))
    return doc


def generate_style_preview(
    kernel_id: str,
    style_name: str,
    style_props: Dict[str, Any],
    sample_text: str = "El veloz murciélago hindú comía feliz cardillo y kiwi"
) -> Optional[str]:
    """Generate a PNG preview of a style rendered by Microsoft Word.
    
    This function:
    1. Creates a minimal DOCX document with sample text using the specified style
    2. Converts it to PDF using Microsoft Word (COM automation)
    3. Extracts the first page as a PNG image
    4. Returns the PNG as a base64-encoded string
    
    Args:
        kernel_id: Kernel ID used to locate template.docx
        style_name: Name of the style (e.g., "Heading 1", "Normal")
        style_props: Style properties to apply:
            - font_name: str (e.g., "Arial")
            - font_size_pt: float (e.g., 14)
            - bold: bool
            - italic: bool
            - underline: bool
            - color_rgb: str (hex like "FF0000" or "000000")
        sample_text: Text to display in preview (default: pangram)
    
    Returns:
        Base64-encoded PNG image string, or None if preview generation fails
    
    Example:
        >>> preview = generate_style_preview("kernel-id", "Heading 1", {
        ...     "font_name": "Arial",
        ...     "font_size_pt": 16,
        ...     "bold": True
        ... })
        >>> if preview:
        ...     # Use as: data:image/png;base64,{preview}
    """
    if not HAS_DOCX:
        _logger.error("[Template] python-docx not installed")
        return None
    
    # Import pdf_converter for Word conversion
    try:
        from app.services.pdf_converter import _convert_to_pdf_word, MS_WORD_AVAILABLE
    except ImportError:
        _logger.error("[Template] pdf_converter not available")
        return None
    
    if not MS_WORD_AVAILABLE:
        _logger.warning("[Template] Microsoft Word not available for preview; trying fallback")
    
    import tempfile
    
    # Create temporary directory manually to handle cleanup errors gracefully
    temp_dir = tempfile.mkdtemp()
    
    try:
        temp_path = Path(temp_dir)
        docx_path = temp_path / "preview.docx"
        
        # -------------------------------------------------------------
        # Step 1: Prepare DOCX (template-based if available)
        # -------------------------------------------------------------
        template_path = get_template_docx_path(kernel_id) if kernel_id else None
        _prepare_preview_docx(
            docx_path,
            template_path=template_path,
            clear_body=True,
            clear_header_footer=False,
            compact_page_setup=True,
        )

        preview_updates = _build_preview_updates(style_props)
        if _has_effective_preview_updates(preview_updates):
            _apply_style_to_docx(docx_path, style_name, preview_updates)

        doc = Document(str(docx_path))

        def _ensure_preview_style(target_style_name: str, style_type: WD_STYLE_TYPE) -> None:
            try:
                _ = doc.styles[target_style_name]
            except KeyError:
                try:
                    doc.styles.add_style(target_style_name, style_type)
                except Exception:
                    pass

        def _apply_paragraph_format(target_para) -> None:
            alignment_value = style_props.get("alignment")
            if alignment_value:
                align_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                target_para.alignment = align_map.get(str(alignment_value).upper().split()[0], None)

            space_before = style_props.get("space_before_pt")
            if space_before is not None:
                try:
                    target_para.paragraph_format.space_before = Pt(float(space_before))
                except (TypeError, ValueError):
                    pass

            space_after = style_props.get("space_after_pt")
            if space_after is not None:
                try:
                    target_para.paragraph_format.space_after = Pt(float(space_after))
                except (TypeError, ValueError):
                    pass

            line_spacing = style_props.get("line_spacing")
            if line_spacing is not None:
                try:
                    target_para.paragraph_format.line_spacing = float(line_spacing)
                except (TypeError, ValueError):
                    pass

            line_spacing_rule = style_props.get("line_spacing_rule")
            if line_spacing_rule:
                rule_map = {
                    "SINGLE": WD_LINE_SPACING.SINGLE,
                    "ONE_POINT_FIVE": WD_LINE_SPACING.ONE_POINT_FIVE,
                    "DOUBLE": WD_LINE_SPACING.DOUBLE,
                    "AT_LEAST": WD_LINE_SPACING.AT_LEAST,
                    "EXACTLY": WD_LINE_SPACING.EXACTLY,
                    "MULTIPLE": WD_LINE_SPACING.MULTIPLE,
                }
                target_para.paragraph_format.line_spacing_rule = rule_map.get(str(line_spacing_rule).upper(), None)

            first_indent = style_props.get("first_line_indent_inches")
            if first_indent is not None:
                try:
                    target_para.paragraph_format.first_line_indent = Inches(float(first_indent))
                except (TypeError, ValueError):
                    pass

            left_indent = style_props.get("left_indent_inches")
            if left_indent is not None:
                try:
                    target_para.paragraph_format.left_indent = Inches(float(left_indent))
                except (TypeError, ValueError):
                    pass

            right_indent = style_props.get("right_indent_inches")
            if right_indent is not None:
                try:
                    target_para.paragraph_format.right_indent = Inches(float(right_indent))
                except (TypeError, ValueError):
                    pass

            if style_props.get("keep_with_next") is not None:
                target_para.paragraph_format.keep_with_next = bool(style_props.get("keep_with_next"))
            if style_props.get("keep_together") is not None:
                target_para.paragraph_format.keep_together = bool(style_props.get("keep_together"))
            if style_props.get("page_break_before") is not None:
                target_para.paragraph_format.page_break_before = bool(style_props.get("page_break_before"))
            if style_props.get("widow_control") is not None:
                target_para.paragraph_format.widow_control = bool(style_props.get("widow_control"))

        def _apply_run_format(target_run) -> None:
            font_name = style_props.get("font_name")
            if font_name:
                target_run.font.name = font_name

            font_size = style_props.get("font_size_pt")
            if font_size:
                target_run.font.size = Pt(float(font_size))

            if style_props.get("bold"):
                target_run.font.bold = True

            if style_props.get("italic"):
                target_run.font.italic = True

            underline_style = style_props.get("underline_style")
            if underline_style:
                underline_map = {
                    "SINGLE": WD_UNDERLINE.SINGLE,
                    "DOUBLE": WD_UNDERLINE.DOUBLE,
                    "DOTTED": WD_UNDERLINE.DOTTED,
                    "DASH": WD_UNDERLINE.DASH,
                    "WAVY": WD_UNDERLINE.WAVY,
                }
                target_run.font.underline = underline_map.get(str(underline_style).upper(), True)
            elif style_props.get("underline"):
                target_run.font.underline = True

            color_rgb = style_props.get("color_rgb")
            if color_rgb:
                try:
                    hex_color = color_rgb.lstrip("#")
                    if len(hex_color) >= 6:
                        r = int(hex_color[0:2], 16)
                        g = int(hex_color[2:4], 16)
                        b = int(hex_color[4:6], 16)
                        target_run.font.color.rgb = RGBColor(r, g, b)
                except (ValueError, IndexError):
                    pass

            highlight_color = style_props.get("highlight_color")
            if highlight_color:
                target_run.font.highlight_color = _resolve_highlight_color(highlight_color)

            if style_props.get("strike"):
                target_run.font.strike = True
            if style_props.get("double_strike"):
                target_run.font.double_strike = True
            if style_props.get("all_caps"):
                target_run.font.all_caps = True
            if style_props.get("small_caps"):
                target_run.font.small_caps = True
            if style_props.get("superscript"):
                target_run.font.superscript = True
            if style_props.get("subscript"):
                target_run.font.subscript = True

        style_type = style_props.get("style_type")
        category = style_props.get("category")
        is_table = style_type == "table" or any(key.startswith("table_") and style_props.get(key) is not None for key in style_props)
        is_list = category == "lists" or any(key.startswith("list_") and style_props.get(key) is not None for key in style_props)
        is_caption = (
            category == "captions"
            or str(style_name or "").strip() == "Caption"
            or str(style_props.get("style_id") or "").strip() == "Caption"
        )

        if is_table:
            _ensure_preview_style(style_name, WD_STYLE_TYPE.TABLE)
            # Use 4 rows to show header and banding effects
            table = doc.add_table(rows=4, cols=2)
            
            # Try to assign the table style - use style_id as fallback
            style_id = style_props.get("style_id")
            style_assigned = False
            applied_table_style = None
            
            # First try by style name (display name)
            try:
                table.style = style_name
                style_assigned = True
                applied_table_style = table.style
                _logger.debug(f"[Template] Table style '{style_name}' applied successfully")
            except Exception as e:
                _logger.debug(f"[Template] Could not apply table style by name '{style_name}': {e}")
            
            # Fallback to style_id if name failed
            if not style_assigned and style_id:
                try:
                    table.style = style_id
                    style_assigned = True
                    applied_table_style = table.style
                    _logger.debug(f"[Template] Table style '{style_id}' applied via style_id")
                except Exception as e:
                    _logger.debug(f"[Template] Could not apply table style by id '{style_id}': {e}")
            
            # Last resort: Table Grid
            if not style_assigned:
                try:
                    table.style = "Table Grid"
                    applied_table_style = table.style
                    _logger.debug("[Template] Fell back to 'Table Grid' style")
                except Exception:
                    _logger.warning("[Template] Could not apply any table style")

            if applied_table_style is None:
                for candidate in (style_id, style_name):
                    if not candidate:
                        continue
                    try:
                        applied_table_style = doc.styles[str(candidate)]
                        break
                    except KeyError:
                        continue

            resolved_look = _resolve_preview_table_look(style_props, applied_table_style)
            preview_runtime_defaults = _extract_preview_table_runtime_defaults(style_props)
            preview_runtime_defaults["look"] = resolved_look
            _apply_table_runtime_defaults_to_preview_table(table, preview_runtime_defaults)

            # Fill table with sample content: first row as header
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    para = cell.paragraphs[0]
                    if row_idx == 0:
                        # Header row
                        text = "COLUMNA 1" if col_idx == 0 else "COLUMNA 2"
                    else:
                        text = f"Fila {row_idx}.{col_idx + 1}"
                    run = para.add_run(text)
                    if not table.style or table.style.name != style_name:
                        _apply_paragraph_format(para)
                        _apply_run_format(run)
        elif is_caption:
            _ensure_preview_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            caption_sample = "Texto de ejemplo"
            caption_prefix = str(style_props.get("caption_label") or "Figura").strip() or "Figura"
            try:
                doc.add_paragraph(f"{caption_prefix} 1. {caption_sample}", style=style_name)
            except Exception:
                para = doc.add_paragraph()
                _apply_paragraph_format(para)
                prefix_run = para.add_run(f"{caption_prefix} 1. ")
                _apply_run_format(prefix_run)
                body_run = para.add_run(caption_sample)
                _apply_run_format(body_run)
        elif is_list:
            _ensure_preview_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            for _ in range(3):
                try:
                    doc.add_paragraph(sample_text, style=style_name)
                except Exception:
                    para = doc.add_paragraph(sample_text)
                    _apply_paragraph_format(para)
                    run = para.runs[0] if para.runs else para.add_run(sample_text)
                    _apply_run_format(run)
        else:
            try:
                doc.add_paragraph(sample_text, style=style_name)
            except Exception:
                para = doc.add_paragraph(sample_text)
                _apply_paragraph_format(para)
                run = para.runs[0] if para.runs else para.add_run(sample_text)
                _apply_run_format(run)

        doc.save(str(docx_path))

        if is_list:
            _apply_list_style_updates(docx_path, style_name, None, preview_updates)
        if is_table:
            _apply_table_style_updates(docx_path, style_name, None, preview_updates)
        
        # -------------------------------------------------------------
        # Step 2: Convert DOCX to PDF (using generic converter with fallback)
        # -------------------------------------------------------------
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
            
        docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
        
        # Import generic converter to enable LibreOffice fallback
        from app.services.pdf_converter import convert_docx_with_diagnostics
        
        # Attempt conversion with 15s timeout
        result = convert_docx_with_diagnostics(docx_b64, timeout_s=15)
        
        pdf_b64 = result.get("pdf_b64")
        if not pdf_b64:
            _logger.warning(f"[Template] Preview conversion failed: {result.get('error')} (used: {result.get('converter_used')})")
            return None
            
        # -------------------------------------------------------------
        # Step 3: Convert PDF to PNG
        # -------------------------------------------------------------
        pdf_bytes = base64.b64decode(pdf_b64)
        png_bytes = _convert_pdf_to_png(pdf_bytes, dpi=150)
        
        if not png_bytes:
            _logger.warning("[Template] PNG conversion failed")
            return None
        
        # -------------------------------------------------------------
        # Step 4: Encode as base64
        # -------------------------------------------------------------
        png_b64 = base64.b64encode(png_bytes).decode("utf-8")
        
        converter_label = result.get('converter_used', 'unknown')
        _logger.info(f"[Template] Generated preview with {converter_label} for '{style_name}'")
        return png_b64

    except Exception as e:
        _logger.error(f"[Template] Error generating style preview: {e}")
        return None
    
    finally:
        # Cleanup temp directory, ignoring errors (e.g. if Word still has file open)
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


def generate_document_table_preview(
    kernel_id: str,
    table_index: int,
    max_rows: int = 4
) -> Dict[str, Any]:
    """Generate a PNG preview of a specific table from the document.
    
    This function:
    1. Copies the template DOCX
    2. Modifies it to keep only the specified table (truncated to max_rows)
    3. Converts to PDF using Word
    4. Returns the PNG as base64
    
    Args:
        kernel_id: Kernel ID to locate template.docx
        table_index: Index of the table to preview (0-based)
        max_rows: Maximum rows to include in preview (default 4)
        
    Returns:
        Dictionary with:
        - success: bool
        - preview_b64: str (if success)
        - error: str (if not success)
        - error_detail: str (if available)
    """
    if not HAS_DOCX:
        error_msg = "python-docx library not installed on server"
        _logger.error(f"[Template] {error_msg}")
        return {"success": False, "error": error_msg}

    normalized_table_index = _validate_table_index_value(table_index)
    if normalized_table_index is None:
        error_msg = _table_index_error_message(table_index)
        _logger.warning(f"[Template] {error_msg}")
        return {"success": False, "error": error_msg}
    table_index = normalized_table_index
    
    try:
        from app.services.pdf_converter import (
            MS_WORD_AVAILABLE,
            PDF_CONVERT_AVAILABLE,
            convert_docx_with_diagnostics,
        )
    except ImportError:
        error_msg = "PDF converter module not available"
        _logger.error(f"[Template] {error_msg}")
        return {"success": False, "error": error_msg}

    import tempfile
    
    template_path = get_template_docx_path(kernel_id)
    if not template_path or not os.path.exists(template_path):
        error_msg = f"No template document found for kernel {kernel_id}"
        _logger.warning(f"[Template] {error_msg}")
        return {"success": False, "error": error_msg}

    preview_pdf_timeout_s = max(5, int(os.getenv("INSPYRO_TEMPLATE_PREVIEW_PDF_TIMEOUT", "10")))

    with open(template_path, "rb") as f:
        source_docx_bytes = f.read()

    detected_tables = _extract_document_tables(source_docx_bytes)
    if table_index >= len(detected_tables):
        error_msg = _table_index_error_message(table_index, len(detected_tables))
        _logger.warning(f"[Template] {error_msg}")
        return {"success": False, "error": error_msg}
    target_table_info = detected_tables[table_index]

    preview_mode = str(os.getenv("INSPYRO_TABLE_PREVIEW_MODE", "fitz_first")).strip().lower()
    if preview_mode in {"fitz_first", "fitz_only"}:
        fallback_preview = _render_table_preview_with_fitz(target_table_info, max_rows=max_rows)
        if fallback_preview:
            return {
                "success": True,
                "preview_b64": fallback_preview,
                "converter_used": "fitz_first",
            }
        if preview_mode == "fitz_only":
            return {"success": False, "error": "fitz preview renderer unavailable"}

    if not MS_WORD_AVAILABLE and not PDF_CONVERT_AVAILABLE:
        # No converter is available; fall back to direct vector rendering.
        fallback_preview = _render_table_preview_with_fitz(target_table_info, max_rows=max_rows)
        if fallback_preview:
            return {
                "success": True,
                "preview_b64": fallback_preview,
                "converter_used": "fitz_fallback_no_converter",
            }
        return {"success": False, "error": "No PDF converter available and fallback renderer failed"}
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        temp_path = Path(temp_dir)
        docx_path = temp_path / "table_preview.docx"
        
        # Copy template
        shutil.copy2(template_path, docx_path)
        
        # Extract and keep only the target table
        with zipfile.ZipFile(docx_path, "r") as zin:
            doc_xml = zin.read("word/document.xml")
            namespace_hints = _collect_docx_namespace_hints_from_zip(zin)
        
        doc_root = ET.fromstring(doc_xml)
        body = doc_root.find("w:body", DOCX_NS)
        
        if body is None:
            error_msg = "Document body not found in template"
            _logger.warning(f"[Template] {error_msg}")
            return {"success": False, "error": error_msg}
        
        # Find all tables
        tables = body.findall(".//w:tbl", DOCX_NS)
        if table_index >= len(tables):
            error_msg = _table_index_error_message(table_index, len(tables))
            _logger.warning(f"[Template] {error_msg}")
            return {"success": False, "error": error_msg}
        
        target_table = tables[table_index]
        
        # Truncate table to max_rows for preview
        rows = target_table.findall("w:tr", DOCX_NS)
        for row in rows[max_rows:]:
            target_table.remove(row)
        
        # Remove everything from body except sectPr, then add the table
        sect_pr = body.find("w:sectPr", DOCX_NS)
        for child in list(body):
            if child.tag != _qn("w", "sectPr"):
                body.remove(child)
        
        # Insert table at beginning
        body.insert(0, target_table)
        
        # Set small margins for compact preview
        if sect_pr is not None:
            # Remove header/footer references to avoid rendering heavy branded
            # assets in previews (faster and less COM failures on some templates).
            for ref_tag in ("headerReference", "footerReference"):
                for ref_node in list(sect_pr.findall(f"w:{ref_tag}", DOCX_NS)):
                    sect_pr.remove(ref_node)

            pg_mar = sect_pr.find("w:pgMar", DOCX_NS)
            if pg_mar is not None:
                pg_mar.set(_qn("w", "top"), "360")      # 0.25 inch
                pg_mar.set(_qn("w", "bottom"), "360")
                pg_mar.set(_qn("w", "left"), "360")
                pg_mar.set(_qn("w", "right"), "360")
        
        # Write modified document
        updated_xml = _serialize_ooxml_part(
            doc_root,
            doc_xml,
            namespace_hints=namespace_hints,
        )
        _write_docx_parts(docx_path, {"word/document.xml": updated_xml})
        
        # Convert to PDF and then PNG
        
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
            
        docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
        
        pdf_result = convert_docx_with_diagnostics(docx_b64, timeout_s=preview_pdf_timeout_s)
        pdf_b64 = pdf_result.get("pdf_b64")
        
        if not pdf_b64:
            error_msg = "PDF conversion failed"
            error_detail = pdf_result.get('error', 'Unknown error')
            # Check for Word specific error if available
            if pdf_result.get('word_error'):
                error_detail += f" (Word error: {pdf_result.get('word_error')})"

            fallback_preview = _render_table_preview_with_fitz(target_table_info, max_rows=max_rows)
            if fallback_preview:
                _logger.info(
                    f"[Template] Table preview fallback renderer used after PDF conversion failure (table={table_index})"
                )
                return {
                    "success": True,
                    "preview_b64": fallback_preview,
                    "converter_used": "fitz_fallback_after_pdf_fail",
                    "error_detail": error_detail,
                }

            _logger.warning(f"[Template] {error_msg}: {error_detail}")
            _logger.warning(f"[Template] PDF converter used: {pdf_result.get('converter_used')}")
            return {
                "success": False,
                "error": error_msg,
                "error_detail": error_detail,
                "converter_used": pdf_result.get('converter_used')
            }
        
        pdf_bytes = base64.b64decode(pdf_b64)
        png_bytes = _convert_pdf_to_png(pdf_bytes, dpi=120)
        if not png_bytes:
            fallback_preview = _render_table_preview_with_fitz(target_table_info, max_rows=max_rows)
            if fallback_preview:
                return {
                    "success": True,
                    "preview_b64": fallback_preview,
                    "converter_used": "fitz_fallback_after_png_fail",
                }
            error_msg = "PNG conversion failed. PyMuPDF may not be installed."
            _logger.warning(f"[Template] {error_msg}")
            return {"success": False, "error": error_msg}

        preview_b64 = base64.b64encode(png_bytes).decode("utf-8")
        
        _logger.info(f"[Template] Successfully generated table preview for table {table_index} (kernel {kernel_id})")
        return {"success": True, "preview_b64": preview_b64}
        
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        _logger.error(f"[Template] Error generating table preview: {error_msg}")
        import traceback
        _logger.error(traceback.format_exc())
        return {"success": False, "error": error_msg}
    
    finally:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


def _normalize_table_border_style(style_value: Any) -> Optional[str]:
    if style_value is None:
        return None
    style = str(style_value).strip()
    if not style:
        return None
    if style.lower() in {"none", "nil"}:
        return "nil"
    return style


def _normalize_ooxml_color(color_value: Any) -> Optional[str]:
    if color_value is None:
        return None
    color = str(color_value).strip().replace("#", "")
    if not color or color.lower() in {"auto", "none"}:
        return None
    return color.upper()


def _border_size_to_ooxml(border_props: Dict[str, Any]) -> Optional[str]:
    if not isinstance(border_props, dict):
        return None
    raw_sz = border_props.get("sz")
    if raw_sz not in (None, ""):
        try:
            return str(max(0, int(float(raw_sz))))
        except (TypeError, ValueError):
            return None
    size_pt = border_props.get("size_pt")
    if size_pt in (None, ""):
        return None
    try:
        return str(max(0, int(round(float(size_pt) * 8))))
    except (TypeError, ValueError):
        return None


def _append_table_borders(parent: ET.Element, container_tag: str, borders: Dict[str, Any]) -> None:
    if not isinstance(borders, dict) or not borders:
        return
    container = ET.SubElement(parent, _qn("w", container_tag))
    side_order = ["top", "left", "bottom", "right", "insideH", "insideV", "start", "end"]
    side_map = {"start": "left", "end": "right"}
    added = False

    normalized_sides: List[str] = []
    for side in side_order:
        mapped = side_map.get(side, side)
        if mapped in borders and mapped not in normalized_sides:
            normalized_sides.append(mapped)
    for side in borders.keys():
        mapped = side_map.get(side, side)
        if mapped not in normalized_sides:
            normalized_sides.append(mapped)

    for side in normalized_sides:
        raw_props = borders.get(side)
        if not isinstance(raw_props, dict):
            continue
        style = _normalize_table_border_style(raw_props.get("style") if "style" in raw_props else raw_props.get("val"))
        size = _border_size_to_ooxml(raw_props)
        color = _normalize_ooxml_color(raw_props.get("color"))
        if style is None and size is None and color is None:
            continue
        if style is None:
            style = "single"

        border = ET.SubElement(container, _qn("w", side))
        border.set(_qn("w", "val"), style)
        if size is not None:
            border.set(_qn("w", "sz"), size)
        if color is not None:
            border.set(_qn("w", "color"), color)
        border.set(_qn("w", "space"), "0")
        added = True

    if not added:
        parent.remove(container)


def _append_cell_margins(parent: ET.Element, margins: Dict[str, Any], container_tag: str = "tblCellMar") -> None:
    if not isinstance(margins, dict) or not margins:
        return
    container = ET.SubElement(parent, _qn("w", container_tag))
    side_order = ["top", "bottom", "left", "right"]
    added = False

    for side in side_order:
        raw_value = margins.get(side)
        if raw_value in (None, ""):
            continue
        try:
            # Extracted margins are in points; convert to twips.
            twips = int(round(float(raw_value) * 20))
        except (TypeError, ValueError):
            continue
        margin = ET.SubElement(container, _qn("w", side))
        margin.set(_qn("w", "w"), str(max(0, twips)))
        margin.set(_qn("w", "type"), "dxa")
        added = True

    if not added:
        parent.remove(container)


def _has_meaningful_font_props(font_props: Dict[str, Any]) -> bool:
    if not isinstance(font_props, dict) or not font_props:
        return False
    return any(value not in (None, "", False) for value in font_props.values())


def _append_rpr_from_font_props(r_pr: ET.Element, font_props: Dict[str, Any]) -> None:
    if not isinstance(font_props, dict) or not font_props:
        return

    font_name = font_props.get("font_name")
    if font_name:
        r_fonts = ET.SubElement(r_pr, _qn("w", "rFonts"))
        _set_explicit_family_on_rfonts(r_fonts, str(font_name))

    size_pt = font_props.get("font_size_pt")
    if size_pt not in (None, ""):
        try:
            size_half_points = int(round(float(size_pt) * 2))
            if size_half_points > 0:
                sz = ET.SubElement(r_pr, _qn("w", "sz"))
                sz.set(_qn("w", "val"), str(size_half_points))
        except (TypeError, ValueError):
            pass

    if font_props.get("bold"):
        ET.SubElement(r_pr, _qn("w", "b"))
    if font_props.get("italic"):
        ET.SubElement(r_pr, _qn("w", "i"))

    underline_style = font_props.get("underline_style")
    underline_enabled = font_props.get("underline")
    if underline_style or underline_enabled:
        u = ET.SubElement(r_pr, _qn("w", "u"))
        if underline_style:
            u.set(_qn("w", "val"), str(underline_style).lower())

    color_rgb = _normalize_ooxml_color(font_props.get("color_rgb"))
    if color_rgb:
        color = ET.SubElement(r_pr, _qn("w", "color"))
        color.set(_qn("w", "val"), color_rgb)

    highlight = font_props.get("highlight_color")
    if highlight:
        hi = ET.SubElement(r_pr, _qn("w", "highlight"))
        hi.set(_qn("w", "val"), str(highlight).lower())

    if font_props.get("strike"):
        ET.SubElement(r_pr, _qn("w", "strike"))
    if font_props.get("double_strike"):
        ET.SubElement(r_pr, _qn("w", "dstrike"))
    if font_props.get("all_caps"):
        ET.SubElement(r_pr, _qn("w", "caps"))
    if font_props.get("small_caps"):
        ET.SubElement(r_pr, _qn("w", "smallCaps"))

    if font_props.get("superscript"):
        va = ET.SubElement(r_pr, _qn("w", "vertAlign"))
        va.set(_qn("w", "val"), "superscript")
    elif font_props.get("subscript"):
        va = ET.SubElement(r_pr, _qn("w", "vertAlign"))
        va.set(_qn("w", "val"), "subscript")


def _append_ppr_from_paragraph_props(p_pr: ET.Element, paragraph_props: Dict[str, Any]) -> None:
    if not isinstance(paragraph_props, dict) or not paragraph_props:
        return

    alignment_value = paragraph_props.get("alignment")
    if alignment_value not in (None, ""):
        jc = ET.SubElement(p_pr, _qn("w", "jc"))
        align_key = str(alignment_value).strip().upper()
        align_map = {
            "LEFT": "left",
            "CENTER": "center",
            "RIGHT": "right",
            "JUSTIFY": "both",
        }
        jc.set(_qn("w", "val"), align_map.get(align_key, str(alignment_value).strip().lower()))

    spacing_attrs: Dict[str, str] = {}
    before_pt = paragraph_props.get("space_before_pt")
    after_pt = paragraph_props.get("space_after_pt")
    if before_pt not in (None, ""):
        try:
            spacing_attrs[_qn("w", "before")] = str(int(round(float(before_pt) * 20)))
        except (TypeError, ValueError):
            pass
    if after_pt not in (None, ""):
        try:
            spacing_attrs[_qn("w", "after")] = str(int(round(float(after_pt) * 20)))
        except (TypeError, ValueError):
            pass

    line_spacing = paragraph_props.get("line_spacing")
    line_rule = paragraph_props.get("line_spacing_rule")
    rule_key = str(line_rule).strip().upper() if line_rule not in (None, "") else ""
    if line_spacing not in (None, "") or rule_key:
        line_rule_map = {
            "SINGLE": ("auto", 240.0),
            "ONE_POINT_FIVE": ("auto", 360.0),
            "DOUBLE": ("auto", 480.0),
            "MULTIPLE": ("auto", None),
            "EXACTLY": ("exact", None),
            "AT_LEAST": ("atLeast", None),
        }
        ooxml_rule, preset_line = line_rule_map.get(rule_key, ("auto", None))
        spacing_attrs[_qn("w", "lineRule")] = ooxml_rule
        line_value: Optional[int] = None
        if line_spacing not in (None, ""):
            try:
                numeric_line = float(line_spacing)
                if ooxml_rule in {"exact", "atLeast"}:
                    line_value = int(round(numeric_line * 20))
                else:
                    line_value = int(round(numeric_line * 240))
            except (TypeError, ValueError):
                line_value = None
        elif preset_line is not None:
            line_value = int(round(preset_line))
        if line_value is not None:
            spacing_attrs[_qn("w", "line")] = str(max(0, line_value))

    if spacing_attrs:
        spacing = ET.SubElement(p_pr, _qn("w", "spacing"))
        for attr_name, attr_value in spacing_attrs.items():
            spacing.set(attr_name, attr_value)

    ind_attrs: Dict[str, str] = {}
    left_indent = paragraph_props.get("left_indent_inches")
    right_indent = paragraph_props.get("right_indent_inches")
    first_line_indent = paragraph_props.get("first_line_indent_inches")
    if left_indent not in (None, ""):
        try:
            ind_attrs[_qn("w", "left")] = str(int(round(float(left_indent) * 1440)))
        except (TypeError, ValueError):
            pass
    if right_indent not in (None, ""):
        try:
            ind_attrs[_qn("w", "right")] = str(int(round(float(right_indent) * 1440)))
        except (TypeError, ValueError):
            pass
    if first_line_indent not in (None, ""):
        try:
            indent_twips = int(round(float(first_line_indent) * 1440))
            if indent_twips < 0:
                ind_attrs[_qn("w", "hanging")] = str(abs(indent_twips))
            else:
                ind_attrs[_qn("w", "firstLine")] = str(indent_twips)
        except (TypeError, ValueError):
            pass
    if ind_attrs:
        ind = ET.SubElement(p_pr, _qn("w", "ind"))
        for attr_name, attr_value in ind_attrs.items():
            ind.set(attr_name, attr_value)

    def _append_on_off(tag_name: str, raw_value: Any) -> None:
        value = _coerce_optional_bool(raw_value)
        if value is None:
            return
        elem = ET.SubElement(p_pr, _qn("w", tag_name))
        if value is False:
            elem.set(_qn("w", "val"), "0")

    _append_on_off("keepNext", paragraph_props.get("keep_with_next"))
    _append_on_off("keepLines", paragraph_props.get("keep_together"))
    _append_on_off("pageBreakBefore", paragraph_props.get("page_break_before"))
    _append_on_off("widowControl", paragraph_props.get("widow_control"))


def _extract_table_format_for_style(table_info: Dict[str, Any]) -> Dict[str, Any]:
    parsed = table_info.get("parsed_properties") if isinstance(table_info, dict) else None
    parsed = parsed if isinstance(parsed, dict) else {}
    return {
        "borders": parsed.get("borders") or table_info.get("borders") or {},
        "shading_color": parsed.get("shading_color") or table_info.get("shading_fill"),
        "cell_margins": parsed.get("cell_margins") or table_info.get("margins") or {},
        "alignment": parsed.get("alignment") or table_info.get("alignment"),
        "width_type": parsed.get("width_type"),
        "width_value": parsed.get("width_value"),
        "layout_type": parsed.get("layout_type"),
        "cell_spacing_pt": parsed.get("cell_spacing_pt"),
        "indent_pt": parsed.get("indent_pt"),
        "look": parsed.get("look") or table_info.get("look") or {},
        "first_row_format": table_info.get("first_row_format") or {},
        "has_distinct_header": bool(table_info.get("has_distinct_header")),
    }


def _apply_table_properties_to_style_tblpr(tbl_pr: ET.Element, table_format: Dict[str, Any]) -> None:
    # Remove existing elements we actively regenerate from direct-format source
    for tag in ("tblBorders", "shd", "tblCellMar", "jc", "tblCellSpacing", "tblInd", "tblW", "tblLayout", "tblLook"):
        for old in list(tbl_pr.findall(f"w:{tag}", DOCX_NS)):
            tbl_pr.remove(old)

    _append_table_borders(tbl_pr, "tblBorders", table_format.get("borders") or {})

    shading_color = _normalize_ooxml_color(table_format.get("shading_color"))
    if shading_color:
        shd = ET.SubElement(tbl_pr, _qn("w", "shd"))
        shd.set(_qn("w", "val"), "clear")
        shd.set(_qn("w", "color"), "auto")
        shd.set(_qn("w", "fill"), shading_color)

    _append_cell_margins(tbl_pr, table_format.get("cell_margins") or {}, "tblCellMar")

    alignment = table_format.get("alignment")
    if alignment:
        jc = ET.SubElement(tbl_pr, _qn("w", "jc"))
        jc.set(_qn("w", "val"), str(alignment).lower())

    spacing_pt = table_format.get("cell_spacing_pt")
    if spacing_pt not in (None, ""):
        try:
            spacing_twips = int(round(float(spacing_pt) * 20))
            spacing = ET.SubElement(tbl_pr, _qn("w", "tblCellSpacing"))
            spacing.set(_qn("w", "w"), str(max(0, spacing_twips)))
            spacing.set(_qn("w", "type"), "dxa")
        except (TypeError, ValueError):
            pass

    indent_pt = table_format.get("indent_pt")
    if indent_pt not in (None, ""):
        try:
            indent_twips = int(round(float(indent_pt) * 20))
            tbl_ind = ET.SubElement(tbl_pr, _qn("w", "tblInd"))
            tbl_ind.set(_qn("w", "w"), str(max(0, indent_twips)))
            tbl_ind.set(_qn("w", "type"), "dxa")
        except (TypeError, ValueError):
            pass



def _apply_first_row_variant(style_elem: ET.Element, table_format: Dict[str, Any]) -> None:
    # Replace firstRow variant only when direct-format extraction has a meaningful signal.
    first_row_format = table_format.get("first_row_format") or {}
    sample_cell = first_row_format.get("sample_cell") or {}
    font_props = first_row_format.get("font_properties") or {}
    shading_fill = first_row_format.get("shading_fill") or sample_cell.get("shading_color")
    borders = first_row_format.get("borders") or sample_cell.get("borders")
    margins = first_row_format.get("margins") or sample_cell.get("margins")
    vertical_align = first_row_format.get("vertical_align") or sample_cell.get("vertical_align")

    has_tc_pr = bool(shading_fill or borders or margins or vertical_align)
    has_r_pr = _has_meaningful_font_props(font_props)
    should_apply = bool(table_format.get("has_distinct_header") or has_tc_pr or has_r_pr)
    if not should_apply:
        return

    for old_pr in list(style_elem.findall(_qn("w", "tblStylePr"))):
        if old_pr.get(_qn("w", "type")) == "firstRow":
            style_elem.remove(old_pr)

    tbl_style_pr = ET.SubElement(style_elem, _qn("w", "tblStylePr"))
    tbl_style_pr.set(_qn("w", "type"), "firstRow")

    if has_tc_pr:
        tc_pr = ET.SubElement(tbl_style_pr, _qn("w", "tcPr"))

        shade = _normalize_ooxml_color(shading_fill)
        if shade:
            shd = ET.SubElement(tc_pr, _qn("w", "shd"))
            shd.set(_qn("w", "val"), "clear")
            shd.set(_qn("w", "color"), "auto")
            shd.set(_qn("w", "fill"), shade)

        _append_table_borders(tc_pr, "tcBorders", borders or {})
        _append_cell_margins(tc_pr, margins or {}, "tcMar")

        if vertical_align:
            v_align = ET.SubElement(tc_pr, _qn("w", "vAlign"))
            v_align.set(_qn("w", "val"), str(vertical_align))

    if has_r_pr:
        r_pr = ET.SubElement(tbl_style_pr, _qn("w", "rPr"))
        _append_rpr_from_font_props(r_pr, font_props)


def create_table_style_from_format(
    kernel_id: str,
    table_index: int,
    new_style_name: str
) -> Optional[Dict[str, Any]]:
    """Create a new table style from a detected table's direct formatting.
    
    This function reads the specified table from the document, extracts its properties
    (borders, shading, etc.), creates a new style in styles.xml, and saves the template.
    
    Args:
        kernel_id: Kernel ID
        table_index: Index of the table in the document
        new_style_name: Name for the new style
        
    Returns:
        The updated template info (dict) or None if failure
    """
    normalized_table_index = _validate_table_index_value(table_index)
    if normalized_table_index is None:
        error_msg = _table_index_error_message(table_index)
        _logger.error(f"[Template] {error_msg}")
        return None
    table_index = normalized_table_index

    template_path = get_template_docx_path(kernel_id)
    if not template_path or not os.path.exists(template_path):
        _logger.error(f"[Template] Template not found for kernel {kernel_id}")
        return None

    try:
        current_template = get_template(kernel_id) or {}
        runtime_defaults = _normalize_table_style_runtime_defaults(
            current_template.get(TABLE_STYLE_RUNTIME_DEFAULTS_KEY)
        )

        # 1. Read DOCX to get tables and properties
        with open(template_path, "rb") as f:
            docx_bytes = f.read()
            
        tables = _extract_document_tables(docx_bytes)
        if not tables or table_index >= len(tables):
            error_msg = _table_index_error_message(table_index, len(tables))
            _logger.error(f"[Template] {error_msg}")
            return None
            
        target_table = tables[table_index]
        table_format = _extract_table_format_for_style(target_table)
        runtime_patch = _extract_runtime_defaults_from_table_format(table_format)
        
        # 2. Read styles.xml
        parts = _read_docx_parts(docx_bytes, ["word/styles.xml"])
        styles_xml = parts.get("word/styles.xml")
        if not styles_xml:
            _logger.error("[Template] No styles.xml found")
            return None
            
        styles_root = ET.fromstring(styles_xml)
        
        # 3. Create new style element
        # Clean name for ID
        clean_name = "".join(c for c in new_style_name if c.isalnum())
        style_id = f"{clean_name}_{str(uuid.uuid4())[:8]}"
        
        new_style = ET.Element(_qn("w", "style"))
        new_style.set(_qn("w", "type"), "table")
        new_style.set(_qn("w", "customStyle"), "1")
        new_style.set(_qn("w", "styleId"), style_id)
        
        # <w:name w:val="..."/>
        name_elem = ET.SubElement(new_style, _qn("w", "name"))
        name_elem.set(_qn("w", "val"), new_style_name)
        
        # <w:basedOn w:val="TableNormal"/>
        based_on = ET.SubElement(new_style, _qn("w", "basedOn"))
        based_on.set(_qn("w", "val"), "TableNormal")
        
        # <w:uiPriority w:val="99"/>
        ui_prio = ET.SubElement(new_style, _qn("w", "uiPriority"))
        ui_prio.set(_qn("w", "val"), "99")
        
        # <w:qFormat/> to show in styles gallery
        ET.SubElement(new_style, _qn("w", "qFormat"))
        
        # <w:tblPr> - Table Properties
        tbl_pr = ET.SubElement(new_style, _qn("w", "tblPr"))
        _apply_table_properties_to_style_tblpr(tbl_pr, table_format)
        _apply_first_row_variant(new_style, table_format)

        # 5. Inject into styles.xml
        styles_root.append(new_style)
        
        # 6. Save and re-extract
        namespace_hints = _collect_docx_namespace_hints_from_bytes(docx_bytes)
        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        
        # Write to DOCX
        _write_docx_parts(Path(template_path), {"word/styles.xml": updated_xml})
        
        # Re-read docx bytes to update cache
        with open(template_path, "rb") as f:
            new_docx_bytes = f.read()
            
        extracted = extract_styles_from_docx(new_docx_bytes)
        runtime_defaults = _apply_runtime_defaults_patch(
            runtime_defaults,
            style_id=style_id,
            style_name=new_style_name,
            patch=runtime_patch,
        )
        extracted[TABLE_STYLE_RUNTIME_DEFAULTS_KEY] = runtime_defaults
        extracted["style_coverage"] = get_style_coverage(extracted)
        save_template(kernel_id, new_docx_bytes, extracted)
        
        _logger.info(f"[Template] Created new table style '{new_style_name}' ({style_id})")
        return extracted

    except Exception as e:
        _logger.error(f"[Template] Error creating table style: {e}")
        import traceback
        _logger.error(traceback.format_exc())
        return None


def apply_table_format_to_style(
    kernel_id: str,
    table_index: int,
    target_style_name: Optional[str],
    target_style_id: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    """Apply the format from a detected table to an existing table style.
    
    This function reads the specified table from the document, extracts its properties
    (borders, shading, etc.), and applies them to the specified existing style.
    
    Args:
        kernel_id: Kernel ID
        table_index: Index of the table in the document
        target_style_name: Name of the existing style to update
        target_style_id: Optional style_id of the existing style to update
        
    Returns:
        The updated template info (dict) or None if failure
    """
    normalized_table_index = _validate_table_index_value(table_index)
    if normalized_table_index is None:
        error_msg = _table_index_error_message(table_index)
        _logger.error(f"[Template] {error_msg}")
        return None
    table_index = normalized_table_index

    template_path = get_template_docx_path(kernel_id)
    if not template_path or not os.path.exists(template_path):
        _logger.error(f"[Template] Template not found for kernel {kernel_id}")
        return None

    try:
        current_template = get_template(kernel_id) or {}
        runtime_defaults = _normalize_table_style_runtime_defaults(
            current_template.get(TABLE_STYLE_RUNTIME_DEFAULTS_KEY)
        )

        # 1. Read DOCX to get tables and properties
        with open(template_path, "rb") as f:
            docx_bytes = f.read()
            
        tables = _extract_document_tables(docx_bytes)
        if not tables or table_index >= len(tables):
            error_msg = _table_index_error_message(table_index, len(tables))
            _logger.error(f"[Template] {error_msg}")
            return None
            
        target_table = tables[table_index]
        table_format = _extract_table_format_for_style(target_table)
        runtime_patch = _extract_runtime_defaults_from_table_format(table_format)
        
        # 2. Read styles.xml
        parts = _read_docx_parts(docx_bytes, ["word/styles.xml"])
        styles_xml = parts.get("word/styles.xml")
        if not styles_xml:
            _logger.error("[Template] No styles.xml found")
            return None
            
        styles_root = ET.fromstring(styles_xml)
        
        # 3. Find existing style by name
        existing_style = _find_style_element(styles_root, target_style_name, target_style_id)
        if existing_style is None:
            _logger.error(f"[Template] Style not found (name={target_style_name!r}, id={target_style_id!r})")
            return None
            
        # Check it's a table style
        style_type = existing_style.get(_qn("w", "type"))
        if style_type != "table":
            _logger.error(f"[Template] Style '{target_style_name}' is not a table style (type: {style_type})")
            return None

        # 3b. Protect header/footer tables before modifying the style
        _freeze_header_footer_table_styles(
            Path(template_path), target_style_name, target_style_id
        )
        
        # 4. Update or create <w:tblPr> - Table Properties
        tbl_pr = existing_style.find(_qn("w", "tblPr"))
        if tbl_pr is None:
            tbl_pr = ET.SubElement(existing_style, _qn("w", "tblPr"))
        _apply_table_properties_to_style_tblpr(tbl_pr, table_format)
        _apply_first_row_variant(existing_style, table_format)

        # 6. Save and re-extract
        namespace_hints = _collect_docx_namespace_hints_from_bytes(docx_bytes)
        updated_xml = _serialize_ooxml_part(
            styles_root,
            styles_xml,
            namespace_hints=namespace_hints,
        )
        
        # Write to DOCX
        _write_docx_parts(Path(template_path), {"word/styles.xml": updated_xml})
        
        # Re-read docx bytes to update cache
        with open(template_path, "rb") as f:
            new_docx_bytes = f.read()
            
        extracted = extract_styles_from_docx(new_docx_bytes)
        style_id = existing_style.get(_qn("w", "styleId")) or target_style_id
        style_name = _find_val(existing_style, "w:name") or target_style_name or style_id
        runtime_defaults = _apply_runtime_defaults_patch(
            runtime_defaults,
            style_id=style_id,
            style_name=style_name,
            patch=runtime_patch,
        )
        extracted[TABLE_STYLE_RUNTIME_DEFAULTS_KEY] = runtime_defaults
        extracted["style_coverage"] = get_style_coverage(extracted)
        save_template(kernel_id, new_docx_bytes, extracted)
        
        applied_label = target_style_name or target_style_id or "<unknown>"
        _logger.info(f"[Template] Applied table format #{table_index + 1} to style '{applied_label}'")
        return extracted

    except Exception as e:
        _logger.error(f"[Template] Error applying table format to style: {e}")
        import traceback
        _logger.error(traceback.format_exc())
        return None
