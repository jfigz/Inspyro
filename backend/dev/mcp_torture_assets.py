"""Canonical assets for the MCP torture probe."""

from __future__ import annotations

from pathlib import Path
from textwrap import dedent
from typing import Any

PROFILE_TOOLSETS = {
    "core": {
        "get_system_info",
        "get_health",
        "list_component_profiles",
        "set_component_profile",
    },
    "authoring": {
        "get_system_info",
        "get_health",
        "list_component_profiles",
        "set_component_profile",
        "notebook_create",
        "notebook_load",
        "list_session_notebooks",
        "notebook_sync_cells",
        "notebook_save",
        "execute_cell",
        "execute_all_cells",
        "get_kernel_status",
        "get_run_status",
        "cancel_run",
        "resume_run",
        "list_cells",
        "get_cell",
        "find_in_notebook",
        "reset_kernel",
        "interrupt_kernel",
        "shutdown_kernel",
        "close_session_notebook",
        "get_variables",
        "get_document_docx",
        "get_document_pdf",
        "export_document_docx",
        "export_document_pdf",
        "check_document_quality",
        "export_clean_document_docx",
        "run_document_workbench",
        "compare_document_versions",
        "manage_document_review",
        "prepare_document_delivery",
        "reconvert_pdf",
        "upload_template",
        "bind_template_to_notebook",
        "get_template_info",
        "delete_template",
        "update_template_style",
        "convert_units",
        "get_units_catalog",
        "check_units_compatible",
    },
    "analysis": {
        "get_system_info",
        "get_health",
        "list_component_profiles",
        "set_component_profile",
        "analyze_dependencies",
        "analyze_impact",
        "run_sensitivity",
        "optimize_design",
        "compare_scenarios",
        "run_code_checks",
    },
    "files": {
        "get_system_info",
        "get_health",
        "list_component_profiles",
        "set_component_profile",
        "list_files",
        "read_file",
        "write_file",
        "create_file",
        "delete_file",
        "rename_file",
    },
    "admin": {
        "get_system_info",
        "get_health",
        "list_component_profiles",
        "set_component_profile",
        "get_metrics",
        "get_pdf_status",
    },
}
PROFILE_TOOLSETS["all"] = set().union(*PROFILE_TOOLSETS.values())

PUBLIC_RESOURCE_URIS = [
    "inspyro://manifest",
    "inspyro://system/info",
    "inspyro://system/health",
    "inspyro://units/catalog",
    "inspyro://pdf/status",
    "inspyro://files/tree",
    "inspyro://session/notebooks",
    "inspyro://guides/start-here",
    "inspyro://guides/client-configuration",
    "inspyro://guides/notebook-workflow",
    "inspyro://guides/docx-quickstart",
    "inspyro://guides/artifact-lifecycle",
    "inspyro://guides/template-workflow",
    "inspyro://guides/analysis-units-workflow",
    "inspyro://guides/error-recovery",
    "inspyro://examples/notebook-docx-report",
]

PUBLIC_RESOURCE_TEMPLATE_URIS = [
    "inspyro://workspace/tree/{path*}",
    "inspyro://workspace/file/{path*}",
    "inspyro://notebooks/{path*}/cells/{cell_id}",
    "inspyro://artifacts/{kernel_id}/{kind}",
    "inspyro://artifacts/{kernel_id}/{kind}/{execution_id}",
    "inspyro://artifacts/token/{kind}/{token}",
    "inspyro://runs/{run_id}",
]

PUBLIC_PROMPT_NAMES = [
    "create_engineering_notebook",
    "debug_cell_error",
    "review_notebook",
    "unit_conversion_help",
    "start_inspyro_session",
    "create_docx_report_notebook",
    "recover_mcp_notebook_session",
]

TOOL_COVERAGE_MATRIX = {
    "get_system_info": {"phase": "discovery", "assertion": "returns workspace metadata"},
    "get_health": {"phase": "discovery", "assertion": "backend is healthy"},
    "list_component_profiles": {"phase": "discovery", "assertion": "lists authoring/analysis/files/admin/all"},
    "set_component_profile": {"phase": "discovery", "assertion": "switches visible tool surface per profile"},
    "notebook_create": {"phase": "authoring", "assertion": "creates main and secondary notebooks"},
    "notebook_load": {"phase": "authoring", "assertion": "reloads notebook with stable ids and source"},
    "list_session_notebooks": {"phase": "authoring", "assertion": "lists notebook sessions and kernel state"},
    "notebook_sync_cells": {"phase": "authoring", "assertion": "syncs canonical notebook spec and mutations"},
    "notebook_save": {"phase": "stress", "assertion": "persists notebook after recovery"},
    "execute_cell": {"phase": "authoring", "assertion": "executes persisted cell and long-running interrupt case"},
    "execute_all_cells": {"phase": "authoring", "assertion": "materializes full notebook and doc artifacts"},
    "get_kernel_status": {"phase": "authoring", "assertion": "reads live kernel status for the active notebook"},
    "get_run_status": {"phase": "authoring", "assertion": "polls a public background run by run_id"},
    "cancel_run": {"phase": "authoring", "assertion": "handles cancellation requests for notebook runs"},
    "resume_run": {"phase": "authoring", "assertion": "handles resume requests for notebook runs"},
    "list_cells": {"phase": "authoring", "assertion": "lists lightweight notebook cells"},
    "get_cell": {"phase": "authoring", "assertion": "reads one persisted notebook cell with source"},
    "find_in_notebook": {"phase": "authoring", "assertion": "searches text inside the notebook"},
    "reset_kernel": {"phase": "stress", "assertion": "covers both soft and hard reset"},
    "interrupt_kernel": {"phase": "stress", "assertion": "interrupts long-running execution"},
    "shutdown_kernel": {"phase": "stress", "assertion": "closes both kernels cleanly"},
    "close_session_notebook": {"phase": "stress", "assertion": "closes explicit notebook session alias cleanly"},
    "get_variables": {"phase": "authoring", "assertion": "extracts analysis payloads from runtime"},
    "get_document_docx": {"phase": "documents", "assertion": "returns DOCX handle and optional inline content"},
    "get_document_pdf": {"phase": "documents", "assertion": "returns PDF handle and optional inline content"},
    "export_document_docx": {"phase": "documents", "assertion": "exports DOCX to stable local path"},
    "export_document_pdf": {"phase": "documents", "assertion": "exports PDF to stable local path"},
    "check_document_quality": {"phase": "documents", "assertion": "runs/reads normalized DOCX quality summaries"},
    "export_clean_document_docx": {"phase": "documents", "assertion": "exports a clean publication-safe DOCX copy"},
    "run_document_workbench": {"phase": "documents", "assertion": "runs compact DOCX Workbench operations"},
    "compare_document_versions": {"phase": "documents", "assertion": "returns compact DOCX diff resources"},
    "manage_document_review": {"phase": "documents", "assertion": "extracts review metadata without binary payloads"},
    "prepare_document_delivery": {"phase": "documents", "assertion": "prepares and optionally exports a delivery DOCX"},
    "reconvert_pdf": {"phase": "documents", "assertion": "regenerates PDF from current DOCX"},
    "upload_template": {"phase": "templates", "assertion": "attaches rich template fixture"},
    "bind_template_to_notebook": {"phase": "templates", "assertion": "persists active template binding beside notebook"},
    "get_template_info": {"phase": "templates", "assertion": "inspects active template metadata"},
    "delete_template": {"phase": "templates", "assertion": "detaches active template once"},
    "update_template_style": {"phase": "templates", "assertion": "mutates Normal and Heading 1"},
    "convert_units": {"phase": "units", "assertion": "converts notebook-related engineering units"},
    "get_units_catalog": {"phase": "units", "assertion": "reads supported units catalog"},
    "check_units_compatible": {"phase": "units", "assertion": "confirms dimensional compatibility"},
    "analyze_dependencies": {"phase": "analysis", "assertion": "builds dependency graph from notebook code"},
    "analyze_impact": {"phase": "analysis", "assertion": "builds impact graph from notebook code"},
    "run_sensitivity": {"phase": "analysis", "assertion": "evaluates parametric sensitivity payloads"},
    "optimize_design": {"phase": "analysis", "assertion": "returns recommended design from notebook formulas"},
    "compare_scenarios": {"phase": "analysis", "assertion": "compares baseline and candidate scenarios"},
    "run_code_checks": {"phase": "analysis", "assertion": "evaluates engineering checks"},
    "list_files": {"phase": "files", "assertion": "lists temp workspace tree"},
    "read_file": {"phase": "files", "assertion": "reads seeded auxiliary files and notebook path rejection"},
    "write_file": {"phase": "files", "assertion": "writes seeded auxiliary inputs and notebook path rejection"},
    "create_file": {"phase": "files", "assertion": "creates directories/files and notebook path rejection"},
    "delete_file": {"phase": "files", "assertion": "deletes scratch file and notebook path rejection"},
    "rename_file": {"phase": "files", "assertion": "renames scratch file and notebook path rejection"},
    "get_metrics": {"phase": "admin", "assertion": "reads backend metrics"},
    "get_pdf_status": {"phase": "admin", "assertion": "reads PDF converter status"},
}

RESOURCE_COVERAGE_MATRIX = {
    "inspyro://manifest": {"phase": "discovery", "assertion": "profile/resource map is readable"},
    "inspyro://system/info": {"phase": "discovery", "assertion": "returns normalized workspace fields"},
    "inspyro://system/health": {"phase": "discovery", "assertion": "returns healthy backend payload"},
    "inspyro://units/catalog": {"phase": "discovery", "assertion": "exposes units catalog resource"},
    "inspyro://pdf/status": {"phase": "discovery", "assertion": "exposes PDF status resource"},
    "inspyro://files/tree": {"phase": "discovery", "assertion": "exposes workspace tree resource"},
    "inspyro://session/notebooks": {"phase": "discovery", "assertion": "exposes live notebook session inventory"},
    "inspyro://guides/start-here": {"phase": "discovery", "assertion": "onboarding guide is readable"},
    "inspyro://guides/client-configuration": {"phase": "discovery", "assertion": "client configuration guide is readable"},
    "inspyro://guides/notebook-workflow": {"phase": "discovery", "assertion": "notebook guide is readable"},
    "inspyro://guides/docx-quickstart": {"phase": "discovery", "assertion": "DOCX guide is readable"},
    "inspyro://guides/artifact-lifecycle": {"phase": "discovery", "assertion": "artifact guide is readable"},
    "inspyro://guides/template-workflow": {"phase": "discovery", "assertion": "template guide is readable"},
    "inspyro://guides/analysis-units-workflow": {"phase": "discovery", "assertion": "analysis guide is readable"},
    "inspyro://guides/error-recovery": {"phase": "discovery", "assertion": "recovery guide is readable"},
    "inspyro://examples/notebook-docx-report": {"phase": "discovery", "assertion": "example flow is readable"},
}

RESOURCE_TEMPLATE_COVERAGE_MATRIX = {
    "inspyro://workspace/tree/{path*}": {"phase": "files", "assertion": "reads tree for temp workspace"},
    "inspyro://workspace/file/{path*}": {"phase": "files", "assertion": "reads seeded loads.json"},
    "inspyro://notebooks/{path*}/cells/{cell_id}": {"phase": "authoring", "assertion": "reads persisted notebook cell snapshot"},
    "inspyro://artifacts/{kernel_id}/{kind}": {"phase": "documents", "assertion": "reads latest session artifact"},
    "inspyro://artifacts/{kernel_id}/{kind}/{execution_id}": {"phase": "documents", "assertion": "reads execution-scoped artifact"},
    "inspyro://artifacts/token/{kind}/{token}": {"phase": "documents", "assertion": "reads portable token artifact"},
    "inspyro://runs/{run_id}": {"phase": "analysis", "assertion": "reads execution run summary"},
}

PROMPT_COVERAGE_MATRIX = {
    "create_engineering_notebook": {"phase": "authoring", "assertion": "prompt routes notebook creation flow"},
    "debug_cell_error": {"phase": "stress", "assertion": "prompt routes negative execution recovery"},
    "review_notebook": {"phase": "analysis", "assertion": "prompt routes notebook review flow"},
    "unit_conversion_help": {"phase": "units", "assertion": "prompt routes unit conversion flow"},
    "start_inspyro_session": {"phase": "discovery", "assertion": "prompt enforces onboarding order"},
    "create_docx_report_notebook": {"phase": "documents", "assertion": "prompt routes DOCX notebook flow"},
    "recover_mcp_notebook_session": {"phase": "stress", "assertion": "prompt routes recovery flow"},
}

PRIMARY_NOTEBOOK_NAME = "mcp_torture_notebook.ipynb"
SECONDARY_NOTEBOOK_NAME = "mcp_torture_secondary.ipynb"
TEMPLATE_FIXTURE_RELATIVE = Path("backend/dev/fixtures/mcp_torture_template.docx")
PRIMARY_NOTEBOOK_DOC_CELL_IDS = [
    "c06_doc_report_cover",
    "c07_doc_report_tables",
    "c08_doc_report_figures",
]
PRIMARY_NOTEBOOK_ANALYSIS_CELL_ID = "c05_analysis_payloads"
PRIMARY_NOTEBOOK_BOOTSTRAP_CELL_ID = "c01_bootstrap_workspace"
PRIMARY_NOTEBOOK_LONG_CELL_ID = "c09_long_running_interrupt"
SECONDARY_NOTEBOOK_CODE_CELL_ID = "s01_quick_runtime"

AUXILIARY_LOADS = {
    "beam_name": "MCP-Torture-Primary",
    "cases": [
        {"name": "service", "w_kN_m": 18.0, "temperature_c": 23.5, "dynamic_factor": 1.00},
        {"name": "wind", "w_kN_m": 22.5, "temperature_c": 18.0, "dynamic_factor": 1.08},
        {"name": "construction", "w_kN_m": 27.0, "temperature_c": 30.0, "dynamic_factor": 1.15},
    ],
}

AUXILIARY_SECTIONS_CSV = dedent(
    """
    name,b_mm,h_mm,area_mm2,I_mm4,Z_mm3
    W350x700,350,700,245000,10004166666.7,8166666.7
    W400x750,400,750,300000,14062500000.0,11250000.0
    W450x800,450,800,360000,19200000000.0,14400000.0
    """
).strip()


def _cell(cell_id: str, cell_type: str, source: str) -> dict[str, Any]:
    return {
        "cell_id": cell_id,
        "cell_type": cell_type,
        "source": dedent(source).strip() + "\n",
    }


PRIMARY_NOTEBOOK_SPEC = [
    _cell(
        "m00_overview",
        "markdown",
        """
        # MCP Torture Notebook

        Notebook canonico para exprimir Inspyro via MCP con unidades,
        analisis, reportes DOCX y recovery.
        """,
    ),
    _cell(
        PRIMARY_NOTEBOOK_BOOTSTRAP_CELL_ID,
        "code",
        """
        import json
        import math
        import statistics
        import time
        from pathlib import Path

        import matplotlib.pyplot as plt
        import numpy as np
        import pandas as pd
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        WORKSPACE_ROOT = Path().resolve()
        INPUT_DIR = WORKSPACE_ROOT / "inputs"
        OUTPUT_DIR = WORKSPACE_ROOT / "outputs"
        OUTPUT_DIR.mkdir(exist_ok=True)

        LOADS = json.loads((INPUT_DIR / "loads.json").read_text(encoding="utf-8"))
        SECTIONS_DF = pd.read_csv(INPUT_DIR / "sections.csv")
        analysis_trace = []

        print(f"workspace={WORKSPACE_ROOT}")
        print(f"load_cases={len(LOADS['cases'])} section_rows={len(SECTIONS_DF)}")
        """,
    ),
    _cell(
        "c02_engineering_units",
        "code",
        """
        if "LOADS" not in globals():
            LOADS = json.loads((INPUT_DIR / "loads.json").read_text(encoding="utf-8"))
        if "SECTIONS_DF" not in globals():
            SECTIONS_DF = pd.read_csv(INPUT_DIR / "sections.csv")

        span = 8.0 * m
        section_b = 0.35 * m
        section_h = 0.70 * m
        cover = 45 * mm
        density = 7850 * kg / m**3
        fy = 420 * MPa
        elastic_modulus = 200000 * MPa
        ambient_temperature = Q_(23.5, degC)
        load_vector = np.array([case["w_kN_m"] for case in LOADS["cases"]], dtype=float) * kN / m
        stiffness_matrix = np.array([[42.0, -21.0], [-21.0, 42.0]], dtype=float) * kN / m
        thermal_profile = np.array([ambient_temperature.to(K).magnitude, 298.15], dtype=float)

        span_m = float(span.to(m).magnitude)
        section_b_mm = float(section_b.to(mm).magnitude)
        section_h_mm = float(section_h.to(mm).magnitude)
        fy_MPa = float(fy.to(MPa).magnitude)
        E_MPa = float(elastic_modulus.to(MPa).magnitude)

        print(load_vector)
        print(density)
        """,
    ),
    _cell(
        "c03_structural_model",
        "code",
        """
        def rectangular_inertia(width_m: float, height_m: float) -> float:
            return width_m * height_m**3 / 12.0


        def section_modulus(width_m: float, height_m: float) -> float:
            return width_m * height_m**2 / 6.0


        def design_moment_kNm(w_kN_m: float, span_m: float, dynamic_factor: float = 1.0) -> float:
            base_moment = w_kN_m * span_m**2 / 8.0
            if dynamic_factor > 1.0:
                return base_moment * dynamic_factor
            return base_moment


        def design_stress_MPa(moment_kNm: float, width_m: float, height_m: float) -> float:
            section_modulus_m3 = section_modulus(width_m, height_m)
            return (moment_kNm * 1000.0) / section_modulus_m3 / 1_000_000.0


        def demand_capacity_ratio(stress_MPa: float, fy_MPa: float, phi: float = 0.9) -> float:
            return stress_MPa / (phi * fy_MPa)


        class BeamScenario:
            def __init__(self, name: str, w_kN_m: float, temperature_c: float, dynamic_factor: float = 1.0):
                self.name = name
                self.w_kN_m = float(w_kN_m)
                self.temperature_c = float(temperature_c)
                self.dynamic_factor = float(dynamic_factor)

            def evaluate(self, width_m: float, height_m: float, fy_MPa: float, E_MPa: float) -> dict[str, float]:
                moment = design_moment_kNm(self.w_kN_m, span_m, self.dynamic_factor)
                shear = self.w_kN_m * span_m / 2.0
                stress = design_stress_MPa(moment, width_m, height_m)
                dcr = demand_capacity_ratio(stress, fy_MPa)
                inertia_m4 = rectangular_inertia(width_m, height_m)
                thermal_factor = 1.0
                if self.temperature_c > 25.0:
                    thermal_factor += (self.temperature_c - 25.0) / 200.0
                deflection_mm = (
                    5.0 * (self.w_kN_m * 1000.0) * span_m**4 / (384.0 * (E_MPa * 1_000_000.0) * inertia_m4)
                ) * 1000.0 * thermal_factor
                return {
                    "name": self.name,
                    "w_kN_m": self.w_kN_m,
                    "dynamic_factor": self.dynamic_factor,
                    "temperature_c": self.temperature_c,
                    "M_max_kNm": moment,
                    "V_max_kN": shear,
                    "sigma_MPa": stress,
                    "dcr": dcr,
                    "deflection_mm": deflection_mm,
                }
        """,
    ),
    _cell(
        "c04_runtime_results",
        "code",
        """
        selected_section = SECTIONS_DF.sort_values("I_mm4", ascending=False).iloc[0]
        section_b_m = float(selected_section["b_mm"]) / 1000.0
        section_h_m = float(selected_section["h_mm"]) / 1000.0

        beam_cases = [
            BeamScenario(
                case["name"],
                case["w_kN_m"],
                case["temperature_c"],
                case.get("dynamic_factor", 1.0),
            )
            for case in LOADS["cases"]
        ]

        scenario_rows = [
            case.evaluate(section_b_m, section_h_m, fy_MPa=fy_MPa, E_MPa=E_MPa)
            for case in beam_cases
        ]
        results_df = pd.DataFrame(scenario_rows)
        governing_row = results_df.loc[results_df["dcr"].idxmax()]
        governing_case_name = str(governing_row["name"])
        governing_moment_kNm = float(governing_row["M_max_kNm"])
        governing_dcr = float(governing_row["dcr"])
        governing_deflection_mm = float(governing_row["deflection_mm"])
        average_temperature_c = statistics.mean(results_df["temperature_c"])

        print(results_df[["name", "M_max_kNm", "sigma_MPa", "dcr"]].round(3))
        analysis_trace.append({"governing_case_name": governing_case_name, "governing_dcr": governing_dcr})
        """,
    ),
    _cell(
        PRIMARY_NOTEBOOK_ANALYSIS_CELL_ID,
        "code",
        """
        analysis_formulas = {
            "M_max_kNm": "w_kN_m * span_m**2 / 8.0 * dynamic_factor",
            "V_max_kN": "w_kN_m * span_m / 2.0",
            "sigma_MPa": "(M_max_kNm * 1000.0) / ((section_b_m * section_h_m**2 / 6.0) * 1_000_000.0)",
            "dcr": "sigma_MPa / (0.9 * fy_MPa)",
            "deflection_mm": "5.0 * (w_kN_m * 1000.0) * span_m**4 / (384.0 * (E_MPa * 1_000_000.0) * (section_b_m * section_h_m**3 / 12.0)) * 1000.0",
            "weight_proxy": "section_b_m * section_h_m * span_m * density_kN_m3",
        }
        analysis_current_values = {
            "w_kN_m": float(governing_row["w_kN_m"]),
            "dynamic_factor": float(governing_row["dynamic_factor"]),
            "span_m": span_m,
            "section_b_m": section_b_m,
            "section_h_m": section_h_m,
            "fy_MPa": fy_MPa,
            "E_MPa": E_MPa,
            "density_kN_m3": float((density * 9.80665 * m / s**2).to(kN / m**3).magnitude),
        }
        analysis_checks = [
            {"name": "stress_limit", "lhs": "sigma_MPa", "op": "<=", "rhs": 0.66 * fy_MPa},
            {"name": "dcr_limit", "lhs": "dcr", "op": "<=", "rhs": 1.0},
            {"name": "deflection_limit", "lhs": "deflection_mm", "op": "<=", "rhs": span_m * 1000.0 / 360.0},
        ]
        analysis_objective = {
            "targets": [
                {"name": "weight_proxy", "goal": "min", "weight": 1.0},
                {"name": "dcr", "goal": "target", "target": 0.90, "weight": 0.4},
            ]
        }
        analysis_variables = [
            {"name": "section_b_m", "min": 0.30, "max": 0.55, "initial": section_b_m},
            {"name": "section_h_m", "min": 0.60, "max": 0.90, "initial": section_h_m},
        ]
        analysis_constraints = [
            {"name": "stress_ok", "lhs": "sigma_MPa", "op": "<=", "rhs": 0.66 * fy_MPa},
            {"name": "dcr_ok", "lhs": "dcr", "op": "<=", "rhs": 1.0},
        ]
        analysis_outputs = ["M_max_kNm", "sigma_MPa", "dcr", "deflection_mm"]
        analysis_baseline = {"name": "baseline", "values": {"w_kN_m": analysis_current_values["w_kN_m"]}}
        analysis_candidates = [
            {"name": "plus_5", "values": {"w_kN_m": analysis_current_values["w_kN_m"] * 1.05}},
            {"name": "minus_5", "values": {"w_kN_m": analysis_current_values["w_kN_m"] * 0.95}},
        ]

        print("analysis payloads ready", governing_case_name, governing_dcr)
        """,
    ),
    _cell(
        "m06_doc_report",
        "markdown",
        """
        ## DOCX report section

        Las siguientes celdas generan un informe DOCX rico para validar
        artefactos, latex, captions, figuras, imagenes y referencias.
        """,
    ),
    _cell(
        "c06_doc_report_cover",
        "code",
        """
        doc_reset(hard=True)

        with build_doc(block_id="cover", order=10) as builder:
            builder.metadata(
                title="MCP Torture Report",
                subject="Notebook-first exhaustive probe",
            )
            builder.heading("MCP Torture Report", level=1)
            builder.text("Informe tecnico generado para llevar Inspyro y su MCP al limite.")
            builder.text(f"Caso gobernante: {governing_case_name}.")
            builder.table_of_contents(depth=2)

        with build_doc(block_id="executive_summary", order=20) as builder:
            builder.heading("Resumen ejecutivo", level=2)
            builder.text(
                f"El momento gobernante es {governing_moment_kNm:.2f} kN m y la demanda/capacidad alcanza {governing_dcr:.3f}."
            )
            builder.math_latex(r"M_{max} = \\frac{wL^2}{8}", label="eq:flexion", number=True)
            builder.caption(
                "Resumen del caso base",
                label="cap:global-summary",
                number=True,
                caption_label="Figura",
            )
        """,
    ),
    _cell(
        "c07_doc_report_tables",
        "code",
        """
        with build_doc(block_id="tables", order=30) as builder:
            builder.heading("Variables y tablas", level=2)
            builder.text("La ecuacion ")
            builder.reference("eq:flexion")
            builder.text(" gobierna el dimensionamiento de flexion.")

            paragraph = builder.document.add_paragraph("La verificacion inline usa ")
            paragraph._p.append(builder.create_math_latex_element(r"\\sigma = \\frac{M y}{I}"))
            paragraph.add_run(" para traducir momento a esfuerzo.")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            builder.table(
                [
                    ["Caso gobernante", governing_case_name],
                    ["Momento maximo", round(governing_moment_kNm, 3)],
                    ["D/C", round(governing_dcr, 4)],
                    ["Flecha", round(governing_deflection_mm, 3)],
                ],
                headers=["Magnitud", "Valor"],
                caption="Resumen manual del caso gobernante",
                label="tbl:manual-summary",
            )

            builder.dataframe(
                results_df[["name", "M_max_kNm", "sigma_MPa", "dcr", "deflection_mm"]].round(4),
                index=False,
                caption="Escenarios evaluados",
                label="tbl:scenarios",
            )
            builder.document.add_page_break()
        """,
    ),
    _cell(
        "c08_doc_report_figures",
        "code",
        """
        x_axis = np.linspace(0.0, span_m, 60)
        response_curve = (analysis_current_values["w_kN_m"] * x_axis * (span_m - x_axis)) / 2.0

        fig, ax = plt.subplots(figsize=(7.0, 4.0))
        ax.plot(x_axis, response_curve, color="#1F4E79", linewidth=2.0)
        ax.set_title("Curva de momento flector")
        ax.set_xlabel("x [m]")
        ax.set_ylabel("M [kN m]")
        ax.grid(True, alpha=0.3)

        figure_path = OUTPUT_DIR / "moment_curve.png"
        fig.savefig(figure_path, dpi=180, bbox_inches="tight")

        with build_doc(block_id="figures", order=40) as builder:
            builder.heading("Figuras y anexos", level=2)
            builder.figure(fig, caption="Curva de momento del caso base", label="fig:moment")
            builder.image(
                str(figure_path),
                width=5.4,
                caption="Imagen exportada a PNG desde el notebook",
                label="fig:png",
            )
            builder.text("La Figura ")
            builder.reference("fig:moment")
            builder.text(" y la Figura ")
            builder.reference("fig:png")
            builder.text(" completan el paquete de evidencia grafica.")
            builder.dataframe(
                SECTIONS_DF,
                index=False,
                caption="Banco de secciones disponibles",
                label="tbl:sections",
            )

        plt.close(fig)
        """,
    ),
    _cell(
        PRIMARY_NOTEBOOK_LONG_CELL_ID,
        "code",
        """
        import time

        for tick in range(30):
            print(f"interruptible_tick={tick}")
            time.sleep(1.0)
        """,
    ),
]

SECONDARY_NOTEBOOK_SPEC = [
    _cell(
        "sm00_overview",
        "markdown",
        """
        # Secondary MCP Torture Notebook

        Notebook liviano para probar aislamiento de kernels y preservacion de outputs.
        """,
    ),
    _cell(
        SECONDARY_NOTEBOOK_CODE_CELL_ID,
        "code",
        """
        from pathlib import Path

        workspace = Path().resolve()
        print(f"secondary_workspace={workspace}")
        quick_values = [n * n for n in range(6)]
        quick_total = sum(quick_values)
        print(f"quick_total={quick_total}")
        """,
    ),
]


def clone_notebook_spec(spec: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "cell_id": str(cell["cell_id"]),
            "cell_type": str(cell["cell_type"]),
            "source": str(cell["source"]),
        }
        for cell in spec
    ]


def build_mutated_primary_spec(cell_id: str, replacement_source: str) -> list[dict[str, Any]]:
    mutated = clone_notebook_spec(PRIMARY_NOTEBOOK_SPEC)
    for cell in mutated:
        if cell["cell_id"] == cell_id:
            cell["source"] = dedent(replacement_source).strip() + "\n"
            break
    return mutated


def ensure_template_fixture(path: Path) -> Path:
    if path.exists():
        return path

    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "MCP Torture Template Header"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "MCP Torture Template Footer"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Cambria"
    heading_style.font.size = Pt(18)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    caption_style = document.styles["Caption"]
    caption_style.font.name = "Georgia"
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.font.color.rgb = RGBColor(0x5B, 0x5B, 0x5B)

    document.core_properties.title = "MCP Torture Template Fixture"
    document.core_properties.subject = "Stable fixture for notebook-first MCP torture probe"
    document.core_properties.author = "Inspyro"

    document.add_heading("Template Fixture Body", level=1)
    document.add_paragraph(
        "Stable template fixture with explicit header, footer, Heading 1, Caption and Table Grid coverage."
    )
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Fixture"
    table.cell(1, 1).text = "MCP Torture"
    document.add_paragraph("Tabla 1. Stable template fixture", style="Caption")
    document.save(path)
    return path
