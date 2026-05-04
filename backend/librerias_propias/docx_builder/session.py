"""Gestión del estado DOCX dentro de notebooks.

Este módulo encapsula toda la lógica relacionada con la vida del documento,
incluyendo la reutilización de celdas, la limpieza de contenido y el registro
de eventos. El objetivo es ofrecer un punto único y predecible para administrar
el estado, reduciendo el riesgo de duplicados o leaks de objetos XML.
"""

from __future__ import annotations

import base64
import copy
import inspect
import io
import json
import logging
import os
import re
import sys
import threading
import uuid
import xml.etree.ElementTree as ET
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, Generator, Iterable, Iterator, List, Optional
from urllib.parse import quote
from urllib.parse import urlsplit
import zipfile

_logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml import OxmlElement, parse_xml  # type: ignore
    from docx.oxml.shared import qn  # type: ignore
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - se valida en tiempo de ejecución
    raise RuntimeError(
        "python-docx es requerido para la nueva API DOCX. Instálalo con 'pip install python-docx'."
    ) from exc


try:
    from .parser import EnhancedMathParser
except ImportError:  # pragma: no cover - fallback cuando no se empaqueta el parser
    EnhancedMathParser = None  # type: ignore

try:
    from .latex_math import LatexMathConverter
except ImportError:  # pragma: no cover - fallback cuando no se empaqueta el conversor LaTeX
    LatexMathConverter = None  # type: ignore

from .utils import (
    iter_relationship_ids_in_xml_element,
    rewrite_relationship_ids_in_xml_element,
    validate_docx_package_bytes,
)
from app.services.docx_sanitizer import sanitize_docx_b64_for_delivery, sanitize_docx_bytes_for_delivery


ISO_EPOCH = datetime(2000, 1, 1, 0, 0, 0)
_TABLE_STYLE_LOOK_KEYS = (
    "firstRow",
    "lastRow",
    "firstColumn",
    "lastColumn",
    "noHBand",
    "noVBand",
)
_DOCX_MAIN_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DOCX_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_NS = {
    "w": _DOCX_MAIN_NS,
    "r": _DOCX_REL_NS,
}
_DOCX_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_DOCX_PROVENANCE_NOTEBOOK_RE = re.compile(r"^<inspyro-notebook:(?P<cell_id>[^>]+)>$")
_DOCX_PROVENANCE_PATH_HINTS = (
    "backend/librerias_propias/docx_builder/",
    "backend/librerias_propias/math_to_docx.py",
    "backend/app/services/notebook_service.py",
)
_DOCX_VISIBLE_INLINE_TAGS = {
    f"{{{_DOCX_MAIN_NS}}}r",
    f"{{{_DOCX_MAIN_NS}}}drawing",
    f"{{{_DOCX_MAIN_NS}}}pict",
    f"{{{_DOCX_MAIN_NS}}}object",
    f"{{{_DOCX_MAIN_NS}}}fldSimple",
    f"{{{_DOCX_MATH_NS}}}oMath",
    f"{{{_DOCX_MATH_NS}}}oMathPara",
}
_DOCX_VISIBLE_CONTAINER_TAGS = {
    qn("w:p"),
    qn("w:tbl"),
}
_PARAGRAPH_HYPERLINK_CHILD_TAGS = {
    f"{{{_DOCX_MAIN_NS}}}r",
    f"{{{_DOCX_MAIN_NS}}}bookmarkStart",
    f"{{{_DOCX_MAIN_NS}}}bookmarkEnd",
    f"{{{_DOCX_MAIN_NS}}}proofErr",
    f"{{{_DOCX_MAIN_NS}}}permStart",
    f"{{{_DOCX_MAIN_NS}}}permEnd",
    f"{{{_DOCX_MAIN_NS}}}fldSimple",
    f"{{{_DOCX_MATH_NS}}}oMath",
    f"{{{_DOCX_MATH_NS}}}oMathPara",
}
_PARAGRAPH_NON_HYPERLINK_MARKER_TAGS = {
    f"{{{_DOCX_MAIN_NS}}}bookmarkStart",
    f"{{{_DOCX_MAIN_NS}}}bookmarkEnd",
    f"{{{_DOCX_MAIN_NS}}}proofErr",
    f"{{{_DOCX_MAIN_NS}}}permStart",
    f"{{{_DOCX_MAIN_NS}}}permEnd",
}


def _frame_location_signature(frame_info: Dict[str, Any]) -> tuple[Optional[str], Optional[str], Optional[int]]:
    return (
        frame_info.get("file_path"),
        frame_info.get("notebook_cell_id"),
        frame_info.get("line"),
    )


def _path_key(path: Iterable[int]) -> str:
    items = [str(int(idx)) for idx in path]
    return "/".join(items)


def _normalize_preview_text(value: Any, *, limit: int = 160) -> Optional[str]:
    text = str(value or "").strip()
    if not text:
        return None
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return f"{compact[: limit - 1].rstrip()}…"


def _normalize_workspace_path(raw_path: Any) -> Optional[str]:
    if raw_path in (None, ""):
        return None
    try:
        return str(Path(str(raw_path)).expanduser().resolve())
    except Exception:
        text = str(raw_path).strip()
        return text or None


def _build_provenance_open_url(provenance_id: str) -> str:
    safe_id = quote(str(provenance_id), safe="")
    relative_path = f"/api/docx/provenance/open?provenance_id={safe_id}"
    backend_url = str(os.getenv("INSPYRO_BACKEND_URL") or "").strip()
    if backend_url:
        try:
            parsed = urlsplit(backend_url if "://" in backend_url else f"http://{backend_url}")
            if parsed.netloc:
                return f"{parsed.scheme or 'http'}://{parsed.netloc}{relative_path}"
        except Exception:
            pass
    backend_port = str(os.getenv("INSPYRO_BACKEND_PORT") or "8000").strip() or "8000"
    backend_host = str(os.getenv("INSPYRO_BACKEND_HOST") or "").strip() or "127.0.0.1"
    if backend_host in {"0.0.0.0", "::", "[::]"}:
        backend_host = "127.0.0.1"
    return f"http://{backend_host}:{backend_port}{relative_path}"


def _is_ipykernel_temp_path(path: Any) -> bool:
    normalized = str(path or "").replace("\\", "/").lower()
    if not normalized:
        return False
    return (
        "/appdata/local/temp/ipykernel_" in normalized
        or "/tmp/ipykernel_" in normalized
        or "/var/folders/" in normalized and "/ipykernel_" in normalized
    )


def _select_useful_callsite(stack: List[Dict[str, Any]]) -> Dict[str, Any]:
    if not stack:
        return {"file_path": None, "notebook_cell_id": None, "line": None}

    for frame in reversed(stack):
        if frame.get("notebook_cell_id"):
            return frame
        file_path = frame.get("file_path")
        if file_path and not _is_ipykernel_temp_path(file_path):
            return frame

    return stack[-1]

# ---------------------------------------------------------------------------
# Default definitions for styles that the builder API requires.
# These are only created when the style does NOT already exist in the
# loaded template - the template version always takes precedence.
# ---------------------------------------------------------------------------
_REQUIRED_STYLE_DEFAULTS = {
    "Heading 1": {"font_size_pt": 16, "bold": True, "space_before_pt": 12, "space_after_pt": 6},
    "Heading 2": {"font_size_pt": 14, "bold": True, "space_before_pt": 10, "space_after_pt": 4},
    "Heading 3": {"font_size_pt": 12, "bold": True, "space_before_pt": 8, "space_after_pt": 4},
    "Heading 4": {"font_size_pt": 11, "bold": True, "italic": True, "space_before_pt": 8, "space_after_pt": 4},
    "Heading 5": {"font_size_pt": 11, "bold": True, "space_before_pt": 6, "space_after_pt": 2},
    "Heading 6": {"font_size_pt": 11, "italic": True, "space_before_pt": 6, "space_after_pt": 2},
    "List Bullet": {"font_size_pt": 11, "left_indent_pt": 36, "first_line_indent_pt": -18, "seed_numbering_style": "List Bullet"},
    "List Number": {"font_size_pt": 11, "left_indent_pt": 36, "first_line_indent_pt": -18, "seed_numbering_style": "List Number"},
    "Code": {"font_size_pt": 10},
    "Caption": {"font_size_pt": 9, "italic": True, "space_after_pt": 6},
}

_SEMANTIC_STYLE_SLOT_NAMES = (
    "body",
    "heading_1",
    "heading_2",
    "heading_3",
    "heading_4",
    "heading_5",
    "heading_6",
    "list_bullet",
    "list_number",
    "caption",
    "code",
    "table_default",
)


def _coerce_optional_int(value: Any) -> Optional[int]:
    if value in (None, ""):
        return None
    try:
        return int(round(float(value)))
    except (TypeError, ValueError):
        return None


def _coerce_boolish(value: Any) -> Optional[bool]:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return None


def _normalize_template_required_style_defaults(data: Any) -> Dict[str, Dict[str, Any]]:
    if not isinstance(data, dict):
        return {}

    normalized: Dict[str, Dict[str, Any]] = {}
    for style_name, raw_entry in data.items():
        style_name_text = str(style_name).strip() if style_name not in (None, "") else ""
        if not style_name_text:
            continue
        entry = raw_entry if isinstance(raw_entry, dict) else {}
        normalized_entry: Dict[str, Any] = {}
        for key, value in entry.items():
            if value in (None, ""):
                continue
            normalized_entry[str(key)] = copy.deepcopy(value)
        if normalized_entry:
            normalized[style_name_text] = normalized_entry
    return normalized


def _normalize_template_semantic_style_slots(data: Any) -> Dict[str, Dict[str, Any]]:
    if not isinstance(data, dict):
        return {}

    normalized: Dict[str, Dict[str, Any]] = {}
    for slot_name, raw_entry in data.items():
        slot_name_text = str(slot_name).strip() if slot_name not in (None, "") else ""
        if not slot_name_text or slot_name_text not in _SEMANTIC_STYLE_SLOT_NAMES:
            continue
        entry = raw_entry if isinstance(raw_entry, dict) else {}
        normalized_entry: Dict[str, Any] = {"slot_name": slot_name_text}
        for key in ("category", "selection_key", "style_id", "style_name", "display_name", "style_type"):
            value = entry.get(key)
            if value in (None, ""):
                continue
            normalized_entry[key] = copy.deepcopy(value)
        normalized[slot_name_text] = normalized_entry
    return normalized


def _normalize_template_table_style_defaults(data: Any) -> Dict[str, Dict[str, Any]]:
    if not isinstance(data, dict):
        return {}

    normalized: Dict[str, Dict[str, Any]] = {}
    for key, raw_entry in data.items():
        entry = raw_entry if isinstance(raw_entry, dict) else {}
        style_id = str(entry.get("style_id") or key).strip() if entry.get("style_id") or key else ""
        if not style_id:
            continue

        normalized_entry: Dict[str, Any] = {"style_id": style_id}
        style_name = entry.get("style_name")
        if style_name not in (None, ""):
            normalized_entry["style_name"] = str(style_name)

        layout_type = entry.get("layout_type")
        if layout_type not in (None, ""):
            normalized_entry["layout_type"] = str(layout_type).strip().lower()

        width_type = entry.get("width_type")
        if width_type not in (None, ""):
            normalized_entry["width_type"] = str(width_type).strip().lower()

        width_value = _coerce_optional_int(entry.get("width_value"))
        if width_value is not None:
            normalized_entry["width_value"] = width_value

        look = entry.get("look")
        if isinstance(look, dict):
            normalized_look: Dict[str, bool] = {}
            for look_key in _TABLE_STYLE_LOOK_KEYS:
                bool_value = _coerce_boolish(look.get(look_key))
                if bool_value is not None:
                    normalized_look[look_key] = bool_value
            if normalized_look:
                normalized_entry["look"] = normalized_look

        if any(key in normalized_entry for key in ("layout_type", "width_type", "width_value", "look")):
            normalized[style_id] = normalized_entry

    return normalized


class DocxSessionError(RuntimeError):
    """Error base para operaciones de sesión."""


class StrictModeError(DocxSessionError):
    """Se lanza cuando el modo estricto está activo y ocurre un fallo."""


class DocxValidationError(DocxSessionError):
    """Se lanza cuando el documento DOCX tiene XML malformado.
    
    Esta excepción indica que el contenido generado no puede ser abierto
    correctamente por Word/LibreOffice y causaría un timeout en la conversión PDF.
    """
    def __init__(self, message: str, details: dict = None):
        super().__init__(message)
        self.details = details or {}
        self.block_id = details.get("block_id") if details else None
        self.errors = details.get("errors", []) if details else []


@dataclass
class CellHandle:
    """Representa el contexto activo de un bloque DOCX."""

    session: "DocxSession"
    block_id: str
    doc: DocumentType
    parser: Optional[EnhancedMathParser]
    latex_converter: Optional[LatexMathConverter]

    def register_element(self, element: Any, *, provenance: Optional[Dict[str, Any]] = None) -> None:
        self.session._register_element(self.block_id, element, provenance=provenance)

    def log_event(self, kind: str, message: str, **data: Any) -> None:
        self.session.log_event(kind, message, **data)


class DocxSession:
    """Administra un documento DOCX compartido dentro de un kernel."""

    _DOC_KEY = "__DOCX_DOC"
    _PARSER_KEY = "__DOCX_PARSER"
    _LATEX_CONVERTER_KEY = "__DOCX_LATEX_CONVERTER"
    _LATEX_CONVERTER_PROBED_KEY = "__DOCX_LATEX_CONVERTER_PROBED"
    _CELL_ITEMS_KEY = "__DOCX_CELL_ITEMS"
    _CELL_CURSOR_KEY = "__DOCX_CELL_CURSOR"
    _CELL_META_KEY = "__DOCX_CELL_META"
    _ACTIVE_CELL_KEY = "__DOCX_ACTIVE_CELL"
    _EVENT_LOG_KEY = "__DOCX_EVENT_LOG"
    _STRICT_MODE_KEY = "__DOCX_STRICT_MODE"
    _LABELS_KEY = "__DOCX_LABELS"
    _EQ_COUNTER_KEY = "__DOCX_EQ_COUNTER"
    _FIG_COUNTER_KEY = "__DOCX_IMAGE_COUNTER"
    _SEQ_COUNTERS_KEY = "__DOCX_SEQUENCE_COUNTERS"
    _BOOKMARK_COUNTER_KEY = "__DOCX_BOOKMARK_COUNTER"
    _DOC_INITIALIZED_KEY = "__DOCX_DOC_INITIALIZED__"
    _CELL_ORDER_KEY = "__DOCX_CELL_ORDER"
    _CELL_ORDER_VALUES_KEY = "__DOCX_CELL_ORDER_VALUES"
    _CELL_SERIALIZED_KEY = "__DOCX_CELL_SERIALIZED"
    _PENDING_PROVENANCE_KEY = "__DOCX_PENDING_PROVENANCE"
    _NOTEBOOK_GROUPS_KEY = "__DOCX_NOTEBOOK_CELL_GROUPS"
    _NOTEBOOK_ACTIVE_KEY = "__DOCX_NOTEBOOK_CELL_ACTIVE"
    _BLOCK_TO_NOTEBOOK_KEY = "__DOCX_BLOCK_TO_NOTEBOOK"
    _TEMPLATE_PATH_KEY = "__DOCX_TEMPLATE_PATH"
    _TEMPLATE_TABLE_STYLE_DEFAULTS_KEY = "__DOCX_TEMPLATE_TABLE_STYLE_DEFAULTS"
    _TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY = "__DOCX_TEMPLATE_REQUIRED_STYLE_DEFAULTS"
    _TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY = "__DOCX_TEMPLATE_SEMANTIC_STYLE_SLOTS"
    _TEMPLATE_SECTION_REFS_KEY = "__DOCX_TEMPLATE_SECTION_REFS"

    _CELL_KEYS = (
        _CELL_ITEMS_KEY,
        _CELL_META_KEY,
        _CELL_CURSOR_KEY,
        _ACTIVE_CELL_KEY,
        _CELL_ORDER_KEY,
    )

    def __init__(self, namespace: Dict[str, Any]):
        if not isinstance(namespace, dict):  # pragma: no cover - sanity check
            raise TypeError("Se esperaba un namespace dict para inicializar DocxSession")
        self.ns = namespace
        self.main_ns = self._get_main_namespace()
        self.doc = self._ensure_document()
        self.parser = self._ensure_parser()
        self.latex_converter = self._ensure_latex_converter()
        self._ensure_structures()
        self._dirty = False  # Flag para reconstrucción diferida

    # ------------------------------------------------------------------
    # Inicialización y utilidades base
    # ------------------------------------------------------------------
    @staticmethod
    def _get_main_namespace() -> Dict[str, Any]:
        main = sys.modules.get("__main__")
        if main and hasattr(main, "__dict__"):
            return getattr(main, "__dict__")  # type: ignore[return-value]
        return {}

    def _ensure_document(self) -> DocumentType:
        doc = self.ns.get(self._DOC_KEY)
        if doc is None:
            template_path = self.ns.get(self._TEMPLATE_PATH_KEY)
            
            if template_path and os.path.exists(template_path):
                try:
                    doc = Document(template_path)
                    self._ensure_template_required_style_defaults(template_path)
                    _logger.info(f"[DOCX Session] Loaded template from: {template_path}")
                except Exception as e:
                    # ✅ Raise exception instead of silent fallback
                    # This ensures users are notified immediately of template issues
                    raise RuntimeError(
                        f"Failed to load DOCX template from '{template_path}': {e}. "
                        f"The template file may be corrupted or inaccessible. "
                        f"Please re-upload the template or clear it to use a blank document."
                    ) from e
            else:
                if template_path:
                    # Template path is set but file doesn't exist
                    _logger.warning(
                        f"[DOCX Session] Template path set to '{template_path}' but file not found. "
                        f"Using blank document. This may indicate the template was deleted externally."
                    )
                self.ns.pop(self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY, None)
                self.ns.pop(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY, None)
                self.ns.pop(self._TEMPLATE_SECTION_REFS_KEY, None)
                if self.main_ns and self.main_ns is not self.ns:
                    self.main_ns.pop(self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY, None)
                    self.main_ns.pop(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY, None)
                    self.main_ns.pop(self._TEMPLATE_SECTION_REFS_KEY, None)
                doc = Document()
                
            self._apply_deterministic_metadata(doc)
            self._ensure_required_styles(doc)
            self.ns[self._DOC_KEY] = doc
        if self.main_ns is not self.ns and self._DOC_KEY not in self.main_ns:
            self.main_ns[self._DOC_KEY] = doc
        return doc

    def set_template_path(self, path: Optional[str]) -> None:
        """Set the template DOCX path for this session.
        
        The template will be used when creating new documents (after reset).
        
        Args:
            path: Absolute path to a .docx template file, or None to clear.
        """
        self.ns[self._TEMPLATE_PATH_KEY] = path
        self.ns.pop(self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY, None)
        self.ns.pop(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY, None)
        self.ns.pop(self._TEMPLATE_SECTION_REFS_KEY, None)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._TEMPLATE_PATH_KEY] = path
            self.main_ns.pop(self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY, None)
            self.main_ns.pop(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY, None)
            self.main_ns.pop(self._TEMPLATE_SECTION_REFS_KEY, None)
        _logger.info(f"[DOCX Session] Template path set to: {path}")

    def get_template_path(self) -> Optional[str]:
        """Get the current template path, if any."""
        return self.ns.get(self._TEMPLATE_PATH_KEY)

    def set_template_table_style_defaults(self, defaults: Optional[Dict[str, Any]]) -> None:
        normalized = _normalize_template_table_style_defaults(defaults)
        self.ns[self._TEMPLATE_TABLE_STYLE_DEFAULTS_KEY] = normalized
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._TEMPLATE_TABLE_STYLE_DEFAULTS_KEY] = copy.deepcopy(normalized)
        _logger.info(
            "[DOCX Session] Template table style defaults set for %d style(s)",
            len(normalized),
        )

    def get_template_table_style_defaults(self) -> Dict[str, Dict[str, Any]]:
        defaults = self.ns.get(self._TEMPLATE_TABLE_STYLE_DEFAULTS_KEY)
        return _normalize_template_table_style_defaults(defaults)

    def set_template_required_style_defaults(self, defaults: Optional[Dict[str, Any]]) -> None:
        normalized = _normalize_template_required_style_defaults(defaults)
        self.ns[self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY] = normalized
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY] = copy.deepcopy(normalized)
        _logger.info(
            "[DOCX Session] Template required style defaults set for %d style(s)",
            len(normalized),
        )

    def get_template_required_style_defaults(self) -> Dict[str, Dict[str, Any]]:
        defaults = self.ns.get(self._TEMPLATE_REQUIRED_STYLE_DEFAULTS_KEY)
        return _normalize_template_required_style_defaults(defaults)

    def set_template_semantic_style_slots(self, slots: Optional[Dict[str, Any]]) -> None:
        normalized = _normalize_template_semantic_style_slots(slots)
        self.ns[self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY] = normalized
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY] = copy.deepcopy(normalized)
        _logger.info(
            "[DOCX Session] Template semantic style slots set for %d slot(s)",
            len(normalized),
        )

    def get_template_semantic_style_slots(self) -> Dict[str, Dict[str, Any]]:
        slots = self.ns.get(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY)
        return _normalize_template_semantic_style_slots(slots)

    def resolve_style_slot(self, slot_name: Optional[str]) -> Optional[str]:
        slot_name_text = str(slot_name).strip() if slot_name not in (None, "") else ""
        if not slot_name_text:
            return None

        slot_entry = self.get_template_semantic_style_slots().get(slot_name_text)
        if not isinstance(slot_entry, dict):
            return None

        requested_style_id = str(slot_entry.get("style_id") or "").strip()
        requested_style_name = str(slot_entry.get("style_name") or "").strip()
        requested_display_name = str(slot_entry.get("display_name") or "").strip()

        for style in self.doc.styles:
            style_id = getattr(style, "style_id", None)
            if requested_style_id and style_id and str(style_id) == requested_style_id:
                return str(getattr(style, "name", "") or requested_style_name or requested_display_name)

        for candidate_name in (requested_style_name, requested_display_name):
            if not candidate_name:
                continue
            try:
                return str(self.doc.styles[candidate_name].name)
            except Exception:
                continue

        return None

    def _ensure_template_required_style_defaults(self, template_path: Optional[str]) -> Dict[str, Dict[str, Any]]:
        cached = self.get_template_required_style_defaults()
        if cached:
            return copy.deepcopy(cached)
        if not template_path or not os.path.exists(template_path):
            return {}

        defaults: Dict[str, Dict[str, Any]] = {}
        try:
            from app.services import template_service as _template_service

            with open(template_path, "rb") as fh:
                extracted = _template_service.extract_styles_from_docx(fh.read())
            built_defaults = _template_service.build_builder_required_style_defaults(extracted)
            if isinstance(built_defaults, dict):
                defaults = copy.deepcopy(built_defaults)
        except Exception as exc:
            _logger.warning(
                "[DOCX Session] Could not derive template font defaults from '%s': %s",
                template_path,
                exc,
            )

        self.set_template_required_style_defaults(defaults)
        return defaults

    def _load_template_section_references(self, template_path: Optional[str]) -> Dict[str, Any]:
        cached = self.ns.get(self._TEMPLATE_SECTION_REFS_KEY)
        if (
            isinstance(cached, dict)
            and cached.get("template_path") == template_path
            and isinstance(cached.get("refs"), dict)
        ):
            return copy.deepcopy(cached["refs"])

        refs: Dict[str, Any] = {
            "header_refs": [],
            "footer_refs": [],
            "title_pg": None,
        }
        if not template_path or not os.path.exists(template_path):
            cache_payload = {"template_path": template_path, "refs": refs}
            self.ns[self._TEMPLATE_SECTION_REFS_KEY] = copy.deepcopy(cache_payload)
            if self.main_ns and self.main_ns is not self.ns:
                self.main_ns[self._TEMPLATE_SECTION_REFS_KEY] = copy.deepcopy(cache_payload)
            return refs

        try:
            with zipfile.ZipFile(template_path, "r") as zf:
                document_xml = zf.read("word/document.xml")
            root = ET.fromstring(document_xml)
            for sect in root.findall(".//w:sectPr", _DOCX_NS):
                header_refs = [
                    ET.tostring(node, encoding="unicode")
                    for node in sect.findall("w:headerReference", _DOCX_NS)
                ]
                footer_refs = [
                    ET.tostring(node, encoding="unicode")
                    for node in sect.findall("w:footerReference", _DOCX_NS)
                ]
                if not header_refs and not footer_refs:
                    continue
                title_pg = sect.find("w:titlePg", _DOCX_NS)
                refs = {
                    "header_refs": header_refs,
                    "footer_refs": footer_refs,
                    "title_pg": (
                        ET.tostring(title_pg, encoding="unicode")
                        if title_pg is not None
                        else None
                    ),
                }
                break
        except Exception as exc:
            _logger.warning(
                "[DOCX Session] Could not load template section refs from '%s': %s",
                template_path,
                exc,
            )

        cache_payload = {"template_path": template_path, "refs": refs}
        self.ns[self._TEMPLATE_SECTION_REFS_KEY] = copy.deepcopy(cache_payload)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._TEMPLATE_SECTION_REFS_KEY] = copy.deepcopy(cache_payload)
        return copy.deepcopy(refs)

    def _document_has_header_footer_references(self) -> bool:
        body = self.doc._element.body
        return any(
            isinstance(getattr(node, "tag", None), str)
            and (
                node.tag.endswith("}headerReference")
                or node.tag.endswith("}footerReference")
            )
            for node in body.iter()
        )

    def _restore_template_header_footer_references_if_missing(self) -> bool:
        template_path = self.get_template_path()
        if not template_path or not os.path.exists(template_path):
            return False
        if self._document_has_header_footer_references():
            return False

        refs = self._load_template_section_references(template_path)
        header_refs = list(refs.get("header_refs") or [])
        footer_refs = list(refs.get("footer_refs") or [])
        title_pg_xml = refs.get("title_pg")
        if not header_refs and not footer_refs:
            return False

        body = self.doc._element.body
        sectPr_nodes = [child for child in body if child.tag.endswith("}sectPr")]
        if sectPr_nodes:
            sectPr = sectPr_nodes[-1]
        else:
            sectPr = OxmlElement("w:sectPr")
            body.append(sectPr)

        for child in list(sectPr):
            if not isinstance(getattr(child, "tag", None), str):
                continue
            if child.tag.endswith("}headerReference") or child.tag.endswith("}footerReference"):
                sectPr.remove(child)

        insert_at = 0
        for idx, child in enumerate(list(sectPr)):
            local_name = child.tag.rsplit("}", 1)[-1] if isinstance(getattr(child, "tag", None), str) else ""
            if local_name not in {"headerReference", "footerReference"}:
                insert_at = idx
                break
        else:
            insert_at = len(list(sectPr))

        restored_count = 0
        for xml_text in [*header_refs, *footer_refs]:
            try:
                sectPr.insert(insert_at, parse_xml(xml_text.encode("utf-8")))
                insert_at += 1
                restored_count += 1
            except Exception as exc:
                _logger.warning(
                    "[DOCX Session] Could not restore template section ref from '%s': %s",
                    template_path,
                    exc,
                )

        has_title_pg = any(
            isinstance(getattr(child, "tag", None), str) and child.tag.endswith("}titlePg")
            for child in sectPr
        )
        if title_pg_xml and not has_title_pg:
            try:
                sectPr.insert(insert_at, parse_xml(title_pg_xml.encode("utf-8")))
            except Exception as exc:
                _logger.warning(
                    "[DOCX Session] Could not restore titlePg from '%s': %s",
                    template_path,
                    exc,
                )

        if restored_count:
            _logger.info(
                "[DOCX Session] Restored %d template header/footer ref(s) before export",
                restored_count,
            )
            return True
        return False

    def _apply_deterministic_metadata(self, doc: DocumentType) -> None:
        cp = doc.core_properties
        cp.author = "inspyro"
        cp.last_modified_by = "inspyro"
        cp.created = ISO_EPOCH
        cp.modified = ISO_EPOCH
        try:
            cp.last_printed = ISO_EPOCH
        except Exception:  # pragma: no cover - atributo opcional
            pass
        try:
            cp.revision = 1
        except Exception:  # pragma: no cover - puede no estar presente
            pass

    def _ensure_required_styles(self, doc) -> None:
        """Create any builder-required style that is missing from the document.

        Only adds styles that do NOT already exist, so templates always
        take precedence.  Each missing style gets a neutral Word default.
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            return

        styles = doc.styles
        existing_names = set()
        seed_doc = None
        template_style_defaults = self._ensure_template_required_style_defaults(self.get_template_path())
        try:
            for s in styles:
                if hasattr(s, "name") and s.name:
                    existing_names.add(s.name)
        except Exception:
            pass

        for style_name, defaults in _REQUIRED_STYLE_DEFAULTS.items():
            if style_name in existing_names:
                continue
            try:
                merged_defaults = dict(defaults)
                template_defaults = template_style_defaults.get(style_name) if isinstance(template_style_defaults, dict) else None
                if isinstance(template_defaults, dict):
                    for key, value in template_defaults.items():
                        if value not in (None, ""):
                            merged_defaults[key] = value
                new_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                self._apply_required_style_defaults(
                    new_style,
                    merged_defaults,
                    pt_cls=Pt,
                    rgb_cls=RGBColor,
                )

                seed_style_name = merged_defaults.get("seed_numbering_style")
                if seed_style_name:
                    if seed_doc is None:
                        seed_doc = Document()
                    if not self._clone_seed_list_numbering(
                        doc,
                        new_style,
                        seed_doc,
                        str(seed_style_name),
                    ):
                        _logger.warning(
                            "[DOCX Session] Could not clone default numbering for '%s'. "
                            "Builder fallback will render explicit markers if needed.",
                            style_name,
                        )

                _logger.debug(
                    "[DOCX Session] Created missing style '%s' with Word defaults",
                    style_name,
                )
                existing_names.add(style_name)
            except Exception as exc:
                _logger.warning(
                    "[DOCX Session] Could not create fallback style '%s': %s",
                    style_name, exc,
                )

    @staticmethod
    def _apply_required_style_defaults(style_obj, defaults: Dict[str, Any], *, pt_cls, rgb_cls) -> None:
        font_name = defaults.get("font_name")
        if font_name:
            style_obj.font.name = str(font_name)

        font_size_pt = defaults.get("font_size_pt")
        if font_size_pt is not None:
            style_obj.font.size = pt_cls(float(font_size_pt))

        font_color_rgb = defaults.get("font_color_rgb")
        if isinstance(font_color_rgb, tuple) and len(font_color_rgb) == 3:
            try:
                style_obj.font.color.rgb = rgb_cls(*font_color_rgb)
            except Exception:
                pass

        if defaults.get("bold") is not None:
            style_obj.font.bold = bool(defaults.get("bold"))

        if defaults.get("italic") is not None:
            style_obj.font.italic = bool(defaults.get("italic"))

        pf = style_obj.paragraph_format
        left_indent_pt = defaults.get("left_indent_pt")
        if left_indent_pt is not None:
            pf.left_indent = pt_cls(float(left_indent_pt))

        first_line_indent_pt = defaults.get("first_line_indent_pt")
        if first_line_indent_pt is not None:
            pf.first_line_indent = pt_cls(float(first_line_indent_pt))

        space_before_pt = defaults.get("space_before_pt")
        if space_before_pt is not None:
            pf.space_before = pt_cls(float(space_before_pt))

        space_after_pt = defaults.get("space_after_pt")
        if space_after_pt is not None:
            pf.space_after = pt_cls(float(space_after_pt))

    def _clone_seed_list_numbering(
        self,
        target_doc: DocumentType,
        style_obj,
        seed_doc: DocumentType,
        seed_style_name: str,
    ) -> bool:
        try:
            seed_style = seed_doc.styles[seed_style_name]
        except KeyError:
            return False

        seed_num_pr = self._extract_style_num_pr(seed_style.element)
        if seed_num_pr is None:
            return False

        seed_num_id_elem = seed_num_pr.find(qn("w:numId"))
        if seed_num_id_elem is None:
            return False

        seed_num_id = seed_num_id_elem.get(qn("w:val"))
        if not seed_num_id:
            return False

        try:
            target_numbering_root = target_doc.part.numbering_part._element
            seed_numbering_root = seed_doc.part.numbering_part._element
        except Exception:
            return False

        seed_num = self._find_numbering_element(seed_numbering_root, "num", "numId", seed_num_id)
        if seed_num is None:
            return False

        seed_abstract_id_elem = seed_num.find(qn("w:abstractNumId"))
        if seed_abstract_id_elem is None:
            return False

        seed_abstract_id = seed_abstract_id_elem.get(qn("w:val"))
        if not seed_abstract_id:
            return False

        seed_abstract = self._find_numbering_element(
            seed_numbering_root,
            "abstractNum",
            "abstractNumId",
            seed_abstract_id,
        )
        if seed_abstract is None:
            return False

        new_abstract_id = self._next_numbering_id(target_numbering_root, "abstractNum", "abstractNumId")
        new_num_id = self._next_numbering_id(target_numbering_root, "num", "numId")

        new_abstract = copy.deepcopy(seed_abstract)
        new_abstract.set(qn("w:abstractNumId"), str(new_abstract_id))

        new_num = copy.deepcopy(seed_num)
        new_num.set(qn("w:numId"), str(new_num_id))
        new_num_abstract_id = new_num.find(qn("w:abstractNumId"))
        if new_num_abstract_id is None:
            return False
        new_num_abstract_id.set(qn("w:val"), str(new_abstract_id))

        target_numbering_root.append(new_abstract)
        target_numbering_root.append(new_num)

        target_p_pr = self._ensure_style_paragraph_properties(style_obj.element)
        existing_num_pr = target_p_pr.find(qn("w:numPr"))
        if existing_num_pr is not None:
            target_p_pr.remove(existing_num_pr)

        new_num_pr = copy.deepcopy(seed_num_pr)
        new_num_pr_num_id = new_num_pr.find(qn("w:numId"))
        if new_num_pr_num_id is None:
            return False
        new_num_pr_num_id.set(qn("w:val"), str(new_num_id))
        target_p_pr.append(new_num_pr)
        return True

    @staticmethod
    def _extract_style_num_pr(style_element):
        p_pr = style_element.find(qn("w:pPr"))
        if p_pr is None:
            return None
        return p_pr.find(qn("w:numPr"))

    @staticmethod
    def _ensure_style_paragraph_properties(style_element):
        p_pr = style_element.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            style_element.append(p_pr)
        return p_pr

    @staticmethod
    def _find_numbering_element(numbering_root, element_name: str, attr_name: str, attr_value: str):
        attr_qn = qn(f"w:{attr_name}")
        for child in numbering_root.findall(qn(f"w:{element_name}")):
            if child.get(attr_qn) == str(attr_value):
                return child
        return None

    @staticmethod
    def _next_numbering_id(numbering_root, element_name: str, attr_name: str) -> int:
        attr_qn = qn(f"w:{attr_name}")
        values: List[int] = []
        for child in numbering_root.findall(qn(f"w:{element_name}")):
            raw_value = child.get(attr_qn)
            if raw_value in (None, ""):
                continue
            try:
                values.append(int(raw_value))
            except (TypeError, ValueError):
                continue
        return (max(values) + 1) if values else 1

    def _ensure_parser(self) -> Optional[EnhancedMathParser]:
        parser = self.ns.get(self._PARSER_KEY)
        if parser is None and EnhancedMathParser is not None:
            parser = EnhancedMathParser()
            self.ns[self._PARSER_KEY] = parser
        if self.main_ns is not self.ns and parser is not None and self._PARSER_KEY not in self.main_ns:
            self.main_ns[self._PARSER_KEY] = parser
        return parser

    def _ensure_latex_converter(self) -> Optional[LatexMathConverter]:
        converter = self.ns.get(self._LATEX_CONVERTER_KEY)
        if converter is None and LatexMathConverter is not None:
            converter = LatexMathConverter()
            self.ns[self._LATEX_CONVERTER_KEY] = converter

        if converter is not None and not self.ns.get(self._LATEX_CONVERTER_PROBED_KEY):
            self.ns[self._LATEX_CONVERTER_PROBED_KEY] = True
            try:
                runtime = converter.describe_runtime()
                if runtime.available:
                    details = f"v{runtime.version}" if runtime.version else "available"
                    location = runtime.texmath_path or runtime.engine or "internal"
                    _logger.info("DOCX LaTeX converter ready via %s (%s)", location, details)
                else:
                    _logger.warning("DOCX LaTeX converter unavailable: %s", runtime.reason)
            except Exception as exc:
                _logger.warning("DOCX LaTeX converter probe failed: %s", exc)

        return converter

    def _ensure_structures(self) -> None:
        self.ns.setdefault(self._CELL_ITEMS_KEY, {})
        self.ns.setdefault(self._CELL_CURSOR_KEY, {})
        self.ns.setdefault(self._CELL_META_KEY, {})
        self.ns.setdefault(self._PENDING_PROVENANCE_KEY, None)
        self.ns.setdefault(self._CELL_ORDER_KEY, [])
        self.ns.setdefault(self._EVENT_LOG_KEY, [])
        self.ns.setdefault(self._LABELS_KEY, {})
        self.ns.setdefault(self._EQ_COUNTER_KEY, 0)
        self.ns.setdefault(self._FIG_COUNTER_KEY, 0)
        self.ns.setdefault(self._SEQ_COUNTERS_KEY, {})
        self.ns.setdefault(self._BOOKMARK_COUNTER_KEY, 1)
        self.ns.setdefault(self._CELL_ORDER_VALUES_KEY, {})
        self.ns.setdefault(self._CELL_SERIALIZED_KEY, {})
        self.ns.setdefault(self._NOTEBOOK_GROUPS_KEY, {})
        self.ns.setdefault(self._NOTEBOOK_ACTIVE_KEY, {})
        self.ns.setdefault(self._BLOCK_TO_NOTEBOOK_KEY, {})
        self.ns.setdefault(self._TEMPLATE_TABLE_STYLE_DEFAULTS_KEY, {})
        self.ns.setdefault(self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY, {})
        if self._DOC_INITIALIZED_KEY not in self.ns:
            self.ns[self._DOC_INITIALIZED_KEY] = False
        if self._STRICT_MODE_KEY not in self.ns:
            self.ns[self._STRICT_MODE_KEY] = False

        if self.main_ns and self.main_ns is not self.ns:
            for key in (
                self._CELL_ITEMS_KEY,
                self._CELL_CURSOR_KEY,
                self._CELL_META_KEY,
                self._PENDING_PROVENANCE_KEY,
                self._EVENT_LOG_KEY,
                self._CELL_ORDER_KEY,
                self._LABELS_KEY,
                self._EQ_COUNTER_KEY,
                self._FIG_COUNTER_KEY,
                self._SEQ_COUNTERS_KEY,
                self._BOOKMARK_COUNTER_KEY,
                self._CELL_ORDER_VALUES_KEY,
                self._CELL_SERIALIZED_KEY,
                self._NOTEBOOK_GROUPS_KEY,
                self._NOTEBOOK_ACTIVE_KEY,
                self._BLOCK_TO_NOTEBOOK_KEY,
                self._TEMPLATE_TABLE_STYLE_DEFAULTS_KEY,
                self._TEMPLATE_SEMANTIC_STYLE_SLOTS_KEY,
                self._STRICT_MODE_KEY,
                self._DOC_INITIALIZED_KEY,
            ):
                if key not in self.main_ns:
                    self.main_ns[key] = self.ns[key]

    # ------------------------------------------------------------------
    # Gestión de celdas
    # ------------------------------------------------------------------
    def _resolve_block_id(self, block_id: Optional[str]) -> str:
        if block_id:
            return str(block_id)
        auto = self.ns.get(self._ACTIVE_CELL_KEY)
        if auto:
            return str(auto)
        # fallback: generar ID estable por orden de creación
        counter = len(self.ns.setdefault(self._CELL_ITEMS_KEY, {})) + 1
        generated = f"cell-{counter}"
        self.ns[self._ACTIVE_CELL_KEY] = generated
        return generated

    def _cell_meta_map(self) -> Dict[str, Dict[str, Any]]:
        return self.ns[self._CELL_META_KEY]

    def _get_block_owner_notebook_cell(self, block_id: str) -> Optional[str]:
        block_map = self.ns.setdefault(self._BLOCK_TO_NOTEBOOK_KEY, {})
        owner = block_map.get(str(block_id))
        if owner in (None, ""):
            return None
        return str(owner)

    def _collect_user_provenance_stack(self) -> List[Dict[str, Any]]:
        frames: List[Dict[str, Any]] = []
        cwd_path = (_normalize_workspace_path(os.getcwd()) or "").replace("\\", "/").lower()
        base_prefix = (_normalize_workspace_path(sys.base_prefix) or "").replace("\\", "/").lower()
        for frame_info in inspect.stack()[2:]:
            filename = str(frame_info.filename or "")
            normalized = filename.replace("\\", "/")
            if any(marker in normalized for marker in _DOCX_PROVENANCE_PATH_HINTS):
                continue

            notebook_match = _DOCX_PROVENANCE_NOTEBOOK_RE.match(filename)
            if notebook_match:
                frame_payload = {
                    "file_path": None,
                    "notebook_cell_id": notebook_match.group("cell_id"),
                    "line": int(frame_info.lineno or 0) or None,
                }
            else:
                if not filename or filename.startswith("<"):
                    continue
                normalized_path = _normalize_workspace_path(filename)
                normalized_cmp = (normalized_path or "").replace("\\", "/").lower()
                if "site-packages" in normalized_cmp or "dist-packages" in normalized_cmp:
                    continue
                if base_prefix and normalized_cmp.startswith(base_prefix) and (not cwd_path or not normalized_cmp.startswith(cwd_path)):
                    continue
                frame_payload = {
                    "file_path": normalized_path,
                    "notebook_cell_id": None,
                    "line": int(frame_info.lineno or 0) or None,
                }

            if frames and _frame_location_signature(frames[-1]) == _frame_location_signature(frame_payload):
                continue
            frames.append(frame_payload)
        return frames

    def _capture_provenance_location(self) -> Dict[str, Any]:
        stack = self._collect_user_provenance_stack()
        callsite = _select_useful_callsite(stack)
        exact = stack[0] if stack else callsite
        return {
            "file_path": callsite.get("file_path"),
            "notebook_cell_id": callsite.get("notebook_cell_id"),
            "line": callsite.get("line"),
            "exact_file_path": exact.get("file_path"),
            "exact_notebook_cell_id": exact.get("notebook_cell_id"),
            "exact_line": exact.get("line"),
            "user_stack": stack,
        }

    def prepare_next_provenance(
        self,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> None:
        self.ns[self._PENDING_PROVENANCE_KEY] = {
            "api_name": str(api_name or "unknown"),
            "element_kind": str(element_kind or "unknown"),
            "text_preview": _normalize_preview_text(text_preview),
            "precision": str(precision or "exact"),
        }
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._PENDING_PROVENANCE_KEY] = copy.deepcopy(self.ns[self._PENDING_PROVENANCE_KEY])

    def _consume_pending_provenance(self) -> Optional[Dict[str, Any]]:
        pending = self.ns.get(self._PENDING_PROVENANCE_KEY)
        self.ns[self._PENDING_PROVENANCE_KEY] = None
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._PENDING_PROVENANCE_KEY] = None
        return copy.deepcopy(pending) if isinstance(pending, dict) else None

    def _build_provenance_entry(
        self,
        block_id: str,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> Dict[str, Any]:
        location = self._capture_provenance_location()
        notebook_cell_id = location.get("notebook_cell_id") or self._get_block_owner_notebook_cell(block_id)
        provenance_id = uuid.uuid4().hex
        return {
            "provenance_id": provenance_id,
            "block_id": str(block_id),
            "notebook_cell_id": notebook_cell_id,
            "file_path": location.get("file_path"),
            "line": int(location["line"]) if location.get("line") else None,
            "exact_notebook_cell_id": location.get("exact_notebook_cell_id"),
            "exact_file_path": location.get("exact_file_path"),
            "exact_line": int(location["exact_line"]) if location.get("exact_line") else None,
            "user_stack": copy.deepcopy(location.get("user_stack") or []),
            "api_name": str(api_name or "unknown"),
            "element_kind": str(element_kind or "unknown"),
            "precision": str(precision or "exact"),
            "text_preview": _normalize_preview_text(text_preview),
            "hyperlink_applied": False,
            "hyperlink_url": _build_provenance_open_url(provenance_id),
            "fragments": [],
        }

    def resolve_element_provenance(
        self,
        block_id: str,
        *,
        default_api_name: str,
        default_element_kind: str,
        default_text_preview: Optional[str] = None,
        default_precision: str = "exact",
        provided: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        base = (
            copy.deepcopy(provided)
            if isinstance(provided, dict)
            else self._consume_pending_provenance()
            or {}
        )
        return self._build_provenance_entry(
            block_id,
            api_name=str(base.get("api_name") or default_api_name),
            element_kind=str(base.get("element_kind") or default_element_kind),
            text_preview=base.get("text_preview") or default_text_preview,
            precision=str(base.get("precision") or default_precision),
        )

    def _node_path_from_root(self, root_node: Any, target_node: Any) -> Optional[List[int]]:
        if root_node is None or target_node is None:
            return None
        if root_node is target_node:
            return []
        path: List[int] = []
        node = target_node
        while node is not None and node is not root_node:
            parent = getattr(node, "getparent", lambda: None)()
            if parent is None:
                return None
            try:
                index = list(parent).index(node)
            except ValueError:
                return None
            path.append(index)
            node = parent
        if node is not root_node:
            return None
        return list(reversed(path))

    def _resolve_node_path(self, root_node: Any, path: Optional[Iterable[int]]) -> Any:
        if root_node is None or path is None:
            return root_node
        node = root_node
        for index in path:
            children = list(node)
            idx = int(index)
            if idx < 0 or idx >= len(children):
                return None
            node = children[idx]
        return node

    @staticmethod
    def _is_visible_node(node: Any) -> bool:
        tag = str(getattr(node, "tag", ""))
        return tag in _DOCX_VISIBLE_INLINE_TAGS or tag in _DOCX_VISIBLE_CONTAINER_TAGS

    def _descendant_visible_nodes(self, node: Any) -> List[Any]:
        if node is None:
            return []
        results: List[Any] = []
        for child in getattr(node, "iter", lambda: [])():
            if child is node:
                continue
            if self._is_visible_node(child):
                results.append(child)
        return results

    def _visible_nodes_for_mutation(self, root_node: Any, mutated_node: Any) -> List[Any]:
        if root_node is None:
            return []
        if mutated_node is None:
            return [root_node]
        if mutated_node is root_node or self._is_visible_node(mutated_node):
            return [mutated_node]

        descendants = self._descendant_visible_nodes(mutated_node)
        if descendants:
            return descendants

        node = mutated_node
        while node is not None:
            tag = str(getattr(node, "tag", ""))
            if tag == qn("w:tc"):
                paragraphs = list(node.findall(f".//{{{_DOCX_MAIN_NS}}}p"))
                return paragraphs or [root_node]
            if tag == qn("w:tr"):
                paragraphs = list(node.findall(f".//{{{_DOCX_MAIN_NS}}}p"))
                return paragraphs or [root_node]
            if self._is_visible_node(node):
                return [node]
            if node is root_node:
                break
            node = getattr(node, "getparent", lambda: None)()

        return [root_node]

    def _preview_from_node(self, node: Any) -> Optional[str]:
        if node is None:
            return None
        texts: List[str] = []
        for child in getattr(node, "iter", lambda: [])():
            tag = str(getattr(child, "tag", ""))
            if tag.endswith("}t") or tag.endswith("}delText") or tag.endswith("}instrText"):
                value = str(getattr(child, "text", "") or "").strip()
                if value:
                    texts.append(value)
        if texts:
            return _normalize_preview_text(" ".join(texts))
        return _normalize_preview_text(getattr(node, "text", None))

    def _ensure_fragment_list(self, entry: Dict[str, Any]) -> List[Dict[str, Any]]:
        fragments = entry.get("fragments")
        if not isinstance(fragments, list):
            fragments = []
            entry["fragments"] = fragments
        return fragments

    def _find_registered_root_index(self, block_id: str, root_element: Any) -> Optional[int]:
        items = self._cell_items_map().get(block_id, [])
        root_node = self._element_node(root_element)
        for idx, element in enumerate(items):
            if element is root_element:
                return idx
            candidate_node = self._element_node(element)
            if candidate_node is not None and root_node is not None and candidate_node is root_node:
                return idx
        return None

    def _build_fragment_record(
        self,
        block_id: str,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> Dict[str, Any]:
        record = self._build_provenance_entry(
            block_id,
            api_name=api_name,
            element_kind=element_kind,
            text_preview=text_preview,
            precision=precision,
        )
        record.pop("fragments", None)
        return record

    @staticmethod
    def _path_is_descendant(path: List[int], prefix: List[int]) -> bool:
        if len(prefix) > len(path):
            return False
        return path[: len(prefix)] == prefix

    def _replace_fragment_paths(self, fragments: List[Dict[str, Any]], path: List[int], record: Dict[str, Any], *, replace: bool) -> None:
        if replace:
            fragments[:] = [
                frag for frag in fragments
                if not self._path_is_descendant(list(frag.get("path") or []), path)
            ]
        else:
            fragments[:] = [
                frag for frag in fragments
                if list(frag.get("path") or []) != path
            ]
        fragments.append(record)

    def ensure_registered_root(
        self,
        block_id: str,
        root_element: Any,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> int:
        existing_index = self._find_registered_root_index(block_id, root_element)
        if existing_index is not None:
            return existing_index
        self._register_element(
            block_id,
            root_element,
            provenance={
                "api_name": api_name,
                "element_kind": element_kind,
                "text_preview": text_preview,
                "precision": precision,
            },
        )
        existing_index = self._find_registered_root_index(block_id, root_element)
        if existing_index is None:
            raise RuntimeError("No se pudo registrar el elemento DOCX raíz para procedencia.")
        return existing_index

    def register_element_before(
        self,
        block_id: str,
        before_element: Any,
        element: Any,
        *,
        provenance: Optional[Dict[str, Any]] = None,
    ) -> None:
        before_index = self._find_registered_root_index(block_id, before_element)
        if before_index is None:
            self._register_element(block_id, element, provenance=provenance)
            return

        items = self._cell_items_map().setdefault(block_id, [])
        block_meta = self._cell_meta_map().setdefault(block_id, {})
        provenance_items = block_meta.setdefault("provenance", [])
        resolved = self.resolve_element_provenance(
            block_id,
            default_api_name=f"document.{type(element).__name__}",
            default_element_kind=getattr(element, "__class__", type(element)).__name__.lower(),
            default_text_preview=getattr(element, "text", None),
            default_precision="fallback",
            provided=provenance,
        )
        items.insert(before_index, element)
        provenance_items.insert(before_index, resolved)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ITEMS_KEY] = items
            self.main_ns[self._CELL_META_KEY] = self._cell_meta_map()

    def record_visible_mutation(
        self,
        block_id: str,
        root_element: Any,
        mutated_target: Any,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
        replace: bool = True,
    ) -> None:
        root_index = self.ensure_registered_root(
            block_id,
            root_element,
            api_name=api_name,
            element_kind=element_kind,
            text_preview=text_preview,
            precision=precision,
        )
        root_node = self._element_node(root_element)
        if root_node is None:
            return

        block_meta = self._cell_meta_map().setdefault(block_id, {})
        provenance_items = block_meta.setdefault("provenance", [])
        while len(provenance_items) <= root_index:
            provenance_items.append(None)
        entry = provenance_items[root_index]
        if not isinstance(entry, dict):
            entry = self._build_fragment_record(
                block_id,
                api_name=api_name,
                element_kind=element_kind,
                text_preview=text_preview,
                precision=precision,
            )
            entry["fragments"] = []
            provenance_items[root_index] = entry

        target_node = self._element_node(mutated_target) if mutated_target is not None and not hasattr(mutated_target, "tag") else mutated_target
        visible_nodes = self._visible_nodes_for_mutation(root_node, target_node)
        fragments = self._ensure_fragment_list(entry)

        for visible_node in visible_nodes:
            path = self._node_path_from_root(root_node, visible_node)
            if path is None:
                continue
            record = self._build_fragment_record(
                block_id,
                api_name=api_name,
                element_kind=element_kind,
                text_preview=text_preview or self._preview_from_node(visible_node) or self._preview_from_node(root_node),
                precision=precision,
            )
            record["path"] = list(path)
            self._replace_fragment_paths(fragments, list(path), record, replace=replace)

        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_META_KEY] = self._cell_meta_map()

    @contextmanager
    def activate_cell(
        self,
        *,
        block_id: Optional[str] = None,
        auto_clear: bool = True,
        strict: Optional[bool] = None,
        order: Optional[int] = None,
        notebook_cell_id: Optional[str] = None,
    ) -> Generator[CellHandle, None, None]:
        if order is None:
            raise ValueError("order es requerido para docblock/doc_begin")
        cid = self._resolve_block_id(block_id)
        explicit_order = int(order)
        previous_cell = self.ns.get(self._ACTIVE_CELL_KEY)
        previous_cursor = self._cell_cursor_map().get(cid, 0)
        previous_strict = self.ns.get(self._STRICT_MODE_KEY, False)

        if strict is not None:
            self.ns[self._STRICT_MODE_KEY] = bool(strict)
        self.ns[self._ACTIVE_CELL_KEY] = cid
        self._cell_cursor_map()[cid] = 0

        if auto_clear:
            self.clear_cell(cid)
        self._apply_explicit_order(cid, explicit_order)
        self._register_block_for_notebook_cell(notebook_cell_id, cid)

        handle = CellHandle(self, cid, self.doc, self.parser, self.latex_converter)
        validation_error = None
        try:
            yield handle
        except Exception:
            # Re-lanzar excepciones de Python sin validar
            raise
        finally:
            self._dedupe_cell(cid)
            self._persist_cell_snapshot(cid)
            
            # Validar documento si strict mode está activo
            current_strict = strict if strict is not None else self.is_strict_mode()
            if current_strict and validation_error is None:
                is_valid, validation_errors = self.validate_document()
                if not is_valid:
                    # Rollback: limpiar el bloque que causó el problema
                    _logger.error(f"[DOCX Validation] Bloque '{cid}' generó XML inválido, haciendo rollback")
                    self.clear_cell(cid)
                    
                    # ✅ Forzar rebuild completo del documento para remover XML corrupto
                    # clear_cell() solo elimina referencias, pero el XML malformado persiste
                    # en doc._element.body. Un rebuild reconstruye el body desde cero.
                    try:
                        _logger.info(f"[DOCX Validation] Forzando rebuild completo para eliminar XML corrupto")
                        self._mark_dirty(reason="validation_rollback")
                        self._ensure_rebuilt()
                    except Exception as rebuild_error:
                        _logger.warning(f"[DOCX Validation] Error en rebuild post-rollback: {rebuild_error}")
                    
                    validation_error = DocxValidationError(
                        f"El bloque '{cid}' generó DOCX con XML malformado. "
                        f"Esto habría causado que Word/LibreOffice se cuelgue. "
                        f"Errores: {'; '.join(validation_errors)}",
                        {"block_id": cid, "errors": validation_errors}
                    )
            
            if validation_error is None:
                self._mark_dirty(reason="cell_update")
            
            # Restaurar estado previo
            if previous_cell is not None:
                self.ns[self._ACTIVE_CELL_KEY] = previous_cell
            else:
                self.ns.pop(self._ACTIVE_CELL_KEY, None)
            self._cell_cursor_map()[cid] = previous_cursor
            if strict is not None:
                self.ns[self._STRICT_MODE_KEY] = previous_strict
        
        # Lanzar error de validación fuera del finally para permitir traceback limpio
        if validation_error is not None:
            raise validation_error

    # ------------------------------------------------------------------
    # Registro y control de elementos
    # ------------------------------------------------------------------
    def _cell_items_map(self) -> Dict[str, List[Any]]:
        return self.ns[self._CELL_ITEMS_KEY]

    def _cell_cursor_map(self) -> Dict[str, int]:
        return self.ns[self._CELL_CURSOR_KEY]

    def _cell_order_values(self) -> Dict[str, Optional[int]]:
        return self.ns[self._CELL_ORDER_VALUES_KEY]

    def _refresh_order_from_values(self) -> None:
        values = self._cell_order_values()
        order_list = [cid for cid, _ in sorted(values.items(), key=lambda kv: (kv[1], kv[0]))]
        self.ns[self._CELL_ORDER_KEY] = order_list
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ORDER_KEY] = list(order_list)
            self.main_ns[self._CELL_ORDER_VALUES_KEY] = dict(values)

    def start_notebook_cell(self, notebook_cell_id: str) -> None:
        key = str(notebook_cell_id)
        active = self.ns.setdefault(self._NOTEBOOK_ACTIVE_KEY, {})
        active[key] = set()
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._NOTEBOOK_ACTIVE_KEY] = active

    def finish_notebook_cell(self, notebook_cell_id: str) -> None:
        key = str(notebook_cell_id)
        active = self.ns.setdefault(self._NOTEBOOK_ACTIVE_KEY, {})
        current = active.pop(key, set())
        groups = self.ns.setdefault(self._NOTEBOOK_GROUPS_KEY, {})
        previous = groups.get(key, [])
        removed = [bid for bid in previous if bid not in current]
        for block_id in removed:
            self.clear_cell(block_id)
        order_values = self._cell_order_values()
        ordered = sorted(current, key=lambda bid: (order_values.get(bid, float("inf")), bid))
        groups[key] = ordered
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._NOTEBOOK_GROUPS_KEY] = groups
            self.main_ns[self._NOTEBOOK_ACTIVE_KEY] = active

    def _register_block_for_notebook_cell(self, notebook_cell_id: Optional[str], block_id: str) -> None:
        if not notebook_cell_id:
            return
        key = str(notebook_cell_id)
        active = self.ns.setdefault(self._NOTEBOOK_ACTIVE_KEY, {})
        current = active.setdefault(key, set())
        current.add(block_id)
        block_map = self.ns.setdefault(self._BLOCK_TO_NOTEBOOK_KEY, {})
        block_map[block_id] = key
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._NOTEBOOK_ACTIVE_KEY] = active
            self.main_ns[self._BLOCK_TO_NOTEBOOK_KEY] = block_map

    def _remove_block_from_notebook_owner(self, block_id: str) -> None:
        block_map = self.ns.setdefault(self._BLOCK_TO_NOTEBOOK_KEY, {})
        owner = block_map.pop(block_id, None)
        if owner:
            groups = self.ns.setdefault(self._NOTEBOOK_GROUPS_KEY, {})
            blocks = groups.get(owner, [])
            groups[owner] = [bid for bid in blocks if bid != block_id]
            if self.main_ns and self.main_ns is not self.ns:
                self.main_ns[self._BLOCK_TO_NOTEBOOK_KEY] = block_map
                self.main_ns[self._NOTEBOOK_GROUPS_KEY] = groups

    def _register_element(
        self,
        block_id: str,
        element: Any,
        *,
        provenance: Optional[Dict[str, Any]] = None,
    ) -> None:
        items = self._cell_items_map()
        cursor_map = self._cell_cursor_map()
        cursor = cursor_map.get(block_id, 0)
        meta_map = self._cell_meta_map()
        block_meta = meta_map.setdefault(block_id, {})
        provenance_items = block_meta.setdefault("provenance", [])

        lst = items.setdefault(block_id, [])
        resolved_provenance = self.resolve_element_provenance(
            block_id,
            default_api_name=f"document.{type(element).__name__}",
            default_element_kind=getattr(element, "__class__", type(element)).__name__.lower(),
            default_text_preview=getattr(element, "text", None),
            default_precision="fallback",
            provided=provenance,
        )
        if cursor < len(lst):
            lst[cursor] = element
            if cursor < len(provenance_items):
                provenance_items[cursor] = resolved_provenance
            else:
                provenance_items.append(resolved_provenance)
        else:
            lst.append(element)
            provenance_items.append(resolved_provenance)
        cursor_map[block_id] = cursor + 1
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_META_KEY] = meta_map

    def _apply_explicit_order(self, block_id: str, order_value: int) -> None:
        values = self._cell_order_values()
        values[str(block_id)] = int(order_value)
        self._refresh_order_from_values()

    def clear_cell(self, block_id: Optional[str] = None) -> bool:
        cid = self._resolve_block_id(block_id)
        items = self._cell_items_map()
        existing = list(items.get(cid, []))
        
        # Diagnóstico de duplicación: verificar snapshots antes de limpiar
        serialized = self.ns.setdefault(self._CELL_SERIALIZED_KEY, {})
        had_serialized = cid in serialized
        serialized_count = len(serialized.get(cid, [])) if had_serialized else 0
        
        for element in existing:
            self._detach_element(element)
        if cid in items:
            items.pop(cid, None)
        cursor_map = self._cell_cursor_map()
        cursor_map.pop(cid, None)
        meta = self.ns[self._CELL_META_KEY]
        meta.pop(cid, None)
        order = self.ns.setdefault(self._CELL_ORDER_KEY, [])
        if cid in order:
            order = [c for c in order if c != cid]
            self.ns[self._CELL_ORDER_KEY] = order
        serialized.pop(cid, None)
        
        # Logging para diagnosticar duplicación
        if existing or had_serialized:
            _logger.info(
                f"[DOCX Clear] Bloque '{cid}': eliminados {len(existing)} elementos vivos "
                f"y {serialized_count} snapshots serializados"
            )
        
        order_values = self._cell_order_values()
        order_values.pop(cid, None)
        self._remove_block_from_notebook_owner(cid)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ITEMS_KEY] = items
            self.main_ns[self._CELL_CURSOR_KEY] = cursor_map
            self.main_ns[self._CELL_META_KEY] = meta
            main_order = self.main_ns.setdefault(self._CELL_ORDER_KEY, [])
            if cid in main_order:
                self.main_ns[self._CELL_ORDER_KEY] = [c for c in main_order if c != cid]
            self.main_ns[self._CELL_ORDER_VALUES_KEY] = order_values
        self._refresh_order_from_values()
        
        # Usar razón específica para permitir optimización en ensure_rebuilt
        self._mark_dirty(reason="cell_update") 
        return bool(existing)


    def _dedupe_cell(self, block_id: str) -> None:
        items = self._cell_items_map()
        lst = items.get(block_id)
        if not lst:
            return
        block_meta = self._cell_meta_map().get(block_id, {})
        provenance_items = list(block_meta.get("provenance") or [])
        unique: List[Any] = []
        unique_provenance: List[Any] = []
        seen = set()
        changed = False
        for idx, element in enumerate(lst):
            anchor = getattr(element, "_p", None)
            if anchor is None:
                anchor = getattr(element, "_tbl", None)
            if anchor is None:
                anchor = element
            key = id(anchor)
            if key in seen:
                changed = True
                continue
            seen.add(key)
            unique.append(element)
            if idx < len(provenance_items):
                unique_provenance.append(provenance_items[idx])
        if changed:
            items[block_id] = unique
            block_meta["provenance"] = unique_provenance

    @staticmethod
    def _element_node(element: Any) -> Any:
        node = getattr(element, "_p", None)
        if node is None:
            node = getattr(element, "_tbl", None)
        if node is None:
            node = getattr(element, "_inline", None)
        if node is None:
            node = getattr(element, "_element", None)
        return node

    def _detach_element(self, element: Any) -> None:
        node = self._element_node(element)
        if node is None:
            return
        parent = getattr(node, "getparent", lambda: None)()
        if parent is None:
            return
        try:
            parent.remove(node)
        except Exception:  # pragma: no cover - fallo silencioso para robustez
            pass

    # ------------------------------------------------------------------
    # Metadatos y numeraciones
    # ------------------------------------------------------------------
    def next_equation_number(self) -> int:
        value = int(self.ns.get(self._EQ_COUNTER_KEY, 0)) + 1
        self.ns[self._EQ_COUNTER_KEY] = value
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._EQ_COUNTER_KEY] = value
        return value

    def next_figure_number(self) -> int:
        value = self.next_sequence_number("Figura")
        self.ns[self._FIG_COUNTER_KEY] = value
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._FIG_COUNTER_KEY] = value
        return value

    def next_table_number(self) -> int:
        return self.next_sequence_number("Tabla")

    def next_sequence_number(self, sequence_name: str) -> int:
        normalized_name = str(sequence_name or "").strip() or "Figura"
        counters = self.ns.setdefault(self._SEQ_COUNTERS_KEY, {})
        value = int(counters.get(normalized_name, 0)) + 1
        counters[normalized_name] = value
        self.ns[self._SEQ_COUNTERS_KEY] = counters
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._SEQ_COUNTERS_KEY] = counters
        return value

    def next_bookmark_id(self) -> int:
        value = int(self.ns.get(self._BOOKMARK_COUNTER_KEY, 1))
        self.ns[self._BOOKMARK_COUNTER_KEY] = value + 1
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._BOOKMARK_COUNTER_KEY] = self.ns[self._BOOKMARK_COUNTER_KEY]
        return value

    def register_label(self, label: str, kind: str, *, number: Optional[int] = None) -> Dict[str, Any]:
        data = self.ns[self._LABELS_KEY]
        info = {
            "type": kind,
            "number": number,
        }
        data[str(label)] = info
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._LABELS_KEY] = data
        return info

    def resolve_label(self, label: str) -> Optional[Dict[str, Any]]:
        data = self.ns.get(self._LABELS_KEY, {})
        return data.get(str(label))

    # ------------------------------------------------------------------
    # Logging y eventos
    # ------------------------------------------------------------------
    def log_event(self, kind: str, message: str, **data: Any) -> None:
        entry = {
            "kind": str(kind),
            "message": str(message),
            "data": data or None,
        }
        log = self.ns[self._EVENT_LOG_KEY]
        log.append(entry)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._EVENT_LOG_KEY] = log

    def get_event_log(self) -> List[Dict[str, Any]]:
        return list(self.ns.get(self._EVENT_LOG_KEY, []))

    def clear_event_log(self) -> None:
        self.ns[self._EVENT_LOG_KEY] = []
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._EVENT_LOG_KEY] = []

    def set_strict_mode(self, value: bool) -> None:
        self.ns[self._STRICT_MODE_KEY] = bool(value)
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._STRICT_MODE_KEY] = self.ns[self._STRICT_MODE_KEY]

    # ------------------------------------------------------------------
    # Validación de documento
    # ------------------------------------------------------------------
    def _wrap_paragraph_with_hyperlink(self, paragraph_node: Any, hyperlink_url: str) -> bool:
        if paragraph_node is None or not hyperlink_url:
            return False
        if any(str(getattr(child, "tag", "")).endswith("}hyperlink") for child in list(paragraph_node)):
            return False

        children = list(paragraph_node)
        content_nodes = []
        insert_index = 0
        seen_content = False
        for child in children:
            tag = str(getattr(child, "tag", ""))
            if tag == qn("w:pPr"):
                insert_index += 1
                continue
            if tag not in _PARAGRAPH_HYPERLINK_CHILD_TAGS:
                return False
            if tag in _PARAGRAPH_NON_HYPERLINK_MARKER_TAGS:
                if not seen_content:
                    insert_index += 1
                continue
            seen_content = True
            content_nodes.append(child)

        if not content_nodes:
            return False

        try:
            from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore

            r_id = self.doc.part.relate_to(str(hyperlink_url), RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            for child in content_nodes:
                paragraph_node.remove(child)
                hyperlink.append(child)
            paragraph_node.insert(insert_index, hyperlink)
            return True
        except Exception as exc:
            _logger.warning("[DOCX Provenance] No se pudo envolver parrafo con hyperlink: %s", exc)
            return False

    @staticmethod
    def _direct_paragraph_child(paragraph_node: Any, node: Any) -> Any:
        current = node
        while current is not None:
            parent = getattr(current, "getparent", lambda: None)()
            if parent is paragraph_node:
                return current
            current = parent
        return None

    @staticmethod
    def _node_is_hyperlinked(node: Any) -> bool:
        current = node
        while current is not None:
            if str(getattr(current, "tag", "")) == qn("w:hyperlink"):
                return True
            current = getattr(current, "getparent", lambda: None)()
        try:
            return any(str(getattr(child, "tag", "")) == qn("w:hyperlink") for child in node.iter())
        except Exception:
            return False

    def _wrap_paragraph_child_groups(
        self,
        paragraph_node: Any,
        child_url_pairs: List[tuple[Any, str]],
    ) -> set[int]:
        if paragraph_node is None or not child_url_pairs:
            return set()
        child_url_map: Dict[int, str] = {}
        for child, url in child_url_pairs:
            if child is None or not url:
                continue
            direct_child = self._direct_paragraph_child(paragraph_node, child)
            if direct_child is None:
                continue
            child_url_map[id(direct_child)] = str(url)
        if not child_url_map:
            return set()

        try:
            from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore
        except Exception:
            return set()

        applied_ids: set[int] = set()
        content_children = [
            child for child in list(paragraph_node)
            if getattr(child, "tag", None) != qn("w:pPr")
        ]
        current_group: List[Any] = []
        current_url: Optional[str] = None

        def flush_group() -> None:
            nonlocal current_group, current_url
            if not current_group or not current_url:
                current_group = []
                current_url = None
                return
            try:
                r_id = self.doc.part.relate_to(str(current_url), RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                insert_at = list(paragraph_node).index(current_group[0])
                for child in current_group:
                    paragraph_node.remove(child)
                    hyperlink.append(child)
                    applied_ids.add(id(child))
                paragraph_node.insert(insert_at, hyperlink)
            except Exception as exc:
                _logger.warning("[DOCX Provenance] No se pudo envolver grupo con hyperlink: %s", exc)
            current_group = []
            current_url = None

        for child in content_children:
            tag = str(getattr(child, "tag", ""))
            if tag not in _PARAGRAPH_HYPERLINK_CHILD_TAGS:
                flush_group()
                continue
            child_url = child_url_map.get(id(child))
            if not child_url:
                flush_group()
                continue
            if current_url == child_url:
                current_group.append(child)
            else:
                flush_group()
                current_url = child_url
                current_group = [child]
        flush_group()
        return applied_ids

    def _apply_hyperlink_to_node(self, node: Any, hyperlink_url: str) -> bool:
        if node is None:
            return False
        node_tag = str(getattr(node, "tag", ""))
        if node_tag == qn("w:p"):
            return self._wrap_paragraph_with_hyperlink(node, hyperlink_url)
        if node_tag == qn("w:tbl"):
            applied = False
            for paragraph_node in node.findall(f".//{{{_DOCX_MAIN_NS}}}p"):
                applied = self._wrap_paragraph_with_hyperlink(paragraph_node, hyperlink_url) or applied
            return applied
        return False

    def _apply_provenance_hyperlinks(self) -> None:
        items_map = self._cell_items_map()
        meta_map = self._cell_meta_map()
        for block_id, items in items_map.items():
            provenance_items = list((meta_map.get(block_id, {}) or {}).get("provenance") or [])
            if not provenance_items:
                continue
            for idx, element in enumerate(items):
                if idx >= len(provenance_items):
                    continue
                provenance_entry = provenance_items[idx]
                if not isinstance(provenance_entry, dict):
                    continue
                node = self._element_node(element)
                fragments = list(provenance_entry.get("fragments") or [])
                records = fragments or [provenance_entry]
                if node is None:
                    continue
                root_node = node
                paragraph_targets: Dict[int, List[tuple[Any, str, Dict[str, Any]]]] = {}
                for record in records:
                    if not isinstance(record, dict):
                        continue
                    record["hyperlink_applied"] = False
                    record["clickable"] = False
                    target_path = record.get("path")
                    target_node = self._resolve_node_path(root_node, target_path) if isinstance(target_path, list) else root_node
                    if target_node is None:
                        continue
                    target_tag = str(getattr(target_node, "tag", ""))
                    if target_tag == qn("w:p"):
                        applied = self._wrap_paragraph_with_hyperlink(target_node, str(record.get("hyperlink_url") or ""))
                        if not applied and self._node_is_hyperlinked(target_node):
                            applied = True
                        record["hyperlink_applied"] = bool(applied)
                        record["clickable"] = bool(applied)
                        continue
                    paragraph_node = target_node if target_tag == qn("w:p") else getattr(target_node, "getparent", lambda: None)()
                    while paragraph_node is not None and str(getattr(paragraph_node, "tag", "")) != qn("w:p"):
                        paragraph_node = getattr(paragraph_node, "getparent", lambda: None)()
                    if paragraph_node is None:
                        applied = self._apply_hyperlink_to_node(target_node, str(record.get("hyperlink_url") or ""))
                        record["hyperlink_applied"] = bool(applied)
                        record["clickable"] = bool(applied)
                        continue
                    paragraph_targets.setdefault(id(paragraph_node), []).append((target_node, str(record.get("hyperlink_url") or ""), record))

                for grouped_targets in paragraph_targets.values():
                    if not grouped_targets:
                        continue
                    paragraph_node = None
                    child_pairs: List[tuple[Any, str]] = []
                    for target_node, url, _record in grouped_targets:
                        paragraph_node = target_node if str(getattr(target_node, "tag", "")) == qn("w:p") else None
                        if paragraph_node is None:
                            current = target_node
                            while current is not None and str(getattr(current, "tag", "")) != qn("w:p"):
                                current = getattr(current, "getparent", lambda: None)()
                            paragraph_node = current
                        child_pairs.append((target_node, url))
                    if paragraph_node is None:
                        continue
                    applied_ids = self._wrap_paragraph_child_groups(paragraph_node, child_pairs)
                    for target_node, _url, record in grouped_targets:
                        direct_child = self._direct_paragraph_child(paragraph_node, target_node)
                        applied = (
                            id(target_node) in applied_ids
                            or (direct_child is not None and id(direct_child) in applied_ids)
                            or self._node_is_hyperlinked(target_node)
                        )
                        record["hyperlink_applied"] = bool(applied)
                        record["clickable"] = bool(applied)

    def build_provenance_manifest(self) -> Dict[str, Any]:
        items_map = self._cell_items_map()
        meta_map = self._cell_meta_map()
        items: List[Dict[str, Any]] = []
        for block_id, elements in items_map.items():
            provenance_items = list((meta_map.get(block_id, {}) or {}).get("provenance") or [])
            if not provenance_items:
                continue
            for idx, _element in enumerate(elements):
                if idx >= len(provenance_items):
                    continue
                entry = provenance_items[idx]
                if not isinstance(entry, dict):
                    continue
                records = list(entry.get("fragments") or []) or [entry]
                for record in records:
                    if not isinstance(record, dict):
                        continue
                    items.append(
                        {
                            "provenance_id": str(record.get("provenance_id") or ""),
                            "block_id": str(entry.get("block_id") or block_id),
                            "notebook_cell_id": record.get("notebook_cell_id"),
                            "file_path": record.get("file_path"),
                            "line": record.get("line"),
                            "exact_notebook_cell_id": record.get("exact_notebook_cell_id"),
                            "exact_file_path": record.get("exact_file_path"),
                            "exact_line": record.get("exact_line"),
                            "user_stack": copy.deepcopy(record.get("user_stack") or []),
                            "api_name": record.get("api_name"),
                            "element_kind": record.get("element_kind"),
                            "precision": record.get("precision"),
                            "text_preview": record.get("text_preview"),
                            "clickable": bool(record.get("clickable") or record.get("hyperlink_applied")),
                            "hyperlink_url": record.get("hyperlink_url"),
                        }
                    )
        return {
            "generated_at": datetime.now(UTC).isoformat().replace("+00:00", "Z"),
            "items": [item for item in items if item.get("provenance_id")],
        }

    def export_provenance_manifest_json(self) -> str:
        self._prepare_for_export()
        return json.dumps(self.build_provenance_manifest(), ensure_ascii=False)

    def validate_document(self) -> tuple:
        """Valida la estructura XML del documento DOCX.
        
        Esta validación detecta problemas que causarían que Word/LibreOffice
        no puedan abrir el documento correctamente (requiriendo "reparación"),
        lo cual bloquearía la conversión PDF.
        
        Returns:
            tuple: (is_valid: bool, errors: List[str])
        """
        try:
            docx_bytes = self.serialize_docx_bytes()
            is_valid, errors = validate_docx_package_bytes(docx_bytes)
        except Exception as exc:
            errors = [f"Error de validación general: {type(exc).__name__}: {exc}"]
            is_valid = False
        is_valid = len(errors) == 0
        if not is_valid:
            _logger.warning(f"[DOCX Validation] Documento invÃ¡lido: {errors}")
        return is_valid, errors

        errors = []
        try:
            # 1. Guardar a buffer temporal
            buffer = io.BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            
            # 2. Verificar que es un archivo ZIP válido
            if not zipfile.is_zipfile(buffer):
                errors.append("El documento no es un archivo ZIP válido (DOCX corrupto)")
                return False, errors
            
            buffer.seek(0)
            with zipfile.ZipFile(buffer, 'r') as zf:
                # 3. Verificar archivos requeridos del formato OOXML
                required_files = ['word/document.xml', '[Content_Types].xml']
                missing = [f for f in required_files if f not in zf.namelist()]
                if missing:
                    errors.append(f"Faltan archivos requeridos: {', '.join(missing)}")
                
                # 4. Validar XML de document.xml
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                if 'word/document.xml' in zf.namelist():
                    try:
                        from lxml import etree
                        xml_content = zf.read('word/document.xml')
                        # Intentar parsear - esto detecta XML malformado
                        root = etree.fromstring(xml_content)
                        
                        # 5. Validación ligera de estructura OOXML
                        # Verificar que el body existe
                        body = root.find('.//w:body', ns)
                        if body is None:
                            errors.append("Estructura DOCX inválida: falta elemento <w:body>")
                            
                    except etree.XMLSyntaxError as e:
                        line = getattr(e, 'lineno', '?')
                        col = getattr(e, 'offset', '?')
                        errors.append(f"XML malformado en document.xml (línea {line}, col {col}): {e.msg}")
                    except Exception as e:
                        errors.append(f"Error parseando document.xml: {type(e).__name__}: {e}")
                        
        except Exception as e:
            errors.append(f"Error de validación general: {type(e).__name__}: {e}")
        
        is_valid = len(errors) == 0
        if not is_valid:
            _logger.warning(f"[DOCX Validation] Documento inválido: {errors}")
        
        return is_valid, errors

    # ------------------------------------------------------------------
    # Ordenamiento y snapshots
    # ------------------------------------------------------------------
    def get_cell_order(self) -> List[str]:
        order = self.ns.get(self._CELL_ORDER_KEY)
        if isinstance(order, list) and order:
            final = [str(cid) for cid in order]
        else:
            self._refresh_order_from_values()
            final = self.ns.get(self._CELL_ORDER_KEY, [])
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ORDER_KEY] = list(final)
        return final

    def set_cell_order(self, order: Iterable[str]) -> List[str]:
        items = self._cell_items_map()
        requested = [str(cid) for cid in order if str(cid) in items]
        remaining = [cid for cid in items.keys() if cid not in requested]
        final_order = requested + remaining
        values = self._cell_order_values()
        for idx, cid in enumerate(final_order):
            values[str(cid)] = idx
        self._refresh_order_from_values()
        self._reorder_body(final_order)
        self._mark_dirty()
        return final_order

    def move_cell(self, block_id: str, direction: str) -> List[str]:
        order = self.get_cell_order()
        cid = str(block_id)
        if cid not in order:
            return order
        idx = order.index(cid)
        if direction == 'up' and idx > 0:
            order[idx - 1], order[idx] = order[idx], order[idx - 1]
        elif direction == 'down' and idx < len(order) - 1:
            order[idx + 1], order[idx] = order[idx], order[idx + 1]
        return self.set_cell_order(order)

    def _reorder_body(self, order: Iterable[str]) -> None:
        if self.doc is None:
            return
        body = getattr(self.doc._element, 'body', None)
        if body is None:
            return
        items = self._cell_items_map()
        ordered_nodes: List[Any] = []
        seen = set()
        for cid in order:
            for element in items.get(str(cid), []):
                node = getattr(element, "_p", None)
                if node is None:
                    node = getattr(element, "_tbl", None)
                if node is None or node in seen:
                    continue
                seen.add(node)
                ordered_nodes.append(node)
        for elements in items.values():
            for element in elements:
                node = getattr(element, "_p", None)
                if node is None:
                    node = getattr(element, "_tbl", None)
                if node is None or node in seen:
                    continue
                seen.add(node)
                ordered_nodes.append(node)
        for node in ordered_nodes:
            parent = getattr(node, 'getparent', lambda: None)()
            if parent is None:
                continue
            if parent is body:
                try:
                    parent.remove(node)
                    body.append(node)
                except Exception:
                    continue
            else:
                try:
                    parent.remove(node)
                    body.append(node)
                except Exception:
                    continue

    def is_strict_mode(self) -> bool:
        return bool(self.ns.get(self._STRICT_MODE_KEY, False))

    def snapshot_cell(self, cell_id: Optional[str] = None, *, include_meta: bool = False) -> Dict[str, Any]:
        if cell_id is None:
            cid = self.ns.get(self._ACTIVE_CELL_KEY)
            if not cid:
                return {"cell_id": None, "elements": [], **({"meta": {}} if include_meta else {})}
            cid = str(cid)
        else:
            cid = str(cell_id)
        items = self._cell_items_map().get(cid, [])
        elements = []
        for idx, element in enumerate(items):
            element_type = getattr(element, "__class__", type(element)).__name__
            descriptor = getattr(element, "text", None)
            entry = {
                "index": idx,
                "type": element_type,
            }
            if descriptor:
                entry["text"] = descriptor[:120]
            elements.append(entry)
        snapshot: Dict[str, Any] = {"cell_id": cid, "elements": elements}
        if include_meta:
            meta = self.ns.get(self._CELL_META_KEY, {}).get(cid, {})
            snapshot["meta"] = meta.copy() if isinstance(meta, dict) else meta
        return snapshot

    def _persist_cell_snapshot(self, cell_id: str) -> None:
        items = list(self._cell_items_map().get(cell_id, []))
        provenance_items = list((self._cell_meta_map().get(cell_id, {}) or {}).get("provenance") or [])
        serialized_map = self.ns.setdefault(self._CELL_SERIALIZED_KEY, {})
        serialized_items: List[Any] = []
        for idx, element in enumerate(items):
            node = self._element_node(element)
            if node is None:
                continue
            provenance_entry = provenance_items[idx] if idx < len(provenance_items) else None
            serialized_items.append(self._serialize_node_snapshot(node, provenance=provenance_entry))
        serialized_map[cell_id] = serialized_items
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_SERIALIZED_KEY] = serialized_map

    def _serialize_node_snapshot(
        self,
        node: Any,
        *,
        provenance: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        payload = {
            "xml": node.xml,
            "relationships": self._snapshot_node_relationships(node),
        }
        if provenance:
            payload["provenance"] = copy.deepcopy(provenance)
        return payload

    def _snapshot_node_relationships(self, node: Any) -> List[Dict[str, Any]]:
        refs = iter_relationship_ids_in_xml_element(node)
        if not refs:
            return []
        rels = getattr(self.doc.part, "rels", {})
        snapshots: List[Dict[str, Any]] = []
        for rid in sorted(refs):
            rel = rels.get(rid)
            if rel is None:
                continue
            snapshot = {
                "rId": rid,
                "reltype": rel.reltype,
                "is_external": rel.is_external,
                "target_ref": rel.target_ref,
            }
            if not rel.is_external:
                snapshot["target_part"] = rel.target_part
            snapshots.append(snapshot)
        return snapshots

    def _node_has_missing_relationships(self, node: Any) -> bool:
        refs = iter_relationship_ids_in_xml_element(node)
        if not refs:
            return False
        rels = getattr(self.doc.part, "rels", {})
        return any(ref not in rels for ref in refs)

    def _restore_node_snapshot(self, fragment: Any) -> tuple[Any, Any, Optional[Dict[str, Any]]]:
        if isinstance(fragment, str):
            xml_text = fragment
            relationships = []
            provenance = None
        else:
            xml_text = str(fragment.get("xml") or "")
            relationships = list(fragment.get("relationships") or [])
            provenance = copy.deepcopy(fragment.get("provenance")) if isinstance(fragment.get("provenance"), dict) else None

        node = parse_xml(xml_text)
        rid_map: Dict[str, str] = {}
        for rel_info in relationships:
            old_rid = str(rel_info.get("rId") or "").strip()
            reltype = str(rel_info.get("reltype") or "").strip()
            if not old_rid or not reltype:
                continue
            try:
                if rel_info.get("is_external"):
                    target_ref = rel_info.get("target_ref")
                    if not target_ref:
                        continue
                    new_rid = self.doc.part.relate_to(str(target_ref), reltype, is_external=True)
                else:
                    target_part = rel_info.get("target_part")
                    if target_part is None:
                        continue
                    new_rid = self.doc.part.relate_to(target_part, reltype)
            except Exception as exc:
                _logger.warning(
                    "[DOCX FastRebuild] No se pudo restaurar relación %s (%s): %s",
                    old_rid,
                    reltype,
                    exc,
                )
                continue
            rid_map[old_rid] = new_rid

        rewrite_relationship_ids_in_xml_element(node, rid_map)

        if node.tag.endswith('}p'):
            return node, Paragraph(node, self.doc), provenance
        if node.tag.endswith('}tbl'):
            return node, Table(node, self.doc), provenance
        return node, node, provenance

    def _ensure_rebuilt(self) -> None:
        """Reconstruye el documento reordenando nodos en memoria (Fast Reorder)."""
        if not self._dirty:
            return

        import time
        start = time.time()
        
        # Estrategia: Reconstuir usando nodos vivos cuando sea posible (O(N) vs O(N*Size))
        self._fast_rebuild_document()
        self._dirty = False
        
        elapsed = time.time() - start
        _logger.info(f"[DOCX FastRebuild] Completado en {elapsed:.4f}s")
        

    def _fast_rebuild_document(self) -> None:
        """
        Reconstruye el documento manipulando el DOM directamente.
        Mueve los elementos a su posición correcta según el orden de celdas.
        """
        self._refresh_order_from_values()
        order = self.ns.get(self._CELL_ORDER_KEY, []) or []
        items_map = self.ns.get(self._CELL_ITEMS_KEY, {})
        meta_map = self.ns.get(self._CELL_META_KEY, {}) or {}
        serialized_map = self.ns.get(self._CELL_SERIALIZED_KEY, {}) or {}

        # 1. Identificar nodos de SectionProperties (sectPr) para preservarlos al final
        body = self.doc._element.body
        sectPr_nodes = [child for child in body if child.tag.endswith("}sectPr")]
        
        # 2. Recolectar TODOS los elementos en el orden correcto
        new_body_elements = []
        new_items_map = {}
        new_meta_map = {}
        
        for cid in order:
            # Opción A: Elementos vivos en memoria (Rápido)
            if cid in items_map and items_map[cid]:
                cell_elements = items_map[cid]
                cell_meta = copy.deepcopy(meta_map.get(cid, {})) if isinstance(meta_map.get(cid, {}), dict) else {}
                has_missing_relationships = any(
                    self._node_has_missing_relationships(self._element_node(el))
                    for el in cell_elements
                )
                # Validar que los elementos sigan vivos (no detached externamente)
                valid_elements = []
                if not has_missing_relationships:
                    for el in cell_elements:
                        # Obtener el nodo XML subyacente
                        node = self._element_node(el)
                        if node is not None:
                            valid_elements.append(node)
                            new_body_elements.append(node)
                
                # Actualizar mapa con lo que encontramos
                if valid_elements:
                    new_items_map[cid] = cell_elements
                    new_meta_map[cid] = cell_meta
                    continue # Éxito con nodos vivos
            
            # Opción B: Fallback a Serialized (Lento, solo si no hay vivos)
            if cid in serialized_map:
                fragments = serialized_map[cid]
                restored_items = []
                restored_provenance = []
                for frag_idx, xml in enumerate(fragments):
                    try:
                        node, restored_item, restored_item_provenance = self._restore_node_snapshot(xml)
                        new_body_elements.append(node)
                        if restored_item is not None:
                            restored_items.append(restored_item)
                            restored_provenance.append(copy.deepcopy(restored_item_provenance) if isinstance(restored_item_provenance, dict) else None)
                    except Exception as exc:
                        _logger.warning(
                            f"[DOCX FastRebuild] Error deserializando fragmento {frag_idx} "
                            f"de celda '{cid}': {type(exc).__name__}: {exc}"
                        )
                if restored_items:
                    new_items_map[cid] = restored_items
                    base_meta = copy.deepcopy(meta_map.get(cid, {})) if isinstance(meta_map.get(cid, {}), dict) else {}
                    base_meta["provenance"] = restored_provenance
                    new_meta_map[cid] = base_meta

        # 3. Aplicar al DOM (Bulk removal & append)
        # Esto es mucho más rápido que remove/append uno por uno si el documento es grande
        # Sin embargo, lxml no tiene "clear_children", así que iteramos.
        # Una optimización es: body[:] = new_body_elements + sectPr_nodes
        # Pero body es un proxy, necesitamos acceder a la lista de hijos
        
        # Remover todos los hijos actuales (excepto sectPr que manejaremos explícitamente)
        body[:] = new_body_elements + sectPr_nodes

        # 4. Actualizar estado
        self.ns[self._CELL_ITEMS_KEY] = new_items_map
        self.ns[self._CELL_META_KEY] = new_meta_map
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ITEMS_KEY] = new_items_map
            self.main_ns[self._CELL_META_KEY] = new_meta_map

    # ------------------------------------------------------------------
    # Dirty flag para reconstrucción diferida
    # ------------------------------------------------------------------
    def _mark_dirty(self, reason: str = "unknown") -> None:
        """Marca el documento como pendiente de reconstrucción."""
        if not self._dirty:
            _logger.debug(f"[DOCX Session] Marcado dirty: {reason}")
        self._dirty = True

    # ------------------------------------------------------------------
    # Reset y exportación
    # ------------------------------------------------------------------
    def reset(self, *, hard: bool = False) -> None:
        if hard:
            self.ns.pop(self._DOC_KEY, None)
            if self.main_ns and self.main_ns is not self.ns:
                self.main_ns.pop(self._DOC_KEY, None)
            self.doc = self._ensure_document()
        else:
            self.clear_document_body()

        for key in (
            self._CELL_ITEMS_KEY,
            self._CELL_CURSOR_KEY,
            self._CELL_META_KEY,
            self._PENDING_PROVENANCE_KEY,
            self._LABELS_KEY,
            self._EVENT_LOG_KEY,
            self._CELL_SERIALIZED_KEY,
            self._CELL_ORDER_VALUES_KEY,
            self._NOTEBOOK_GROUPS_KEY,
            self._NOTEBOOK_ACTIVE_KEY,
            self._BLOCK_TO_NOTEBOOK_KEY,
        ):
            if key == self._EVENT_LOG_KEY:
                self.ns[key] = []
            elif key == self._PENDING_PROVENANCE_KEY:
                self.ns[key] = None
            else:
                self.ns[key] = {}
            if self.main_ns and self.main_ns is not self.ns:
                self.main_ns[key] = self.ns[key]

        self.ns[self._CELL_ORDER_KEY] = []
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._CELL_ORDER_KEY] = []

        self.ns[self._EQ_COUNTER_KEY] = 0
        self.ns[self._FIG_COUNTER_KEY] = 0
        self.ns[self._SEQ_COUNTERS_KEY] = {}
        self.ns[self._BOOKMARK_COUNTER_KEY] = 1
        self.ns[self._DOC_INITIALIZED_KEY] = False
        if self.main_ns and self.main_ns is not self.ns:
            self.main_ns[self._EQ_COUNTER_KEY] = 0
            self.main_ns[self._FIG_COUNTER_KEY] = 0
            self.main_ns[self._SEQ_COUNTERS_KEY] = {}
            self.main_ns[self._BOOKMARK_COUNTER_KEY] = 1
            self.main_ns[self._DOC_INITIALIZED_KEY] = False

    def clear_document_body(self) -> None:
        document_element = self.doc._element  # type: ignore[attr-defined]
        body = document_element.body
        existing = list(body)
        sectPr_nodes = [child for child in existing if child.tag.endswith("}sectPr")]
        if not sectPr_nodes:
            sectPr_nodes = [OxmlElement("w:sectPr")]
        for child in existing:
            body.remove(child)
        for sectPr in sectPr_nodes:
            body.append(sectPr)

    def _prepare_for_export(self) -> None:
        if self._dirty:
            _logger.info("[DOCX Export] Reconstruyendo documento (dirty flag activo)")
        self._ensure_rebuilt()
        self._apply_provenance_hyperlinks()
        self._restore_template_header_footer_references_if_missing()

    def serialize_docx_bytes(self) -> bytes:
        self._prepare_for_export()
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    def serialize_docx_bytes_for_delivery(self) -> bytes:
        return sanitize_docx_bytes_for_delivery(self.serialize_docx_bytes())

    def export_docx_base64(self) -> str:
        import time
        start = time.time()

        self._prepare_for_export()
        rebuild_time = time.time() - start

        buffer = io.BytesIO()
        self.doc.save(buffer)
        save_time = time.time() - start - rebuild_time
        
        result = base64.b64encode(buffer.getvalue()).decode("utf-8")
        encode_time = time.time() - start - rebuild_time - save_time
        total_time = time.time() - start
        
        docx_size = len(buffer.getvalue())
        _logger.info(
            f"[DOCX Export] Completado en {total_time:.2f}s "
            f"(rebuild: {rebuild_time:.2f}s, save: {save_time:.2f}s, encode: {encode_time:.2f}s) "
            f"| Tamaño: {docx_size / 1024:.1f} KB"
        )
        
        return result

    def export_docx_base64_for_delivery(self) -> str:
        return sanitize_docx_b64_for_delivery(self.export_docx_base64()) or ""



# ----------------------------------------------------------------------
# Helpers de acceso global
# ----------------------------------------------------------------------
_SESSION_CACHE: Dict[int, DocxSession] = {}
_SESSION_LOCK = threading.Lock()


def get_session(namespace: Optional[Dict[str, Any]] = None) -> DocxSession:
    if namespace is None:
        frame = sys._getframe(1)
        namespace = frame.f_globals
    key = id(namespace)
    with _SESSION_LOCK:
        session = _SESSION_CACHE.get(key)
        if session is None:
            session = DocxSession(namespace)
            _SESSION_CACHE[key] = session
    return session


def reset_session_cache() -> None:
    with _SESSION_LOCK:
        _SESSION_CACHE.clear()
