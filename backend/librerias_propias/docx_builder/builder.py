"""Builder fluido para crear documentos DOCX en notebooks."""

from __future__ import annotations

import warnings
import weakref
import io
import re
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Tuple

from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.shared import qn  # type: ignore
from docx.shared import Inches, Pt, RGBColor

try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
except ImportError:  # pragma: no cover - versiones antiguas de python-docx
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    WD_BREAK = None  # type: ignore

from docx.table import Table
from docx.text.paragraph import Paragraph

from .latex_math import (
    LatexMathError,
    LatexMathUnavailableError,
)
from .proxies import wrap_docx_proxy
from .session import CellHandle, DocxSession, StrictModeError
from .utils import ImageConversionError, image_to_bytes

try:
    from librerias_propias.inspyro_units.formatting import build_docx_unit_runs
    from librerias_propias.inspyro_units.normalization import validate_unit_expression
except Exception:  # pragma: no cover - fallback defensivo
    def build_docx_unit_runs(unit_text: str) -> list[dict]:
        text = str(unit_text).strip() if unit_text is not None else ""
        return [{"text": text, "italic": True}] if text else []

    def validate_unit_expression(raw_unit: str) -> tuple[bool, str]:
        text = str(raw_unit or "").strip()
        return (bool(text), text)


AlignmentValue = Optional[str]

# ==========================================================================
# LÃ­mites y configuraciÃ³n de validaciones
# ==========================================================================
MAX_LIST_ITEMS = 500  # MÃ¡ximo de items en una lista antes de warning
MAX_TABLE_ROWS = 1000  # MÃ¡ximo de filas en tabla antes de warning
MAX_TABLE_COLS = 50  # MÃ¡ximo de columnas en tabla antes de warning
MAX_SECTIONS_PER_BLOCK = 1  # MÃ¡ximo de secciones por bloque (en notebooks)
MAX_HEADER_FOOTER_CALLS = 2  # MÃ¡ximo de llamadas a header/footer por bloque
MAX_IMAGE_DIMENSION = 20  # MÃ¡ximo ancho/alto en pulgadas
MAX_STYLES_PER_BLOCK = 10  # Máximo de estilos nuevos por bloque

# Deteccion agresiva de cantidades en texto libre (magnitude + unidad).
_UNIT_QUANTITY_RE = re.compile(
    r"(?P<magnitude>[+-]?(?:\d+\.\d+|\d+|\.\d+)(?:[eE][+-]?\d+)?)"
    r"(?P<space>\s+)"
    r"(?P<unit>[A-Za-z°µ_][A-Za-z0-9_°µ/*^().+\-\u00B2\u00B3\u2070-\u2079\u207B]*)"
)
_UNIT_TRAILING_PUNCT = ",.;:!?)]"


def _split_unit_token(unit_token: str) -> tuple[str, str]:
    if not unit_token:
        return "", ""
    split_at = len(unit_token)
    while split_at > 0 and unit_token[split_at - 1] in _UNIT_TRAILING_PUNCT:
        split_at -= 1
    return unit_token[:split_at], unit_token[split_at:]


_UNIT_VALIDATION_CACHE: dict[str, bool] = {}


def _is_valid_unit_token(unit_token: str) -> bool:
    token = str(unit_token or "").strip()
    if not token:
        return False
    cached = _UNIT_VALIDATION_CACHE.get(token)
    if cached is not None:
        return cached
    try:
        valid, _ = validate_unit_expression(token)
        result = bool(valid)
    except Exception:
        result = False
    _UNIT_VALIDATION_CACHE[token] = result
    return result


class DocxWarning(UserWarning):
    """Warning especÃ­fico para la API DOCX."""
    pass


def _emit_warning(message: str, category: type = DocxWarning) -> None:
    """Emite un warning que serÃ¡ visible en el notebook."""
    warnings.warn(f"[DOCX API] {message}", category, stacklevel=4)


def _normalize_alignment(value: AlignmentValue):
    if value is None or WD_ALIGN_PARAGRAPH is None:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(str(value).lower())


def _set_run_picture_alt_text(run: Any, alt_text: Optional[str]) -> None:
    text = str(alt_text or "").strip()
    if not text:
        return
    try:
        for docpr in run._r.xpath(".//wp:docPr"):
            docpr.set("descr", text)
            if not docpr.get("title"):
                docpr.set("title", text[:128])
    except Exception:
        return


def _mark_row_as_repeating_header(row: Any) -> None:
    try:
        tr_pr = row._tr.get_or_add_trPr()
        existing = tr_pr.find(qn("w:tblHeader"))
        if existing is None:
            existing = OxmlElement("w:tblHeader")
            tr_pr.append(existing)
        existing.set(qn("w:val"), "true")
    except Exception:
        return


def _set_cell_padding(cell: Any, padding_twips: int) -> None:
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for side in ("top", "left", "bottom", "right"):
            node = tc_mar.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(int(padding_twips)))
            node.set(qn("w:type"), "dxa")
    except Exception:
        return


def _set_cell_vertical_alignment(cell: Any, align: Optional[str]) -> None:
    if not align:
        return
    value = str(align).strip().lower()
    if value not in {"top", "center", "bottom"}:
        return
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        node = tc_pr.find(qn("w:vAlign"))
        if node is None:
            node = OxmlElement("w:vAlign")
            tc_pr.append(node)
        node.set(qn("w:val"), value)
    except Exception:
        return


def _apply_paragraph_style(paragraph: Paragraph, style_name: str, *, fallback_style: str = "Normal") -> bool:
    try:
        paragraph.style = style_name
        return True
    except KeyError:
        try:
            paragraph.style = fallback_style
        except KeyError:
            pass
        return False


def _style_has_numbering(style_obj: Any) -> bool:
    if style_obj is None:
        return False
    try:
        element = style_obj.element
        p_pr = element.find(qn("w:pPr"))
        if p_pr is None:
            return False
        return p_pr.find(qn("w:numPr")) is not None
    except Exception:
        return False


def _apply_visible_run_fallback(
    run,
    *,
    font_name: Optional[str] = None,
    font_size_pt: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
) -> None:
    if font_name:
        run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(float(font_size_pt))
    if bold is not None:
        run.bold = bool(bold)
    if italic is not None:
        run.italic = bool(italic)
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass


_TABLE_STYLE_LOOK_DEFAULTS = {
    "firstRow": True,
    "lastRow": False,
    "firstColumn": False,
    "lastColumn": False,
    "noHBand": False,
    "noVBand": True,
}


def _coerce_boolish(value: Any) -> Optional[bool]:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return None


def _coerce_optional_int(value: Any) -> Optional[int]:
    if value in (None, ""):
        return None
    try:
        return int(round(float(value)))
    except (TypeError, ValueError):
        return None


def _resolve_table_runtime_defaults(
    defaults_map: Optional[Dict[str, Dict[str, Any]]],
    table_style: Any,
    requested_style: Optional[str],
) -> Optional[Dict[str, Any]]:
    if not isinstance(defaults_map, dict) or not defaults_map:
        return None

    style_id = getattr(table_style, "style_id", None) if table_style is not None else None
    style_name = getattr(table_style, "name", None) if table_style is not None else None

    if style_id and style_id in defaults_map:
        return defaults_map[style_id]

    for candidate_name in (style_name, requested_style):
        if candidate_name in (None, ""):
            continue
        for entry in defaults_map.values():
            if not isinstance(entry, dict):
                continue
            if entry.get("style_name") == candidate_name:
                return entry
    return None


def _apply_table_runtime_defaults(table: Table, runtime_defaults: Optional[Dict[str, Any]]) -> None:
    if not runtime_defaults or not isinstance(runtime_defaults, dict):
        return

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for tag in ("tblLook", "tblLayout", "tblW"):
        existing = tbl_pr.find(qn(f"w:{tag}"))
        if existing is not None:
            tbl_pr.remove(existing)

    look = runtime_defaults.get("look")
    if isinstance(look, dict) and look:
        tbl_look = OxmlElement("w:tblLook")
        for key, default_value in _TABLE_STYLE_LOOK_DEFAULTS.items():
            bool_value = _coerce_boolish(look.get(key))
            resolved_bool = default_value if bool_value is None else bool_value
            tbl_look.set(qn(f"w:{key}"), "1" if resolved_bool else "0")
        tbl_pr.append(tbl_look)

    layout_type = runtime_defaults.get("layout_type")
    if layout_type not in (None, ""):
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), str(layout_type).strip().lower())
        tbl_pr.append(tbl_layout)

    width_type = runtime_defaults.get("width_type")
    width_value = runtime_defaults.get("width_value")
    if width_type not in (None, "") or width_value not in (None, ""):
        tbl_w = OxmlElement("w:tblW")
        resolved_width_type = str(width_type or "auto").strip().lower()
        tbl_w.set(qn("w:type"), resolved_width_type)
        if resolved_width_type == "auto":
            tbl_w.set(qn("w:w"), "0")
        else:
            tbl_w.set(qn("w:w"), str(_coerce_optional_int(width_value) or 0))
        tbl_pr.append(tbl_w)


class DocBuilder:
    """Interfaz encadenable para componer documentos DOCX."""

    # MÃ©todos del Document que crean elementos rastreables
    _TRACKABLE_METHODS = ('add_paragraph', 'add_table', 'add_picture', 'add_heading', 'add_page_break')

    def __init__(self, session: DocxSession, handle: CellHandle, on_close: Callable[[], None]):
        self._session = session
        self._handle = handle
        self._context_finalizer = weakref.finalize(self, on_close)
        self._log_finalizer = weakref.finalize(self, handle.log_event, "lifecycle", "builder-finalized")
        
        # Contadores para detectar uso problemÃ¡tico
        self._section_count = 0
        self._header_count = 0
        self._footer_count = 0
        self._style_count = 0
        self._total_elements = 0
        self._registered_labels: set = set()
        
        # Almacenar mÃ©todos originales y aplicar monkey-patching
        self._original_methods: dict = {}
        self._patch_document_methods()

    def _patch_document_methods(self) -> None:
        """Intercepta los mÃ©todos add_* del Document para registrar elementos automÃ¡ticamente."""
        doc = self._handle.doc
        
        for method_name in self._TRACKABLE_METHODS:
            if hasattr(doc, method_name):
                original = getattr(doc, method_name)
                self._original_methods[method_name] = original
                
                # Crear wrapper que registra el elemento
                def make_wrapper(orig_method, tracked_name):
                    def wrapper(*args, **kwargs):
                        result = orig_method(*args, **kwargs)
                        # Registrar el elemento creado
                        if result is not None:
                            default_preview = None
                            if args:
                                first_arg = args[0]
                                if isinstance(first_arg, str):
                                    default_preview = first_arg
                            pending_provenance = self._session.ns.get(self._session._PENDING_PROVENANCE_KEY)
                            self._handle.register_element(
                                result,
                                provenance=(
                                    None
                                    if isinstance(pending_provenance, dict)
                                    else {
                                        "api_name": f"document.{tracked_name}",
                                        "element_kind": (
                                            "table" if tracked_name == "add_table"
                                            else "picture" if tracked_name == "add_picture"
                                            else "paragraph"
                                        ),
                                        "text_preview": default_preview,
                                        "precision": "fallback",
                                    }
                                ),
                            )
                        return result
                    return wrapper
                
                setattr(doc, method_name, make_wrapper(original, method_name))

    def _restore_document_methods(self) -> None:
        """Restaura los mÃ©todos originales del Document."""
        doc = self._handle.doc
        for method_name, original in self._original_methods.items():
            if hasattr(doc, method_name):
                setattr(doc, method_name, original)
        self._original_methods.clear()

    # ------------------------------------------------------------------
    # Propiedades pÃºblicas (Acceso de bajo nivel)
    # ------------------------------------------------------------------
    @property
    def document(self):
        """Acceso rastreado al objeto Document de python-docx subyacente.
        
        Permite el uso de toda la funcionalidad de python-docx que no estÃ©
        expuesta explÃ­citamente en el builder.
        
        Los elementos y mutaciones visibles se instrumentan para conservar
        procedencia exacta/callsite mientras el acceso permanezca dentro del
        proxy. Para omitir este tracking usar ``document_raw``.
        """
        return wrap_docx_proxy(
            self._handle.doc,
            session=self._session,
            handle=self._handle,
            root_element=self._handle.doc,
        )

    @property
    def document_raw(self):
        """Escape hatch al Document real sin garantÃ­a de procedencia fina."""
        return self._handle.doc

    def create_math_element(self, expression: str):
        """Parsea una expresiÃ³n matemÃ¡tica y devuelve el elemento XML (OMML).
        
        Ãštil para insertar ecuaciones inline dentro de pÃ¡rrafos de texto
        o en estructuras personalizadas donde math() no es suficiente.
        
        Args:
            expression: FÃ³rmula en sintaxis natural (no LaTeX).
            
        Returns:
            Elemento lxml (OMML) listo para usar con append().
            Retorna None si el parser no estÃ¡ disponible.
            
        Raises:
            ValueError: Si hay un error de sintaxis en la expresiÃ³n.
        """
        parser = self._handle.parser
        if parser is None:
            return None
            
        result = parser.parse_expression(expression)
        if not getattr(result, "success", False):
            error_msg = getattr(result, "error_message", "error al parsear expresiÃ³n")
            raise ValueError(f"Error parseando '{expression}': {error_msg}")
            
        return result.omml_element

    def create_math_latex_element(self, expression: str):
        """Convierte una expresion LaTeX matematica inline a OMML."""
        converter = self._handle.latex_converter
        if converter is None:
            raise LatexMathUnavailableError("Conversor LaTeX -> OMML no disponible.")
        try:
            return converter.create_omml_element(expression, inline=True)
        except LatexMathError:
            raise
        except Exception as exc:
            raise ValueError(f"Error convirtiendo LaTeX '{expression}': {exc}") from exc

    def resolve_style_slot(self, slot_name: str) -> Optional[str]:
        """Resolve the active Word style name for a semantic slot."""
        return self._session.resolve_style_slot(slot_name)

    # ------------------------------------------------------------------
    # GestiÃ³n bÃ¡sica
    # ------------------------------------------------------------------
    def close(self) -> None:
        # Restaurar mÃ©todos originales del Document
        self._restore_document_methods()
        
        if self._context_finalizer.alive:
            try:
                self._context_finalizer()  # type: ignore[call-arg]
            except Exception:
                self._handle.log_event("warning", "on_close fallÃ³")
        if self._log_finalizer.alive:
            self._log_finalizer()  # type: ignore[call-arg]


    def __enter__(self) -> "DocBuilder":
        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        self.close()

    # ------------------------------------------------------------------
    # Helpers internos
    # ------------------------------------------------------------------
    def _prepare_next_provenance(
        self,
        *,
        api_name: str,
        element_kind: str,
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> None:
        self._session.prepare_next_provenance(
            api_name=api_name,
            element_kind=element_kind,
            text_preview=text_preview,
            precision=precision,
        )

    def _new_paragraph(
        self,
        *,
        api_name: Optional[str] = None,
        element_kind: str = "paragraph",
        text_preview: Optional[str] = None,
        precision: str = "exact",
    ) -> Paragraph:
        # NOTE: add_paragraph() is monkey-patched in _patch_document_methods()
        # to auto-register elements â€” do NOT call register_element() again here.
        if api_name:
            self._prepare_next_provenance(
                api_name=api_name,
                element_kind=element_kind,
                text_preview=text_preview,
                precision=precision,
            )
        paragraph = self._handle.doc.add_paragraph()
        return paragraph

    def _fail(self, kind: str, exc: Exception) -> None:
        message = f"{kind} error: {exc}"
        # Evitar colisiÃ³n con el argumento posicional `kind`
        self._handle.log_event("error", message, event_kind=kind, error=str(exc))
        if self._session.is_strict_mode():
            raise StrictModeError(message) from exc
        p = self._new_paragraph()
        p.add_run(f"[{kind} error] {exc}")

    def _bookmark(self, paragraph: Paragraph, label: str, *, container: Any = None) -> None:
        bookmark_id = self._session.next_bookmark_id()
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), str(label))
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        if container is None:
            paragraph._p.insert(0, start)
            paragraph._p.append(end)
            return
        container.append(start)
        container.append(end)

    def _register_requested_label(self, label: Optional[str]) -> None:
        if not label:
            return
        if label in self._registered_labels:
            _emit_warning(
                f"Label '{label}' ya fue usado en este bloque. "
                f"Las referencias pueden ser ambiguas."
            )
        self._registered_labels.add(label)

    @staticmethod
    def _normalize_caption_position(position: Optional[str], *, default: str) -> str:
        normalized = str(position or default).strip().lower() or default
        if normalized not in {"above", "below"}:
            _emit_warning(
                f"caption_position inválido ({position!r}). "
                f"Usando '{default}'."
            )
            return default
        return normalized

    @staticmethod
    def _normalize_sequence_name(caption_label: Optional[str], *, default: str) -> str:
        normalized = str(caption_label or "").strip()
        return normalized or default

    @staticmethod
    def _field_sequence_identifier(sequence_name: str) -> str:
        normalized = re.sub(r"\s+", "_", str(sequence_name or "").strip())
        normalized = re.sub(r"[^0-9A-Za-z_À-ÿ]", "_", normalized)
        return normalized or "Figura"

    def _append_simple_field(
        self,
        paragraph: Paragraph,
        *,
        instruction: str,
        display_text: str,
        bookmark_label: Optional[str] = None,
        fallback_font_size_pt: Optional[float] = None,
        fallback_italic: Optional[bool] = None,
        fallback_bold: Optional[bool] = None,
    ) -> None:
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), str(instruction))
        bookmark_id = self._session.next_bookmark_id() if bookmark_label else None
        if bookmark_label and bookmark_id is not None:
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), str(bookmark_id))
            start.set(qn("w:name"), str(bookmark_label))
            fld.append(start)

        run = OxmlElement("w:r")
        if any(value is not None for value in (fallback_font_size_pt, fallback_italic, fallback_bold)):
            r_pr = OxmlElement("w:rPr")
            if fallback_bold:
                r_pr.append(OxmlElement("w:b"))
            if fallback_italic:
                r_pr.append(OxmlElement("w:i"))
            if fallback_font_size_pt is not None:
                size_value = max(1, int(round(float(fallback_font_size_pt) * 2)))
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(size_value))
                r_pr.append(sz)
                sz_cs = OxmlElement("w:szCs")
                sz_cs.set(qn("w:val"), str(size_value))
                r_pr.append(sz_cs)
            run.append(r_pr)
        text = OxmlElement("w:t")
        text.text = str(display_text)
        run.append(text)
        fld.append(run)
        if bookmark_label and bookmark_id is not None:
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), str(bookmark_id))
            fld.append(end)
        paragraph._p.append(fld)

    def _insert_caption_paragraph(
        self,
        text: str,
        *,
        sequence_name: Optional[str] = None,
        number_value: Optional[int] = None,
        label: Optional[str] = None,
        api_name: str = "caption",
    ) -> Paragraph:
        paragraph = self._new_paragraph(
            api_name=api_name,
            element_kind="caption",
            text_preview=text,
        )
        target_style = self.resolve_style_slot("caption") or "Caption"
        style_applied = _apply_paragraph_style(paragraph, target_style)
        if not style_applied:
            _emit_warning(f"Estilo '{target_style}' no encontrado. Usando estilo por defecto.")

        if sequence_name and number_value is not None:
            prefix_run = paragraph.add_run(f"{sequence_name} ")
            if not style_applied:
                _apply_visible_run_fallback(prefix_run, font_size_pt=9, italic=True)
            self._append_simple_field(
                paragraph,
                instruction=f"SEQ {self._field_sequence_identifier(sequence_name)} \\* ARABIC",
                display_text=str(number_value),
                bookmark_label=label,
                fallback_font_size_pt=9 if not style_applied else None,
                fallback_italic=True if not style_applied else None,
            )
            suffix_text = f". {text}" if text else ""
            if suffix_text:
                suffix_run = paragraph.add_run(suffix_text)
                if not style_applied:
                    _apply_visible_run_fallback(suffix_run, font_size_pt=9, italic=True)
            return paragraph

        run = paragraph.add_run(str(text))
        if not style_applied:
            _apply_visible_run_fallback(run, font_size_pt=9, italic=True)
        if label:
            self._bookmark(paragraph, label)
        return paragraph

    def _append_text_run(
        self,
        paragraph: Paragraph,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        superscript: bool = False,
    ) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bool(bold)
        run.italic = bool(italic)
        run.underline = bool(underline)
        if superscript:
            try:
                run.font.superscript = True
            except Exception:
                pass

    def _render_quantity_runs(
        self,
        paragraph: Paragraph,
        *,
        magnitude: str,
        unit_token: str,
        trailing: str = "",
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
    ) -> None:
        self._append_text_run(
            paragraph,
            magnitude,
            bold=bold,
            italic=italic,
            underline=underline,
        )
        # Thin space between magnitude and unit for typographic quality.
        self._append_text_run(
            paragraph,
            "\u2009",
            bold=bold,
            italic=italic,
            underline=underline,
        )
        for unit_run in build_docx_unit_runs(unit_token):
            unit_text = str(unit_run.get("text", ""))
            if not unit_text:
                continue
            self._append_text_run(
                paragraph,
                unit_text,
                bold=bold,
                italic=bool(unit_run.get("italic", True) or italic),
                underline=underline,
                superscript=bool(unit_run.get("superscript")),
            )
        if trailing:
            self._append_text_run(
                paragraph,
                trailing,
                bold=bold,
                italic=italic,
                underline=underline,
            )

    def _render_text_with_aggressive_units(
        self,
        paragraph: Paragraph,
        raw_text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
    ) -> bool:
        text = str(raw_text)
        if not text:
            return False

        segments: list[dict[str, Any]] = []
        last_index = 0

        for match in _UNIT_QUANTITY_RE.finditer(text):
            start, end = match.span()
            unit_token = match.group("unit") or ""
            unit_core, trailing = _split_unit_token(unit_token)
            if not unit_core:
                continue
            if not _is_valid_unit_token(unit_core):
                continue

            if start > last_index:
                segments.append({"kind": "plain", "text": text[last_index:start]})

            segments.append(
                {
                    "kind": "quantity",
                    "magnitude": match.group("magnitude") or "",
                    "unit_token": unit_core,
                    "unit_runs": build_docx_unit_runs(unit_core),
                    "trailing": trailing,
                }
            )
            last_index = end

        if not any(segment.get("kind") == "quantity" for segment in segments):
            return False

        if last_index < len(text):
            segments.append({"kind": "plain", "text": text[last_index:]})

        for segment in segments:
            kind = segment.get("kind")
            if kind == "plain":
                self._append_text_run(
                    paragraph,
                    str(segment.get("text", "")),
                    bold=bold,
                    italic=italic,
                    underline=underline,
                )
                continue

            magnitude = str(segment.get("magnitude", ""))
            unit_token = str(segment.get("unit_token", ""))
            trailing = str(segment.get("trailing", ""))
            unit_runs = segment.get("unit_runs")
            if not isinstance(unit_runs, list):
                unit_runs = build_docx_unit_runs(unit_token)

            self._append_text_run(
                paragraph,
                magnitude,
                bold=bold,
                italic=italic,
                underline=underline,
            )
            self._append_text_run(
                paragraph,
                "\u2009",
                bold=bold,
                italic=italic,
                underline=underline,
            )
            for unit_run in unit_runs:
                unit_text = str(unit_run.get("text", ""))
                if not unit_text:
                    continue
                self._append_text_run(
                    paragraph,
                    unit_text,
                    bold=bold,
                    italic=bool(unit_run.get("italic", True) or italic),
                    underline=underline,
                    superscript=bool(unit_run.get("superscript")),
                )
            if trailing:
                self._append_text_run(
                    paragraph,
                    trailing,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                )

        return True

    # ------------------------------------------------------------------
    # Bloques de contenido
    # ------------------------------------------------------------------
    def heading(self, text: str, *, level: int = 1, style: Optional[str] = None) -> "DocBuilder":
        """Inserta un encabezado.
        
        text: Texto del encabezado.
        level: Nivel del encabezado (1-6).
        style: Estilo personalizado (opcional).
        """
        # Validaciones
        if not text:
            _emit_warning("heading() llamado con texto vacÃ­o.")
        
        try:
            lvl = int(level)
            if lvl < 1 or lvl > 6:
                _emit_warning(f"level debe estar entre 1 y 6 (recibido: {lvl}). Ajustando.")
            lvl = max(1, min(6, lvl))
        except (ValueError, TypeError):
            _emit_warning(f"level invÃ¡lido ({level}). Usando nivel 1.")
            lvl = 1
        
        try:
            paragraph = self._new_paragraph(
                api_name="heading",
                element_kind="heading",
                text_preview=text,
            )
            target_style = style or self.resolve_style_slot(f"heading_{lvl}") or f"Heading {lvl}"
            style_applied = _apply_paragraph_style(paragraph, target_style)
            if not style_applied:
                _emit_warning(f"Estilo '{target_style}' no encontrado. Usando estilo por defecto.")
            run = paragraph.add_run(str(text))
            if not style_applied:
                heading_sizes = {1: 16, 2: 14, 3: 12, 4: 11, 5: 11, 6: 11}
                _apply_visible_run_fallback(
                    run,
                    font_size_pt=heading_sizes.get(lvl, 11),
                    bold=(lvl <= 5),
                    italic=(lvl in {4, 6}),
                )
        except Exception as exc:
            self._fail("heading", exc)
        return self

    def text(
        self,
        text: str,
        *,
        style: Optional[str] = None,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        align: AlignmentValue = None,
    ) -> "DocBuilder":
        """Inserta un pÃ¡rrafo de texto.
        
        text: Contenido del pÃ¡rrafo.
        style: Estilo de pÃ¡rrafo (opcional).
        bold: Texto en negrita.
        italic: Texto en cursiva.
        underline: Texto subrayado.
        align: AlineaciÃ³n ('left', 'center', 'right', 'justify').
        """
        # ValidaciÃ³n de alineaciÃ³n
        if align is not None:
            valid_aligns = {"left", "center", "right", "justify"}
            if str(align).lower() not in valid_aligns:
                _emit_warning(
                    f"AlineaciÃ³n '{align}' no vÃ¡lida. "
                    f"Opciones: {', '.join(valid_aligns)}. Ignorando."
                )
                align = None
        
        try:
            paragraph = self._new_paragraph(
                api_name="text",
                element_kind="paragraph",
                text_preview=text,
            )
            target_style = style or self.resolve_style_slot("body")
            if target_style:
                try:
                    paragraph.style = target_style
                except KeyError:
                    _emit_warning(f"Estilo '{target_style}' no encontrado. Usando estilo por defecto.")
                    try:
                        paragraph.style = "Normal"
                    except KeyError:
                        pass
            alignment = _normalize_alignment(align)
            if alignment is not None:
                paragraph.alignment = alignment
            plain_text = str(text)
            try:
                rendered_quantity = self._render_text_with_aggressive_units(
                    paragraph,
                    plain_text,
                    bold=bool(bold),
                    italic=bool(italic),
                    underline=bool(underline),
                )
            except Exception:
                rendered_quantity = False

            if not rendered_quantity:
                self._append_text_run(
                    paragraph,
                    plain_text,
                    bold=bool(bold),
                    italic=bool(italic),
                    underline=bool(underline),
                )
        except Exception as exc:
            self._fail("text", exc)
        return self

    def list(self, items: Sequence[Any], *, ordered: bool = False) -> "DocBuilder":
        """Inserta una lista con viÃ±etas o numerada.
        
        items: Secuencia de strings o tuplas (nivel, texto) para multinivel.
        ordered: True para lista numerada, False para viÃ±etas.
        """
        # Validaciones
        if not items:
            _emit_warning("list() llamado con items vacÃ­o. No se insertÃ³ ningÃºn elemento.")
            return self
        
        item_count = len(items)
        if item_count > MAX_LIST_ITEMS:
            _emit_warning(
                f"Lista con {item_count} items (mÃ¡ximo recomendado: {MAX_LIST_ITEMS}). "
                f"Listas muy largas pueden afectar el rendimiento."
            )
        
        style_name = (
            self.resolve_style_slot("list_number" if ordered else "list_bullet")
            or ("List Number" if ordered else "List Bullet")
        )
        ordered_counters: Dict[int, int] = {}
        for item in items:
            try:
                level = 0
                text = item
                if isinstance(item, (list, tuple)) and len(item) >= 2:
                    level = int(item[0])
                    text = item[1]
                    # Validar nivel de indentaciÃ³n razonable
                    if level > 9:
                        _emit_warning(f"Nivel de lista muy profundo ({level}). MÃ¡ximo recomendado: 9.")
                        level = min(level, 9)
                paragraph = self._new_paragraph(
                    api_name="list",
                    element_kind="list-item",
                    text_preview=text,
                )
                style_applied = _apply_paragraph_style(paragraph, style_name)
                if not style_applied:
                    _emit_warning(f"Estilo '{style_name}' no encontrado. Usando estilo 'Normal'.")
                use_explicit_markers = (not style_applied) or (not _style_has_numbering(paragraph.style))
                if use_explicit_markers and style_applied:
                    _emit_warning(
                        f"Estilo '{style_name}' no tiene numeración/viñeta válida. "
                        f"Se aplicará fallback visible."
                    )
                if use_explicit_markers:
                    if ordered:
                        for stale_level in list(ordered_counters):
                            if stale_level > level:
                                ordered_counters.pop(stale_level, None)
                        ordered_counters[level] = ordered_counters.get(level, 0) + 1
                    base_indent_pt = 36 + (18 * max(level, 0))
                    paragraph.paragraph_format.left_indent = Pt(base_indent_pt)
                    paragraph.paragraph_format.first_line_indent = Pt(-18)
                    marker = "\u2022"
                    if ordered:
                        marker = f"{ordered_counters[level]}."
                    run = paragraph.add_run(f"{marker} {text}")
                    _apply_visible_run_fallback(run, font_size_pt=11)
                else:
                    paragraph.add_run(str(text))
                    if level > 0:
                        paragraph.paragraph_format.left_indent = Pt(14 * level)
            except Exception as exc:
                self._fail("list", exc)
                break
        return self

    def code(self, text: str, *, language: Optional[str] = None) -> "DocBuilder":
        """Inserta un bloque de cÃ³digo con fuente monoespaciada.
        
        text: CÃ³digo a insertar.
        language: Lenguaje (informativo, no afecta el formato).
        """
        if not text:
            _emit_warning("code() llamado con texto vacÃ­o.")
            return self
        
        # Advertir si el cÃ³digo es muy largo
        lines = str(text).count('\n') + 1
        if lines > 100:
            _emit_warning(
                f"Bloque de cÃ³digo con {lines} lÃ­neas. "
                f"Bloques muy largos pueden afectar la legibilidad."
            )
        
        try:
            paragraph = self._new_paragraph(
                api_name="code",
                element_kind="code",
                text_preview=text,
            )
            target_style = self.resolve_style_slot("code") or "Code"
            style_applied = _apply_paragraph_style(paragraph, target_style)
            if not style_applied:
                _emit_warning(f"Estilo '{target_style}' no encontrado. Usando estilo 'Normal'.")
            run = paragraph.add_run(str(text))
            if not style_applied:
                _apply_visible_run_fallback(run, font_size_pt=10)
        except Exception as exc:
            self._fail("code", exc)
        return self

    def link(self, text: str, url: str) -> "DocBuilder":
        """Inserta un hipervÃ­nculo.
        
        text: Texto visible del enlace.
        url: URL de destino.
        """
        # Validaciones
        if not text:
            _emit_warning("link() llamado con texto vacÃ­o. Usando URL como texto.")
            text = url
        
        if not url:
            _emit_warning("link() llamado con URL vacÃ­a. No se creÃ³ el enlace.")
            return self
        
        # Validar formato bÃ¡sico de URL
        url_str = str(url).strip()
        if not url_str.startswith(('http://', 'https://', 'mailto:', 'ftp://')):
            _emit_warning(
                f"URL '{url_str[:50]}...' no parece vÃ¡lida. "
                f"Las URLs deben comenzar con http://, https://, mailto:, etc."
            )
        
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore

            paragraph = self._new_paragraph(
                api_name="link",
                element_kind="link",
                text_preview=text,
            )
            hyperlink = OxmlElement("w:hyperlink")
            r_id = paragraph.part.relate_to(url_str, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink.set(qn("r:id"), r_id)

            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0000FF")
            r_pr.append(u)
            r_pr.append(color)
            t = OxmlElement("w:t")
            t.text = str(text)
            r.append(r_pr)
            r.append(t)
            hyperlink.append(r)
            paragraph._p.append(hyperlink)
        except Exception as exc:
            self._fail("link", exc)
        return self

    def page_break(self) -> "DocBuilder":
        """Inserta un salto de pÃ¡gina.
        
        Este es el mÃ©todo recomendado para separar contenido en pÃ¡ginas
        diferentes, especialmente en notebooks donde section() puede
        causar problemas.
        """
        try:
            paragraph = self._new_paragraph(
                api_name="page_break",
                element_kind="page-break",
                precision="fallback",
            )
            run = paragraph.add_run()
            if WD_BREAK is not None:
                run.add_break(WD_BREAK.PAGE)
            else:  # pragma: no cover - fallback antiguo
                run.add_break()
        except Exception as exc:
            self._fail("page-break", exc)
        return self

    def metadata(
        self,
        *,
        title: Optional[str] = None,
        subject: Optional[str] = None,
        keywords: Optional[Iterable[str]] = None,
    ) -> "DocBuilder":
        """Configura las propiedades del documento.
        
        title: TÃ­tulo del documento.
        subject: Asunto/descripciÃ³n.
        keywords: Lista de palabras clave.
        
        Nota: Llamadas mÃºltiples sobrescriben valores anteriores.
        """
        # Validar que al menos un parÃ¡metro fue proporcionado
        if title is None and subject is None and keywords is None:
            _emit_warning("metadata() llamado sin parÃ¡metros. No se realizÃ³ ningÃºn cambio.")
            return self
        
        try:
            props = self._handle.doc.core_properties
            if title is not None:
                if len(str(title)) > 255:
                    _emit_warning("title muy largo (>255 caracteres). Puede truncarse.")
                props.title = str(title)
            if subject is not None:
                if len(str(subject)) > 255:
                    _emit_warning("subject muy largo (>255 caracteres). Puede truncarse.")
                props.subject = str(subject)
            if keywords is not None:
                kw_list = list(keywords)
                if len(kw_list) > 50:
                    _emit_warning(f"Muchas keywords ({len(kw_list)}). Considere reducir la lista.")
                props.keywords = ",".join(str(k) for k in kw_list)
        except Exception as exc:
            self._fail("metadata", exc)
        return self

    def style(
        self,
        name: str,
        *,
        base: str = "Normal",
        font: Optional[str] = None,
        size_pt: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        spacing: Optional[Dict[str, float]] = None,
    ) -> "DocBuilder":
        """Crea o modifica un estilo de pÃ¡rrafo.
        
        name: Nombre del estilo (se crea si no existe).
        base: Estilo base (default: 'Normal').
        font: Nombre de la fuente.
        size_pt: TamaÃ±o en puntos.
        bold: Negrita.
        italic: Cursiva.
        spacing: Dict con space_before_pt, space_after_pt, line_spacing.
        """
        # Validaciones
        if not name or not name.strip():
            _emit_warning("style() requiere un nombre vÃ¡lido. No se creÃ³ ningÃºn estilo.")
            return self
        
        self._style_count += 1
        if self._style_count > MAX_STYLES_PER_BLOCK:
            _emit_warning(
                f"Muchos estilos creados en un solo bloque ({self._style_count}). "
                f"Considere definir estilos en un bloque de configuraciÃ³n separado."
            )
        
        if size_pt is not None:
            size_f = float(size_pt)
            if size_f <= 0:
                _emit_warning(f"size_pt debe ser positivo ({size_f}). Ignorando.")
                size_pt = None
            elif size_f > 100:
                _emit_warning(f"size_pt muy grande ({size_f}). Â¿EstÃ¡ seguro?")
        
        try:
            styles = self._handle.doc.styles
            is_new = False
            try:
                style_obj = styles[name]
                # Estilo ya existe, se estÃ¡ redefiniendo
                self._handle.log_event(
                    "info", 
                    f"Estilo '{name}' ya existe, se estÃ¡ modificando"
                )
            except KeyError:
                from docx.enum.style import WD_STYLE_TYPE  # type: ignore

                style_obj = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                is_new = True
                base_style = None
                if base:
                    try:
                        base_style = styles[base]
                    except KeyError:
                        _emit_warning(f"Estilo base '{base}' no encontrado. Usando 'Normal'.")
                        base_style = None
                if base_style is None:
                    base_style = styles["Normal"]
                style_obj.base_style = base_style
            
            if font is not None:
                style_obj.font.name = str(font)
            if size_pt is not None:
                style_obj.font.size = Pt(float(size_pt))
            if bold is not None:
                style_obj.font.bold = bool(bold)
            if italic is not None:
                style_obj.font.italic = bool(italic)
            if spacing:
                pf = style_obj.paragraph_format
                if "space_before_pt" in spacing:
                    pf.space_before = Pt(float(spacing["space_before_pt"]))
                if "space_after_pt" in spacing:
                    pf.space_after = Pt(float(spacing["space_after_pt"]))
                if "line_spacing" in spacing:
                    ls = float(spacing["line_spacing"])
                    if ls <= 0:
                        _emit_warning(f"line_spacing debe ser positivo ({ls}).")
                    else:
                        pf.line_spacing = ls
        except Exception as exc:
            self._fail("style", exc)
        return self

    # ------------------------------------------------------------------
    # MatemÃ¡tica y referencias
    # ------------------------------------------------------------------
    def _register_equation_label(self, label: Optional[str]) -> None:
        if not label:
            return
        if label in self._registered_labels:
            _emit_warning(
                f"Label '{label}' ya fue usado en este bloque. "
                f"Las referencias pueden ser ambiguas."
            )
        self._registered_labels.add(label)

    def _finalize_equation_paragraph(
        self,
        paragraph: Paragraph,
        *,
        label: Optional[str],
        number: bool,
    ) -> None:
        number_value = None
        if number:
            number_value = self._session.next_equation_number()
            paragraph.add_run(f"  ({number_value})")
        if label:
            self._bookmark(paragraph, label)
            self._session.register_label(label, "equation", number=number_value)

    def math(self, expression: str, *, label: Optional[str] = None, number: bool = False) -> "DocBuilder":
        """Inserta una ecuaciÃ³n matemÃ¡tica.
        
        expression: ExpresiÃ³n matemÃ¡tica (sintaxis LaTeX-like).
        label: Etiqueta para referencias cruzadas.
        number: Si True, numera la ecuaciÃ³n.
        """
        # Validaciones
        if not expression or not expression.strip():
            _emit_warning("math() llamado con expresiÃ³n vacÃ­a.")
            return self

        self._register_equation_label(label)
        
        paragraph = self._new_paragraph(
            api_name="math",
            element_kind="equation",
            text_preview=expression,
        )
        parser = self._handle.parser
        try:
            if parser is None:
                _emit_warning(
                    "Parser matemÃ¡tico no disponible. "
                    "La ecuaciÃ³n se insertarÃ¡ como texto plano."
                )
                paragraph.add_run(str(expression))
                return self
            result = parser.parse_expression(expression)
            if not getattr(result, "success", False):
                error_msg = getattr(result, "error_message", "error al parsear expresiÃ³n")
                _emit_warning(f"Error en ecuaciÃ³n: {error_msg}. Se insertarÃ¡ como texto.")
                raise ValueError(error_msg)
            paragraph._p.append(result.omml_element)
            self._finalize_equation_paragraph(paragraph, label=label, number=number)
        except Exception as exc:
            self._fail("equation", exc)
        return self

    def math_latex(self, expression: str, *, label: Optional[str] = None, number: bool = False) -> "DocBuilder":
        """Inserta una ecuacion matematica usando LaTeX math-only."""

        if not expression or not expression.strip():
            _emit_warning("math_latex() llamado con expresion vacia.")
            return self

        self._register_equation_label(label)

        paragraph = self._new_paragraph(
            api_name="math_latex",
            element_kind="equation",
            text_preview=expression,
        )
        try:
            converter = self._handle.latex_converter
            if converter is None:
                raise LatexMathUnavailableError("Conversor LaTeX -> OMML no disponible.")
            omml_element = converter.create_omml_element(expression, inline=False)
            paragraph._p.append(omml_element)
            self._finalize_equation_paragraph(paragraph, label=label, number=number)
        except LatexMathError as exc:
            self._fail("equation_latex", exc)
        except Exception as exc:
            self._fail("equation_latex", exc)
        return self

    def reference(self, label: str) -> "DocBuilder":
        """Inserta una referencia cruzada a una etiqueta.
        
        label: Etiqueta de la ecuaciÃ³n, figura o caption a referenciar.
        """
        # Validaciones
        if not label or not label.strip():
            _emit_warning("reference() llamado con label vacÃ­o.")
            return self
        
        info = self._session.resolve_label(label)
        
        if info is None:
            _emit_warning(
                f"Label '{label}' no encontrado. "
                f"AsegÃºrese de que la ecuaciÃ³n/figura con ese label se haya creado antes. "
                f"Se mostrarÃ¡ [{label}] como placeholder."
            )
        
        paragraph = self._new_paragraph(
            api_name="reference",
            element_kind="reference",
            text_preview=label,
        )
        try:
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), f"REF {label} \\h")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = f"({info.get('number', '?')})" if info else f"[{label}]"
            run.append(text)
            fld.append(run)
            paragraph._p.append(fld)
        except Exception as exc:
            self._fail("reference", exc)
        return self

    # ------------------------------------------------------------------
    # ImÃ¡genes y datos
    # ------------------------------------------------------------------
    def image(
        self,
        image: Any,
        *,
        width: Optional[float] = None,
        height: Optional[float] = None,
        align: AlignmentValue = "center",
        caption: Optional[str] = None,
        label: Optional[str] = None,
        alt_text: Optional[str] = None,
        caption_position: str = "below",
        caption_label: str = "Figura",
    ) -> "DocBuilder":
        """Inserta una imagen en el documento.
        
        image: Ruta (str), bytes, PIL.Image, numpy.ndarray o matplotlib.Figure.
        width: Ancho en pulgadas (opcional).
        height: Alto en pulgadas (opcional).
        align: AlineaciÃ³n ('left', 'center', 'right').
        """
        # Validaciones de dimensiones
        if width is not None:
            width_f = float(width)
            if width_f <= 0:
                _emit_warning(f"width debe ser positivo ({width_f}). Ignorando parÃ¡metro.")
                width = None
            elif width_f > MAX_IMAGE_DIMENSION:
                _emit_warning(
                    f"width muy grande ({width_f} pulgadas). "
                    f"MÃ¡ximo recomendado: {MAX_IMAGE_DIMENSION} pulgadas."
                )
        
        if height is not None:
            height_f = float(height)
            if height_f <= 0:
                _emit_warning(f"height debe ser positivo ({height_f}). Ignorando parÃ¡metro.")
                height = None
            elif height_f > MAX_IMAGE_DIMENSION:
                _emit_warning(
                    f"height muy grande ({height_f} pulgadas). "
                    f"MÃ¡ximo recomendado: {MAX_IMAGE_DIMENSION} pulgadas."
                )
        
        self._register_requested_label(label)
        resolved_caption_position = self._normalize_caption_position(caption_position, default="below")
        sequence_name = self._normalize_sequence_name(caption_label, default="Figura")

        try:
            number = self._session.next_figure_number()
            final_label = label or f"fig:{number}"
            self._session.register_label(final_label, "figure", number=number)
            if caption and resolved_caption_position == "above":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=number,
                    label=final_label,
                    api_name="image",
                )

            data = image_to_bytes(image)
            paragraph = self._new_paragraph(
                api_name="image",
                element_kind="image",
                text_preview=caption or label or "image",
            )
            run = paragraph.add_run()
            stream = io.BytesIO(data)
            if width is not None:
                run.add_picture(stream, width=Inches(float(width)))
            elif height is not None:
                run.add_picture(stream, height=Inches(float(height)))
            else:
                run.add_picture(stream)
            _set_run_picture_alt_text(run, alt_text or caption or label)
            alignment = _normalize_alignment(align)
            if alignment is not None:
                paragraph.alignment = alignment
            if caption and resolved_caption_position == "below":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=number,
                    label=final_label,
                    api_name="image",
                )
        except ImageConversionError as exc:
            self._fail("image", exc)
        except Exception as exc:  # pragma: no cover - errores genÃ©ricos
            self._fail("image", exc)
        return self

    def figure(
        self,
        figure: Any,
        *,
        caption: Optional[str] = None,
        label: Optional[str] = None,
        width: Optional[float] = None,
        height: Optional[float] = None,
        alt_text: Optional[str] = None,
        dpi: int = 200,
        caption_position: str = "below",
        caption_label: str = "Figura",
    ) -> "DocBuilder":
        """Inserta una figura de Matplotlib con caption opcional.
        
        figure: matplotlib.figure.Figure a insertar.
        caption: Texto del caption (opcional).
        label: Etiqueta para referencias cruzadas (opcional).
        width: Ancho en pulgadas.
        height: Alto en pulgadas.
        dpi: ResoluciÃ³n (default: 200).
        """
        # Validaciones
        if dpi <= 0:
            _emit_warning(f"dpi debe ser positivo ({dpi}). Usando 200.")
            dpi = 200
        elif dpi > 600:
            _emit_warning(
                f"dpi muy alto ({dpi}) puede generar imÃ¡genes muy grandes. "
                f"MÃ¡ximo recomendado: 300-400 para documentos."
            )
        
        if width is not None and float(width) > MAX_IMAGE_DIMENSION:
            _emit_warning(f"width muy grande ({width} pulgadas).")
        
        if height is not None and float(height) > MAX_IMAGE_DIMENSION:
            _emit_warning(f"height muy grande ({height} pulgadas).")
        
        self._register_requested_label(label)
        resolved_caption_position = self._normalize_caption_position(caption_position, default="below")
        sequence_name = self._normalize_sequence_name(caption_label, default="Figura")
        
        try:
            number_value = self._session.next_figure_number()
            final_label = label or f"fig:{number_value}"
            self._session.register_label(final_label, "figure", number=number_value)
            if caption and resolved_caption_position == "above":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=number_value,
                    label=final_label,
                    api_name="figure",
                )

            data = image_to_bytes(figure, dpi=dpi)
            paragraph = self._new_paragraph(
                api_name="figure",
                element_kind="figure",
                text_preview=caption or label or "figure",
            )
            run = paragraph.add_run()
            stream = io.BytesIO(data)
            if width is not None:
                run.add_picture(stream, width=Inches(float(width)))
            elif height is not None:
                run.add_picture(stream, height=Inches(float(height)))
            else:
                run.add_picture(stream)
            _set_run_picture_alt_text(run, alt_text or caption or label)
            if caption and resolved_caption_position == "below":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=number_value,
                    label=final_label,
                    api_name="figure",
                )
        except ImageConversionError as exc:
            self._fail("figure", exc)
        except Exception as exc:  # pragma: no cover
            self._fail("figure", exc)
        return self

    def caption(
        self,
        text: str,
        *,
        label: Optional[str] = None,
        number: bool = False,
        caption_label: str = "Figura",
    ) -> "DocBuilder":
        """Inserta una leyenda/caption.
        
        text: Texto del caption.
        label: Etiqueta para referencias cruzadas (opcional).
        """
        if not text:
            _emit_warning("caption() llamado con texto vacÃ­o.")
        
        self._register_requested_label(label)
        
        try:
            if number:
                sequence_name = self._normalize_sequence_name(caption_label, default="Figura")
                number_value = self._session.next_sequence_number(sequence_name)
                self._insert_caption_paragraph(
                    str(text),
                    sequence_name=sequence_name,
                    number_value=number_value,
                    label=label,
                    api_name="caption",
                )
                if label:
                    self._session.register_label(label, "caption", number=number_value)
                return self

            self._insert_caption_paragraph(str(text), label=label, api_name="caption")
            if label:
                self._session.register_label(label, "caption")
        except Exception as exc:
            self._fail("caption", exc)
        return self

    def table(
        self,
        data: Sequence[Sequence[Any]],
        *,
        headers: Optional[Sequence[Any]] = None,
        style: Optional[str] = None,
        autofit: bool = True,
        caption: Optional[str] = None,
        label: Optional[str] = None,
        caption_position: str = "above",
        caption_label: str = "Tabla",
        repeat_header_row: bool = True,
        column_widths: Optional[Sequence[float]] = None,
        cell_padding_twips: Optional[int] = None,
        vertical_align: Optional[str] = "center",
    ) -> "DocBuilder":
        """Inserta una tabla en el documento.
        
        data: Secuencia de filas, cada fila es una secuencia de valores.
        headers: Encabezados opcionales de la tabla.
        style: Estilo de la tabla (default: "Table Grid").
        autofit: Ajustar automÃ¡ticamente el ancho de columnas.
        """
        # Validaciones
        rows = list(data)
        row_count = len(rows)
        
        if row_count == 0:
            _emit_warning("table() llamado con data vacÃ­o. No se insertÃ³ ninguna tabla.")
            return self
        
        cols = max((len(r) for r in rows), default=0)
        if headers:
            cols = max(cols, len(headers))
        
        if cols == 0:
            _emit_warning("table() llamado con filas vacÃ­as. No se insertÃ³ ninguna tabla.")
            return self
        
        if row_count > MAX_TABLE_ROWS:
            _emit_warning(
                f"Tabla con {row_count} filas (mÃ¡ximo recomendado: {MAX_TABLE_ROWS}). "
                f"Tablas muy grandes pueden afectar el rendimiento. "
                f"Considere usar max_rows con dataframe() o dividir los datos."
            )
        
        if cols > MAX_TABLE_COLS:
            _emit_warning(
                f"Tabla con {cols} columnas (mÃ¡ximo recomendado: {MAX_TABLE_COLS}). "
                f"Tablas muy anchas pueden no visualizarse correctamente."
            )

        self._register_requested_label(label)
        resolved_caption_position = self._normalize_caption_position(caption_position, default="above")
        sequence_name = self._normalize_sequence_name(caption_label, default="Tabla")
        caption_number = None
        final_label = None
        if caption:
            caption_number = self._session.next_table_number()
            final_label = label or f"tbl:{caption_number}"
            self._session.register_label(final_label, "table", number=caption_number)
            if resolved_caption_position == "above":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=caption_number,
                    label=final_label,
                    api_name="table",
                )
        
        try:
            self._prepare_next_provenance(
                api_name="table",
                element_kind="table",
                text_preview=caption or label or (headers[0] if headers else None),
            )
            table = self._handle.doc.add_table(rows=(row_count + (1 if headers else 0)), cols=max(1, cols))
            target_style = style or self.resolve_style_slot("table_default") or "Table Grid"
            try:
                table.style = target_style
            except KeyError:
                _emit_warning(f"Estilo de tabla '{target_style}' no encontrado. Usando estilo por defecto.")
            idx = 0
            if headers:
                for j, value in enumerate(headers):
                    table.cell(0, j).text = str(value)
                idx = 1
            if repeat_header_row:
                header_index = 0
                if header_index < len(table.rows):
                    _mark_row_as_repeating_header(table.rows[header_index])
            for i, row in enumerate(rows):
                for j in range(cols):
                    cell_value = row[j] if j < len(row) else ""
                    table.cell(idx + i, j).text = str(cell_value)
            table.autofit = bool(autofit)
            if column_widths:
                for row in table.rows:
                    for j, width_value in enumerate(column_widths[:cols]):
                        try:
                            row.cells[j].width = Inches(float(width_value))
                        except Exception:
                            continue
            if cell_padding_twips is not None or vertical_align:
                for row in table.rows:
                    for cell in row.cells:
                        if cell_padding_twips is not None:
                            _set_cell_padding(cell, int(cell_padding_twips))
                        _set_cell_vertical_alignment(cell, vertical_align)

            runtime_defaults = _resolve_table_runtime_defaults(
                self._session.get_template_table_style_defaults(),
                getattr(table, "style", None),
                target_style,
            )
            _apply_table_runtime_defaults(table, runtime_defaults)
            if caption and resolved_caption_position == "below":
                self._insert_caption_paragraph(
                    str(caption),
                    sequence_name=sequence_name,
                    number_value=caption_number,
                    label=final_label,
                    api_name="table",
                )
        except Exception as exc:
            self._fail("table", exc)
        return self

    def dataframe(
        self,
        df: Any,
        *,
        style: Optional[str] = None,
        index: bool = False,
        number_format: Optional[Dict[str, str]] = None,
        max_rows: Optional[int] = None,
        caption: Optional[str] = None,
        label: Optional[str] = None,
        caption_position: str = "above",
        caption_label: str = "Tabla",
        repeat_header_row: bool = True,
        column_widths: Optional[Sequence[float]] = None,
        cell_padding_twips: Optional[int] = None,
        vertical_align: Optional[str] = "center",
    ) -> "DocBuilder":
        """Convierte un DataFrame de Pandas a tabla DOCX.
        
        df: pandas.DataFrame a convertir.
        style: Estilo de la tabla.
        index: Incluir el Ã­ndice del DataFrame como columna.
        number_format: Dict con formato por columna (ej: {'col': '.2f'}).
        max_rows: Limitar el nÃºmero de filas mostradas.
        """
        try:
            import pandas as pd  # type: ignore
        except ImportError:
            self._fail("dataframe", ImportError(
                "pandas no estÃ¡ instalado. InstÃ¡lalo con: pip install pandas"
            ))
            return self

        if not isinstance(df, pd.DataFrame):
            self._fail("dataframe", TypeError(
                f"Se esperaba un pandas.DataFrame, se recibiÃ³ {type(df).__name__}"
            ))
            return self
        
        # Validaciones del DataFrame
        if df.empty:
            _emit_warning("dataframe() llamado con DataFrame vacÃ­o. No se insertÃ³ ninguna tabla.")
            return self
        
        original_rows = len(df)
        original_cols = len(df.columns)
        
        if original_rows > MAX_TABLE_ROWS and max_rows is None:
            _emit_warning(
                f"DataFrame con {original_rows} filas (mÃ¡ximo recomendado: {MAX_TABLE_ROWS}). "
                f"Considere usar max_rows para limitar las filas mostradas."
            )
        
        if original_cols > MAX_TABLE_COLS:
            _emit_warning(
                f"DataFrame con {original_cols} columnas (mÃ¡ximo recomendado: {MAX_TABLE_COLS}). "
                f"Tablas muy anchas pueden no visualizarse correctamente."
            )
        
        try:
            prepared = df.copy()
            if max_rows is not None:
                max_rows_int = int(max_rows)
                if max_rows_int <= 0:
                    _emit_warning("max_rows debe ser positivo. Usando todas las filas.")
                else:
                    prepared = prepared.head(max_rows_int)
                    if original_rows > max_rows_int:
                        self._handle.log_event(
                            "info",
                            f"DataFrame truncado de {original_rows} a {max_rows_int} filas"
                        )
            
            rows: List[List[Any]] = []
            cols = list(prepared.columns)
            if index:
                cols = [prepared.index.name or "index"] + cols
            for idx_value, row in prepared.iterrows():
                values = [row[col] for col in prepared.columns]
                if index:
                    values = [idx_value] + values
                rendered: List[str] = []
                for col_name, value in zip(cols, values):
                    if (
                        number_format
                        and col_name in number_format
                        and isinstance(value, (int, float))
                    ):
                        try:
                            rendered.append(format(value, number_format[col_name]))
                        except ValueError:
                            _emit_warning(f"Formato invÃ¡lido '{number_format[col_name]}' para columna '{col_name}'")
                            rendered.append(str(value))
                    else:
                        rendered.append(str(value))
                rows.append(rendered)
            headers = [str(c) for c in cols]
            self.table(
                rows,
                headers=headers,
                style=style,
                caption=caption,
                label=label,
                caption_position=caption_position,
                caption_label=caption_label,
                repeat_header_row=repeat_header_row,
                column_widths=column_widths,
                cell_padding_twips=cell_padding_twips,
                vertical_align=vertical_align,
            )
        except Exception as exc:
            self._fail("dataframe", exc)
        return self

    # ------------------------------------------------------------------
    # Estructura
    # ------------------------------------------------------------------
    def section(
        self,
        *,
        orientation: str = "portrait",
        page_size: Optional[Tuple[float, float]] = None,
        margins: Optional[Dict[str, float]] = None,
    ) -> "DocBuilder":
        """Crea una nueva secciÃ³n con configuraciÃ³n de pÃ¡gina.
        
        âš ï¸ ADVERTENCIA: Este mÃ©todo tiene limitaciones en modo notebook.
        Los nodos sectPr creados por add_section() NO se registran en el
        sistema de tracking, lo que puede causar problemas de rendimiento
        o loops infinitos al reconstruir el documento.
        
        RECOMENDACIÃ“N: Use page_break() para separar contenido en notebooks.
        Reserve section() para scripts standalone o uso Ãºnico por documento.
        """
        # ValidaciÃ³n: detectar uso excesivo en notebooks
        self._section_count += 1
        if self._section_count > MAX_SECTIONS_PER_BLOCK:
            _emit_warning(
                f"MÃºltiples llamadas a section() en el mismo bloque ({self._section_count}). "
                f"Esto puede causar problemas de rendimiento o loops en notebooks. "
                f"Considere usar page_break() en su lugar."
            )
            self._handle.log_event(
                "warning", 
                f"section() llamado {self._section_count} veces en bloque '{self._handle.block_id}'",
                section_count=self._section_count
            )
        
        # Advertencia general sobre uso en notebooks
        if self._section_count == 1:
            _emit_warning(
                "section() puede causar problemas en modo notebook. "
                "Los cambios de secciÃ³n NO se preservan correctamente al re-ejecutar celdas. "
                "Use page_break() para separar contenido de forma segura."
            )
        
        try:
            section = self._handle.doc.add_section()
            try:
                from docx.enum.section import WD_ORIENT  # type: ignore

                section.orientation = (
                    WD_ORIENT.LANDSCAPE if str(orientation).lower().startswith("land") else WD_ORIENT.PORTRAIT
                )
            except Exception:
                pass
            if page_size:
                # Validar dimensiones razonables
                w, h = float(page_size[0]), float(page_size[1])
                if w > MAX_IMAGE_DIMENSION or h > MAX_IMAGE_DIMENSION:
                    _emit_warning(f"TamaÃ±o de pÃ¡gina muy grande: {w}x{h} pulgadas")
                section.page_width = Inches(w)
                section.page_height = Inches(h)
            if margins:
                section.top_margin = Inches(float(margins.get("top", section.top_margin.inches)))
                section.bottom_margin = Inches(float(margins.get("bottom", section.bottom_margin.inches)))
                section.left_margin = Inches(float(margins.get("left", section.left_margin.inches)))
                section.right_margin = Inches(float(margins.get("right", section.right_margin.inches)))
        except Exception as exc:
            self._fail("section", exc)
        return self

    def table_of_contents(self, *, depth: int = 3, hyperlinks: bool = True) -> "DocBuilder":
        """Inserta una tabla de contenidos.
        
        depth: Niveles de encabezados a incluir (1-9).
        hyperlinks: Incluir enlaces clickeables (default: True).
        
        Nota: La TOC debe actualizarse manualmente en Word (Ctrl+A, F9).
        """
        # Validaciones
        try:
            depth_int = int(depth)
            if depth_int < 1 or depth_int > 9:
                _emit_warning(f"depth debe estar entre 1 y 9 (recibido: {depth_int}). Ajustando.")
                depth_int = max(1, min(9, depth_int))
        except (ValueError, TypeError):
            _emit_warning(f"depth invÃ¡lido ({depth}). Usando 3.")
            depth_int = 3
        
        paragraph = self._new_paragraph()
        try:
            field = OxmlElement("w:fldSimple")
            instr = f"TOC \\o \"1-{depth_int}\""
            if hyperlinks:
                instr += " \\h"
            field.set(qn("w:instr"), instr)
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "Tabla de contenidos (actualizar en Word)"
            run.append(text)
            field.append(run)
            paragraph._p.append(field)
        except Exception as exc:
            self._fail("toc", exc)
        return self

    # ------------------------------------------------------------------
    # Utilidades avanzadas
    # ------------------------------------------------------------------
    def snapshot(self) -> List[str]:
        items = self._session._cell_items_map()  # acceso controlado
        cid = self._handle.block_id
        elements = items.get(cid, [])
        return [f"{idx}: {getattr(el, '__class__', type(el)).__name__}" for idx, el in enumerate(elements)]

    def log(self, kind: str, message: str, **data: Any) -> "DocBuilder":
        self._handle.log_event(kind, message, **data)
        return self

    def header(self, *, text: Optional[str] = None, image: Any = None) -> "DocBuilder":
        """Configura el encabezado de la secciÃ³n actual.
        
        Nota: Llamadas mÃºltiples sobrescriben el contenido anterior.
        El header se aplica a la Ãºltima secciÃ³n del documento.
        """
        self._header_count += 1
        if self._header_count > MAX_HEADER_FOOTER_CALLS:
            _emit_warning(
                f"header() llamado {self._header_count} veces en el mismo bloque. "
                f"Cada llamada sobrescribe el header anterior. "
                f"Considere consolidar el contenido en una sola llamada."
            )
        
        # Validar que al menos un parÃ¡metro fue proporcionado
        if text is None and image is None:
            _emit_warning("header() llamado sin text ni image. No se realizÃ³ ningÃºn cambio.")
            return self
        
        try:
            section = self._handle.doc.sections[-1]
            header = section.header
            header.is_linked_to_previous = False
            if text is not None:
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.text = str(text)
            if image is not None:
                data = image_to_bytes(image)
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(io.BytesIO(data))
        except Exception as exc:
            self._fail("header", exc)
        return self

    def footer(self, *, text: Optional[str] = None) -> "DocBuilder":
        """Configura el pie de pÃ¡gina de la secciÃ³n actual.
        
        Nota: Llamadas mÃºltiples sobrescriben el contenido anterior.
        El footer se aplica a la Ãºltima secciÃ³n del documento.
        """
        self._footer_count += 1
        if self._footer_count > MAX_HEADER_FOOTER_CALLS:
            _emit_warning(
                f"footer() llamado {self._footer_count} veces en el mismo bloque. "
                f"Cada llamada sobrescribe el footer anterior. "
                f"Considere consolidar el contenido en una sola llamada."
            )
        
        # Validar que el parÃ¡metro fue proporcionado
        if text is None:
            _emit_warning("footer() llamado sin text. No se realizÃ³ ningÃºn cambio.")
            return self
        
        try:
            section = self._handle.doc.sections[-1]
            footer = section.footer
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = str(text)
        except Exception as exc:
            self._fail("footer", exc)
        return self
