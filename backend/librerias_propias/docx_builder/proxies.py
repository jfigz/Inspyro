from __future__ import annotations

from typing import Any, Iterable, TYPE_CHECKING

from docx.document import Document as DocumentObject  # type: ignore
from docx.table import Table, _Cell, _Row  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.text.run import Run  # type: ignore

try:
    from docx.text.parfmt import ParagraphFormat  # type: ignore
except Exception:  # pragma: no cover - compat
    ParagraphFormat = None  # type: ignore

try:
    from docx.text.font import Font  # type: ignore
except Exception:  # pragma: no cover - compat
    Font = None  # type: ignore

if TYPE_CHECKING:  # pragma: no cover
    from .session import CellHandle, DocxSession


_INTERNAL_ATTRS = {
    "_obj",
    "_session",
    "_handle",
    "_block_id",
    "_root_element",
}


def unwrap_proxy(value: Any) -> Any:
    if hasattr(value, "_unwrap_raw"):
        return value._unwrap_raw()
    if isinstance(value, list):
        return [unwrap_proxy(item) for item in value]
    if isinstance(value, tuple):
        return tuple(unwrap_proxy(item) for item in value)
    return value


def _is_xml_node(value: Any) -> bool:
    return hasattr(value, "tag") and hasattr(value, "getparent")


def _prefer_defined(primary: Any, fallback: Any) -> Any:
    return primary if primary is not None else fallback


class _BaseProxy:
    def __init__(self, obj: Any, session: "DocxSession", handle: "CellHandle", root_element: Any) -> None:
        object.__setattr__(self, "_obj", obj)
        object.__setattr__(self, "_session", session)
        object.__setattr__(self, "_handle", handle)
        object.__setattr__(self, "_block_id", handle.block_id)
        object.__setattr__(self, "_root_element", root_element)

    def _unwrap_raw(self) -> Any:
        return self._obj

    def _wrap(self, value: Any, *, root_element: Any | None = None) -> Any:
        return wrap_docx_proxy(
            value,
            session=self._session,
            handle=self._handle,
            root_element=self._root_element if root_element is None else root_element,
        )

    def _mark_mutation(
        self,
        *,
        target: Any = None,
        api_name: str,
        element_kind: str,
        text_preview: str | None = None,
        replace: bool = True,
    ) -> None:
        self._session.record_visible_mutation(
            self._block_id,
            self._root_element,
            target,
            api_name=api_name,
            element_kind=element_kind,
            text_preview=text_preview,
            replace=replace,
        )

    def __repr__(self) -> str:  # pragma: no cover - debug helper
        return repr(self._obj)


class XmlNodeProxy(_BaseProxy):
    _MUTATING_METHODS = {"append", "insert", "remove", "set"}

    def __iter__(self):
        for child in list(self._obj):
            yield self._wrap(child)

    def __getattr__(self, name: str) -> Any:
        attr = getattr(self._obj, name)
        if callable(attr):
            def wrapped(*args, **kwargs):
                raw_args = [unwrap_proxy(arg) for arg in args]
                raw_kwargs = {key: unwrap_proxy(value) for key, value in kwargs.items()}
                if name == "remove":
                    result = attr(*raw_args, **raw_kwargs)
                    self._mark_mutation(
                        target=self._obj,
                        api_name="document.xml.remove",
                        element_kind="xml",
                        replace=True,
                    )
                    return self._wrap(result)
                result = attr(*raw_args, **raw_kwargs)
                if name in self._MUTATING_METHODS:
                    target = raw_args[-1] if raw_args and name in {"append", "insert"} else self._obj
                    self._mark_mutation(
                        target=target,
                        api_name=f"document.xml.{name}",
                        element_kind="xml",
                        replace=(name in {"set"}),
                    )
                return self._wrap(result)
            return wrapped
        return self._wrap(attr)


class ParagraphFormatProxy(_BaseProxy):
    def __getattr__(self, name: str) -> Any:
        return getattr(self._obj, name)

    def __setattr__(self, name: str, value: Any) -> None:
        if name in _INTERNAL_ATTRS:
            object.__setattr__(self, name, value)
            return
        setattr(self._obj, name, unwrap_proxy(value))
        paragraph = _prefer_defined(
            getattr(self._obj, "_element", None),
            getattr(self._obj, "_pPr", None),
        )
        self._mark_mutation(
            target=paragraph,
            api_name=f"document.paragraph_format.{name}",
            element_kind="paragraph",
            replace=True,
        )


class FontProxy(_BaseProxy):
    def __getattr__(self, name: str) -> Any:
        return getattr(self._obj, name)

    def __setattr__(self, name: str, value: Any) -> None:
        if name in _INTERNAL_ATTRS:
            object.__setattr__(self, name, value)
            return
        setattr(self._obj, name, unwrap_proxy(value))
        run = _prefer_defined(
            getattr(self._obj, "_element", None),
            getattr(self._obj, "_rPr", None),
        )
        self._mark_mutation(
            target=run,
            api_name=f"document.run.font.{name}",
            element_kind="run",
            replace=True,
        )


class RunProxy(_BaseProxy):
    @property
    def text(self) -> str:
        return self._obj.text

    @text.setter
    def text(self, value: Any) -> None:
        self._obj.text = "" if value is None else str(value)
        self._mark_mutation(
            target=self._obj._r,
            api_name="document.run.text",
            element_kind="run",
            text_preview=self._obj.text,
            replace=True,
        )

    @property
    def bold(self) -> Any:
        return self._obj.bold

    @bold.setter
    def bold(self, value: Any) -> None:
        self._obj.bold = value
        self._mark_mutation(target=self._obj._r, api_name="document.run.bold", element_kind="run", replace=True)

    @property
    def italic(self) -> Any:
        return self._obj.italic

    @italic.setter
    def italic(self, value: Any) -> None:
        self._obj.italic = value
        self._mark_mutation(target=self._obj._r, api_name="document.run.italic", element_kind="run", replace=True)

    @property
    def underline(self) -> Any:
        return self._obj.underline

    @underline.setter
    def underline(self, value: Any) -> None:
        self._obj.underline = value
        self._mark_mutation(target=self._obj._r, api_name="document.run.underline", element_kind="run", replace=True)

    @property
    def font(self) -> Any:
        if Font is None:
            return self._obj.font
        return FontProxy(self._obj.font, self._session, self._handle, self._root_element)

    @property
    def _r(self) -> Any:
        return XmlNodeProxy(self._obj._r, self._session, self._handle, self._root_element)

    @property
    def _element(self) -> Any:
        return XmlNodeProxy(self._obj._element, self._session, self._handle, self._root_element)

    def add_break(self, *args, **kwargs):
        result = self._obj.add_break(*args, **kwargs)
        self._mark_mutation(target=self._obj._r, api_name="document.run.add_break", element_kind="run", replace=True)
        return result

    def add_picture(self, *args, **kwargs):
        result = self._obj.add_picture(*args, **kwargs)
        self._mark_mutation(target=self._obj._r, api_name="document.run.add_picture", element_kind="picture", replace=True)
        return result

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name))


class ParagraphProxy(_BaseProxy):
    def add_run(self, text: Any = "", *args, **kwargs) -> RunProxy:
        result = self._obj.add_run("" if text is None else str(text), *args, **kwargs)
        self._mark_mutation(
            target=result._r,
            api_name="document.add_run",
            element_kind="run",
            text_preview=result.text,
            replace=False,
        )
        return RunProxy(result, self._session, self._handle, self._root_element)

    def insert_paragraph_before(self, text: Any = "", *args, **kwargs):
        result = self._obj.insert_paragraph_before("" if text is None else str(text), *args, **kwargs)
        provenance = {
            "api_name": "document.insert_paragraph_before",
            "element_kind": "paragraph",
            "text_preview": None if text is None else str(text),
            "precision": "exact",
        }
        if self._root_element is self._obj:
            self._session.register_element_before(
                self._block_id,
                self._obj,
                result,
                provenance=provenance,
            )
        else:
            self._mark_mutation(
                target=result._p,
                api_name="document.insert_paragraph_before",
                element_kind="paragraph",
                text_preview=None if text is None else str(text),
                replace=False,
            )
        return ParagraphProxy(result, self._session, self._handle, result if self._root_element is self._obj else self._root_element)

    @property
    def text(self) -> str:
        return self._obj.text

    @text.setter
    def text(self, value: Any) -> None:
        self._obj.text = "" if value is None else str(value)
        self._mark_mutation(
            target=self._obj._p,
            api_name="document.paragraph.text",
            element_kind="paragraph",
            text_preview=self._obj.text,
            replace=True,
        )

    @property
    def style(self) -> Any:
        return self._obj.style

    @style.setter
    def style(self, value: Any) -> None:
        self._obj.style = unwrap_proxy(value)
        self._mark_mutation(target=self._obj._p, api_name="document.paragraph.style", element_kind="paragraph", replace=True)

    @property
    def alignment(self) -> Any:
        return self._obj.alignment

    @alignment.setter
    def alignment(self, value: Any) -> None:
        self._obj.alignment = unwrap_proxy(value)
        self._mark_mutation(target=self._obj._p, api_name="document.paragraph.alignment", element_kind="paragraph", replace=True)

    @property
    def paragraph_format(self) -> Any:
        if ParagraphFormat is None:
            return self._obj.paragraph_format
        return ParagraphFormatProxy(self._obj.paragraph_format, self._session, self._handle, self._root_element)

    @property
    def _p(self) -> Any:
        return XmlNodeProxy(self._obj._p, self._session, self._handle, self._root_element)

    @property
    def _element(self) -> Any:
        return XmlNodeProxy(self._obj._element, self._session, self._handle, self._root_element)

    @property
    def runs(self) -> list[RunProxy]:
        return [RunProxy(run, self._session, self._handle, self._root_element) for run in self._obj.runs]

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name))


class CellProxy(_BaseProxy):
    def add_paragraph(self, text: Any = "", *args, **kwargs) -> ParagraphProxy:
        result = self._obj.add_paragraph("" if text is None else str(text), *args, **kwargs)
        self._mark_mutation(
            target=result._p,
            api_name="document.cell.add_paragraph",
            element_kind="paragraph",
            text_preview=result.text,
            replace=False,
        )
        return ParagraphProxy(result, self._session, self._handle, self._root_element)

    @property
    def text(self) -> str:
        return self._obj.text

    @text.setter
    def text(self, value: Any) -> None:
        self._obj.text = "" if value is None else str(value)
        self._mark_mutation(
            target=self._obj._tc,
            api_name="document.cell.text",
            element_kind="cell",
            text_preview=self._obj.text,
            replace=True,
        )

    @property
    def paragraphs(self) -> list[ParagraphProxy]:
        return [ParagraphProxy(p, self._session, self._handle, self._root_element) for p in self._obj.paragraphs]

    @property
    def _tc(self) -> Any:
        return XmlNodeProxy(self._obj._tc, self._session, self._handle, self._root_element)

    @property
    def _element(self) -> Any:
        return XmlNodeProxy(self._obj._element, self._session, self._handle, self._root_element)

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name), root_element=self._root_element)


class RowProxy(_BaseProxy):
    @property
    def cells(self) -> list[CellProxy]:
        return [CellProxy(cell, self._session, self._handle, self._root_element) for cell in self._obj.cells]

    @property
    def _tr(self) -> Any:
        return XmlNodeProxy(self._obj._tr, self._session, self._handle, self._root_element)

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name), root_element=self._root_element)


class TableProxy(_BaseProxy):
    @property
    def rows(self) -> list[RowProxy]:
        return [RowProxy(row, self._session, self._handle, self._root_element) for row in self._obj.rows]

    @property
    def style(self) -> Any:
        return self._obj.style

    @style.setter
    def style(self, value: Any) -> None:
        self._obj.style = unwrap_proxy(value)
        self._mark_mutation(target=self._obj._tbl, api_name="document.table.style", element_kind="table", replace=True)

    @property
    def alignment(self) -> Any:
        return self._obj.alignment

    @alignment.setter
    def alignment(self, value: Any) -> None:
        self._obj.alignment = unwrap_proxy(value)
        self._mark_mutation(target=self._obj._tbl, api_name="document.table.alignment", element_kind="table", replace=True)

    @property
    def _tbl(self) -> Any:
        return XmlNodeProxy(self._obj._tbl, self._session, self._handle, self._root_element)

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name), root_element=self._root_element)


class DocumentProxy(_BaseProxy):
    def add_paragraph(self, text: Any = "", *args, **kwargs) -> ParagraphProxy:
        self._session.prepare_next_provenance(
            api_name="document.add_paragraph",
            element_kind="paragraph",
            text_preview=None if text is None else str(text),
            precision="exact",
        )
        result = self._obj.add_paragraph("" if text is None else str(text), *args, **kwargs)
        return ParagraphProxy(result, self._session, self._handle, result)

    def add_table(self, *args, **kwargs) -> TableProxy:
        self._session.prepare_next_provenance(
            api_name="document.add_table",
            element_kind="table",
            precision="exact",
        )
        result = self._obj.add_table(*args, **kwargs)
        return TableProxy(result, self._session, self._handle, result)

    def add_picture(self, *args, **kwargs):
        self._session.prepare_next_provenance(
            api_name="document.add_picture",
            element_kind="picture",
            precision="exact",
        )
        result = self._obj.add_picture(*args, **kwargs)
        return self._wrap(result)

    def add_heading(self, text: Any = "", *args, **kwargs) -> ParagraphProxy:
        self._session.prepare_next_provenance(
            api_name="document.add_heading",
            element_kind="heading",
            text_preview=None if text is None else str(text),
            precision="exact",
        )
        result = self._obj.add_heading("" if text is None else str(text), *args, **kwargs)
        return ParagraphProxy(result, self._session, self._handle, result)

    def add_page_break(self):
        self._session.prepare_next_provenance(
            api_name="document.add_page_break",
            element_kind="page-break",
            precision="exact",
        )
        return self._obj.add_page_break()

    @property
    def paragraphs(self) -> list[ParagraphProxy]:
        return [ParagraphProxy(p, self._session, self._handle, p) for p in self._obj.paragraphs]

    @property
    def tables(self) -> list[TableProxy]:
        return [TableProxy(table, self._session, self._handle, table) for table in self._obj.tables]

    @property
    def sections(self) -> Any:
        return self._obj.sections

    @property
    def _element(self) -> Any:
        return XmlNodeProxy(self._obj._element, self._session, self._handle, self._obj)

    def __getattr__(self, name: str) -> Any:
        return self._wrap(getattr(self._obj, name))


def wrap_docx_proxy(value: Any, *, session: "DocxSession", handle: "CellHandle", root_element: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, _BaseProxy):
        return value
    if isinstance(value, DocumentObject):
        return DocumentProxy(value, session, handle, root_element)
    if isinstance(value, Paragraph):
        return ParagraphProxy(value, session, handle, root_element)
    if isinstance(value, Run):
        return RunProxy(value, session, handle, root_element)
    if isinstance(value, Table):
        return TableProxy(value, session, handle, root_element)
    if isinstance(value, _Row):
        return RowProxy(value, session, handle, root_element)
    if isinstance(value, _Cell):
        return CellProxy(value, session, handle, root_element)
    if ParagraphFormat is not None and isinstance(value, ParagraphFormat):
        return ParagraphFormatProxy(value, session, handle, root_element)
    if Font is not None and isinstance(value, Font):
        return FontProxy(value, session, handle, root_element)
    if _is_xml_node(value):
        return XmlNodeProxy(value, session, handle, root_element)
    if isinstance(value, list):
        return [wrap_docx_proxy(item, session=session, handle=handle, root_element=root_element) for item in value]
    if isinstance(value, tuple):
        return tuple(wrap_docx_proxy(item, session=session, handle=handle, root_element=root_element) for item in value)
    return value

